import os
import io
import base64
from io import BytesIO
import re
import json
from reportlab.platypus import FrameBreak
import math
import datetime
import time
import random
from reportlab.pdfgen import canvas
import secrets
import requests
import tempfile
import warnings
from reportlab.lib.utils import ImageReader
import hashlib
import shutil
import plotly.subplots as sp
import matplotlib.image as mpimg
from datetime import datetime
from pathlib import Path
from collections import defaultdict, OrderedDict
from concurrent.futures import ThreadPoolExecutor
import numpy as np
import pandas as pd
from scipy import stats
from scipy.stats import jarque_bera, kurtosis, skew
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import matplotlib.ticker as mticker
from matplotlib.path import Path
from matplotlib.patches import PathPatch, Rectangle
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib.lines import Line2D
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import distinctipy
import xgboost as xgb
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler, PolynomialFeatures
from sklearn.model_selection import train_test_split
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.linear_model import LinearRegression, Ridge, Lasso
from sklearn.tree import DecisionTreeRegressor
from sklearn.ensemble import RandomForestRegressor
from sklearn.svm import SVR
import statsmodels.api as sm
from statsmodels.tsa.arima.model import ARIMA
from statsmodels.stats.diagnostic import het_breuschpagan, acorr_ljungbox
from statsmodels.stats.stattools import durbin_watson, omni_normtest
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.tsa.stattools import adfuller
from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
from docx2pdf import convert
from pdf2docx import Converter
import PyPDF2
from pypdf import PdfReader, PdfWriter
import fitz  # PyMuPDF
import img2pdf
from PIL import Image, ImageEnhance
from reportlab.lib import colors
from reportlab.lib.colors import Color, HexColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, letter, legal, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Flowable, Frame, Indenter, Image as ReportLabImage)
from reportlab.graphics import renderPDF
from reportlab.graphics.shapes import Drawing, Line, String, Rect
from reportlab.graphics.charts.lineplots import LinePlot
from reportlab.graphics.charts.linecharts import HorizontalLineChart
from reportlab.graphics.charts.legends import Legend
from reportlab.graphics.widgets.markers import makeMarker
import streamlit as st
import streamlit.components.v1 as components
from streamlit_lottie import st_lottie
from streamlit_option_menu import option_menu
from streamlit_cookies_manager import EncryptedCookieManager
import base64
import io
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from datetime import datetime
from openpyxl.formatting.rule import ColorScaleRule
import itertools
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from datetime import datetime as dt
from reportlab.lib.units import inch
from reportlab.graphics.shapes import Line
from datetime import datetime, timedelta
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import green, red, black
import calendar
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors
import warnings
import statistics
import plotly.express as fx
import plotly.graph_objs as go
import plotly.io as pio
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, PageBreak,Paragraph, Spacer,HRFlowable
from openpyxl import Workbook
def price():
 def calculate_effective_invoice(df, region, month, year):
    df['Date'] = pd.to_datetime(df['Date'])
    month_start = pd.Timestamp(year=year, month=month, day=1)
    prev_month_data = df[(df['Region(District)'] == region) & (df['Date'] < month_start)].sort_values('Date', ascending=False)
    last_available_invoice = None
    if not prev_month_data.empty:
        last_available_invoice = prev_month_data.iloc[0]['Inv.']
    month_data = df[(df['Region(District)'] == region) & (df['Date'].dt.month == month) & (df['Date'].dt.year == year)].copy()
    if month_data.empty and last_available_invoice is not None:
        month_data = pd.DataFrame([{'Date': month_start,'Inv.': last_available_invoice,'Region(District)': region}])
    elif month_data.empty and last_available_invoice is None:
        return None
    month_data = month_data.sort_values('Date')
    last_day = pd.Timestamp(year, month, 1) + pd.offsets.MonthEnd(1)
    days_in_month = last_day.day
    first_period = pd.date_range(start=f"{year}-{month:02d}-01", end=f"{year}-{month:02d}-10")
    middle_period = pd.date_range(start=f"{year}-{month:02d}-11", end=f"{year}-{month:02d}-20")
    last_period = pd.date_range(start=f"{year}-{month:02d}-21", end=f"{year}-{month:02d}-{days_in_month}")
    def calculate_period_invoice(period_dates, data, weight):
        if data[data['Date'] <= period_dates[-1]].empty and last_available_invoice is not None:
            return last_available_invoice * weight
        period_data = data[data['Date'].dt.date.isin(period_dates.date)]
        if period_data.empty:
            prev_data = data[data['Date'] < period_dates[0]]
            if prev_data.empty and last_available_invoice is not None:
                return last_available_invoice * weight
            elif not prev_data.empty:
                return prev_data.iloc[-1]['Inv.'] * weight
            return 0
        invoice_values = []
        if period_data.iloc[0]['Date'].date() > period_dates[0].date():
            prev_data = data[data['Date'] < period_dates[0]]
            initial_invoice = last_available_invoice if prev_data.empty else prev_data.iloc[-1]['Inv.']
            days_until_first_change = (period_data.iloc[0]['Date'].date() - period_dates[0].date()).days
            if days_until_first_change > 0:
                invoice_values.append((initial_invoice, days_until_first_change))
        for idx, row in period_data.iterrows():
            next_change = period_data[period_data['Date'] > row['Date']].iloc[0]['Date'] if not period_data[period_data['Date'] > row['Date']].empty else period_dates[-1]
            days_effective = (min(next_change, period_dates[-1]).date() - row['Date'].date()).days + 1
            invoice_values.append((row['Inv.'], days_effective))
        total_days = sum(days for _, days in invoice_values)
        weighted_invoice = sum(invoice * (days / total_days) for invoice, days in invoice_values)
        return weighted_invoice * weight
    first_period_invoice = calculate_period_invoice(first_period, month_data, 0.20)
    middle_period_invoice = calculate_period_invoice(middle_period, month_data, 0.30)
    last_period_invoice = calculate_period_invoice(last_period, month_data, 0.50)
    effective_invoice = first_period_invoice + middle_period_invoice + last_period_invoice
    return {'effective_invoice': round(effective_invoice, 2),'first_period_invoice': round(first_period_invoice / 0.20, 2) if first_period_invoice != 0 else 0,'middle_period_invoice': round(middle_period_invoice / 0.30, 2) if middle_period_invoice != 0 else 0,'last_period_invoice': round(last_period_invoice / 0.50, 2) if last_period_invoice != 0 else 0,'first_period_contribution': round(first_period_invoice, 2),'middle_period_contribution': round(middle_period_invoice, 2),'last_period_contribution': round(last_period_invoice, 2),'last_available_invoice': last_available_invoice}
 def create_effective_invoice_analysis(story, df, region, current_date, styles):
    normal_style = styles['Normal']
    month_style = ParagraphStyle('MonthStyle',parent=styles['Heading3'],textColor=colors.green,spaceAfter=1)
    metric_style = ParagraphStyle('MetricStyle',parent=styles['Normal'],fontSize=12,textColor=colors.brown,spaceAfter=1)
    current_month = current_date.month
    current_year = current_date.year
    last_month = current_month - 1 if current_month > 1 else 12
    last_month_year = current_year if current_month > 1 else current_year - 1
    current_month_effective = calculate_effective_invoice(df, region, current_month, current_year)
    last_month_effective = calculate_effective_invoice(df, region, last_month, last_month_year)
    story.append(Paragraph("Effective Invoice Analysis:-", month_style))
    table_data = [['Period', 'First 10 days (20%)', 'Middle 10 days (30%)', 'Last 10 days (50%)', 'Total Effective Invoice']]
    if current_month_effective:
        current_row = ['Current Month',f"Rs.{current_month_effective['first_period_invoice']:,.0f}\n(Cont: Rs.{current_month_effective['first_period_contribution']:,.0f})",f"Rs.{current_month_effective['middle_period_invoice']:,.0f}\n(Cont: Rs.{current_month_effective['middle_period_contribution']:,.0f})",f"Rs.{current_month_effective['last_period_invoice']:,.0f}\n(Cont: Rs.{current_month_effective['last_period_contribution']:,.0f})",f"Rs.{current_month_effective['effective_invoice']:,.2f}"]
        table_data.append(current_row)
    if last_month_effective:
        last_row = ['Last Month',f"Rs.{last_month_effective['first_period_invoice']:,.0f}\n(Cont: Rs.{last_month_effective['first_period_contribution']:,.0f})",f"Rs.{last_month_effective['middle_period_invoice']:,.0f}\n(Cont: Rs.{last_month_effective['middle_period_contribution']:,.0f})",f"Rs.{last_month_effective['last_period_invoice']:,.0f}\n(Cont: Rs.{last_month_effective['last_period_contribution']:,.0f})",f"Rs.{last_month_effective['effective_invoice']:,.2f}"]
        table_data.append(last_row)
    if current_month_effective or last_month_effective:
        t = Table(table_data, colWidths=[80, 110, 110, 110, 100])
        t.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),('FONTSIZE', (0, 0), (-1, 0), 9),('FONTSIZE', (0, 1), (-1, -1), 8),('GRID', (0, 0), (-1, -1), 1, colors.black),('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('BACKGROUND', (0, 1), (-1, 1), colors.lightgrey),('LEFTPADDING', (0, 0), (-1, -1), 3),('RIGHTPADDING', (0, 0), (-1, -1), 3),('TOPPADDING', (0, 0), (-1, -1), 3),('BOTTOMPADDING', (0, 0), (-1, -1), 3),]))
        story.append(t)
        story.append(Spacer(1, 6))
 def create_effective_nod_analysis(story, df, region, current_date, styles):
    normal_style = styles['Normal']
    month_style = ParagraphStyle('MonthStyle',parent=styles['Heading3'],textColor=colors.green,spaceAfter=1)
    metric_style = ParagraphStyle('MetricStyle',parent=styles['Normal'],fontSize=12,textColor=colors.brown,spaceAfter=1)
    current_month = current_date.month
    current_year = current_date.year
    last_month = current_month - 1 if current_month > 1 else 12
    last_month_year = current_year if current_month > 1 else current_year - 1
    current_month_effective = calculate_effective_nod(df, region, current_month, current_year)
    last_month_effective = calculate_effective_nod(df, region, last_month, last_month_year)
    story.append(Paragraph("Effective NOD Analysis:-", month_style))
    table_data = [['Period', 'First 10 days (20%)', 'Middle 10 days (30%)', 'Last 10 days (50%)', 'Total Effective NOD']]
    if current_month_effective:
        current_row = ['Current Month',f"Rs.{current_month_effective['first_period_nod']:,.0f}\n(Cont: Rs.{current_month_effective['first_period_contribution']:,.0f})",f"Rs.{current_month_effective['middle_period_nod']:,.0f}\n(Cont: Rs.{current_month_effective['middle_period_contribution']:,.0f})",f"Rs.{current_month_effective['last_period_nod']:,.0f}\n(Cont: Rs.{current_month_effective['last_period_contribution']:,.0f})",f"Rs.{current_month_effective['effective_nod']:,.2f}"]
        table_data.append(current_row)
    if last_month_effective:
        last_row = ['Last Month',f"Rs.{last_month_effective['first_period_nod']:,.0f}\n(Cont: Rs.{last_month_effective['first_period_contribution']:,.0f})",f"Rs.{last_month_effective['middle_period_nod']:,.0f}\n(Cont: Rs.{last_month_effective['middle_period_contribution']:,.0f})",f"Rs.{last_month_effective['last_period_nod']:,.0f}\n(Cont: Rs.{last_month_effective['last_period_contribution']:,.0f})",f"Rs.{last_month_effective['effective_nod']:,.2f}"]
        table_data.append(last_row)
    if current_month_effective or last_month_effective:
        t = Table(table_data, colWidths=[80, 110, 110, 110, 100])
        t.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),('FONTSIZE', (0, 0), (-1, 0), 9),('FONTSIZE', (0, 1), (-1, -1), 8),('GRID', (0, 0), (-1, -1), 1, colors.black),('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('BACKGROUND', (0, 1), (-1, 1), colors.lightgrey),('LEFTPADDING', (0, 0), (-1, -1), 3),('RIGHTPADDING', (0, 0), (-1, -1), 3),('TOPPADDING', (0, 0), (-1, -1), 3),('BOTTOMPADDING', (0, 0), (-1, -1), 3),]))
        story.append(t)
        story.append(Spacer(1, 6))
 def calculate_effective_nod(df, region, month, year):
    df['Date'] = pd.to_datetime(df['Date'])
    month_start = pd.Timestamp(year=year, month=month, day=1)
    prev_month_data = df[(df['Region(District)'] == region) & (df['Date'] < month_start)].sort_values('Date', ascending=False)
    last_available_nod = None
    if not prev_month_data.empty:
        last_available_nod = prev_month_data.iloc[0]['Net']
    month_data = df[(df['Region(District)'] == region) & (df['Date'].dt.month == month) & (df['Date'].dt.year == year)].copy()
    if month_data.empty and last_available_nod is not None:
        month_data = pd.DataFrame([{'Date': month_start,'Net': last_available_nod,'Region(District)': region}])
    elif month_data.empty and last_available_nod is None:
        return None
    month_data = month_data.sort_values('Date')
    last_day = pd.Timestamp(year, month, 1) + pd.offsets.MonthEnd(1)
    days_in_month = last_day.day
    first_period = pd.date_range(start=f"{year}-{month:02d}-01", end=f"{year}-{month:02d}-10")
    middle_period = pd.date_range(start=f"{year}-{month:02d}-11", end=f"{year}-{month:02d}-20")
    last_period = pd.date_range(start=f"{year}-{month:02d}-21", end=f"{year}-{month:02d}-{days_in_month}")
    def calculate_period_nod(period_dates, data, weight):
        if data[data['Date'] <= period_dates[-1]].empty and last_available_nod is not None:
            return last_available_nod * weight
        period_data = data[data['Date'].dt.date.isin(period_dates.date)]
        if period_data.empty:
            prev_data = data[data['Date'] < period_dates[0]]
            if prev_data.empty and last_available_nod is not None:
                return last_available_nod * weight
            elif not prev_data.empty:
                return prev_data.iloc[-1]['Net'] * weight
            return 0
        nod_values = []
        current_period_start = period_dates[0]
        if period_data.iloc[0]['Date'].date() > period_dates[0].date():
            prev_data = data[data['Date'] < period_dates[0]]
            initial_nod = last_available_nod if prev_data.empty else prev_data.iloc[-1]['Net']
            days_until_first_change = (period_data.iloc[0]['Date'].date() - period_dates[0].date()).days
            if days_until_first_change > 0:
                nod_values.append((initial_nod, days_until_first_change))
        for idx, row in period_data.iterrows():
            next_change = period_data[period_data['Date'] > row['Date']].iloc[0]['Date'] if not period_data[period_data['Date'] > row['Date']].empty else period_dates[-1]
            days_effective = (min(next_change, period_dates[-1]).date() - row['Date'].date()).days 
            nod_values.append((row['Net'], days_effective))
        total_days = sum(days for _, days in nod_values)
        weighted_nod = sum(nod * (days / total_days) for nod, days in nod_values)
        return weighted_nod * weight
    first_period_nod = calculate_period_nod(first_period, month_data, 0.20)
    middle_period_nod = calculate_period_nod(middle_period, month_data, 0.30)
    last_period_nod = calculate_period_nod(last_period, month_data, 0.50)
    effective_nod = first_period_nod + middle_period_nod + last_period_nod
    return {'effective_nod': round(effective_nod, 2),'first_period_nod': round(first_period_nod / 0.20, 2) if first_period_nod != 0 else 0,'middle_period_nod': round(middle_period_nod / 0.30, 2) if middle_period_nod != 0 else 0,'last_period_nod': round(last_period_nod / 0.50, 2) if last_period_nod != 0 else 0,'first_period_contribution': round(first_period_nod, 2),'middle_period_contribution': round(middle_period_nod, 2),'last_period_contribution': round(last_period_nod, 2),'last_available_nod': last_available_nod}
 def get_competitive_brands_wsp_data():
    include_competitive_brands = st.checkbox("Include Competitive Brands WSP Data")
    competitive_brands_wsp = {}
    if include_competitive_brands:
        competitive_brands_file = st.file_uploader("Upload Competitive Brands WSP Data File", type=['xlsx'],help="Upload an Excel file with multiple sheets, each representing a different brand's WSP data")
        if competitive_brands_file is not None:
            try:
                xls = pd.ExcelFile(competitive_brands_file)
                required_columns = ['Region(District)', 'D1-3', 'D4-6', 'D7-9', 'D10-12', 'D13-15','D16-18','D19-21','D22-24','D25-27','D28-30','D1-3 J','D4-6 J','D7-8 J']
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(competitive_brands_file, sheet_name=sheet_name)
                    missing_columns = [col for col in required_columns if col not in df.columns]
                    if missing_columns:
                        st.warning(f"Sheet '{sheet_name}' is missing columns: {missing_columns}")
                        continue
                    competitive_brands_wsp[sheet_name] = df
                if not competitive_brands_wsp:
                    st.error("No valid brand sheets found in the uploaded file.")
                    return None
                return competitive_brands_wsp
            except Exception as e:
                st.error(f"Could not read competitive brands WSP file: {e}")
                return None
    return None
 def get_start_data_point(df, reference_date):
    first_day_data = df[(df['Date'].dt.year == reference_date.year) & (df['Date'].dt.month == reference_date.month) & (df['Date'].dt.day == 1)]
    if not first_day_data.empty:
        return first_day_data.iloc[0]
    prev_month = reference_date.replace(day=1) - timedelta(days=1)
    last_data_of_prev_month = df[(df['Date'].dt.year == prev_month.year) & (df['Date'].dt.month == prev_month.month)]
    if not last_data_of_prev_month.empty:
        return last_data_of_prev_month.iloc[-1]
    return None
 def get_start_data_point_current_month(df, reference_date):
    nov_30_data = df[(df['Date'].dt.year == reference_date.year) & (df['Date'].dt.month == 11) & (df['Date'].dt.day == 30)]
    if not nov_30_data.empty:
        return nov_30_data.iloc[-1]
    first_day_data = df[(df['Date'].dt.year == reference_date.year) & (df['Date'].dt.month == 12) & (df['Date'].dt.day == 1)]
    if not first_day_data.empty:
        return first_day_data.iloc[0]
    nov_data = df[(df['Date'].dt.year == reference_date.year) & (df['Date'].dt.month == 11)]
    if not nov_data.empty:
        return nov_data.iloc[-1]
    return None
 def create_wsp_progression(story, wsp_df, region, styles, brand_name=None, is_last_brand=False, company_wsp_df=None):
    if region in ['UK (Dehradun)', 'UK (Haridwar)']:
        return
        
    normal_style = styles['Normal']
    month_style = ParagraphStyle('MonthStyle', parent=styles['Heading3'], textColor=colors.green, spaceAfter=1)
    large_price_style = ParagraphStyle('LargePriceStyle', parent=styles['Normal'], fontSize=14, spaceAfter=1)
    total_change_style = ParagraphStyle('TotalChangeStyle', parent=styles['Normal'], fontSize=12, textColor=colors.brown, alignment=TA_LEFT, spaceAfter=1, fontName='Helvetica-Bold')
    
    if wsp_df is None:
        return
        
    region_wsp = wsp_df[wsp_df['Region(District)'] == region]
    if region_wsp.empty:
        story.append(Paragraph(f"No WSP data available for {region}" + (f" - {brand_name}" if brand_name else ""), normal_style))
        story.append(Spacer(1, 0))
        return
        
    # Define December and January columns
    dec_columns = ['D1-3', 'D4-6', 'D7-9', 'D10-12', 'D13-15', 'D16-18', 'D19-21', 'D22-24', 'D25-27', 'D28-30']
    jan_columns = ['D1-3 J', 'D4-6 J', 'D7-8 J']
    
    # Get the values and handle NaN
    dec_values = region_wsp[dec_columns].values.flatten().tolist()
    jan_values = region_wsp[jan_columns].values.flatten().tolist()
    
    # Check if all values are NaN
    if all(pd.isna(val) for val in dec_values + jan_values):
        story.append(Paragraph(f"No data available for {region}" + (f" - {brand_name}" if brand_name else ""), normal_style))
        story.append(Spacer(1, 0))
        return
    
    dec_labels = ['01-03 Dec', '04-06 Dec', '07-09 Dec', '10-12 Dec', '13-15 Dec', '16-18 Dec', '19-21 Dec', '22-24 Dec', '25-27 Dec', '28-30 Dec']
    jan_labels = ['01-03 Jan', '04-06 Jan', '07-09 Jan']
    
    header_text = f"WSP Progression from December 2024 to January 2025" + \
                  (f" - {brand_name}" if brand_name else "")
    story.append(Paragraph(header_text + ":-", month_style))
    
    # December Progression
    dec_progression_parts = []
    last_valid_value = None
    for i in range(len(dec_values)):
        if pd.isna(dec_values[i]):
            if last_valid_value is not None:
                dec_progression_parts.append(f"{last_valid_value:.0f}")
            else:
                dec_progression_parts.append("No data")
        else:
            dec_progression_parts.append(f"{dec_values[i]:.0f}")
            last_valid_value = dec_values[i]
        
        if i < len(dec_values) - 1 and not (pd.isna(dec_values[i]) or pd.isna(dec_values[i+1])):
            change = float(dec_values[i+1]) - float(dec_values[i])
            if change > 0:
                dec_progression_parts.append(f'<sup><font color="green" size="7">+{change:.0f}</font></sup>→')
            elif change < 0:
                dec_progression_parts.append(f'<sup><font color="red" size="7">{change:.0f}</font></sup>→')
            else:
                dec_progression_parts.append(f'<sup><font size="8">00</font></sup>→')
        elif i < len(dec_values) - 1:
            dec_progression_parts.append("→")
    
    dec_full_progression = " ".join(dec_progression_parts)
    dec_week_progression_text = "- ".join(dec_labels)
    
    story.append(Paragraph("December 2024:", normal_style))
    story.append(Paragraph(dec_full_progression, large_price_style))
    story.append(Paragraph(dec_week_progression_text, normal_style))
    
    # January Progression
    jan_progression_parts = []
    last_valid_value = None
    for i in range(len(jan_values)):
        if pd.isna(jan_values[i]):
            if last_valid_value is not None:
                jan_progression_parts.append(f"{last_valid_value:.0f}")
            else:
                jan_progression_parts.append("No data")
        else:
            jan_progression_parts.append(f"{jan_values[i]:.0f}")
            last_valid_value = jan_values[i]
        
        if i < len(jan_values) - 1 and not (pd.isna(jan_values[i]) or pd.isna(jan_values[i+1])):
            change = float(jan_values[i+1]) - float(jan_values[i])
            if change > 0:
                jan_progression_parts.append(f'<sup><font color="green" size="7">+{change:.0f}</font></sup>→')
            elif change < 0:
                jan_progression_parts.append(f'<sup><font color="red" size="7">{change:.0f}</font></sup>→')
            else:
                jan_progression_parts.append(f'<sup><font size="8">00</font></sup>→')
        elif i < len(jan_values) - 1:
            jan_progression_parts.append("→")
    
    jan_full_progression = " ".join(jan_progression_parts)
    jan_week_progression_text = "- ".join(jan_labels)
    
    story.append(Paragraph("January 2025:", normal_style))
    story.append(Paragraph(jan_full_progression, large_price_style))
    story.append(Paragraph(jan_week_progression_text, normal_style))
    
    # Calculate and display changes only for valid data
    if len(dec_values) > 1 and not all(pd.isna(val) for val in dec_values):
        valid_dec_values = [val for val in dec_values if not pd.isna(val)]
        if len(valid_dec_values) >= 2:
            dec_change = valid_dec_values[-1] - valid_dec_values[0]
            dec_change_text = f"Net Change in WSP{' - ' + brand_name if brand_name else ''} in December: {dec_change:+.0f} Rs."
            story.append(Paragraph(dec_change_text, total_change_style))
    
    # Calculate January change if data is available
    if len(jan_values) > 0 and not all(pd.isna(val) for val in dec_values + jan_values):
        valid_dec_values = [val for val in dec_values if not pd.isna(val)]
        valid_jan_values = [val for val in jan_values if not pd.isna(val)]
        
        if valid_dec_values and valid_jan_values:
            total_jan_change = valid_jan_values[-1] - valid_dec_values[-1]
            jan_change_text = f"Net Change in WSP{' - ' + brand_name if brand_name else ''} in January: {total_jan_change:+.0f} Rs."
            story.append(Paragraph(jan_change_text, total_change_style))
    
    # Calculate total change if both start and end data are available
    valid_all_values = [val for val in dec_values + jan_values if not pd.isna(val)]
    if len(valid_all_values) >= 2:
        total_change = valid_all_values[-1] - valid_all_values[0]
        total_change_text = f"Total Change in WSP{' - ' + brand_name if brand_name else ''} from 1st Dec.: {total_change:+.0f} Rs."
        story.append(Paragraph(total_change_text, total_change_style))
    
    if company_wsp_df is not None and brand_name is not None:
        company_region_wsp = company_wsp_df[company_wsp_df['Region(District)'] == region]
        if not company_region_wsp.empty and not region_wsp.empty:
            company_w1_dec_wsp = company_region_wsp['D1-3'].values[0]
            competitive_w1_dec_wsp = region_wsp['D1-3'].values[0]
            if not pd.isna(company_w1_dec_wsp) and not pd.isna(competitive_w1_dec_wsp):
                wsp_difference = company_w1_dec_wsp - competitive_w1_dec_wsp
                wsp_diff_text = f"The Difference in WSP between JKLC and {brand_name} was {wsp_difference:+.0f} Rs. at the start of Decemeber."
                story.append(Paragraph(wsp_diff_text, total_change_style))
            company_w3j_jan_wsp = company_region_wsp['D7-8 J'].values[0]
            competitive_w3j_jan_wsp = region_wsp['D7-8 J'].values[0]
            if not pd.isna(company_w3j_jan_wsp) and not pd.isna(competitive_w3j_jan_wsp):
                wsp_diff_1 = company_w3j_jan_wsp-competitive_w3j_jan_wsp
                wsp_diff_text_1 = f"The Difference in WSP between JKLC and {brand_name} currently is {wsp_diff_1:+.0f} Rs."
                story.append(Paragraph(wsp_diff_text_1, total_change_style))
            
    story.append(Spacer(1, 0))
    if not is_last_brand:
        story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=colors.black, spaceBefore=1, spaceAfter=1))
 def create_comprehensive_metric_progression(story, region_df, current_date, last_month, metric_column, title, styles, is_secondary_metric=False):
    if is_secondary_metric:
        box_style = ParagraphStyle(
            f'{title}BoxStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.darkgreen,
            borderColor=colors.lightgrey,
            borderWidth=1,
            borderPadding=5,
            backColor=colors.whitesmoke,
            spaceAfter=0  # Reduced from 1 to 0
        )
        normal_style = ParagraphStyle(f'{title}NormalStyle', parent=styles['Normal'], fontSize=10)
        total_change_style = ParagraphStyle(
            f'{title}TotalChangeStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.brown,
            alignment=TA_LEFT,
            spaceAfter=0  # Reduced from 1 to 0
        )
    else:
        month_style = ParagraphStyle('MonthStyle', parent=styles['Heading3'], textColor=colors.green, spaceAfter=1)
        normal_style = styles['Normal']
        large_price_style = ParagraphStyle('LargePriceStyle', parent=styles['Normal'], fontSize=14, spaceAfter=1)
        total_change_style = ParagraphStyle(
            'TotalChangeStyle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.brown,
            alignment=TA_LEFT,
            spaceAfter=0,  # Reduced from 1 to 0
            fontName='Helvetica-Bold'
        )

    start_data_point = get_start_data_point(region_df, last_month)
    if start_data_point is None:
        if not is_secondary_metric:
            story.append(Paragraph("No data available for this period", normal_style))
        return

    progression_df = region_df[(region_df['Date'] >= start_data_point['Date']) & 
                             (region_df['Date'] <= current_date)].copy().sort_values('Date')

    if progression_df.empty:
        if not is_secondary_metric:
            story.append(Paragraph("No data available for this period", normal_style))
        return

    if not is_secondary_metric:
        story.append(Paragraph(f"{title} Progression from {last_month.strftime('%B %Y')} to {current_date.strftime('%B %Y')}:-", month_style))
        
        metric_values = progression_df[metric_column].apply(lambda x: f"{x:.0f}").tolist()
        dates = progression_df['Date'].dt.strftime('%d-%b').tolist()
        metric_progression_parts = []
        
        for i in range(len(metric_values)):
            metric_progression_parts.append(metric_values[i])
            if i < len(metric_values) - 1:
                change = float(metric_values[i+1]) - float(metric_values[i])
                if change > 0:
                    metric_progression_parts.append(f'<sup><font color="green" size="7">+{change:.0f}</font></sup>→')
                elif change < 0:
                    metric_progression_parts.append(f'<sup><font color="red" size="7">{change:.0f}</font></sup>→')
                else:
                    metric_progression_parts.append(f'<sup><font size="8">00</font></sup>→')

        full_progression = " ".join(metric_progression_parts)
        date_progression_text = " ----- ".join(dates)

        progression_table = Table([
            [Paragraph(full_progression, large_price_style)],
            [Paragraph(date_progression_text, normal_style)]
        ], colWidths=[400])
        progression_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 1),   # Reduced padding
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1)  # Reduced padding
        ]))
        story.append(progression_table)

    if len(progression_df[metric_column]) > 1:
        dec_data = progression_df[progression_df['Date'].dt.month == 12]
        jan_data = progression_df[progression_df['Date'].dt.month == 1]
        
        changes_text = []
        
        if not dec_data.empty:
            dec_change = dec_data[metric_column].iloc[-1] - dec_data[metric_column].iloc[0]
            changes_text.append(f"Net Change in {title} for December: {dec_change:+.0f} Rs.")

        if not jan_data.empty:
            dec_last_value = dec_data[metric_column].iloc[-1] if not dec_data.empty else progression_df[metric_column].iloc[0]
            jan_change = jan_data[metric_column].iloc[-1] - dec_last_value
            changes_text.append(f"Net Change in {title} for January: {jan_change:+.0f} Rs.")

        total_change = progression_df[metric_column].iloc[-1] - progression_df[metric_column].iloc[0]
        changes_text.append(f"Total Change in {title} from 1st Dec.: {total_change:+.0f} Rs.")

        if is_secondary_metric:
            box_content = [Paragraph(f"<b>{title} Changes</b>", box_style)]
            for text in changes_text:
                box_content.append(Paragraph(text, total_change_style))
            
            box_table = Table([[content] for content in box_content], 
                            colWidths=[200],
                            style=[
                                ('BOX', (0, 0), (-1, -1), 1, colors.lightgrey),
                                ('BACKGROUND', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                                ('TOPPADDING', (0, 0), (-1, -1), 3),  # Reduced padding
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),  # Reduced padding
                            ])
            story.append(box_table)
        else:
            changes_table = Table([[Paragraph(text, total_change_style)] for text in changes_text],
                                colWidths=[400])
            changes_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 1),  # Reduced padding
                ('BOTTOMPADDING', (0, 0), (-1, -1), 1)  # Reduced padding
            ]))
            story.append(changes_table)

 def save_regional_price_trend_report(df):
    company_wsp_df = get_wsp_data()
    competitive_brands_wsp = get_competitive_brands_wsp_data()
    return generate_regional_price_trend_report(df, company_wsp_df, competitive_brands_wsp)
 def generate_regional_price_trend_report(df, company_wsp_df=None, competitive_brands_wsp=None):
    try:
        region_order = ['GJ (Ahmedabad)', 'GJ (Surat)','RJ(Jaipur)', 'RJ(Udaipur)','HY (Gurgaon)','PB (Bhatinda)','Delhi','CG (Raipur)','ORR (Khorda)', 'ORR (Sambalpur)', 'UP (Gaziabad)','UK (Haridwar)','UK (Dehradun)', 'M.P.(East)[Balaghat]', 'M.P.(West)[Indore]', 'M.H.(East)[Nagpur Urban]']
        required_columns = ['Date', 'Region(District)', 'Inv.', 'Net', 'RD', 'STS']
        for col in required_columns:
            if col not in df.columns:
                raise ValueError(f"Missing required column: {col}")
        
        df['Date'] = pd.to_datetime(df['Date'], format='%d-%b %Y')
        df['region_order'] = df['Region(District)'].map({region: idx for idx, region in enumerate(region_order)})
        df = df.sort_values(['region_order', 'Date'])
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=8, leftMargin=1, topMargin=5, bottomMargin=1)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('TitleStyle', parent=styles['Title'], fontSize=20, textColor=colors.darkblue, alignment=TA_CENTER, spaceAfter=1)
        region_style = ParagraphStyle('RegionStyle', parent=styles['Heading2'], textColor=colors.blue, spaceAfter=1, fontSize=12)
        
        story = []
        story.append(Paragraph("Regional Price Trend Analysis Report", title_style))
        story.append(Paragraph("Comprehensive Price Movement Insights", ParagraphStyle('SubtitleStyle', parent=styles['Normal'], fontSize=12, textColor=colors.red, alignment=TA_CENTER, spaceAfter=1)))
        story.append(Spacer(1, 0))
        
        current_date = datetime.now()
        last_month = current_date.replace(day=1) - timedelta(days=1)
        regions = [region for region in region_order if region in df['Region(District)'].unique()]
        
        for i, region in enumerate(regions):
            region_story = []
            region_df = df[df['Region(District)'] == region].copy()
            
            # Add separator line between regions (except for the first region)
            if i > 0:
                region_story.append(HRFlowable(width="100%", thickness=2, lineCap='round', color=colors.grey, spaceBefore=10, spaceAfter=10))
            
            # Region header
            region_story.append(Paragraph(f"{region}", region_style))
            region_story.append(Spacer(1, 1))
            
            # Create main metrics table (Invoice and NOD progression)
            metrics_data = []
            
            # Left column: Invoice and NOD progression
            left_column = []
            create_comprehensive_metric_progression(left_column, region_df, current_date, last_month, 'Inv.', 'Invoice Price', styles)
            create_comprehensive_metric_progression(left_column, region_df, current_date, last_month, 'Net', 'NOD', styles)
            
            # Right column: RD and STS boxes
            right_column = []
            create_comprehensive_metric_progression(right_column, region_df, current_date, last_month, 'RD', 'RD', styles, is_secondary_metric=True)
            create_comprehensive_metric_progression(right_column, region_df, current_date, last_month, 'STS', 'STS', styles, is_secondary_metric=True)
            
            # Create a table for the two-column layout
            layout_table = Table([[left_column, right_column]], colWidths=[400, 200])
            layout_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            region_story.append(layout_table)
            
            # WSP Progression
            brand_count = 1 if company_wsp_df is not None and not company_wsp_df.empty else 0
            if competitive_brands_wsp:
                brand_count += len(competitive_brands_wsp)
            is_last_brand = (brand_count == 1)
            
            create_wsp_progression(region_story, company_wsp_df, region, styles, is_last_brand=is_last_brand, company_wsp_df=company_wsp_df)
            
            if competitive_brands_wsp:
                brand_names = list(competitive_brands_wsp.keys())
                for i, (brand, brand_wsp_df) in enumerate(competitive_brands_wsp.items()):
                    is_last_brand = (i == len(brand_names) - 1)
                    create_wsp_progression(region_story, brand_wsp_df, region, styles, brand_name=brand, is_last_brand=is_last_brand, company_wsp_df=company_wsp_df)
            
            story.append(KeepTogether(region_story))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    except Exception as e:
        print(f"Error generating report: {e}")
        raise
 def get_wsp_data():
    include_wsp = st.checkbox("Include WSP (Wholesale Price) Data")
    if include_wsp:
        wsp_file = st.file_uploader("Upload WSP Data File", type=['csv', 'xlsx'])
        if wsp_file is not None:
            try:
                if wsp_file.name.endswith('.csv'):
                    wsp_df = pd.read_csv(wsp_file)
                else:
                    wsp_df = pd.read_excel(wsp_file)
                required_columns = ['Region(District)', 'D1-3', 'D4-6', 'D7-9', 'D10-12', 'D13-15','D16-18','D19-21','D22-24','D25-27','D28-30','D1-3 J','D4-6 J','D7-8 J']
                for col in required_columns:
                    if col not in wsp_df.columns:
                        st.error(f"Missing required WSP column: {col}")
                        return None
                return wsp_df
            except Exception as e:
                st.error(f"Could not read WSP file: {e}")
                return None
    return None
 def save_regional_price_trend_report(df):
    wsp_df = get_wsp_data()
    competitive_brands_wsp_df = get_competitive_brands_wsp_data()
    return generate_regional_price_trend_report(df, wsp_df,competitive_brands_wsp_df)
 def convert_dataframe_to_pdf(df, filename):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    title_style = styles['Heading1']
    data = [df.columns.tolist()]  # Header row
    for _, row in df.iterrows():
        data.append([str(val) for val in row.tolist()])
    table = Table(data)
    table.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.grey),('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),('ALIGN', (0,0), (-1,-1), 'CENTER'),('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),('FONTSIZE', (0,0), (-1,0), 12),('BOTTOMPADDING', (0,0), (-1,0), 12),('BACKGROUND', (0,1), (-1,-1), colors.beige),('GRID', (0,0), (-1,-1), 1, colors.black)]))
    content = []
    content.append(table)
    doc.build(content)
    buffer.seek(0)
    return buffer
 def save_processed_dataframe(df, start_date=None, download_format='xlsx'):
    region_order = ['GJ (Ahmedabad)', 'GJ (Surat)', 'RJ(Jaipur)', 'RJ(Udaipur)', 'HY (Gurgaon)', 'PB (Bhatinda)','Delhi','CG (Raipur)', 'ORR (Khorda)', 'ORR (Sambalpur)', 'UP (Gaziabad)','UK (Haridwar)','UK (Dehradun)', 'M.P.(East)[Balaghat]', 'M.P.(West)[Indore]', 'M.H.(East)[Nagpur Urban]']
    if 'processed_dataframe' in st.session_state:
        df = st.session_state['processed_dataframe']
    df_to_save = df.copy()
    df_to_save['region_order'] = df_to_save['Region(District)'].map({region: idx for idx, region in enumerate(region_order)})
    df_to_save = df_to_save.sort_values(['region_order', 'Date'])
    df_to_save = df_to_save.drop(columns=['region_order'])
    if 'Date' in df_to_save.columns:
        df_to_save['Date'] = pd.to_datetime(df_to_save['Date'], format='%d-%b %Y')
        if start_date:
            df_to_save = df_to_save[df_to_save['Date'] >= start_date]
            df_to_save['Date'] = df_to_save['Date'].dt.strftime('%d-%b %Y')
    output = io.BytesIO()
    if download_format == 'xlsx':
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df_to_save.to_excel(writer, sheet_name='Sheet1', index=False)
            workbook = writer.book
            worksheet = writer.sheets['Sheet1']
            worksheet.repeat_rows(0)
            worksheet.set_page_view()
            dark_blue = '#2C3E50'
            white = '#FFFFFF'
            light_gray = '#F2F2F2'
            format_header = workbook.add_format({'bold': True, 'font_size': 14,'bg_color': dark_blue,'font_color': white,'align': 'center','valign': 'vcenter','border': 1,'border_color': '#000000','text_wrap': True})
            format_general = workbook.add_format({'font_size': 12,'valign': 'vcenter','align': 'center'})
            format_alternating = workbook.add_format({'font_size': 12,'bg_color': light_gray,'valign': 'vcenter','align': 'center'})
            worksheet.set_row(0, 30, format_header)
            for row_num in range(1, len(df_to_save) + 1):
                if row_num % 2 == 0:
                    worksheet.set_row(row_num, None, format_alternating)
                else:
                    worksheet.set_row(row_num, None, format_general)
            for col_num, col_name in enumerate(df_to_save.columns):
                max_len = max(df_to_save[col_name].astype(str).map(len).max(),len(str(col_name)))
                worksheet.set_column(col_num, col_num, max_len + 2, format_general)
            if 'MoM Change' in df_to_save.columns:
                mom_change_col_index = df_to_save.columns.get_loc('MoM Change')
                format_negative = workbook.add_format({'bg_color': '#FFC7CE','font_size': 12,'align': 'center','valign': 'vcenter'})
                format_zero = workbook.add_format({'bg_color': '#D9D9D9','font_size': 12,'align': 'center','valign': 'vcenter'})
                format_positive = workbook.add_format({'bg_color': '#C6EFCE','font_size': 12,'align': 'center','valign': 'vcenter'})
                worksheet.conditional_format(1, mom_change_col_index, len(df_to_save), mom_change_col_index, {'type': 'cell', 'criteria': '<', 'value': 0, 'format': format_negative})
                worksheet.conditional_format(1, mom_change_col_index, len(df_to_save), mom_change_col_index, {'type': 'cell','criteria': '=','value': 0,'format': format_zero})
                worksheet.conditional_format(1, mom_change_col_index, len(df_to_save), mom_change_col_index, {'type': 'cell','criteria': '>','value': 0, 'format': format_positive})
            writer.close()
    elif download_format == 'pdf':
        output = convert_dataframe_to_pdf(df_to_save, 'processed_price_tracker.pdf')
    output.seek(0)
    return output
 def parse_date(date_str):
    try:
        date_formats = ['%d-%b %Y','%d-%b-%Y','%d-%B %Y','%Y-%m-%d','%m/%d/%Y','%d/%m/%Y',]
        for fmt in date_formats:
            try:
                return pd.to_datetime(date_str, format=fmt)
            except ValueError:
                continue
        return pd.to_datetime(date_str, format='mixed', dayfirst=True)
    except Exception as e:
        st.warning(f"Could not parse date: {date_str}. Error: {e}")
        return pd.NaT
 def process_excel_file(uploaded_file, requires_editing):
    warnings.simplefilter("ignore")
    df = pd.read_excel(uploaded_file)
    if not requires_editing:
        if 'Date' in df.columns:
            df['Date'] = df['Date'].apply(parse_date)
        return df
    df = df.iloc[1:] 
    df = df.iloc[:, 1:]
    new_header = df.iloc[0]
    df = df[1:]
    df.columns = new_header
    df = df[~df.iloc[:, 1].str.contains('Date', na=False, case=False)]
    df.iloc[:, 1] = df.iloc[:, 1].apply(parse_date)
    df = df.loc[:, df.columns.notnull()] 
    df = df[df.iloc[:, 0] != "JKLC Price Tracker Mar'24 - till 03-12-24"]
    mask = df.iloc[:, 0].notna()
    current_value = None
    for i in range(len(df)):     
        if mask.iloc[i]:         
            current_value = df.iloc[i, 0]     
        else:         
            if current_value is not None:             
                df.iloc[i, 0] = current_value 
    df = df.rename(columns={df.columns[0]: 'Region(District)'})
    df = df.reset_index(drop=True)
    return df
 def generate_wsp_comparison_report(company_wsp_df, competitive_brands_wsp=None):
    try:
        region_order = ['GJ (Ahmedabad)', 'GJ (Surat)', 'RJ(Jaipur)', 'RJ(Udaipur)', 
                       'HY (Gurgaon)', 'PB (Bhatinda)', 'Delhi', 'CG (Raipur)', 
                       'ORR (Khorda)', 'ORR (Sambalpur)', 'UP (Gaziabad)', 'UK (Haridwar)', 
                       'UK (Dehradun)', 'M.P.(East)[Balaghat]', 'M.P.(West)[Indore]', 
                       'M.H.(East)[Nagpur Urban]']
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=8, leftMargin=8, 
                              topMargin=8, bottomMargin=8)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle('TitleStyle', parent=styles['Title'], 
                                   fontSize=20, textColor=colors.darkblue, 
                                   alignment=TA_CENTER, spaceAfter=1)
        region_style = ParagraphStyle('RegionStyle', parent=styles['Heading2'], 
                                    textColor=colors.blue, spaceAfter=1, fontSize=12)
        normal_style = styles['Normal']
        month_style = ParagraphStyle('MonthStyle', parent=styles['Heading3'], 
                                   textColor=colors.green, spaceAfter=1)
        
        story = []
        
        # Report header
        story.append(Paragraph("WSP Comparison Report", title_style))
        story.append(Paragraph("December 2024 - January 2025", 
                             ParagraphStyle('SubtitleStyle', parent=styles['Normal'], 
                                          fontSize=12, textColor=colors.red, 
                                          alignment=TA_CENTER, spaceAfter=1)))
        story.append(Spacer(1, 12))
        
        # Process each region
        for region in region_order:
            if region in ['UK (Dehradun)', 'UK (Haridwar)']:
                continue
                
            region_story = []
            region_story.append(Paragraph(f"{region}", region_style))
            region_story.append(Spacer(1, 6))
            
            # Company WSP data
            if company_wsp_df is not None and not company_wsp_df.empty:
                region_story.append(Paragraph("JKLC WSP Progression:", month_style))
                company_data = create_wsp_summary(company_wsp_df, region)
                if company_data:
                    region_story.extend(company_data)
                    region_story.append(Spacer(1, 6))
            
            # Competitive brands WSP data
            if competitive_brands_wsp:
                for brand, brand_wsp_df in competitive_brands_wsp.items():
                    region_story.append(Paragraph(f"{brand} WSP Progression:", month_style))
                    competitor_data = create_wsp_summary(brand_wsp_df, region)
                    if competitor_data:
                        region_story.extend(competitor_data)
                        region_story.append(Spacer(1, 6))
            
            # Add comparison summary if both company and competitor data exist
            if company_wsp_df is not None and competitive_brands_wsp:
                comparison_data = create_comparison_summary(company_wsp_df, competitive_brands_wsp, region)
                if comparison_data:
                    region_story.extend(comparison_data)
            
            story.append(KeepTogether(region_story))
            story.append(Paragraph("<pagebreak/>", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    except Exception as e:
        print(f"Error generating WSP comparison report: {e}")
        raise

 def create_wsp_summary(wsp_df, region):
    styles = getSampleStyleSheet()
    normal_style = styles['Normal']
    summary_style = ParagraphStyle('SummaryStyle', parent=styles['Normal'], 
                                 textColor=colors.brown, spaceAfter=0)
    
    region_wsp = wsp_df[wsp_df['Region(District)'] == region]
    if region_wsp.empty:
        return None
        
    summary = []
    
    # December WSP calculation
    dec_columns = ['D1-3', 'D4-6', 'D7-9', 'D10-12', 'D13-15', 'D16-18', 
                  'D19-21', 'D22-24', 'D25-27', 'D28-30']
    dec_values = region_wsp[dec_columns].values.flatten().tolist()
    
    # January WSP calculation
    jan_columns = ['D1-3 J', 'D4-6 J', 'D7-8 J']
    jan_values = region_wsp[jan_columns].values.flatten().tolist()
    
    # Calculate changes
    dec_change = float(dec_values[-1]) - float(dec_values[0])
    jan_change = float(jan_values[-1]) - float(dec_values[-1])
    total_change = float(jan_values[-1]) - float(dec_values[0])
    
    # Add WSP values
    summary.append(Paragraph(f"December Start: Rs. {dec_values[0]:.0f}", normal_style))
    summary.append(Paragraph(f"December End: Rs. {dec_values[-1]:.0f}", normal_style))
    summary.append(Paragraph(f"January Latest: Rs. {jan_values[-1]:.0f}", normal_style))
    
    # Add changes
    summary.append(Paragraph(f"December Net Change: {dec_change:+.0f} Rs.", summary_style))
    summary.append(Paragraph(f"January Net Change: {jan_change:+.0f} Rs.", summary_style))
    summary.append(Paragraph(f"Total Net Change: {total_change:+.0f} Rs.", summary_style))
    
    return summary

 def create_comparison_summary(company_wsp_df, competitive_brands_wsp, region):
    styles = getSampleStyleSheet()
    comparison_style = ParagraphStyle('ComparisonStyle', parent=styles['Normal'], 
                                    textColor=colors.blue, spaceAfter=0)
    
    summary = []
    summary.append(Spacer(1, 6))
    summary.append(Paragraph("Comparative Analysis:", comparison_style))
    
    company_region_wsp = company_wsp_df[company_wsp_df['Region(District)'] == region]
    if company_region_wsp.empty:
        return None
    
    company_dec_start = company_region_wsp['D1-3'].values[0]
    
    for brand, brand_wsp_df in competitive_brands_wsp.items():
        competitor_region_wsp = brand_wsp_df[brand_wsp_df['Region(District)'] == region]
        if not competitor_region_wsp.empty:
            competitor_dec_start = competitor_region_wsp['D1-3'].values[0]
            difference = company_dec_start - competitor_dec_start
            summary.append(Paragraph(
                f"JKLC vs {brand} (December Start): {difference:+.0f} Rs.", 
                comparison_style))
    
    return summary
 def download_wsp_comparison_report():
    st.subheader("WSP Comparison Report Generator")
    
    # Get company WSP data
    company_wsp_df = get_wsp_data()
    
    # Get competitive brands WSP data
    competitive_brands_wsp = get_competitive_brands_wsp_data()
    
    if company_wsp_df is not None or competitive_brands_wsp is not None:
        # Create a download button
        if st.button("Generate WSP Comparison Report"):
            try:
                # Generate the report
                report_buffer = generate_wsp_comparison_report(company_wsp_df, competitive_brands_wsp)
                
                # Create the download button
                current_date = datetime.now().strftime("%d%b%Y")
                st.download_button(
                    label="📥 Download WSP Comparison Report",
                    data=report_buffer,
                    file_name=f"WSP_Comparison_Report_{current_date}.pdf",
                    mime="application/pdf",
                    key='download_wsp_report'
                )
                
                st.success("Report generated successfully! Click the download button above to save it.")
                
            except Exception as e:
                st.error(f"Error generating report: {e}")
    else:
        st.info("Please upload the WSP data files to generate the comparison report.")
 def main():
    st.title("📊 Price Tracker Analysis Tool")
    st.markdown("""
    ### Welcome to the Price Tracker Analysis Tool
    
    **Instructions:**
    1. Upload your Excel price tracking file
    2. Choose whether the file needs initial editing
    3. Add new data, analyze regions, and download processed files
    """)
    uploaded_file = st.file_uploader("Please upload the Price Tracker file", type=['xlsx'], help="Upload an Excel file containing price tracking data")
    if uploaded_file is not None:
        requires_editing = st.radio("Does this file require initial editing?", ["No", "Yes"],help="Select 'Yes' if the uploaded file needs preprocessing")
        try:
            df = process_excel_file(uploaded_file, requires_editing == "Yes")
            required_columns = ['Region(District)', 'Date', 'Inv.', 'RD', 'STS', 'Reglr']
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                st.error(f"Missing required columns: {', '.join(missing_columns)}")
                st.stop()
            col1, col2 = st.columns([1,3])
            with col1:
                st.subheader("🔄 Data Entry")
                price_changed = st.radio("Do you want to add new data?", ["No", "Yes"])
                if price_changed == "Yes":
                    unique_regions = df['Region(District)'].unique()
                    if len(unique_regions) == 0:
                        st.warning("No regions found in the dataframe.")
                    else:
                        selected_regions = st.multiselect("Select Region(s)", unique_regions)
                        data_entries = []
                        for selected_region in selected_regions:
                            st.markdown(f"### Data Entry for {selected_region}")
                            region_df = df[df['Region(District)'] == selected_region]
                            from datetime import datetime
                            date_input = st.text_input(f"Enter Date for {selected_region}", value=datetime.now().strftime("%d-%b %Y"),placeholder="DD-Mon YYYY, e.g., 01-Jan 2024",key=f"date_{selected_region}")
                            inv_input = st.number_input(f"Enter Inv. value for {selected_region}",value=0.0,format="%.2f",key=f"inv_{selected_region}")
                            rd_input = st.number_input(f"Enter RD value for {selected_region}",value=0.0, format="%.2f",key=f"rd_{selected_region}")
                            sts_input = st.number_input(f"Enter STS value for {selected_region}",value=0.0, format="%.2f", key=f"sts_{selected_region}")
                            reglr_input = st.number_input(f"Enter Reglr value for {selected_region}",value=0.0,format="%.2f",key=f"reglr_{selected_region}")
                            net_input = inv_input - rd_input - sts_input - reglr_input
                            st.write(f"Calculated Net value for {selected_region}: {net_input}")
                            last_net_value = region_df['Net'].iloc[-1] if 'Net' in region_df.columns and not region_df['Net'].empty else 0
                            mom_change = net_input - last_net_value
                            st.write(f"Calculated MoM Change for {selected_region}: {mom_change}")
                            remarks_input = st.text_area(f"Enter Remarks for {selected_region} (Optional)",key=f"remarks_{selected_region}")
                            new_row = {'Region(District)': selected_region,'Date': parse_date(date_input).strftime('%d-%b %Y'),'Inv.': inv_input,'RD': rd_input,'STS': sts_input,'Reglr': reglr_input,'Net': net_input,'MoM Change': mom_change,'Remarks': remarks_input}
                            data_entries.append(new_row)
                            st.markdown("---")
                        if st.button("Add New Rows to Dataframe"):
                            if not data_entries:
                                st.warning("No new entries to add.")
                                return
                            updated_df = df.copy()
                            new_rows_df = pd.DataFrame(data_entries)
                            for col in df.columns:
                                if col not in new_rows_df.columns:
                                    new_rows_df[col] = None
                            new_rows_df = new_rows_df.reindex(columns=df.columns)
                            for region in new_rows_df['Region(District)'].unique():
                                region_new_rows = new_rows_df[new_rows_df['Region(District)'] == region]
                                region_existing_indices = updated_df[updated_df['Region(District)'] == region].index
                                if not region_existing_indices.empty:
                                    last_region_index = region_existing_indices[-1]
                                    before_region = updated_df.iloc[:last_region_index+1]
                                    after_region = updated_df.iloc[last_region_index+1:]
                                    updated_df = pd.concat([before_region,region_new_rows,after_region]).reset_index(drop=True)
                                else:
                                    updated_df = pd.concat([updated_df, region_new_rows]).reset_index(drop=True)
                            df = updated_df
                            st.session_state['processed_dataframe'] = df
                            st.success(f"{len(data_entries)} new rows added successfully!")
            with col2:
                st.subheader("📈 Region Analysis")
                unique_regions = df['Region(District)'].unique()
                selected_region_analysis = st.selectbox("Select Region for Analysis", unique_regions,key="region")
                region_analysis_df = df[df['Region(District)'] == selected_region_analysis]
                region_analysis_df['Date'] = pd.to_datetime(region_analysis_df['Date'], format='%d-%b %Y')
                current_month = dt.now().month
                current_year = dt.now().year
                last_month = current_month - 1 if current_month > 1 else 12
                last_month_year = current_year if current_month > 1 else current_year - 1
                last_month_data = region_analysis_df[(region_analysis_df['Date'].dt.month == last_month) & (region_analysis_df['Date'].dt.year == last_month_year)]
                current_month_data = region_analysis_df[(region_analysis_df['Date'].dt.month == current_month) & (region_analysis_df['Date'].dt.year == current_year)]
                display_columns = ['Date', 'Inv.', 'RD', 'STS', 'Reglr', 'Net', 'MoM Change']
                st.markdown(f"### Monthly Data for {selected_region_analysis}")
                st.markdown("#### Last Month Data")
                if not last_month_data.empty:
                      last_month_display = last_month_data[display_columns].copy()
                      last_month_display['Date'] = last_month_display['Date'].dt.strftime('%d-%b %Y')
                      last_month_display.set_index('Date', inplace=True)
                      last_month_display['Inv.']= last_month_display['Inv.'].abs().round(0).astype(int)
                      last_month_display['RD'] = last_month_display['RD'].abs().round(0).astype(int)
                      last_month_display['STS'] = last_month_display['STS'].abs().round(0).astype(int)
                      last_month_display['Reglr'] = last_month_display['Reglr'].abs().round(0).astype(int)
                      last_month_display['Net'] = last_month_display['Net'].abs().round(0).astype(int)
                      last_month_display['MoM Change'] = last_month_display['MoM Change'].round(0).astype(int)
                      st.dataframe(last_month_display.style.background_gradient(cmap='Blues'), use_container_width=True)
                      col_last_1, col_last_2 = st.columns(2)
                      with col_last_1:
                       st.metric(f"Total No. of Price Change in (Last Month)", len(last_month_data))
                      with col_last_2:
                       st.metric("Total Change in NOD(Last Month)(in Rs.)", last_month_data['MoM Change'].sum())
                else:
                     st.info(f"No data found for last month in {selected_region_analysis}")
                st.markdown("#### Current Month Data")
                if not current_month_data.empty:
                     current_month_display = current_month_data[display_columns].copy()
                     current_month_display['Date'] = current_month_display['Date'].dt.strftime('%d-%b %Y')
                     current_month_display.set_index('Date', inplace=True)
                     current_month_display['Inv.']= current_month_display['Inv.'].abs().round(0).astype(int)
                     current_month_display['RD'] = current_month_display['RD'].abs().round(0).astype(int)
                     current_month_display['STS'] = current_month_display['STS'].abs().round(0).astype(int)
                     current_month_display['Reglr'] = current_month_display['Reglr'].abs().round(0).astype(int)
                     current_month_display['Net'] = current_month_display['Net'].abs().round(0).astype(int)
                     current_month_display['MoM Change'] = current_month_display['MoM Change'].round(0).astype(int)
                     st.dataframe(current_month_display.style.background_gradient(cmap='Blues'), use_container_width=True)
                     col_curr_1, col_curr_2 = st.columns(2)
                     with col_curr_1:
                        st.metric("Total No. of Price Change in (Current Month)", len(current_month_data))
                     with col_curr_2:
                         st.metric("Total Change in NOD(Current Month)(in Rs.)", current_month_data['MoM Change'].sum())
                else:
                      st.info(f"No data found for current month in {selected_region_analysis}")
                region_analysis_df = df[df['Region(District)'] == selected_region_analysis]
                col_metrics_1, col_metrics_2 = st.columns(2)
                with col_metrics_1:
                    st.metric("Total Price Changes", len(region_analysis_df))
                st.markdown("### Graph Date Range")
                col_start_month, col_start_year = st.columns(2)
                with col_start_month:
                  start_month = st.selectbox("Select Start Month", ['January', 'February', 'March', 'April', 'May', 'June','July', 'August', 'September', 'October', 'November', 'December'],index=8)
                with col_start_year:
                  start_year = st.number_input("Select Start Year", min_value=2000, max_value=2030, value=2024)
                start_date = pd.to_datetime(f'01-{start_month[:3].lower()} {start_year}', format='%d-%b %Y')
                region_analysis_df = df[df['Region(District)'] == selected_region_analysis]
                region_analysis_df['Date'] = pd.to_datetime(region_analysis_df['Date'], format='%d-%b %Y')
                filtered_df = region_analysis_df[region_analysis_df['Date'] >= start_date].copy()
                if filtered_df.empty:
                    st.warning(f"No data available for {selected_region_analysis} from {start_month} {start_year}")
                else:
                    graph_type = st.selectbox("Select Metric for Analysis", ['Net', 'Inv.', 'RD', 'STS', 'Reglr', 'MoM Change'])
                filtered_df = filtered_df.sort_values('Date')
                fig = go.Figure()
                fig.add_trace(go.Scatter(x=filtered_df['Date'],y=filtered_df[graph_type], mode='lines+markers+text',text=filtered_df[graph_type].abs().round(0).astype(int),textposition='top center',name=f'{graph_type} Value',line=dict(color='#1E90FF',width=3),marker=dict(size=10,color='#4169E1',symbol='circle',line=dict(color='#FFFFFF',width=2)),hovertemplate=('<b>Date</b>: %{x|%d %B %Y}<br>' +f'<b>{graph_type}</b>: %{{y:.2f}}<br>' +'<extra></extra>')))
                fig.update_layout(title=f'{graph_type} Value Trend for {selected_region_analysis}',xaxis_title='Date',yaxis_title=f'{graph_type} Value',height=400)
                st.plotly_chart(fig, use_container_width=True)
                graph_download_format = st.selectbox("Download Graph as", ['PNG', 'PDF'])
                if st.button("Download Graph"):
                        if graph_download_format == 'PNG':
                            img_bytes = pio.to_image(fig, format='png')
                            st.download_button(label="Download Graph as PNG",data=img_bytes,file_name=f'{selected_region_analysis}_{graph_type}_trend.png',mime='image/png')
                        else:
                            pdf_bytes = pio.to_image(fig, format='pdf')
                            st.download_button(label="Download Graph as PDF",data=pdf_bytes,file_name=f'{selected_region_analysis}_{graph_type}_trend.pdf',mime='application/pdf')
                st.markdown("#### Effective NOD Analysis")
                current_month = dt.now().month
                current_year = dt.now().year
                current_month_effective = calculate_effective_nod(df, selected_region_analysis, current_month, current_year)
                last_month_effective = calculate_effective_nod(df, selected_region_analysis, current_month - 1 if current_month > 1 else 12,current_year if current_month > 1 else current_year - 1)
                col_eff_1, col_eff_2 = st.columns(2)
                with col_eff_1:
                     st.markdown("##### Current Month Effective NOD(Estimated)")
                     if current_month_effective:
                       st.metric("Effective NOD", f"₹{current_month_effective['effective_nod']:,.2f}")
                       with st.expander("View Breakdown"):
                            st.markdown(f"""
                        - First 10 days (20%): ₹{current_month_effective['first_period_nod']:,.2f}
                          * Contribution: ₹{current_month_effective['first_period_contribution']:,.2f}
                        - Middle 10 days (30%): ₹{current_month_effective['middle_period_nod']:,.2f}
                            * Contribution: ₹{current_month_effective['middle_period_contribution']:,.2f}
                        - Last 10 days (50%): ₹{current_month_effective['last_period_nod']:,.2f}
                            * Contribution: ₹{current_month_effective['last_period_contribution']:,.2f}
                        """)
                     else:
                        st.info("No data available for current month")
                with col_eff_2:
                     st.markdown("##### Last Month Effective NOD")
                     if last_month_effective:
                         st.metric("Effective NOD", f"₹{last_month_effective['effective_nod']:,.2f}")
                         with st.expander("View Breakdown"):
                             st.markdown(f"""
                        - First 10 days (20%): ₹{last_month_effective['first_period_nod']:,.2f}
                             * Contribution: ₹{last_month_effective['first_period_contribution']:,.2f}
                        - Middle 10 days (30%): ₹{last_month_effective['middle_period_nod']:,.2f}
                             * Contribution: ₹{last_month_effective['middle_period_contribution']:,.2f}
                        - Last 10 days (50%): ₹{last_month_effective['last_period_nod']:,.2f}
                             * Contribution: ₹{last_month_effective['last_period_contribution']:,.2f}
                        """)
                     else:
                         st.info("No data available for last month")
                if current_month_effective or last_month_effective:
                   st.markdown("##### Effective NOD Composition")
                   fig = go.Figure()
                   if current_month_effective:
                      current_month_name = dt.now().strftime('%B')
                      fig.add_trace(go.Bar(name=current_month_name,x=['First 10 Days', 'Middle 10 Days', 'Last 10 Days'],y=[current_month_effective['first_period_contribution'],current_month_effective['middle_period_contribution'],current_month_effective['last_period_contribution']],text=[f"₹{val:,.0f}" for val in [current_month_effective['first_period_contribution'],current_month_effective['middle_period_contribution'],current_month_effective['last_period_contribution']]],textposition='auto',))
                   if last_month_effective:
                      last_month_name = (dt.now().replace(day=1) - timedelta(days=1)).strftime('%B')
                      fig.add_trace(go.Bar(name=last_month_name,x=['First 10 Days', 'Middle 10 Days', 'Last 10 Days'],y=[last_month_effective['first_period_contribution'],last_month_effective['middle_period_contribution'],last_month_effective['last_period_contribution']],text=[f"₹{val:,.0f}" for val in [last_month_effective['first_period_contribution'],last_month_effective['middle_period_contribution'],last_month_effective['last_period_contribution']]],textposition='auto',))
                   fig.update_layout(title='Effective NOD Composition by Period',xaxis_title='Period',yaxis_title='Contribution to Effective NOD (₹)',barmode='group',height=400)
                   st.plotly_chart(fig, use_container_width=True)
                st.markdown("#### Effective Invoice Analysis")
                current_month = dt.now().month
                current_year = dt.now().year
                current_month_effective_invoice = calculate_effective_invoice(df, selected_region_analysis, current_month, current_year)
                last_month_effective_invoice = calculate_effective_invoice(df, selected_region_analysis,current_month - 1 if current_month > 1 else 12,current_year if current_month > 1 else current_year - 1)
                col_eff_inv_1, col_eff_inv_2 = st.columns(2)
                with col_eff_inv_1:
                  st.markdown("##### Current Month Effective Invoice(Estimated)")
                  if current_month_effective_invoice:
                   st.metric("Effective Invoice", f"₹{current_month_effective_invoice['effective_invoice']:,.2f}")
                   with st.expander("View Breakdown"):
                      st.markdown(f"""
            - First 10 days (20%): ₹{current_month_effective_invoice['first_period_invoice']:,.2f}
              * Contribution: ₹{current_month_effective_invoice['first_period_contribution']:,.2f}
            - Middle 10 days (30%): ₹{current_month_effective_invoice['middle_period_invoice']:,.2f}
              * Contribution: ₹{current_month_effective_invoice['middle_period_contribution']:,.2f}
            - Last 10 days (50%): ₹{current_month_effective_invoice['last_period_invoice']:,.2f}
              * Contribution: ₹{current_month_effective_invoice['last_period_contribution']:,.2f}
            """)
                  else:
                     st.info("No data available for current month")
                with col_eff_inv_2:
                  st.markdown("##### Last Month Effective Invoice")
                  if last_month_effective_invoice:
                    st.metric("Effective Invoice", f"₹{last_month_effective_invoice['effective_invoice']:,.2f}")
                    with st.expander("View Breakdown"):
                         st.markdown(f"""
            - First 10 days (20%): ₹{last_month_effective_invoice['first_period_invoice']:,.2f}
              * Contribution: ₹{last_month_effective_invoice['first_period_contribution']:,.2f}
            - Middle 10 days (30%): ₹{last_month_effective_invoice['middle_period_invoice']:,.2f}
              * Contribution: ₹{last_month_effective_invoice['middle_period_contribution']:,.2f}
            - Last 10 days (50%): ₹{last_month_effective_invoice['last_period_invoice']:,.2f}
              * Contribution: ₹{last_month_effective_invoice['last_period_contribution']:,.2f}
            """)
                  else:
                     st.info("No data available for last month")
                if current_month_effective_invoice or last_month_effective_invoice:
                   st.markdown("##### Effective Invoice Composition")
                   fig = go.Figure()
                   if current_month_effective_invoice:
                      current_month_name = dt.now().strftime('%B')
                      fig.add_trace(go.Bar(name=current_month_name,x=['First 10 Days', 'Middle 10 Days', 'Last 10 Days'],y=[current_month_effective_invoice['first_period_contribution'],current_month_effective_invoice['middle_period_contribution'],current_month_effective_invoice['last_period_contribution']],text=[f"₹{val:,.0f}" for val in [current_month_effective_invoice['first_period_contribution'],current_month_effective_invoice['middle_period_contribution'],current_month_effective_invoice['last_period_contribution']]],textposition='auto',))
                   if last_month_effective_invoice:
                      last_month_name = (dt.now().replace(day=1) - timedelta(days=1)).strftime('%B')
                      fig.add_trace(go.Bar(name=last_month_name,x=['First 10 Days', 'Middle 10 Days', 'Last 10 Days'],y=[last_month_effective_invoice['first_period_contribution'],last_month_effective_invoice['middle_period_contribution'],last_month_effective_invoice['last_period_contribution']],text=[f"₹{val:,.0f}" for val in [last_month_effective_invoice['first_period_contribution'],last_month_effective_invoice['middle_period_contribution'],last_month_effective_invoice['last_period_contribution']]],textposition='auto',))
                   fig.update_layout(title='Effective Invoice Composition by Period',xaxis_title='Period',yaxis_title='Contribution to Effective Invoice (₹)',barmode='group',height=400)
                   st.plotly_chart(fig, use_container_width=True)
                st.markdown("### Remarks")
                remarks_df = region_analysis_df[['Date', 'Remarks']].dropna(subset=['Remarks'])
                remarks_df = remarks_df.sort_values('Date', ascending=False)
                if not remarks_df.empty:
                        for _, row in remarks_df.iterrows():
                            st.markdown(f"""<div style="background-color:#f0f2f6;border-left: 5px solid #4a4a4a;padding: 10px;margin-bottom: 10px;border-radius: 5px;"><strong>{row['Date'].strftime('%d-%b %Y')}</strong>: {row['Remarks']}</div>""", unsafe_allow_html=True)
                else:
                        st.info("No remarks found for this region.")
            st.markdown("## 📥 Download Options")
            download_options = st.radio("Download File From:", ["Entire Dataframe", "Specific Month", "Regional Price Trend Report","DOWNLOAD WSP REPORT"], horizontal=True)
            start_date = None
            if download_options =="Entire Dataframe":
                if st.button("Download Processed File"):
                 try:
                    output = save_processed_dataframe(df, start_date, selected_format)
                    st.download_button(label=f"Click to Download {download_format}",data=output,file_name=f'processed_price_tracker.{selected_format}',mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' if selected_format == 'xlsx' else 'application/pdf')
                 except Exception as e:
                    st.error(f"Error during download: {e}")
            if download_options == "Specific Month":
                col1, col2 = st.columns(2)
                with col1:
                    month_input = st.selectbox("Select Month", ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December'])
                with col2:
                    year_input = st.number_input("Select Year", min_value=2000, max_value=2030, value=2024)
                start_date = pd.to_datetime(f'01-{month_input[:3].lower()} {year_input}', format='%d-%b %Y')
                download_format = st.selectbox("Select Download Format", ['Excel (.xlsx)', 'PDF (.pdf)'])
                format_map = {'Excel (.xlsx)': 'xlsx', 'PDF (.pdf)': 'pdf'}
                selected_format = format_map[download_format]
                if st.button("Download Processed File"):
                 try:
                    output = save_processed_dataframe(df, start_date, selected_format)
                    st.download_button(label=f"Click to Download {download_format}",data=output,file_name=f'processed_price_tracker.{selected_format}',mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' if selected_format == 'xlsx' else 'application/pdf')
                 except Exception as e:
                    st.error(f"Error during download: {e}")
            if download_options == "Regional Price Trend Report":
                output = save_regional_price_trend_report(df)
                st.download_button(label="Download Regional Price Trend Report (PDF)",data=output,file_name="regional_price_trend_report.pdf",mime="application/pdf")
            if download_options == "DOWNLOAD WSP REPORT":
               download_wsp_comparison_report()
        except Exception as e:
            st.error(f"An error occurred: {e}")
            st.exception(e)
 if __name__ == "__main__":
    main()
def price_input():
 warnings.filterwarnings('ignore', category=DeprecationWarning)
 def parse_date(date_str):
    if pd.isna(date_str):
        return None
    date_formats = ['%d/%m/%Y','%m/%d/%Y','%Y-%m-%d','%d-%m-%Y','%d.%m.%Y',]
    date_str = str(date_str).strip()
    for fmt in date_formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    try:
        return pd.to_datetime(date_str)
    except:
        st.warning(f"Could not parse date: {date_str}")
        return None
 def preprocess_dataframe(df):
    df = df.dropna(subset=['Owner: Full Name', 'Brand: Name', 'checkin date'])
    df['checkin date'] = df['checkin date'].apply(parse_date)
    df = df.dropna(subset=['checkin date'])
    df['checkin ate'] = df['checkin date'].dt.strftime('%d/%m/%Y')
    return df
 def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)
 st.markdown("""<style>.reportview-container {background-color: #f0f2f6;}.sidebar .sidebar-content {background-color: #ffffff;box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);}.stButton>button {background-color: #1F4E78;color: white;border: none;padding: 10px 20px;border-radius: 5px;transition: background-color 0.3s ease;}.stButton>button:hover {background-color: #2C6BA3;}.stMultiSelect, .stSelectbox {width: 100%;}h1, h2, h3 {color: #1F4E78;}</style>""", unsafe_allow_html=True)
 def normalize_brand_name(brand):
    if pd.isna(brand): 
        return ""
    return str(brand).lower().strip()
 def create_price_report(df, selected_owners):
    wb = Workbook()
    ws = wb.active
    ws.title = "Price Reports"
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    border = Border(left=Side(style='thin'),right=Side(style='thin'),top=Side(style='thin'),bottom=Side(style='thin'))
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 30  
    ws.column_dimensions['C'].width = 15  
    ws.column_dimensions['D'].width = 15  
    ws.column_dimensions['E'].width = 15  
    ws.column_dimensions['F'].width = 50  
    headers = ["Regional Head", "Brand Name", "Total Reports", "First Report", "Last Report", "Report Dates"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal='center')
    current_row = 2
    for owner in sorted(selected_owners):
        owner_data = df[df['Owner: Full Name'] == owner]
        if len(owner_data) == 0:
            continue
        for brand in sorted(owner_data['Brand: Name'].unique()):
            brand_data = owner_data[owner_data['Brand: Name'] == brand]
            unique_dates = sorted(list(set([parse_date(date) for date in brand_data['checkin date']])))
            if not unique_dates:
                continue
            ws.cell(row=current_row, column=1, value=owner)
            ws.cell(row=current_row, column=2, value=brand)
            ws.cell(row=current_row, column=3, value=len(unique_dates))
            ws.cell(row=current_row, column=4, value=unique_dates[0].strftime('%d/%m/%Y'))
            ws.cell(row=current_row, column=5, value=unique_dates[-1].strftime('%d/%m/%Y'))
            ws.cell(row=current_row, column=6, value=", ".join(d.strftime('%d/%m/%Y') for d in unique_dates))
            for col in range(1, 7):
                cell = ws.cell(row=current_row, column=col)
                cell.border = border
                cell.alignment = Alignment(horizontal='left' if col in [1, 2, 6] else 'center')
            current_row += 1
    ws.auto_filter.ref = f"A1:F{current_row-1}"
    filename = f"price_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    wb.save(filename)
    return filename
 def create_price_report1(df, selected_owners):
    df = df.dropna(subset=['Brand: Name'])
    wb = Workbook()
    ws = wb.active
    ws.title = "Price Reports"
    TARGET_BRANDS = ['jk', 'wonder', 'shree', 'platinum', 'ambuja', 'ultratech']
    COLORS = {'header_bg': "1F4E78",'header_text': "FFFFFF",'alt_row': "F5F9FF",'border_color': "C5D9F1",'low_visits': "FF0000",'high_visits': "00FF00"}
    header_fill = PatternFill(start_color=COLORS['header_bg'], end_color=COLORS['header_bg'], fill_type="solid")
    alt_row_fill = PatternFill(start_color=COLORS['alt_row'], end_color=COLORS['alt_row'], fill_type="solid")
    header_font = Font(name='Calibri', size=11, color=COLORS['header_text'], bold=True)
    normal_font = Font(name='Calibri', size=10)
    border = Border(left=Side(style='thin', color=COLORS['border_color']),right=Side(style='thin', color=COLORS['border_color']),top=Side(style='thin', color=COLORS['border_color']),bottom=Side(style='thin', color=COLORS['border_color']))
    all_brands = df['Brand: Name'].unique()
    matched_brands = []
    for brand in all_brands:
        normalized = normalize_brand_name(brand)
        if any(target in normalized for target in TARGET_BRANDS):
            matched_brands.append(brand)
    matched_brands.sort()
    total_columns = 3 + len(matched_brands) + 1
    column_letter_end = chr(64 + total_columns)
    ws.merge_cells(f'A1:{column_letter_end}1')
    title_cell = ws['A1']
    title_cell.value = "Regional Price Report"
    title_cell.font = Font(name='Calibri', size=16, bold=True, color=COLORS['header_bg'])
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    timestamp_cell = ws['A2']
    timestamp_cell.value = f"Generated on: {datetime.now().strftime('%d %B %Y, %H:%M')}"
    timestamp_cell.font = Font(name='Calibri', size=9, italic=True)
    ws.merge_cells(f'A2:{column_letter_end}2')
    ws.insert_rows(3)
    headers = ["Regional Head", "First Report", "Last Report"]
    headers.extend(matched_brands)
    headers.append("Total Visits")
    for col in range(1, len(headers) + 1):
        if col == 1:  ws.column_dimensions[chr(64 + col)].width = 30
        elif col in [2, 3]: ws.column_dimensions[chr(64 + col)].width = 15
        else: ws.column_dimensions[chr(64 + col)].width = 12
    header_row = 4
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    current_row = header_row + 1
    all_visit_counts = []
    for owner in sorted(selected_owners):
        owner_data = df[df['Owner: Full Name'] == owner]
        if len(owner_data) == 0:
            continue
        row_data = {brand: 0 for brand in matched_brands}
        first_report_date = None
        last_report_date = None
        for brand in matched_brands:
            brand_data = owner_data[owner_data['Brand: Name'] == brand]
            if len(brand_data) > 0:
                unique_dates = sorted(list(set([parse_date(date) for date in brand_data['checkin date']])))
                if unique_dates:
                    row_data[brand] = len(unique_dates)
                    if row_data[brand] > 0:
                        all_visit_counts.append(len(unique_dates))
                    if first_report_date is None or unique_dates[0] < first_report_date:
                        first_report_date = unique_dates[0]
                    if last_report_date is None or unique_dates[-1] > last_report_date:
                        last_report_date = unique_dates[-1]
        total_visits = sum(row_data.values())
        if total_visits > 0:
            if current_row % 2 == 0:
                for col in range(1, len(headers) + 1):
                    ws.cell(row=current_row, column=col).fill = alt_row_fill
            ws.cell(row=current_row, column=1, value=owner)
            ws.cell(row=current_row, column=2, value=first_report_date.strftime('%d/%m/%Y') if first_report_date else "N/A")
            ws.cell(row=current_row, column=3, value=last_report_date.strftime('%d/%m/%Y') if last_report_date else "N/A")
            for col, brand in enumerate(matched_brands, 4):
                cell = ws.cell(row=current_row, column=col, value=row_data[brand])
                cell.alignment = Alignment(horizontal='center')
            ws.cell(row=current_row, column=len(headers), value=total_visits)
            for col in range(1, len(headers) + 1):
                cell = ws.cell(row=current_row, column=col)
                cell.border = border
                cell.font = normal_font
            current_row += 1
    median_visits = statistics.median(all_visit_counts) if all_visit_counts else 0
    brand_cols_start = 4
    brand_cols_end = len(headers) - 1
    for row in range(header_row + 1, current_row):
        for col in range(brand_cols_start, brand_cols_end + 1):
            cell = ws.cell(row=row, column=col)
            if cell.value and cell.value < median_visits:
                cell.font = Font(color="FF0000", bold=True)
    for col in range(brand_cols_start, brand_cols_end + 1):
        color_scale = ColorScaleRule(start_type='min', start_color='FF0000',mid_type='percentile', mid_value=50, mid_color='FFFF00',end_type='max', end_color='00FF00')
        col_letter = chr(64 + col)
        ws.conditional_formatting.add(f'{col_letter}{header_row+1}:{col_letter}{current_row-1}', color_scale)
    ws.auto_filter.ref = f"A{header_row}:{column_letter_end}{header_row}"
    ws.freeze_panes = ws[f'A{header_row+1}']
    filename = f"price_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    wb.save(filename)
    return filename
 def parse_pasted_owners(pasted_text):
        """Parse pasted text into a list of owner names"""
        if not pasted_text:
            return []
        # Split by newlines and clean up each name
        owners = [name.strip() for name in pasted_text.split('\n') if name.strip()]
        return owners

 def main():
        st.title('📊 Price Report Generator')
        st.markdown("""
        ### Upload Your Excel File
        Please upload an Excel file containing the following columns:
        - `Owner: Full Name`
        - `Brand: Name`
        - `checkin date`
        """)
        
        uploaded_file = st.file_uploader("Choose an Excel file", type=['xlsx', 'xls'], 
                                       help="Upload your price report Excel file")
        
        if uploaded_file is not None:
            try:
                df = pd.read_excel(uploaded_file)
                required_columns = ['Owner: Full Name', 'Brand: Name', 'checkin date']
                if not all(col in df.columns for col in required_columns):
                    st.error("Invalid file format. Please ensure your file contains the required columns.")
                    return
                
                st.sidebar.header('🔍 Select Owners')
                
                # Add text area for pasting owners
                st.sidebar.subheader("Paste Owners")
                pasted_owners = st.sidebar.text_area(
                    "Paste owner names (one per line)",
                    help="Copy and paste owner names from your Excel file, one name per line",
                    height=150
                )
                
                # Get all unique owners from the DataFrame
                owners = sorted(df['Owner: Full Name'].astype(str).unique().tolist())
                
                # Parse pasted owners and find matches
                if pasted_owners:
                    parsed_owners = parse_pasted_owners(pasted_owners)
                    # Find valid owners from the pasted list
                    valid_owners = [owner for owner in parsed_owners if owner in owners]
                    invalid_owners = [owner for owner in parsed_owners if owner not in owners]
                    
                    if invalid_owners:
                        st.sidebar.warning("Some pasted names were not found in the data:")
                        st.sidebar.write("\n".join(invalid_owners))
                else:
                    valid_owners = []

                # MultiSelect widget with pasted owners pre-selected
                selected_owners = st.sidebar.multiselect(
                    "Choose Regional Heads",
                    options=owners,
                    default=valid_owners if valid_owners else owners[:5],
                    help="Select the regional heads for your report"
                )
                
                st.sidebar.markdown("---")
                report_type = st.sidebar.radio(
                    "Choose Report Type",
                    ["Date-Based Report", "Owner-Based Report"],
                    help="Select the type of report you want to generate"
                )

                if st.sidebar.button('🚀 Generate Report', type='primary'):
                    if not selected_owners:
                        st.warning("Please select at least one owner.")
                    else:
                        with st.spinner('Generating Report...'):
                            try:
                                if report_type == "Date-Based Report":
                                    filename = create_price_report(df, selected_owners)
                                    st.success(f"Date-Based Report Generated: {filename}")
                                    with open(filename, 'rb') as file:
                                        st.download_button(
                                            label="Download Date-Based Report",
                                            data=file,
                                            file_name=filename,
                                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                        )
                                else:
                                    filename = create_price_report1(df, selected_owners)
                                    st.success(f"Owner-Based Report Generated: {filename}")
                                    with open(filename, 'rb') as file:
                                        st.download_button(
                                            label="Download Owner-Based Report",
                                            data=file,
                                            file_name=filename,
                                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                        )
                            except Exception as e:
                                st.error(f"An error occurred: {e}")
            except Exception as e:
                st.error(f"Error processing the file: {e}")

 if __name__ == "__main__":
        main()
def geo():
 def fill_second_column(df):
    processed_df = df.copy()
    first_col = processed_df.columns[0]
    second_col = processed_df.columns[1]
    processed_df[first_col] = processed_df[first_col].ffill()
    first_col_mask = processed_df[first_col].notna()
    processed_df.loc[first_col_mask, second_col] = processed_df.loc[first_col_mask, second_col].ffill()
    return processed_df
 def fill_third_column_comprehensively(df):
    processed_df = df.copy()
    if len(processed_df.columns) < 3:
        return processed_df
    third_col = processed_df.columns[2]
    processed_df[third_col] = processed_df[third_col].ffill()
    processed_df[third_col] = processed_df[third_col].bfill()
    return processed_df
 def process_excel_file(uploaded_file):
        xls = pd.ExcelFile(uploaded_file)
        sheet_names = xls.sheet_names
        channel_non_total_dfs = {}
        for sheet_name in sheet_names:
            df = pd.read_excel(uploaded_file, sheet_name=sheet_name, header=3)
            df = df[~df.iloc[:, 0].astype(str).str.lower().str.contains(r'trade|non trade')]
            non_total_df = df[~df.iloc[:, 0].astype(str).str.contains("Total", case=False, na=False)]
            non_total_df = fill_second_column(non_total_df)
            if len(non_total_df.columns) > 1:
                channel_non_total_df = non_total_df[~non_total_df.iloc[:, 1].astype(str).str.contains("Total", case=False, na=False)]
            else:
                channel_non_total_df = non_total_df
            if len(channel_non_total_df.columns) >= 3:
                channel_non_total_df = fill_third_column_comprehensively(channel_non_total_df)
            channel_non_total_dfs[sheet_name] = channel_non_total_df
        return channel_non_total_dfs
 def get_download_link(df_dict, filename):
    output_buffer = io.BytesIO()
    with pd.ExcelWriter(output_buffer, engine='openpyxl') as writer:
        for sheet_name, df in df_dict.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    output_buffer.seek(0)
    excel_file = output_buffer.read()
    b64 = base64.b64encode(excel_file).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="Channel_Non_Total_{filename}">Download Channel Non-Total Sheets</a>'
    return href
 def cross_month_analyze_sheets(processed_files):
    st.header("🔍 Cross-Month Sheet Analyzer")
    all_months = list(processed_files.keys())
    st.subheader("Select Files to Compare")
    col1, col2 = st.columns(2)
    with col1:
        month1 = st.selectbox("Select First Month", all_months)
    with col2:
        remaining_months = [m for m in all_months if m != month1]
        month2 = st.selectbox("Select Second Month", remaining_months)
    month1_sheets = list(processed_files[month1].keys())
    month2_sheets = list(processed_files[month2].keys())
    common_sheets = list(set(month1_sheets) & set(month2_sheets))
    if not common_sheets:
        st.warning("No common sheets found between the selected months.")
        return
    selected_sheet = st.selectbox("Select a Common Sheet to Analyze", common_sheets)
    df1 = processed_files[month1][selected_sheet]
    df2 = processed_files[month2][selected_sheet]
    if len(df1.columns) < 5 or len(df2.columns) < 5:
        st.warning("Selected sheets do not have enough columns for analysis.")
        return
    first_four_cols = df1.columns[:4]
    column_filters = {}
    filtered_dfs = {month1: df1.copy(), month2: df2.copy()}
    for i, col in enumerate(first_four_cols):
        unique_values = pd.concat([filtered_dfs[month1][col],filtered_dfs[month2][col]]).dropna().unique()
        column_filters[col] = st.multiselect(f"Select values for {col}", list(unique_values))
        for month in [month1, month2]:
            if column_filters[col]:
                filtered_dfs[month] = filtered_dfs[month][filtered_dfs[month][col].isin(column_filters[col])]
    if st.button("Compare Filtered Data"):
        if len(filtered_dfs[month1]) > 0 and len(filtered_dfs[month2]) > 0:
            st.success(f"Found {len(filtered_dfs[month1])} rows for {month1} and {len(filtered_dfs[month2])} rows for {month2}")
            display_cols = st.multiselect("Select columns to display (from 5th column onwards)", list(df1.columns[4:]),default=list(df1.columns[4:]))
            cols_to_show = list(first_four_cols) + display_cols
            comparison_df = pd.DataFrame()
            for col in cols_to_show:
                comparison_df[f"{month1}_{col}"] = filtered_dfs[month1][col].reset_index(drop=True)
            for col in cols_to_show:
                comparison_df[f"{month2}_{col}"] = filtered_dfs[month2][col].reset_index(drop=True)
            st.dataframe(comparison_df)
            if st.download_button(label="Download Comparison Data",data=comparison_df.to_csv(index=False),file_name=f"{month1}_{month2}_{selected_sheet}_comparison.csv",mime='text/csv'):
                st.success("Comparison data downloaded successfully!")
        else:
            st.warning("No matching rows found after filtering")
 def main():
    st.markdown("""
    <style>
    .big-font {
        font-size:20px !important;
        font-weight: bold;
    }
    .highlight {
        background-color: #f0f2f6;
        padding: 20px;
        border-radius: 10px;
    }
    </style>
    """, unsafe_allow_html=True)
    st.title("📊 Multi-Month Excel Channel Non-Total Sheet Processor")
    st.markdown("""
    <div class="big-font">
    Upload, Process, Analyze, and Compare Sheets Across Multiple Months
    </div>
    """, unsafe_allow_html=True)
    num_months = st.number_input("How many months of files do you want to process?", min_value=1, max_value=12, value=1)
    if 'processed_files' not in st.session_state:
        st.session_state.processed_files = {}
    tabs = st.tabs([f"Month {i+1}" for i in range(num_months)])
    for i in range(num_months):
        with tabs[i]:
            uploaded_files = st.file_uploader(f"Choose Excel files for Month {i+1}", type=['xlsx', 'xls'], accept_multiple_files=True,key=f"file_uploader_{i}")
            if uploaded_files:
                for uploaded_file in uploaded_files:
                    month_filename = f"Month {i+1} - {uploaded_file.name}"
                    
                    # Check if file is not already processed
                    if month_filename not in st.session_state.processed_files:
                        try:
                            # Process the file
                            channel_non_total_dfs = process_excel_file(uploaded_file)
                            
                            # Store processed files in session state
                            st.session_state.processed_files[month_filename] = channel_non_total_dfs
                        except Exception as e:
                            st.error(f"Error processing {month_filename}: {str(e)}")
    tab1, tab2 = st.tabs(["File Processing", "Cross-Month Analyzer"])
    with tab1:
        if st.session_state.processed_files:
            st.header("📁 Processed Files")
            file_tabs = st.tabs(list(st.session_state.processed_files.keys()))
            for i, (filename, file_dfs) in enumerate(st.session_state.processed_files.items()):
                with file_tabs[i]:
                    st.subheader(f"Sheets in {filename}")
                    current_sheets = list(file_dfs.keys())
                    sheets_to_keep = st.multiselect(f"Select sheets to KEEP from {filename}",current_sheets,default=current_sheets)
                    filtered_dfs = {sheet: df for sheet, df in file_dfs.items() if sheet in sheets_to_keep}
                    for sheet_name, df in filtered_dfs.items():
                        with st.expander(f"Sheet: {sheet_name}"):
                            st.dataframe(df)
                    if st.button(f"Download Processed Sheets for {filename}"):
                        download_link = get_download_link(filtered_dfs, filename)
                        st.markdown(download_link, unsafe_allow_html=True)
                    if st.button(f"Remove {filename} from Processed Files"):
                        del st.session_state.processed_files[filename]
                        st.experimental_rerun()
    with tab2:
        if len(st.session_state.processed_files) > 1:
            cross_month_analyze_sheets(st.session_state.processed_files)
        else:
            st.info("Please upload and process files from at least two months to use the cross-month analyzer.")
    st.markdown("---")
    st.markdown("""
    ### 🤔 How to Use
    1. Select number of months to process
    2. Upload Excel files for each month
    3. Automatically process Channel Non-Total sheets
    4. Preview sheets for each file
    5. Use the Cross-Month Analyzer to:
       - Select files from different months
       - Find common sheets
       - Apply filters across both months
       - Compare rows side by side
    """)
 if __name__ == "__main__":
    main()
def pro():
 from openpyxl import load_workbook
 from sklearn.ensemble import RandomForestRegressor
 import calendar
 import warnings
 warnings.filterwarnings('ignore')
 def read_excel_skip_hidden(uploaded_file):
    with open("temp.xlsx", "wb") as f:
        f.write(uploaded_file.getvalue())
    wb = load_workbook(filename="temp.xlsx")
    ws = wb.active
    hidden_rows = [i + 1 for i in range(ws.max_row) if ws.row_dimensions[i + 1].hidden]
    df = pd.read_excel(uploaded_file, skiprows=hidden_rows)
    return df
 def prepare_features(df, current_month_data=None):
    features = pd.DataFrame()
    for month in ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct']:
        features[f'sales_{month}'] = df[f'Monthly Achievement({month})']
    features['prev_sep'] = df['Total Sep 2023']
    features['prev_oct'] = df['Total Oct 2023']
    features['prev_nov'] = df['Total Nov 2023']
    for month in ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct']:
        features[f'month_target_{month}'] = df[f'Month Tgt ({month})']
        features[f'monthly_achievement_rate_{month}'] = (features[f'sales_{month}'] / features[f'month_target_{month}'])
        features[f'ags_target_{month}'] = df[f'AGS Tgt ({month})']
        features[f'ags_achievement_rate_{month}'] = (features[f'sales_{month}'] / features[f'ags_target_{month}'])
    features['month_target_nov'] = df['Month Tgt (Nov)']
    features['ags_target_nov'] = df['AGS Tgt (Nov)']
    features['avg_monthly_achievement_rate'] = features[[f'monthly_achievement_rate_{m}' 
        for m in ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct']]].mean(axis=1)
    features['avg_ags_achievement_rate'] = features[[f'ags_achievement_rate_{m}' 
        for m in ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct']]].mean(axis=1)
    weights = np.array([0.05, 0.1, 0.1, 0.15, 0.2, 0.2, 0.2])
    features['weighted_monthly_achievement_rate'] = np.average(features[[f'monthly_achievement_rate_{m}' for m in ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct']]],weights=weights, axis=1)
    features['weighted_ags_achievement_rate'] = np.average(features[[f'ags_achievement_rate_{m}' for m in ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct']]],weights=weights, axis=1)
    features['avg_monthly_sales'] = features[[f'sales_{m}' for m in ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct']]].mean(axis=1)
    months = ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct']
    for i in range(1, len(months)):
        features[f'growth_{months[i]}'] = features[f'sales_{months[i]}'] / features[f'sales_{months[i-1]}']
    features['yoy_sep_growth'] = features['sales_Sep'] / features['prev_sep']
    features['yoy_oct_growth'] = features['sales_Oct'] / features['prev_oct']
    if current_month_data:
        features['current_month_yoy_growth'] = current_month_data['current_year'] / current_month_data['previous_year']
        features['projected_full_month'] = (current_month_data['current_year'] / current_month_data['days_passed']) * current_month_data['total_days']
        features['current_month_daily_rate'] = current_month_data['current_year'] / current_month_data['days_passed']
        features['yoy_weighted_growth'] = (features['yoy_sep_growth'] * 0.3 + features['yoy_oct_growth'] * 0.4 + features['current_month_yoy_growth'] * 0.3)
    else:
        features['yoy_weighted_growth'] = (features['yoy_sep_growth'] * 0.4 + features['yoy_oct_growth'] * 0.6)
    features['target_achievement_rate'] = features['sales_Oct'] / features['month_target_Oct']
    return features
 def display_historical_data(df_filtered, current_month_data=None):
    if current_month_data:
        historical_data = pd.DataFrame({'Period': ['October 2023', 'November 2023', 'October 2024', f'November 2024 (First {current_month_data["days_passed"]} days)',f'November 2023 (First {current_month_data["days_passed"]} days)'],'Sales': [df_filtered['Total Oct 2023'].iloc[0],df_filtered['Total Nov 2023'].iloc[0],df_filtered['Monthly Achievement(Oct)'].iloc[0],current_month_data['current_year'],current_month_data['previous_year']]})
        historical_data['Growth'] = ['Base',f"{(historical_data['Sales'][1] / historical_data['Sales'][0] - 1) * 100:.1f}% MoM",f"{(historical_data['Sales'][2] / historical_data['Sales'][0] - 1) * 100:.1f}% YoY",f"Current Progress ({(current_month_data['current_year'] / current_month_data['previous_year'] - 1) * 100:.1f}% YoY)",'Previous Year Baseline']
    else:
        historical_data = pd.DataFrame({'Period': ['October 2023', 'November 2023', 'October 2024'],'Sales': [df_filtered['Total Oct 2023'].iloc[0],df_filtered['Total Nov 2023'].iloc[0],df_filtered['Monthly Achievement(Oct)'].iloc[0]]})
        historical_data['Growth'] = ['Base',f"{(historical_data['Sales'][1] / historical_data['Sales'][0] - 1) * 100:.1f}% MoM",f"{(historical_data['Sales'][2] / historical_data['Sales'][0] - 1) * 100:.1f}% YoY"]
    return historical_data
 def calculate_trend_prediction(features, growth_weights, current_month_data=None):
    if current_month_data:
        adjusted_weights = {k: v * 0.7 for k, v in growth_weights.items()}
        adjusted_weights['current_month'] = 0.3
        base_weighted_growth = sum(features[month] * weight for month, weight in adjusted_weights.items() if month != 'current_month') / sum(adjusted_weights.values())
        current_month_growth = features['current_month_yoy_growth'].iloc[0]
        weighted_growth = (base_weighted_growth * 0.7 + current_month_growth * 0.3)
    else:
        weighted_growth = sum(features[month] * weight for month, weight in growth_weights.items()) / sum(growth_weights.values())
    return features['sales_Oct'] * weighted_growth
 def predict_november_sales(df, selected_zone, selected_brand, growth_weights, method_weights, current_month_data=None):
    df_filtered = df[(df['Zone'] == selected_zone) & (df['Brand'] == selected_brand)]
    if len(df_filtered) == 0:
        st.error("No data available for the selected combination of Zone and Brand")
        return None, None
    features = prepare_features(df_filtered, current_month_data)
    historical_data = display_historical_data(df_filtered, current_month_data)
    exclude_columns = ['month_target_nov', 'ags_target_nov','avg_monthly_achievement_rate', 'avg_ags_achievement_rate','weighted_monthly_achievement_rate', 'weighted_ags_achievement_rate','avg_monthly_sales', 'yoy_sep_growth', 'yoy_oct_growth','yoy_weighted_growth']
    feature_cols = [col for col in features.columns if col not in exclude_columns]
    rf_model_monthly = RandomForestRegressor(n_estimators=100, random_state=42)
    rf_model_ags = RandomForestRegressor(n_estimators=100, random_state=42)
    rf_model_monthly.fit(features[feature_cols],features['month_target_nov'] * features['weighted_monthly_achievement_rate'])
    rf_model_ags.fit(features[feature_cols],features['ags_target_nov'] * features['weighted_ags_achievement_rate'])
    rf_prediction_monthly = rf_model_monthly.predict(features[feature_cols])
    rf_prediction_ags = rf_model_ags.predict(features[feature_cols])
    rf_prediction = (rf_prediction_monthly + rf_prediction_ags) / 2
    yoy_prediction = features['prev_nov'] * features['yoy_weighted_growth']
    trend_prediction = calculate_trend_prediction(features, growth_weights, current_month_data)
    if current_month_data:
        target_based_prediction = (features['avg_monthly_sales'] * features['target_achievement_rate'] * (1 + (features['current_month_yoy_growth'] - 1) * 0.3))
    else:
        target_based_prediction = features['avg_monthly_sales'] * features['target_achievement_rate']
    final_prediction = (method_weights['rf'] * rf_prediction +method_weights['yoy'] * yoy_prediction +method_weights['trend'] * trend_prediction +method_weights['target'] * target_based_prediction)
    predictions = pd.DataFrame({'Zone': df_filtered['Zone'],'Brand': df_filtered['Brand'],'RF_Prediction': rf_prediction,'YoY_Prediction': yoy_prediction,'Trend_Prediction': trend_prediction,'Target_Based_Prediction': target_based_prediction,'Final_Prediction': final_prediction})
    return predictions, historical_data
 def main():
    st.markdown("""
        <style>
        .stApp {
            max-width: 1200px;
            margin: 0 auto;
        }
        .metric-card {
            background-color: #f0f2f6;
            padding: 20px;
            border-radius: 10px;
            margin: 10px 0;
        }
        .metric-value {
            font-size: 24px;
            font-weight: bold;
            color: #0066cc;
        }
        </style>
    """, unsafe_allow_html=True)
    st.title("Sales Forecasting Model")
    uploaded_file = st.file_uploader("Upload Excel File", type=['xlsx'])
    if uploaded_file:
        df = read_excel_skip_hidden(uploaded_file)
        col1, col2 = st.columns(2)
        with col1:
            zones = sorted(df['Zone'].unique())
            selected_zone = st.selectbox('Select Zone:', zones)
        with col2:
            brands = sorted(df[df['Zone'] == selected_zone]['Brand'].unique())
            selected_brand = st.selectbox('Select Brand:', brands)
        st.subheader("Current Month Analysis (Optional)")
        use_current_month = st.checkbox("Include current month data")
        current_month_data = None
        if use_current_month:
            col1, col2, col3 = st.columns(3)
            with col1:
                days_passed = st.number_input("Days passed:", min_value=1, max_value=31, value=1)
            with col2:
                current_year_sales = st.number_input("Current year sales:", min_value=0.0, value=0.0)
            with col3:
                previous_year_sales = st.number_input("Previous year sales:", min_value=0.0, value=0.0)
            if days_passed > 0 and current_year_sales > 0 and previous_year_sales > 0:
                total_days = calendar.monthrange(2024, 11)[1]
                current_month_data = {'days_passed': days_passed,'total_days': total_days,'current_year': current_year_sales,'previous_year': previous_year_sales}
        st.subheader("Growth Weights")
        col1, col2 = st.columns(2)
        growth_weights = {}
        with col1:
            growth_weights['growth_May'] = st.slider('May/Apr:', 0.0, 1.0, 0.05, 0.05)
            growth_weights['growth_June'] = st.slider('June/May:', 0.0, 1.0, 0.10, 0.05)
            growth_weights['growth_July'] = st.slider('July/June:', 0.0, 1.0, 0.15, 0.05)
        with col2:
            growth_weights['growth_Aug'] = st.slider('Aug/July:', 0.0, 1.0, 0.20, 0.05)
            growth_weights['growth_Sep'] = st.slider('Sep/Aug:', 0.0, 1.0, 0.25, 0.05)
            growth_weights['growth_Oct'] = st.slider('Oct/Sep:', 0.0, 1.0, 0.25, 0.05)
        st.subheader("Method Weights")
        col1, col2 = st.columns(2)
        method_weights = {}
        with col1:
            method_weights['rf'] = st.slider('Random Forest:', 0.0, 1.0, 0.4, 0.05)
            method_weights['yoy'] = st.slider('Year over Year:', 0.0, 1.0, 0.1, 0.05)
        with col2:
            method_weights['trend'] = st.slider('Trend:', 0.0, 1.0, 0.4, 0.05)
            method_weights['target'] = st.slider('Target:', 0.0, 1.0, 0.1, 0.05)
        if abs(sum(growth_weights.values()) - 1.0) > 0.01:
            st.error("Growth weights should sum to 1")
        elif abs(sum(method_weights.values()) - 1.0) > 0.01:
            st.error("Method weights should sum to 1")
        else:
            predictions, historical_data = predict_november_sales(df, selected_zone, selected_brand, growth_weights, method_weights, current_month_data)
            if predictions is not None and historical_data is not None:
                st.subheader("Historical Sales Data")
                st.dataframe(historical_data.style.format({'Sales': '₹{:,.2f}'}))
                st.subheader("November 2024 Predictions")
                st.dataframe(predictions.style.format({'RF_Prediction': '₹{:,.2f}','YoY_Prediction': '₹{:,.2f}','Trend_Prediction': '₹{:,.2f}','Target_Based_Prediction': '₹{:,.2f}','Final_Prediction': '₹{:,.2f}'}))
                st.subheader("Summary Metrics")
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Average Prediction", f"₹{predictions['Final_Prediction'].mean():,.2f}")
                if current_month_data:
                    current_performance = (current_month_data['current_year'] / current_month_data['previous_year'] - 1) * 100
                    last_year_nov = df[(df['Zone'] == selected_zone) & (df['Brand'] == selected_brand)]['Total Nov 2023'].iloc[0]
                    nov_target = df[(df['Zone'] == selected_zone) & (df['Brand'] == selected_brand)]['Month Tgt (Nov)'].iloc[0]
                    growth_multiplier = current_month_data['current_year'] / current_month_data['previous_year']
                    projected_full_month = last_year_nov * growth_multiplier
                    st.subheader("Current Month Analysis")
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Days Completed", f"{current_month_data['days_passed']} / {current_month_data['total_days']}")
                    with col2:
                        st.metric("Current Performance vs Last Year", f"{current_performance:,.1f}%")
                    with col3:
                        st.metric("Current Year Daily Rate", f"₹{(current_month_data['current_year'] / current_month_data['days_passed']):,.2f}")
                    with st.expander("Detailed Analysis"):
                        col1, col2 = st.columns(2)
                        with col1:
                            st.metric("Previous Year Daily Rate", f"₹{(current_month_data['previous_year'] / current_month_data['days_passed']):,.2f}")
                            st.metric("Projected Full Month", f"₹{projected_full_month:,.2f}")
                        with col2:
                            st.metric("Last Year November Total", f"₹{last_year_nov:,.2f}")
                        remaining_days = current_month_data['total_days'] - current_month_data['days_passed']
                        if remaining_days > 0:
                            st.subheader("Target Analysis")
                            current_daily = current_month_data['current_year'] / current_month_data['days_passed']
                            required_additional_prediction = predictions['Final_Prediction'].iloc[0] - current_month_data['current_year']
                            required_daily_prediction = required_additional_prediction / remaining_days
                            required_additional_target = nov_target - current_month_data['current_year']
                            required_daily_target = required_additional_target / remaining_days
                            col1, col2 = st.columns(2)
                            with col1:
                                st.markdown("**For Predicted Amount:**")
                                st.metric("Required Daily Rate",f"₹{required_daily_prediction:,.2f}")
                                st.metric("Required Growth in Daily Rate",f"{((required_daily_prediction/current_daily) - 1) * 100:,.1f}%")
                            with col2:
                                st.markdown("**For Monthly Target:**")
                                st.metric("Target Amount",f"₹{nov_target:,.2f}")
                                st.metric("Required Daily Rate",f"₹{required_daily_target:,.2f}")
                            st.subheader("Target Achievement Analysis")
                            target_achievement_projected = (projected_full_month / nov_target) * 100
                            col1, col2 = st.columns(2)
                            with col1:
                                st.metric("Projected Achievement at Current Rate",f"{target_achievement_projected:.1f}%")
                            shortfall_or_excess = projected_full_month - nov_target
                            with col2:
                                if shortfall_or_excess < 0:
                                    st.metric("Projected Shortfall",f"₹{abs(shortfall_or_excess):,.2f}")
                                else:
                                    st.metric("Projected Excess",f"₹{shortfall_or_excess:,.2f}")
 if __name__ == "__main__":
    main()
def load_lottie_url(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()
import zipfile
def create_file_management_tab():
    st.markdown("""
        <style>
        .main-header {
            text-align: center;
            padding: 1rem;
            background-color: #f0f2f6;
            border-radius: 10px;
            margin-bottom: 2rem;
        }
        .sub-header {
            color: #0f1116;
            text-align: center;
            padding: 0.5rem;
            margin-bottom: 1rem;
            border-bottom: 2px solid #e6e6e6;
        }
        .success-message {
            padding: 1rem;
            border-radius: 5px;
            background-color: #d1e7dd;
            color: #0f5132;
            text-align: center;
        }
        .warning-message {
            padding: 1rem;
            border-radius: 5px;
            background-color: #fff3cd;
            color: #664d03;
            text-align: center;
        }
        .error-message {
            padding: 1rem;
            border-radius: 5px;
            background-color: #f8d7da;
            color: #842029;
            text-align: center;
        }
        .tool-container {
            background-color: white;
            padding: 1.5rem;
            border-radius: 10px;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
            margin-bottom: 1rem;
        }
        .download-button {
            width: 100%;
            margin-top: 1rem;
            text-align: center;
        }
        </style>
    """, unsafe_allow_html=True)
    if 'current_password' not in st.session_state:
        st.session_state.current_password = None
    if 'current_pdf_name' not in st.session_state:
        st.session_state.current_pdf_name = None
    st.markdown('<div class="main-header"><h1>📁 File Management System</h1></div>', unsafe_allow_html=True)
    row1_col1, row1_col2 = st.columns(2)
    row2_col1, row2_col2 = st.columns(2)
    with row1_col1:
        st.markdown('<div class="tool-container">', unsafe_allow_html=True)
        st.markdown('<div class="sub-header"><h3>📦 Create ZIP</h3></div>', unsafe_allow_html=True)
        uploaded_files = st.file_uploader("Choose files to zip",accept_multiple_files=True,key="zip_files")
        if uploaded_files:
            st.markdown("**Selected files:**")
            for file in uploaded_files:
                st.markdown(f"• {file.name}")
        folder_name = st.text_input("📁 Enter folder name", "my_folder")
        if st.button("🔒 Create ZIP", use_container_width=True):
            if uploaded_files:
                try:
                    os.makedirs(folder_name, exist_ok=True)
                    for file in uploaded_files:
                        file_path = os.path.join(folder_name, file.name)
                        with open(file_path, "wb") as f:
                            f.write(file.getbuffer())
                    zip_path = f"{folder_name}.zip"
                    with zipfile.ZipFile(zip_path, 'w') as zipf:
                        for root, dirs, files in os.walk(folder_name):
                            for file in files:
                                file_path = os.path.join(root, file)
                                arcname = os.path.relpath(file_path, folder_name)
                                zipf.write(file_path, arcname)
                    with open(zip_path, "rb") as f:
                        st.markdown('<div class="download-button">', unsafe_allow_html=True)
                        st.download_button(label="⬇️ Download ZIP",data=f,file_name=zip_path,mime="application/zip",use_container_width=True)
                        st.markdown('</div>', unsafe_allow_html=True)
                    for file in os.listdir(folder_name):
                        os.remove(os.path.join(folder_name, file))
                    os.rmdir(folder_name)
                    os.remove(zip_path)
                    st.markdown('<div class="success-message">✅ ZIP created successfully!</div>', unsafe_allow_html=True)
                except Exception as e:
                    st.markdown(f'<div class="error-message">❌ Error: {str(e)}</div>', unsafe_allow_html=True)
            else:
                st.markdown('<div class="warning-message">⚠️ Please upload files first!</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    with row1_col2:
        st.markdown('<div class="tool-container">', unsafe_allow_html=True)
        st.markdown('<div class="sub-header"><h3>📤 Extract ZIP</h3></div>', unsafe_allow_html=True)
        uploaded_zip = st.file_uploader("Upload ZIP file", type=['zip'], key="unzip_file")
        if uploaded_zip:
            if st.button("📂 Extract Files", use_container_width=True):
                try:
                    zip_bytes = BytesIO(uploaded_zip.read())
                    with zipfile.ZipFile(zip_bytes, 'r') as zip_ref:
                        file_list = zip_ref.namelist()
                        st.markdown("**Extracted files:**")
                        for file_name in file_list:
                            with zip_ref.open(file_name) as file:
                                file_bytes = BytesIO(file.read())
                                st.download_button(label=f"⬇️ {file_name}",data=file_bytes,file_name=file_name,mime="application/octet-stream",use_container_width=True)
                    st.markdown('<div class="success-message">✅ Files extracted successfully!</div>', unsafe_allow_html=True)
                except Exception as e:
                    st.markdown(f'<div class="error-message">❌ Error: {str(e)}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    with row2_col1:
        st.markdown('<div class="tool-container">', unsafe_allow_html=True)
        st.markdown('<div class="sub-header"><h3>🔒 Protect PDF</h3></div>', unsafe_allow_html=True)
        pdf_file = st.file_uploader("Choose PDF file", type=['pdf'], key="pdf_file")
        if pdf_file:
            if pdf_file.name != st.session_state.current_pdf_name:
                st.session_state.current_pdf_name = pdf_file.name
                if st.session_state.get('password_option', '') == "Generate Random 4-digit Password":
                    st.session_state.current_password = str(random.randint(1000, 9999))
            password_option = st.radio("Password Option",["Generate Random 4-digit Password", "Enter Custom Password"],key="password_option",horizontal=True)
            if password_option == "Generate Random 4-digit Password":
                if not st.session_state.current_password:
                    st.session_state.current_password = str(random.randint(1000, 9999))
                st.info(f"🔑 Generated Password: **{st.session_state.current_password}**")
                password = st.session_state.current_password
            else:
                password = st.text_input("🔑 Enter password", type="password")
                st.session_state.current_password = password
            if st.button("🔒 Protect PDF", use_container_width=True):
                try:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    pdf_writer = PyPDF2.PdfWriter()
                    for page in pdf_reader.pages:
                        pdf_writer.add_page(page)
                    pdf_writer.encrypt(password)
                    output_pdf = BytesIO()
                    pdf_writer.write(output_pdf)
                    output_pdf.seek(0)
                    st.download_button(label="⬇️ Download Protected PDF",data=output_pdf,file_name=f"protected_{pdf_file.name}",mime="application/pdf",use_container_width=True)
                    st.markdown('<div class="success-message">✅ PDF protected successfully!</div>', unsafe_allow_html=True)
                    if password_option == "Generate Random 4-digit Password":
                        st.info("📝 Make sure to save the password!")
                except Exception as e:
                    st.markdown(f'<div class="error-message">❌ Error: {str(e)}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    with row2_col2:
        st.markdown('<div class="tool-container">', unsafe_allow_html=True)
        st.markdown('<div class="sub-header"><h3>📊 Extract Tables from PDF</h3></div>', unsafe_allow_html=True)
        uploaded_file = st.file_uploader("Upload PDF file", type=['pdf'], key="table_pdf")
        if uploaded_file:
            try:
                doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                tables_found = []
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    tables = page.find_tables()
                    if tables.tables:
                        tables_found.extend([{'page': page_num + 1,'table': table} for table in tables.tables])
                if tables_found:
                    st.markdown(f"📋 Found **{len(tables_found)}** tables in the PDF")
                    selected_tables = st.multiselect("Select tables to extract",options=range(len(tables_found)),format_func=lambda x: f"Table {x+1} (Page {tables_found[x]['page']})")
                    if selected_tables and st.button("📥 Extract Selected Tables", use_container_width=True):
                        output = BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            for i in selected_tables:
                                table = tables_found[i]['table']
                                df = pd.DataFrame(table.extract())
                                df.to_excel(writer,sheet_name=f"Table_{i+1}_Page_{tables_found[i]['page']}",index=False)
                        st.download_button(label="⬇️ Download Excel File",data=output.getvalue(),file_name="extracted_tables.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.document",use_container_width=True)
                        st.markdown('<div class="success-message">✅ Tables extracted successfully!</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div class="warning-message">⚠️ No tables found in the PDF</div>', unsafe_allow_html=True)
            except Exception as e:
                st.markdown(f'<div class="error-message">❌ Error: {str(e)}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
def create_privacy_section():
    st.markdown("""
        <style>
        .privacy-header {
            text-align: center;
            padding: 1.5rem;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border-radius: 10px;
            margin-bottom: 2rem;
        }
        .section-card {
            background-color: white;
            padding: 1.5rem;
            border-radius: 10px;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
            margin-bottom: 1.5rem;
        }
        .section-header {
            color: #4a5568;
            border-bottom: 2px solid #e2e8f0;
            padding-bottom: 0.5rem;
            margin-bottom: 1rem;
        }
        .value-card {
            background-color: #f7fafc;
            padding: 1rem;
            border-radius: 8px;
            border-left: 4px solid #4299e1;
            margin-bottom: 1rem;
        }
        .highlight-text {
            color: #2b6cb0;
            font-weight: bold;
        }
        .feature-list {
            list-style-type: none;
            padding-left: 0;
        }
        .feature-item {
            padding: 0.5rem 0;
            display: flex;
            align-items: center;
            gap: 0.5rem;
        }
        </style>
    """, unsafe_allow_html=True)
    st.markdown("""<div class="privacy-header"><h1>🔒 Data Privacy & Terms of Service</h1><p>Your privacy and security are our top priorities</p></div>""", unsafe_allow_html=True)
    st.markdown("""
        <div class="section-card">
            <h2 class="section-header">📱 About Our Data Central</h2>
            <p>Our application is designed to provide secure, efficient, and user-friendly file management solutions. 
            We understand the importance of your data and have built this system with privacy and security at its core.</p>
            <h3>🎯 Our Core Values</h3>
            <div class="value-card">
                <p><span class="highlight-text">Privacy First:</span> Your data privacy is non-negotiable.</p>
            </div>
            <div class="value-card">
                <p><span class="highlight-text">Transparency:</span> We're clear about how our system works.</p>
            </div>
            <div class="value-card">
                <p><span class="highlight-text">Security:</span> Industry-standard security measures protect your files.</p>
            </div>
            <div class="value-card">
                <p><span class="highlight-text">User Control:</span> You maintain complete control over your files.</p>
            </div>
        </div>
    """, unsafe_allow_html=True)
    st.markdown("""
        <div class="section-card">
            <h2 class="section-header">🛡️ Our Data Privacy Commitments</h2>
            <ul class="feature-list">
                <li class="feature-item">
                    ✨ <strong>No Data Storage:</strong> We do not store any of your uploaded files or processed data on our servers
                </li>
                <li class="feature-item">
                    ✨ <strong>Temporary Processing:</strong> Files are only held in temporary memory during processing and are immediately deleted afterward
                </li>
                <li class="feature-item">
                    ✨ <strong>No Data Collection:</strong> We don't collect personal information or tracking data
                </li>
                <li class="feature-item">
                    ✨ <strong>Secure Processing:</strong> All file processing happens locally in your browser session
                </li>
                <li class="feature-item">
                    ✨ <strong>No Third-Party Sharing:</strong> Your data is never shared with third parties
                </li>
            </ul>
        </div>
    """, unsafe_allow_html=True)
    st.markdown("""
        <div class="section-card">
            <h2 class="section-header">🔐 Security Features</h2>
            <ul class="feature-list">
                <li class="feature-item">
                    🛠️ <strong>Secure File Processing:</strong> All file operations are performed in-memory
                </li>
                <li class="feature-item">
                    🛠️ <strong>Client-Side Processing:</strong> Files are processed locally on your device
                </li>
                <li class="feature-item">
                    🛠️ <strong>Industry-Standard Encryption:</strong> For PDF password protection
                </li>
                <li class="feature-item">
                    🛠️ <strong>Secure Downloads:</strong> Direct file downloads without server storage
                </li>
            </ul>
        </div>
    """, unsafe_allow_html=True)
    st.markdown("""
        <div class="section-card">
            <h2 class="section-header">📜 Terms of Service</h2>
            <h4>Usage Agreement</h4>
            <p>By using our File Management System, you agree to:</p>
            <ul class="feature-list">
                <li class="feature-item">
                    📋 Use the service for legal purposes only
                </li>
                <li class="feature-item">
                    📋 Not attempt to circumvent any security features
                </li>
                <li class="feature-item">
                    📋 Accept responsibility for the files you process
                </li>
                <li class="feature-item">
                    📋 Understand that we provide no warranty for the service
                </li>
            </ul>
            <h4>Limitations of Liability</h4>
            <p>We strive to provide a reliable service but cannot guarantee:</p>
            <ul class="feature-list">
                <li class="feature-item">
                    ⚠️ Uninterrupted service availability
                </li>
                <li class="feature-item">
                    ⚠️ Perfect accuracy in table extraction
                </li>
                <li class="feature-item">
                    ⚠️ Compatibility with all file formats
                </li>
            </ul>
        </div>
    """, unsafe_allow_html=True)
    st.markdown("""
        <div class="section-card">
            <h2 class="section-header">✅ Acceptable Use Policy</h2>
            <p>To maintain the security and reliability of our service, we require users to:</p>
            <ul class="feature-list">
                <li class="feature-item">
                    ✔️ Only upload files you have the right to process
                </li>
                <li class="feature-item">
                    ✔️ Respect file size limitations
                </li>
                <li class="feature-item">
                    ✔️ Not attempt to upload malicious files
                </li>
                <li class="feature-item">
                    ✔️ Use the service in a manner that doesn't disrupt other users
                </li>
            </ul>
        </div>
    """, unsafe_allow_html=True)
    st.markdown("""
        <div class="section-card">
            <h2 class="section-header">📞 Contact & Support</h2>
            <p>If you have any questions about our privacy policy or terms of service, please contact us:</p>
            <ul class="feature-list">
                <li class="feature-item">
                    📧 Email: prasoon.bajpai@lc.jkmail.com
                </li>
                <li class="feature-item">
                    ⏰ Response Time: Within 24 hours
                </li>
            </ul>
        </div>
    """, unsafe_allow_html=True)
def integrate_privacy_section():
    # Add this to your main app's sidebar or as a separate page
    st.sidebar.markdown("---")
    if st.sidebar.button("📜 View Privacy Policy & Terms"):
        create_privacy_section()
def process_pdf(input_pdf, operations):
    from PyPDF2 import PdfReader, PdfWriter
    from io import BytesIO
    writer = PdfWriter()
    reader = PdfReader(input_pdf)
    if "extract" in operations:
        selected_pages = operations["extract"]["pages"]  
        for page_num in selected_pages:
            if 0 <= page_num - 1 < len(reader.pages): 
                writer.add_page(reader.pages[page_num - 1])
    else:
        for page in reader.pages:
            writer.add_page(page)
    if "merge" in operations and operations["merge"]["files"]:
        # Add pages from additional PDFs
        for additional_pdf in operations["merge"]["files"]:
            merge_reader = PdfReader(additional_pdf)
            for page in merge_reader.pages:
                writer.add_page(page)
    if len(writer.pages) == 0:
        return BytesIO(input_pdf.read())
    pdf_width = float(writer.pages[0].mediabox.width)
    pdf_height = float(writer.pages[0].mediabox.height)
    transformed_writer = PdfWriter()
    for i in range(len(writer.pages)):
        page = writer.pages[i]
        if "resize" in operations:
            scale = operations["resize"]["scale"] / 100
            page.scale(scale, scale)
        if "crop" in operations:
            left = operations["crop"]["left"] * pdf_width / 100
            bottom = operations["crop"]["bottom"] * pdf_height / 100
            right = operations["crop"]["right"] * pdf_width / 100
            top = operations["crop"]["top"] * pdf_height / 100
            page.cropbox.lower_left = (left, bottom)
            page.cropbox.upper_right = (right, top)  
        if "rotate" in operations:
            angle = operations["rotate"]["angle"]
            page.rotate(angle)
        transformed_writer.add_page(page)
    output = BytesIO()
    transformed_writer.write(output)
    return output
def add_watermark(pdf_writer, watermark_options):
    watermark_buffer = BytesIO()
    c = Canvas(watermark_buffer)
    first_page = pdf_writer.pages[0]
    page_width = float(first_page.mediabox.width)
    page_height = float(first_page.mediabox.height)  
    if watermark_options["type"] == "text":
        text = watermark_options["text"]
        color = watermark_options["color"]
        font_size = watermark_options["size"]
        opacity = watermark_options["opacity"]
        angle = watermark_options["angle"]
        position = watermark_options["position"]
        r = int(color[1:3], 16) / 255
        g = int(color[3:5], 16) / 255
        b = int(color[5:7], 16) / 255
        c.setFillColor(Color(r, g, b, alpha=opacity))
        c.setFont("Helvetica", font_size)
        if position == "center":
            x, y = page_width/2, page_height/2
        elif position == "top-left":
            x, y = 50, page_height-50
        elif position == "top-right":
            x, y = page_width-50, page_height-50
        elif position == "bottom-left":
            x, y = 50, 50
        elif position == "bottom-right":
            x, y = page_width-50, 50
        c.saveState()
        c.translate(x, y)
        c.rotate(angle)
        c.drawString(-len(text)*font_size/4, 0, text)
        c.restoreState()  
    else:
        image = Image.open(watermark_options["image"])
        opacity = watermark_options["opacity"]
        angle = watermark_options["angle"]
        position = watermark_options["position"]
        size = watermark_options["size"]
        img_width = page_width * size / 100
        img_height = img_width * image.height / image.width
        if position == "center":
            x, y = (page_width-img_width)/2, (page_height-img_height)/2
        elif position == "top-left":
            x, y = 0, page_height-img_height
        elif position == "top-right":
            x, y = page_width-img_width, page_height-img_height
        elif position == "bottom-left":
            x, y = 0, 0
        elif position == "bottom-right":
            x, y = page_width-img_width, 0
        img_buffer = BytesIO()
        if image.mode != 'RGBA':
            image = image.convert('RGBA')
        image.putalpha(int(opacity * 255))
        image.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        c.saveState()
        c.translate(x + img_width/2, y + img_height/2)
        c.rotate(angle)
        c.translate(-img_width/2, -img_height/2)
        c.drawImage(ImageReader(img_buffer), 0, 0, width=img_width, height=img_height)
        c.restoreState()
    c.save()
    watermark_buffer.seek(0)
    watermark_pdf = PdfReader(watermark_buffer)
    selected_pages = watermark_options.get("pages", "all")
    for i, page in enumerate(pdf_writer.pages):
        if selected_pages == "all" or (i+1) in selected_pages:
            page.merge_page(watermark_pdf.pages[0])
    return pdf_writer
def get_pdf_preview(pdf_file, page_num=0):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    page = doc[page_num]
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    return img
def get_image_size_metrics(original_image_bytes, processed_image_bytes):
    original_size = len(original_image_bytes) / 1024 
    processed_size = len(processed_image_bytes) / 1024
    size_change = ((original_size - processed_size) / original_size) * 100  
    return {'original_size': original_size,'processed_size': processed_size,'size_change': size_change}
def process_image(image, operations):
    if "resize" in operations:
        width = operations["resize"]["width"]
        height = operations["resize"]["height"]
        image = image.resize((width, height), Image.Resampling.LANCZOS)
    if "compress" in operations:
        quality = operations["compress"]["quality"]
        return image, quality
    if "crop" in operations:
        left = operations["crop"]["left"]
        top = operations["crop"]["top"]
        right = operations["crop"]["right"]
        bottom = operations["crop"]["bottom"]
        image = image.crop((left, top, right, bottom))
    if "rotate" in operations:
        angle = operations["rotate"]["angle"]
        image = image.rotate(angle, expand=True)
    if "brightness" in operations:
        factor = operations["brightness"]["factor"]
        enhancer = ImageEnhance.Brightness(image)
        image = enhancer.enhance(factor)
    if "contrast" in operations:
        factor = operations["contrast"]["factor"]
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(factor)
    return image, None
def convert_uploadedfile_to_image(uploaded_file):
    if uploaded_file is None:
        return None
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, uploaded_file.name)
    with open(temp_file_path, 'wb') as f:
        f.write(uploaded_file.getvalue())
    return temp_file_path
TEMPLATES = {"Classic Professional": {"background_type": "Color","background_color": "#FFFFFF","design_elements": ["Border"],"border_color": "#000000","border_width": 1,"title_font": "Helvetica-Bold","title_color": "#000000","accent_color": "#4A4A4A","layout": "centered"},
             "Modern Minimal": {"background_type": "Color","background_color": "#FFFFFF","design_elements": ["Accent Bar"],"accent_color": "#2C3E50","title_font": "Helvetica","title_color": "#2C3E50","layout": "left-aligned"},
             "Corporate Blue": {"background_type": "Gradient","gradient_start": "#E8F0FE","gradient_end": "#FFFFFF","design_elements": ["Corner Lines", "Header Bar"],"accent_color": "#1B4F72","title_font": "Helvetica-Bold","title_color": "#1B4F72","layout": "centered"},
              "Creative Bold": {"background_type": "Color","background_color": "#FFFFFF","design_elements": ["Diagonal Lines", "Side Bar"],"accent_color": "#FF5733","title_font": "Helvetica-Bold","title_color": "#2C3E50","layout": "asymmetric"},
            "Executive Elite": {"background_type": "Color","background_color": "#F5F5F5","design_elements": ["Gold Accents", "Double Border"],"accent_color": "#D4AF37","border_color": "#000000","title_font": "Times-Bold","title_color": "#000000","layout": "centered"}}
def register_custom_fonts():
    try:
        custom_fonts = [("Roboto-Regular.ttf", "Roboto"),("Montserrat-Regular.ttf", "Montserrat"),("OpenSans-Regular.ttf", "OpenSans")]
        for font_file, font_name in custom_fonts:
            if not font_name in pdfmetrics.getRegisteredFontNames():
                font_path = f"fonts/{font_file}"
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
    except Exception as e:
        st.warning(f"Some custom fonts couldn't be loaded: {str(e)}")
def draw_design_elements(c, options, width, height):
    margin = cm
    if "Border" in options["design_elements"]:
        c.setStrokeColor(options["border_color"])
        c.setLineWidth(options["border_width"])
        c.rect(margin, margin, width - 2*margin, height - 2*margin)
    if "Double Border" in options["design_elements"]:
        c.setStrokeColor(options["border_color"])
        c.setLineWidth(options["border_width"])
        c.rect(margin, margin, width - 2*margin, height - 2*margin)
        inner_margin = margin + 0.5*cm
        c.rect(inner_margin, inner_margin, width - 2*inner_margin, height - 2*inner_margin)
    if "Corner Lines" in options["design_elements"]:
        c.setStrokeColor(options["accent_color"])
        c.setLineWidth(2)
        corner_size = 3*cm
        for x, y in [(margin, height-margin), (width-margin, height-margin),
                     (margin, margin), (width-margin, margin)]:
            c.saveState()
            c.translate(x, y)
            if y > height/2:
                c.rotate(180)
            c.lines([(0, 0, corner_size, 0), (0, 0, 0, -corner_size)])
            c.restoreState()
    if "Diagonal Lines" in options["design_elements"]:
        c.setStrokeColor(options["accent_color"])
        c.setLineWidth(1)
        spacing = 1*cm
        for i in range(int(height/(2*spacing))):
            y = i * 2*spacing
            c.line(0, y, 2*cm, y + 2*cm)
    if "Side Bar" in options["design_elements"]:
        c.setFillColor(options["accent_color"])
        c.rect(0, 0, 2*cm, height, fill=1)
    if "Header Bar" in options["design_elements"]:
        c.setFillColor(options["accent_color"])
        c.rect(0, height-3*cm, width, 3*cm, fill=1)
    if "Accent Bar" in options["design_elements"]:
        c.setFillColor(options["accent_color"])
        bar_width = 0.5*cm
        c.rect(margin, height/2, width - 2*margin, bar_width, fill=1)
    if "Gold Accents" in options["design_elements"]:
        c.setStrokeColor(options["accent_color"])
        c.setLineWidth(1)
        pattern_size = 1*cm
        for i in range(4):
            x = margin + i * pattern_size
            c.line(x, height-margin, x + pattern_size, height-margin-pattern_size)
            c.line(x, margin, x + pattern_size, margin+pattern_size)
def draw_watermark(c, options, width, height):
    if options.get("watermark_text"):
        c.saveState()
        c.translate(width/2, height/2)
        c.rotate(45)
        c.setFont(options["watermark_font"], options["watermark_size"])
        c.setFillColor(colors.Color(0, 0, 0, alpha=0.1))
        c.drawCentredString(0, 0, options["watermark_text"])
        c.restoreState()
def create_front_page(options):
    buffer = io.BytesIO()
    page_size = {"A4": A4,"A4 Landscape": landscape(A4),"Letter": letter,"Letter Landscape": landscape(letter),"Legal": legal}[options["page_size"]]
    c = canvas.Canvas(buffer, pagesize=page_size)
    width, height = page_size
    if options.get("template"):
        template = TEMPLATES[options["template"]]
        options = {**template, **options}
    if options["background_type"] == "Color":
        c.setFillColor(options["background_color"])
        c.rect(0, 0, width, height, fill=True)
    elif options["background_type"] == "Gradient":
        steps = 100
        for i in range(steps):
            r = options["gradient_start"].red + (options["gradient_end"].red - options["gradient_start"].red) * i / steps
            g = options["gradient_start"].green + (options["gradient_end"].green - options["gradient_start"].green) * i / steps
            b = options["gradient_start"].blue + (options["gradient_end"].blue - options["gradient_start"].blue) * i / steps
            c.setFillColor((r, g, b))
            c.rect(0, height * i / steps, width, height / steps, fill=True)
    elif options["background_type"] == "Pattern":
        pattern_size = 1*cm
        c.setStrokeColor(colors.Color(0, 0, 0, alpha=0.1))
        for x in range(0, int(width), int(pattern_size)):
            for y in range(0, int(height), int(pattern_size)):
                if (x + y) % (2 * int(pattern_size)) == 0:
                    c.rect(x, y, pattern_size, pattern_size, fill=True)
    draw_design_elements(c, options, width, height)
    if options.get("watermark_text"):
        draw_watermark(c, options, width, height)
    if options.get("logo"):
        try:
            logo_path = convert_uploadedfile_to_image(options["logo"])
            if logo_path:
                logo_img = Image.open(logo_path)
                aspect = logo_img.height / logo_img.width
                logo_width = options["logo_width"]
                logo_height = logo_width * aspect
                if options["layout"] == "centered":
                    x = (width - logo_width) / 2
                    y = height - logo_height - 3*cm
                elif options["layout"] == "left-aligned":
                    x = 3*cm
                    y = height - logo_height - 3*cm
                elif options["layout"] == "asymmetric":
                    x = width - logo_width - 3*cm
                    y = height - logo_height - 3*cm
                c.drawImage(logo_path, x, y, width=logo_width, height=logo_height)
                os.unlink(logo_path)
                os.rmdir(os.path.dirname(logo_path))
        except Exception as e:
            st.error(f"Error processing logo: {str(e)}")
    c.setFont(options["title_font"], options["title_size"])
    c.setFillColor(options["title_color"])
    title_lines = options["title"].split('\n')
    if options["layout"] == "centered":
        title_height = (height + len(title_lines) * options["title_size"]) / 2
    elif options["layout"] == "left-aligned":
        title_height = height - 5*cm
    else:  # asymmetric
        title_height = (height + len(title_lines) * options["title_size"]) / 1.5
    for line in title_lines:
        title_width = c.stringWidth(line, options["title_font"], options["title_size"])
        if options["layout"] == "centered":
            x = (width - title_width) / 2
        elif options["layout"] == "left-aligned":
            x = 3*cm
        else:  # asymmetric
            x = width - title_width - 3*cm
        c.drawString(x, title_height, line)
        title_height -= options["title_size"] * 1.2
    if options["subtitle"]:
        c.setFont(options["subtitle_font"], options["subtitle_size"])
        c.setFillColor(options["subtitle_color"])
        subtitle_width = c.stringWidth(options["subtitle"], options["subtitle_font"], options["subtitle_size"])
        if options["layout"] == "centered":
            x = (width - subtitle_width) / 2
        elif options["layout"] == "left-aligned":
            x = 3*cm
        else: 
            x = width - subtitle_width - 3*cm
        c.drawString(x, title_height - cm, options["subtitle"])
    y_position = title_height - 4*cm
    for text_block in options["text_blocks"]:
        if text_block["text"]:
            c.setFont(text_block["font"], text_block["size"])
            c.setFillColor(text_block["color"])
            lines = text_block["text"].split('\n')
            for line in lines:
                text_width = c.stringWidth(line, text_block["font"], text_block["size"])
                if options["layout"] == "centered":
                    x = (width - text_width) / 2
                elif options["layout"] == "left-aligned":
                    x = 3*cm
                else:  # asymmetric
                    x = width - text_width - 3*cm
                c.drawString(x, y_position, line)
                y_position -= text_block["size"] * 1.5
    if options["show_date"] or options["footer_text"]:
        footer_font = options.get("footer_font", "Helvetica")
        footer_size = options.get("footer_size", 10)
        c.setFont(footer_font, footer_size)
        c.setFillColor(options.get("footer_color", colors.black))
        footer_elements = []
        if options["show_date"]:
            date_format = options.get("date_format", "%B %d, %Y")
            footer_elements.append(datetime.now().strftime(date_format))
        if options["footer_text"]:
            footer_elements.append(options["footer_text"])
        footer_text = " | ".join(footer_elements)
        footer_width = c.stringWidth(footer_text, footer_font, footer_size)
        if options["layout"] == "centered":
            x = (width - footer_width) / 2
        elif options["layout"] == "left-aligned":
            x = 3*cm
        else:  # asymmetric
            x = width - footer_width - 3*cm
        c.drawString(x, 2*cm, footer_text)
    c.save()
    buffer.seek(0)
    return buffer
def front_page_creator():
    st.header("📄 Professional Front Page Creator")
    st.subheader("Choose a Template")
    template = st.selectbox("Select a Template", 
        ["Custom"] + list(TEMPLATES.keys()),
        help="Choose a pre-designed template or create your own custom design")
    with st.expander("Preview Template", expanded=False):
        st.write("Template Preview would appear here")
    with st.container():
            st.subheader("Basic Settings")
            col1, col2 = st.columns(2)
            with col1:
                page_size = st.selectbox(
                    "Page Size", 
                    ["A4", "A4 Landscape", "Letter", "Letter Landscape", "Legal"],
                    help="Choose the size and orientation of your front page")                
                title = st.text_area(
                    "Title",
                    placeholder="Enter title (can be multiple lines)\nUse new lines for multi-line titles",
                    help="Main title of your front page. Use new lines for multiple lines of text")
                subtitle = st.text_input(
                    "Subtitle",
                    placeholder="Enter subtitle (optional)",
                    help="Optional subtitle that appears below the main title")
            with col2:
                title_font = st.selectbox(
                    "Title Font",
                    ["Helvetica", "Helvetica-Bold", "Times-Roman", "Times-Bold", 
                     "Courier", "Courier-Bold", "Roboto", "Montserrat", "OpenSans"],
                    help="Choose the font for your title")
                title_size = st.slider("Title Size",20, 72, 48,help="Adjust the size of your title text")
                title_color = st.color_picker("Title Color","#000000",help="Choose the color for your title")
            st.subheader("Layout Settings")
            col3, col4 = st.columns(2)
            with col3:
                layout_style = st.selectbox("Layout Style",["centered", "left-aligned", "asymmetric"],help="Choose how your content is aligned on the page")
                content_spacing = st.slider("Content Spacing",1.0, 3.0, 1.5,0.1,help="Adjust the spacing between content elements")
            with col4:
                margins = st.slider("Page Margins (cm)",1.0, 5.0, 2.5,0.5,help="Adjust the margins around your content")
            st.subheader("Background Settings")
            background_type = st.radio("Background Type",["Color", "Gradient", "Pattern", "None"],help="Choose the type of background for your front page")
            if background_type == "Color":
                background_color = st.color_picker("Background Color","#FFFFFF",help="Choose a solid color for your background")
            elif background_type == "Gradient":
                col5, col6 = st.columns(2)
                with col5:
                    gradient_start = st.color_picker("Gradient Start Color","#FFFFFF",help="Choose the starting color for your gradient")
                    gradient_direction = st.selectbox("Gradient Direction",["Top to Bottom", "Left to Right", "Diagonal"],help="Choose the direction of your gradient")
                with col6:
                    gradient_end = st.color_picker("Gradient End Color","#E0E0E0",help="Choose the ending color for your gradient")
            elif background_type == "Pattern":
                col7, col8 = st.columns(2)
                with col7:
                    pattern_type = st.selectbox("Pattern Type",["Dots", "Lines", "Grid", "Chevron"],help="Choose the type of pattern")
                    pattern_color = st.color_picker("Pattern Color","#E0E0E0",help="Choose the color for your pattern")
                with col8:
                    pattern_opacity = st.slider("Pattern Opacity",0.0, 1.0, 0.1,0.1,help="Adjust the opacity of the pattern")
                    pattern_size = st.slider("Pattern Size",0.5, 3.0, 1.0,0.1,help="Adjust the size of the pattern elements")
            st.subheader("Logo/Image Settings")
            logo = st.file_uploader("Upload Logo/Image",type=["png", "jpg", "jpeg"],help="Upload your organization's logo or an image")
            if logo:
                col9, col10 = st.columns(2)
                with col9:
                    logo_width = st.slider("Logo Width",50, 400, 200,help="Adjust the width of your logo")
                    logo_opacity = st.slider("Logo Opacity",0.1, 1.0, 1.0,0.1,help="Adjust the opacity of your logo")
                with col10:
                    logo_position = st.selectbox("Logo Position",["Top Center", "Top Left", "Top Right", "Bottom Center", "Bottom Left", "Bottom Right"],help="Choose where to place your logo")
                    logo_padding = st.slider("Logo Padding (cm)",0.5, 5.0, 2.0,0.5,help="Adjust the space around your logo")
            st.subheader("Design Elements")
            design_elements = st.multiselect("Add Design Elements",["Border", "Double Border", "Corner Lines", "Diagonal Lines", "Side Bar", "Header Bar", "Accent Bar", "Gold Accents"],default=["Border"],help="Choose decorative elements to enhance your design")
            if any(design_elements):
                col11, col12 = st.columns(2)
                with col11:
                    accent_color = st.color_picker("Accent Color","#000000",help="Choose the color for decorative elements")
                if "Border" in design_elements or "Double Border" in design_elements:
                    with col12:
                        border_color = st.color_picker("Border Color","#000000",help="Choose the color for the border")
                        border_width = st.slider("Border Width",0.5, 5.0, 1.0,0.5,help="Adjust the thickness of the border")
            st.subheader("Additional Text Blocks")
            num_blocks = st.number_input("Number of Additional Text Blocks", 0, 5, 0)
            text_blocks = []
            for i in range(num_blocks):
             st.markdown(f"#### Text Block {i+1}")
             col13, col14 = st.columns(2)
             with col13:
                block_text = st.text_input(f"Text for Block {i+1}")
                block_font = st.selectbox(f"Font for Block {i+1}", ["Helvetica", "Times-Roman", "Courier"])
             with col14:
                block_size = st.slider(f"Size for Block {i+1}", 8, 36, 12)
                block_color = st.color_picker(f"Color for Block {i+1}", "#000000")
             text_blocks.append({"text": block_text,"font": block_font,"size": block_size,"color": block_color})
            st.subheader("Watermark Settings")
            add_watermark = st.checkbox("Add Watermark",help="Add a watermark to your front page")
            if add_watermark:
                col15, col16 = st.columns(2)
                with col15:
                    watermark_text = st.text_input("Watermark Text",placeholder="Enter watermark text",help="Text to use as watermark")
                    watermark_font = st.selectbox("Watermark Font",["Helvetica", "Times-Roman", "Courier"],help="Choose the font for your watermark")
                with col16:
                    watermark_size = st.slider("Watermark Size",20, 100, 60,help="Adjust the size of your watermark")
                    watermark_opacity = st.slider("Watermark Opacity",0.0, 1.0, 0.1,0.1,help="Adjust the opacity of your watermark")
            st.subheader("Footer Options")
            col15, col16 = st.columns(2)
            with col15:
                show_date = st.checkbox("Show Date",help="Include the current date in the footer")
                if show_date:
                    date_format = st.selectbox("Date Format",["%B %d, %Y", "%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"],help="Choose how the date should be displayed")
            with col16:
                footer_text = st.text_input("Custom Footer Text",placeholder="Enter custom footer text",help="Add custom text to the footer")
                footer_alignment = st.selectbox("Footer Alignment",["Center", "Left", "Right"],help="Choose how the footer is aligned")
            if st.button("Generate Front Page", help="Click to create your front page"):
                if not title:
                    st.error("Please enter a title for your front page.")
                    return
                try:
                    options = {"template": template if template != "Custom" else None,"page_size": page_size,"title": title,"subtitle": subtitle,"title_font": title_font,
                        "title_size": title_size,"title_color": colors.HexColor(title_color),"subtitle_font": title_font,"subtitle_size": int(title_size * 0.6),"subtitle_color": colors.HexColor(title_color),
                        "layout": layout_style,"content_spacing": content_spacing,"margins": margins,"background_type": background_type,"background_color": colors.HexColor(background_color) if background_type == "Color" else None,
                        "gradient_start": colors.HexColor(gradient_start) if background_type == "Gradient" else None,"gradient_end": colors.HexColor(gradient_end) if background_type == "Gradient" else None,
                        "gradient_direction": gradient_direction if background_type == "Gradient" else None,"pattern_type": pattern_type if background_type == "Pattern" else None,
                        "pattern_color": colors.HexColor(pattern_color) if background_type == "Pattern" else None,"pattern_opacity": pattern_opacity if background_type == "Pattern" else None,
                        "pattern_size": pattern_size if background_type == "Pattern" else None,"logo": logo,"logo_width": logo_width if logo else None,"logo_position": logo_position if logo else None,"logo_opacity": logo_opacity if logo else None,
                        "logo_padding": logo_padding if logo else None,"design_elements": design_elements,"accent_color": colors.HexColor(accent_color) if any(design_elements) else None,
                        "border_color": colors.HexColor(border_color) if "Border" in design_elements or "Double Border" in design_elements else None,"border_width": border_width if "Border" in design_elements or "Double Border" in design_elements else None,
                        "watermark_text": watermark_text if add_watermark else None,"watermark_font": watermark_font if add_watermark else None,"watermark_size": watermark_size if add_watermark else None,
                        "watermark_opacity": watermark_opacity if add_watermark else None,"show_date": show_date,"date_format": date_format if show_date else "%B %d, %Y",
                        "text_blocks": text_blocks,"footer_text": footer_text,"footer_alignment": footer_alignment,}
                    pdf_buffer = create_front_page(options)
                    filename = f"{title.split()[0].lower()}_front_page.pdf"
                    st.download_button(
                        label="📥 Download Front Page",
                        data=pdf_buffer,
                        file_name=filename,
                        mime="application/pdf",
                        help="Download your generated front page as a PDF")
                    st.success("✨ PDF generated successfully! Click the download button above to save your front page.")
                    st.info("💡 Tip: After downloading, you may want to preview the PDF to ensure everything looks perfect.")
                except Exception as e:
                    st.error(f"Error generating front page: {str(e)}")
                    st.info("Please try adjusting your settings or contact support if the problem persists.")
def excel_editor_and_analyzer():
    st.title("🧩 Data Central")
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["Excel Editor","File Converter","Data Analyzer","Front Page Creator","File Management","Terms & Conditions"])
    with tab1:
        excel_editor()
    with tab2:
        file_converter()
    with tab3:
        data_analyzer()
    with tab4:
        front_page_creator()
    with tab5:
        create_file_management_tab()
    with tab6:
        create_privacy_section()
def file_converter():
    st.header("🔄 Universal File Converter")
    st.markdown("""
        <style>
        .converter-card {
            background-color: #f8f9fa;
            padding: 1.5rem;
            border-radius: 0.5rem;
            margin: 1rem 0;
            border: 1px solid #e9ecef;
        }
        .stButton>button {
            width: 100%;
            margin-top: 1rem;
        }
        .success-message {
            color: #28a745;
            padding: 0.75rem;
            border-radius: 0.25rem;
            margin: 1rem 0;
            background-color: #d4edda;
            border: 1px solid #c3e6cb;
        }
        .info-message {
            color: #0c5460;
            padding: 0.75rem;
            border-radius: 0.25rem;
            margin: 1rem 0;
            background-color: #d1ecf1;
            border: 1px solid #bee5eb;
        }
        .conversion-stats {
            padding: 1rem;
            background-color: #fff;
            border-radius: 0.25rem;
            box-shadow: 0 0.125rem 0.25rem rgba(0,0,0,0.075);
            margin-top: 1rem;
        }
        </style>
    """, unsafe_allow_html=True)
    converter_type = st.selectbox(
        "Select Conversion Type",
        [
            "Excel ↔️ CSV Converter",
            "Word ↔️ PDF Converter",
            "Image to PDF Converter",
            "PDF Editor",
            "Image Editor"])
    if converter_type == "Excel ↔️ CSV Converter":
        st.markdown("### Excel ↔️ CSV Converter")
        conversion_direction = st.radio(
            "Select conversion direction:",
            ["CSV to Excel", "Excel to CSV"],
            horizontal=True)
        if conversion_direction == "CSV to Excel":
            with st.container():
                st.markdown('<div class="converter-card">', unsafe_allow_html=True)
                uploaded_file = st.file_uploader("Upload CSV file", type="csv", key="csv_to_excel")
                if uploaded_file is not None:
                    try:
                        col1, col2 = st.columns(2)
                        with col1:
                            separator = st.selectbox(
                                "Select delimiter",
                                options=[",", ";", "|", "\t"],
                                index=0)
                        with col2:
                            encoding = st.selectbox(
                                "Select encoding",
                                options=["utf-8", "iso-8859-1", "cp1252"],
                                index=0)
                        df = pd.read_csv(uploaded_file, sep=separator, encoding=encoding)
                        st.markdown("#### Preview")
                        st.dataframe(df.head(), use_container_width=True)
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Rows", df.shape[0])
                        with col2:
                            st.metric("Columns", df.shape[1])
                        with col3:
                            st.metric("Size", f"{uploaded_file.size / 1024:.2f} KB")
                        output = BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            df.to_excel(writer, index=False)
                        excel_data = output.getvalue()
                        st.download_button(
                            label="📥 Download Excel File",
                            data=excel_data,
                            file_name=f"{uploaded_file.name.split('.')[0]}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
                st.markdown('</div>', unsafe_allow_html=True)
        else: 
            with st.container():
                st.markdown('<div class="converter-card">', unsafe_allow_html=True)
                uploaded_file = st.file_uploader("Upload Excel file", type=["xlsx", "xls"], key="excel_to_csv")
                if uploaded_file is not None:
                    try:
                        df = pd.read_excel(uploaded_file)
                        st.markdown("#### Preview")
                        st.dataframe(df.head(), use_container_width=True)
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Rows", df.shape[0])
                        with col2:
                            st.metric("Columns", df.shape[1])
                        with col3:
                            st.metric("Size", f"{uploaded_file.size / 1024:.2f} KB")
                        csv_data = BytesIO()
                        df.to_csv(csv_data, index=False)
                        st.download_button(
                            label="📥 Download CSV File",
                            data=csv_data.getvalue(),
                            file_name=f"{uploaded_file.name.split('.')[0]}.csv",
                            mime="text/csv")
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
                st.markdown('</div>', unsafe_allow_html=True)
    elif converter_type == "Word ↔️ PDF Converter":
        st.markdown("### Word ↔️ PDF Converter")
        conversion_direction = st.radio(
            "Select conversion direction:",
            ["Word to PDF", "PDF to Word"],
            horizontal=True)
        with st.container():
            st.markdown('<div class="converter-card">', unsafe_allow_html=True)
            if conversion_direction == "Word to PDF":
                uploaded_file = st.file_uploader("Upload Word file", type=["docx", "doc"], key="word_to_pdf")
                if uploaded_file is not None:
                    try:
                        doc = Document(uploaded_file)
                        output = BytesIO()
                        pdf = SimpleDocTemplate(output, pagesize=letter)
                        story = []
                        for paragraph in doc.paragraphs:
                            story.append(Paragraph(paragraph.text))
                        pdf.build(story)
                        st.download_button(
                            label="📥 Download PDF File",
                            data=output.getvalue(),
                            file_name=f"{uploaded_file.name.split('.')[0]}.pdf",
                            mime="application/pdf")
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
            else:
                uploaded_file = st.file_uploader("Upload PDF file", type=["pdf"], key="pdf_to_word")
                if uploaded_file is not None:
                    try:
                        pdf_reader = PdfReader(uploaded_file)
                        doc = Document()
                        for page in pdf_reader.pages:
                            text = page.extract_text()
                            doc.add_paragraph(text)
                        docx_output = BytesIO()
                        doc.save(docx_output)
                        st.download_button(label="📥 Download Word File",data=docx_output.getvalue(),file_name=f"{uploaded_file.name.split('.')[0]}.docx",mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
            st.markdown('</div>', unsafe_allow_html=True)
    elif converter_type == "Image to PDF Converter":
        st.markdown("### Image to PDF Converter")
        with st.container():
            st.markdown('<div class="converter-card">', unsafe_allow_html=True)
            uploaded_files = st.file_uploader(
                "Upload images (you can select multiple files)",
                type=["png", "jpg", "jpeg"],
                accept_multiple_files=True,
                key="image_to_pdf")
            if uploaded_files:
                try:
                    if len(uploaded_files) > 0:
                        st.markdown("#### Preview")
                        cols = st.columns(min(3, len(uploaded_files)))
                        for idx, file in enumerate(uploaded_files[:3]):
                            cols[idx].image(file, use_column_width=True)
                        if len(uploaded_files) > 3:
                            st.info(f"+ {len(uploaded_files) - 3} more images")
                    output = BytesIO()
                    pdf = Canvas(output, pagesize=letter)
                    for image_file in uploaded_files:
                        img = Image.open(image_file)
                        img_width, img_height = img.size
                        aspect = img_height / float(img_width)
                        if aspect > 1:
                            img_width = letter[0] - 40
                            img_height = img_width * aspect
                        else:
                            img_height = letter[1] - 40
                            img_width = img_height / aspect
                        pdf.drawImage(ImageReader(img), 20, letter[1] - img_height - 20,width=img_width, height=img_height)
                        pdf.showPage()
                    pdf.save()
                    st.download_button(
                        label="📥 Download PDF File",
                        data=output.getvalue(),
                        file_name="converted_images.pdf",
                        mime="application/pdf")
                except Exception as e:
                    st.error(f"Error: {str(e)}")
            st.markdown('</div>', unsafe_allow_html=True)
    elif converter_type == "PDF Editor":
      st.markdown("### PDF Editor") 
      with st.container():
        st.markdown('<div class="converter-card">', unsafe_allow_html=True) 
        uploaded_file = st.file_uploader("Upload PDF file", type=["pdf"], key="pdf_editor")
        if uploaded_file is not None:
            col1, col2 = st.columns(2)
            pdf_reader = PdfReader(uploaded_file)
            total_pages = len(pdf_reader.pages)
            with col1:
                st.markdown("#### Original PDF")
                preview_page = st.number_input("Preview page", 1, total_pages, 1) - 1
                uploaded_file.seek(0)
                original_preview = get_pdf_preview(uploaded_file, preview_page)
                st.image(original_preview, use_column_width=True)
                operations = st.multiselect("Select operations to perform",["Extract Pages", "Merge PDFs", "Rotate Pages", "Add Watermark","Resize", "Crop"])
            try:
                pdf_operations = {}
                if "Extract Pages" in operations:
                    st.markdown("#### Extract Pages")
                    total_pages = len(pdf_reader.pages)
                    all_pages = list(range(1, total_pages + 1))
                    selected_pages = st.multiselect("Select pages to extract",options=all_pages,default=[1],help="You can select multiple non-consecutive pages")
                    if selected_pages:
                          selected_pages.sort()
                          pdf_operations["extract"] = {"pages": selected_pages}
                          st.info(f"Selected pages: {', '.join(map(str, selected_pages))}")
                if "Merge PDFs" in operations:
                    st.markdown("#### Merge PDFs")
                    additional_pdfs = st.file_uploader(
                        "Upload PDFs to merge",
                        type=["pdf"],
                        accept_multiple_files=True,
                        key="merge_pdfs")
                    if additional_pdfs:
                        pdf_operations["merge"] = {"files": additional_pdfs}
                if "Rotate Pages" in operations:
                    st.markdown("#### Rotate Pages")
                    rotation = st.selectbox("Rotation angle", [90, 180, 270])
                    pdf_operations["rotate"] = {"angle": rotation}
                if "Add Watermark" in operations:
                    st.markdown("#### Add Watermark")
                    watermark_type = st.radio("Watermark Type", ["Text", "Image"])
                    watermark_options = {
                        "type": watermark_type.lower(),
                        "position": st.selectbox(
                            "Position",
                            ["center", "top-left", "top-right", "bottom-left", "bottom-right"]
                        ),
                        "angle": st.slider("Rotation Angle", -180, 180, 45),
                        "opacity": st.slider("Opacity", 0.1, 1.0, 0.3)}
                    page_selection = st.radio("Apply watermark to", ["All Pages", "Selected Pages"])
                    if page_selection == "Selected Pages":
                        selected_pages = st.multiselect("Select pages",range(1, total_pages + 1))
                        watermark_options["pages"] = selected_pages
                    else:
                        watermark_options["pages"] = "all"
                    if watermark_type == "Text":
                        watermark_options.update({"text": st.text_input("Watermark text"),"color": st.color_picker("Color", "#000000"),"size": st.slider("Size", 20, 100, 40)})
                    else:
                        watermark_image = st.file_uploader("Upload watermark image",type=["png", "jpg", "jpeg"])
                        if watermark_image:
                            watermark_options.update({"image": watermark_image,"size": st.slider("Size (% of page width)", 10, 100, 30)})
                    if (watermark_type == "Text" and watermark_options["text"]) or (watermark_type == "Image" and watermark_image):
                        pdf_operations["watermark"] = watermark_options
                if "Resize" in operations:
                    st.markdown("#### Resize PDF")
                    scale = st.slider("Scale percentage", 1, 200, 100,help="100% is original size")
                    pdf_operations["resize"] = {"scale": scale}
                if "Crop" in operations:
                    st.markdown("#### Crop PDF")
                    st.info("Values are in percentage of original size")
                    crop_col1, crop_col2 = st.columns(2)
                    with crop_col1:
                        left = st.number_input("Left", 0, 100, 0)
                        right = st.number_input("Right", 0, 100, 100)
                    with crop_col2:
                        top = st.number_input("Top", 0, 100, 100)
                        bottom = st.number_input("Bottom", 0, 100, 0)
                    pdf_operations["crop"] = {"left": left,"right": right,"top": top,"bottom": bottom}
                if pdf_operations:
                    output = BytesIO()
                    uploaded_file.seek(0)
                    output = process_pdf(uploaded_file, pdf_operations)
                    if "watermark" in pdf_operations:
                        output.seek(0)
                        pdf_writer = PdfWriter()
                        temp_reader = PdfReader(output)
                        for page in temp_reader.pages:
                            pdf_writer.add_page(page)
                        pdf_writer = add_watermark(pdf_writer, pdf_operations["watermark"])
                        final_output = BytesIO()
                        pdf_writer.write(final_output)
                        output = final_output
                    with col2:
                        st.markdown("#### Processed PDF")
                        output.seek(0)
                        processed_preview = get_pdf_preview(output, preview_page)
                        st.image(processed_preview, use_column_width=True)
                    original_size = len(uploaded_file.getvalue()) / 1024  # KB
                    output.seek(0)
                    new_size = len(output.getvalue()) / 1024 
                    metric_col1, metric_col2, metric_col3 = st.columns(3)
                    with metric_col1:
                        st.metric("Original Size", f"{original_size:.1f} KB")
                    with metric_col2:
                        st.metric("New Size", f"{new_size:.1f} KB")
                    with metric_col3:
                        reduction = ((original_size - new_size) / original_size) * 100
                        st.metric("Size Change", f"{reduction:.1f}%")
                    st.download_button(label="📥 Download Modified PDF",data=output.getvalue(),file_name=f"modified_{uploaded_file.name}",mime="application/pdf")       
            except Exception as e:
                st.error(f"Error: {str(e)}")
      st.markdown('</div>', unsafe_allow_html=True)
    elif converter_type == "Image Editor":
        st.markdown("### Image Editor")
        with st.container():
            st.markdown('<div class="converter-card">', unsafe_allow_html=True)
            uploaded_file = st.file_uploader("Upload image",type=["png", "jpg", "jpeg"],key="image_editor")
            if uploaded_file is not None:
                try:
                    original_bytes = uploaded_file.getvalue()
                    image = Image.open(uploaded_file)
                    st.markdown("#### Original Image")
                    st.image(image, use_column_width=True)
                    operations = {}
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.checkbox("Resize"):
                            st.markdown("##### Resize Settings")
                            orig_width, orig_height = image.size
                            width = st.number_input("Width", min_value=1, value=orig_width)
                            height = st.number_input("Height", min_value=1, value=orig_height)
                            operations["resize"] = {"width": width, "height": height}
                        if st.checkbox("Crop"):
                            st.markdown("##### Crop Settings")
                            width, height = image.size
                            left = st.number_input("Left", 0, width-1, 0)
                            top = st.number_input("Top", 0, height-1, 0)
                            right = st.number_input("Right", left+1, width, width)
                            bottom = st.number_input("Bottom", top+1, height, height)
                            operations["crop"] = {"left": left, "top": top, "right": right, "bottom": bottom}
                    with col2:
                        if st.checkbox("Rotate"):
                            angle = st.slider("Rotation Angle", -180, 180, 0)
                            operations["rotate"] = {"angle": angle}
                        if st.checkbox("Adjust"):
                            brightness = st.slider("Brightness", 0.0, 2.0, 1.0)
                            contrast = st.slider("Contrast", 0.0, 2.0, 1.0)
                            operations["brightness"] = {"factor": brightness}
                            operations["contrast"] = {"factor": contrast}
                        if st.checkbox("Compress"):
                            quality = st.slider("Quality", 1, 100, 85)
                            operations["compress"] = {"quality": quality}
                    if operations:
                        processed_image, quality = process_image(image, operations)
                        st.markdown("#### Processed Image")
                        st.image(processed_image, use_column_width=True)
                        output = BytesIO()
                        if quality is not None:
                            processed_image.save(output, format=image.format, quality=quality, optimize=True)
                        else:
                            processed_image.save(output, format=image.format, optimize=True)
                        processed_bytes = output.getvalue()
                        col1, col2, col3 = st.columns(3)
                        with col1:
                         st.metric("Width", f"{processed_image.width}px")
                        with col2:
                         st.metric("Height", f"{processed_image.height}px")
                        with col3:
                         st.metric("Size", f"{len(processed_bytes)/1024:.1f} KB")
                        st.markdown("#### Size Comparison")
                        metrics = get_image_size_metrics(original_bytes, processed_bytes)
                        col1, col2, col3 = st.columns(3)
                        with col1:
                         st.metric("Original Size", f"{metrics['original_size']:.1f} KB")
                        with col2:
                         st.metric("New Size", f"{metrics['processed_size']:.1f} KB")
                        with col3:
                         st.metric(
                            "Size Change", 
                            f"{metrics['size_change']:.1f}%",
                            delta_color="inverse")
                        st.download_button(
                        label="📥 Download Processed Image",
                        data=processed_bytes,
                        file_name=f"processed_{uploaded_file.name}",
                        mime=f"image/{image.format.lower()}")
                except Exception as e:
                    st.error(f"Error: {str(e)}")
            st.markdown('</div>', unsafe_allow_html=True)
    with st.expander("ℹ️ Need Help?"):
        st.markdown("""
        ### Usage Instructions
        1. Select the type of conversion you want to perform
        2. Upload your file(s) in the supported format
        3. Configure any additional settings if available
        4. Click the download button to save your converted file
        ### Supported Formats
        - Excel: .xlsx, .xls
        - CSV: .csv
        - Word: .docx, .doc
        - PDF: .pdf
        - Images: .png, .jpg, .jpeg
        ### Common Issues
        - If you're having trouble with CSV encoding, try different encoding options
        - Large files may take longer to process
        """)
def excel_editor():
    st.header("Excel Editor")
    def create_excel_structure_html(sheet, max_rows=5):
        html = "<table class='excel-table'>"
        merged_cells = sheet.merged_cells.ranges
        for idx, row in enumerate(sheet.iter_rows(max_row=max_rows)):
            html += "<tr>"
            for cell in row:
                merged = False
                for merged_range in merged_cells:
                    if cell.coordinate in merged_range:
                        if cell.coordinate == merged_range.start_cell.coordinate:
                            rowspan = min(merged_range.max_row - merged_range.min_row + 1, max_rows - idx)
                            colspan = merged_range.max_col - merged_range.min_col + 1
                            html += f"<td rowspan='{rowspan}' colspan='{colspan}'>{cell.value}</td>"
                        merged = True
                        break
                if not merged:
                    html += f"<td>{cell.value}</td>"
            html += "</tr>"
        html += "</table>"
        return html
    def get_merged_column_groups(sheet):
        merged_groups = {}
        for merged_range in sheet.merged_cells.ranges:
            if merged_range.min_row == 1:  # Only consider merged cells in the first row (header)
                main_col = sheet.cell(1, merged_range.min_col).value
                merged_groups[main_col] = list(range(merged_range.min_col, merged_range.max_col + 1))
        return merged_groups
    uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx")
    if uploaded_file is not None:
        # Read Excel file
        excel_file = openpyxl.load_workbook(uploaded_file)
        sheet = excel_file.active
        st.subheader("Original Excel Structure (First 5 Rows)")
        excel_html = create_excel_structure_html(sheet, max_rows=5)
        st.markdown(excel_html, unsafe_allow_html=True)
        merged_groups = get_merged_column_groups(sheet)
        column_headers = []
        column_indices = OrderedDict()  # To store the column indices for each header
        for col in range(1, sheet.max_column + 1):
            cell_value = sheet.cell(1, col).value
            if cell_value is not None:
                column_headers.append(cell_value)
                if cell_value not in column_indices:
                    column_indices[cell_value] = []
                column_indices[cell_value].append(col - 1)  # pandas uses 0-based index
            else:
                prev_header = column_headers[-1]
                column_headers.append(prev_header)
                column_indices[prev_header].append(col - 1)
        df = pd.read_excel(uploaded_file, header=None, names=column_headers)
        df = df.iloc[1:]  # Remove the first row as it's now our header
        st.subheader("Select columns to delete")
        all_columns = list(column_indices.keys())  # Use OrderedDict keys to maintain order
        cols_to_delete = st.multiselect("Choose columns to remove", all_columns)
        if cols_to_delete:
            columns_to_remove = []
            for col in cols_to_delete:
                columns_to_remove.extend(column_indices[col])
            df = df.drop(df.columns[columns_to_remove], axis=1)
            st.success(f"Deleted columns: {', '.join(cols_to_delete)}")
        st.subheader("Delete rows")
        num_rows = st.number_input("Enter the number of rows to delete from the start", min_value=0, max_value=len(df)-1, value=0)
        if num_rows > 0:
            df = df.iloc[num_rows:]
            st.success(f"Deleted first {num_rows} rows")
        st.subheader("Edit Data")
        st.write("You can edit individual cell values directly in the table below:")
        df_dict = df.where(pd.notnull(df), None).to_dict('records')
        edited_data = st.data_editor(df_dict)
        edited_df = pd.DataFrame(edited_data)
        st.subheader("Edited Data")
        st.dataframe(edited_df)
        def get_excel_download_link(df):
            output = BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False)
            excel_data = output.getvalue()
            b64 = base64.b64encode(excel_data).decode()
            return f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="edited_file.xlsx">Download Edited Excel File</a>'
        st.markdown(get_excel_download_link(edited_df), unsafe_allow_html=True)
        if st.button("Upload Edited File to Home"):
            st.session_state.edited_df = edited_df
            st.session_state.edited_file_name = "edited_" + uploaded_file.name
            st.success("Edited file has been uploaded to Home. Please switch to the Home tab to see the uploaded file.")
    else:
        st.info("Please upload an Excel file to begin editing.")
def data_analyzer():
    st.header("Advanced Data Analyzer")
    uploaded_file = st.file_uploader("Choose an Excel file for analysis", type="xlsx", key="analyser")
    if uploaded_file is not None:
        df = pd.read_excel(uploaded_file)
        st.write("Dataset Information:")
        st.write(f"Number of rows: {df.shape[0]}")
        st.write(f"Number of columns: {df.shape[1]}")
        numeric_columns = df.select_dtypes(include=['float64', 'int64']).columns
        categorical_columns = df.select_dtypes(include=['object']).columns
        analysis_type = st.selectbox("Select analysis type", ["Univariate Analysis", "Bivariate Analysis", "Regression Analysis", "Machine Learning Models", "Advanced Statistics"])
        if analysis_type == "Univariate Analysis":
            univariate_analysis(df, numeric_columns, categorical_columns)
        elif analysis_type == "Bivariate Analysis":
            bivariate_analysis(df, numeric_columns)
        elif analysis_type == "Regression Analysis":
            regression_analysis(df, numeric_columns, categorical_columns)
        elif analysis_type == "Machine Learning Models":
            machine_learning_models(df, numeric_columns, categorical_columns)
        elif analysis_type == "Advanced Statistics":
            advanced_statistics(df, numeric_columns)
def univariate_analysis(df, numeric_columns, categorical_columns):
    st.subheader("Univariate Analysis")
    column = st.selectbox("Select a column for analysis", numeric_columns.tolist() + categorical_columns.tolist())
    if column in numeric_columns:
        st.write(df[column].describe())
        col1, col2 = st.columns(2)
        with col1:
            fig = go.Figure()
            fig.add_trace(go.Histogram(x=df[column], name="Histogram"))
            fig.update_layout(title=f"Histogram for {column}")
            st.plotly_chart(fig, use_container_width=True)
        with col2:
            fig = go.Figure()
            fig.add_trace(go.Box(y=df[column], name="Box Plot"))
            fig.update_layout(title=f"Box Plot for {column}")
            st.plotly_chart(fig, use_container_width=True)
        col3, col4 = st.columns(2)
        with col3:
            fig = go.Figure()
            fig.add_trace(go.Violin(y=df[column], box_visible=True, line_color='black', meanline_visible=True, fillcolor='lightseagreen', opacity=0.6, x0=column))
            fig.update_layout(title=f"Violin Plot for {column}")
            st.plotly_chart(fig, use_container_width=True)
        with col4:
            fig = px.line(df, y=column, title=f"Line Plot for {column}")
            st.plotly_chart(fig, use_container_width=True)
        st.subheader("Additional Statistics")
        col5, col6, col7 = st.columns(3)
        with col5:
            st.metric("Skewness", f"{skew(df[column]):.4f}")
        with col6:
            st.metric("Kurtosis", f"{kurtosis(df[column]):.4f}")
        with col7:
            st.metric("Coefficient of Variation", f"{df[column].std() / df[column].mean():.4f}")
    else:
        st.write(df[column].value_counts())
        col1, col2 = st.columns(2)
        with col1:
            fig = px.bar(df[column].value_counts(), title=f"Bar Plot for {column}")
            st.plotly_chart(fig, use_container_width=True)
        with col2:
            fig = px.pie(df, names=column, title=f"Pie Chart for {column}")
            st.plotly_chart(fig, use_container_width=True)
def bivariate_analysis(df, numeric_columns):
    st.subheader("Bivariate Analysis")
    x_col = st.selectbox("Select X-axis variable", numeric_columns)
    y_col = st.selectbox("Select Y-axis variable", numeric_columns)
    chart_type = st.selectbox("Select chart type", ["Scatter", "Line", "Bar", "Box", "Violin", "3D Scatter", "Heatmap"])
    if chart_type == "Scatter":
        fig = px.scatter(df, x=x_col, y=y_col, title=f"{y_col} vs {x_col}")
    elif chart_type == "Line":
        fig = px.line(df, x=x_col, y=y_col, title=f"{y_col} vs {x_col}")
    elif chart_type == "Bar":
        fig = px.bar(df, x=x_col, y=y_col, title=f"{y_col} vs {x_col}")
    elif chart_type == "Box":
        fig = px.box(df, x=x_col, y=y_col, title=f"{y_col} vs {x_col}")
    elif chart_type == "Violin":
        fig = px.violin(df, x=x_col, y=y_col, title=f"{y_col} vs {x_col}")
    elif chart_type == "3D Scatter":
        z_col = st.selectbox("Select Z-axis variable", numeric_columns)
        fig = px.scatter_3d(df, x=x_col, y=y_col, z=z_col, title=f"3D Scatter Plot")
    elif chart_type == "Heatmap":
        corr_matrix = df[numeric_columns].corr()
        fig = px.imshow(corr_matrix, title="Correlation Heatmap")
    st.plotly_chart(fig, use_container_width=True)
    correlation = df[[x_col, y_col]].corr().iloc[0, 1]
    st.write(f"Correlation between {x_col} and {y_col}: {correlation:.4f}")
    st.subheader("Correlation Interpretation")
    st.write("""
    The correlation coefficient ranges from -1 to 1:
    - 1: Perfect positive correlation
    - 0: No correlation
    - -1: Perfect negative correlation
    Interpretation:
    - 0.00 to 0.19: Very weak correlation
    - 0.20 to 0.39: Weak correlation
    - 0.40 to 0.59: Moderate correlation
    - 0.60 to 0.79: Strong correlation
    - 0.80 to 1.00: Very strong correlation
    """)
    st.latex(r'''
    r = \frac{\sum_{i=1}^{n} (x_i - \bar{x})(y_i - \bar{y})}{\sqrt{\sum_{i=1}^{n} (x_i - \bar{x})^2} \sqrt{\sum_{i=1}^{n} (y_i - \bar{y})^2}}
    ''')
    st.write("Where:")
    st.write("- r is the correlation coefficient")
    st.write("- x_i and y_i are individual sample points")
    st.write("- x̄ and ȳ are the sample means")
def regression_analysis(df, numeric_columns, categorical_columns):
    st.subheader("Regression Analysis")
    regression_type = st.selectbox("Select regression type", ["Simple Linear", "Multiple Linear", "Polynomial", "Ridge", "Lasso"])
    y_col = st.selectbox("Select dependent variable", numeric_columns)
    x_cols = st.multiselect("Select independent variables", numeric_columns.tolist() + categorical_columns.tolist())
    if len(x_cols) == 0:
        st.warning("Please select at least one independent variable.")
        return
    X = df[x_cols]
    y = df[y_col]
    X = pd.get_dummies(X, drop_first=True)
    if regression_type == "Polynomial":
        degree = st.slider("Select polynomial degree", 1, 5, 2)
        poly = PolynomialFeatures(degree=degree)
        X = poly.fit_transform(X)
    X = sm.add_constant(X)
    try:
        if regression_type == "Ridge":
            alpha = st.slider("Select alpha for Ridge regression", 0.0, 10.0, 1.0)
            model = sm.OLS(y, X).fit_regularized(alpha=alpha, L1_wt=0)
        elif regression_type == "Lasso":
            alpha = st.slider("Select alpha for Lasso regression", 0.0, 10.0, 1.0)
            model = sm.OLS(y, X).fit_regularized(alpha=alpha, L1_wt=1)
        else:
            model = sm.OLS(y, X).fit()
        st.write(model.summary())
        fig = px.scatter(x=y, y=model.predict(X), labels={'x': 'Actual', 'y': 'Predicted'}, title="Actual vs Predicted Values")
        fig.add_trace(go.Scatter(x=[y.min(), y.max()], y=[y.min(), y.max()], mode='lines', name='y=x'))
        st.plotly_chart(fig, use_container_width=True)
        residuals = model.resid
        fig = px.scatter(x=model.predict(X), y=residuals, labels={'x': 'Predicted', 'y': 'Residuals'}, title="Residual Plot")
        fig.add_hline(y=0, line_dash="dash", line_color="red")
        st.plotly_chart(fig, use_container_width=True)
        st.subheader("Statistical Tests")
        jb_statistic, jb_p_value = jarque_bera(residuals)
        st.write(f"Jarque-Bera Test for Normality: statistic = {jb_statistic:.4f}, p-value = {jb_p_value:.4f}")
        st.write(f"{'Reject' if jb_p_value < 0.05 else 'Fail to reject'} the null hypothesis of normality at 5% significance level.")
        _, bp_p_value, _, _ = het_breuschpagan(residuals, model.model.exog)
        st.write(f"Breusch-Pagan Test for Heteroscedasticity: p-value = {bp_p_value:.4f}")
        st.write(f"{'Reject' if bp_p_value < 0.05 else 'Fail to reject'} the null hypothesis of homoscedasticity at 5% significance level.")
        dw_statistic = durbin_watson(residuals)
        st.write(f"Durbin-Watson Test for Autocorrelation: {dw_statistic:.4f}")
        st.write("Values close to 2 suggest no autocorrelation, while values toward 0 or 4 suggest positive or negative autocorrelation.")
        vif_data = pd.DataFrame()
        vif_data["Variable"] = X.columns
        vif_data["VIF"] = [variance_inflation_factor(X.values, i) for i in range(X.shape[1])]
        st.write("Variance Inflation Factors (VIF) for Multicollinearity:")
        st.write(vif_data)
        st.write("VIF > 5 suggests high multicollinearity.")
        st.subheader("Regression Formulas")
        if regression_type == "Simple Linear":
            st.latex(r'y = \beta_0 + \beta_1x + \epsilon')
        elif regression_type == "Multiple Linear":
            st.latex(r'y = \beta_0 + \beta_1x_1 + \beta_2x_2 + ... + \beta_nx_n + \epsilon')
        elif regression_type == "Polynomial":
            st.latex(r'y = \beta_0 + \beta_1x + \beta_2x^2 + ... + \beta_nx^n + \epsilon')
        elif regression_type == "Ridge":
            st.latex(r'\min_{\beta} \sum_{i=1}^n (y_i - \beta_0 - \sum_{j=1}^p \beta_jx_{ij})^2 + \lambda \sum_{j=1}^p \beta_j^2')
        elif regression_type == "Lasso":
            st.latex(r'\min_{\beta} \sum_{i=1}^n (y_i - \beta_0 - \sum_{j=1}^p \beta_jx_{ij})^2 + \lambda \sum_{j=1}^p |\beta_j|')
        st.write("Where:")
        st.write("- y is the dependent variable")
        st.write("- x, x_1, x_2, ..., x_n are independent variables")
        st.write("- β_0, β_1, β_2, ..., β_n are regression coefficients")
        st.write("- ε is the error term")
        st.write("- λ is the regularization parameter (for Ridge and Lasso)")
    except Exception as e:
        st.error(f"An error occurred during regression analysis: {str(e)}")
        st.write("This error might be due to multicollinearity, insufficient data, or other issues in the dataset.")
        st.write("Try selecting different variables or using a different regression type.")
def machine_learning_models(df, numeric_columns, categorical_columns):
    st.subheader("Machine Learning Models")
    model_type = st.selectbox("Select model type", ["Supervised", "Unsupervised"])
    if model_type == "Supervised":
        supervised_models(df, numeric_columns, categorical_columns)
    else:
        unsupervised_models(df, numeric_columns)
def supervised_models(df, numeric_columns, categorical_columns):
    st.write("Supervised Learning Models")
    y_col = st.selectbox("Select target variable", numeric_columns)
    x_cols = st.multiselect("Select features", numeric_columns.tolist() + categorical_columns.tolist())
    if len(x_cols) == 0:
        st.warning("Please select at least one feature.")
        return
    X = df[x_cols]
    y = df[y_col]
    X = pd.get_dummies(X, drop_first=True)
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)
    models = {
        "Linear Regression": LinearRegression(),
        "Decision Tree": DecisionTreeRegressor(),
        "Random Forest": RandomForestRegressor(),
        "SVR": SVR()
    }
    selected_model = st.selectbox("Select a model", list(models.keys()))
    try:
        model = models[selected_model]
        model.fit(X_train_scaled, y_train)
        y_pred = model.predict(X_test_scaled)
        mse = mean_squared_error(y_test, y_pred)
        r2 = r2_score(y_test, y_pred)
        st.write(f"Mean Squared Error: {mse:.4f}")
        st.write(f"R-squared Score: {r2:.4f}")
        fig = px.scatter(x=y_test, y=y_pred, labels={'x': 'Actual', 'y': 'Predicted'}, title="Actual vs Predicted Values")
        fig.add_trace(go.Scatter(x=[y_test.min(), y_test.max()], y=[y_test.min(), y_test.max()], mode='lines', name='y=x'))
        st.plotly_chart(fig, use_container_width=True)
        if selected_model in ["Decision Tree", "Random Forest"]:
            feature_importance = pd.DataFrame({
                'feature': X.columns,
                'importance': model.feature_importances_
            }).sort_values('importance', ascending=False)
            st.write("Feature Importance:")
            fig = px.bar(feature_importance, x='feature', y='importance', title="Feature Importance")
            st.plotly_chart(fig, use_container_width=True)
        st.subheader("Model Formulas and Explanations")
        if selected_model == "Linear Regression":
            st.latex(r'y = \beta_0 + \beta_1x_1 + \beta_2x_2 + ... + \beta_nx_n')
            st.write("Linear Regression finds the best-fitting linear relationship between the target variable and the features.")
        elif selected_model == "Decision Tree":
            st.write("Decision Trees make predictions by learning decision rules inferred from the data features.")
            st.image("https://scikit-learn.org/stable/_images/iris_dtc.png", caption="Example of a Decision Tree")
        elif selected_model == "Random Forest":
            st.write("Random Forest is an ensemble of Decision Trees, where each tree is trained on a random subset of the data and features.")
            st.image("https://scikit-learn.org/stable/_images/plot_forest_importances_faces_001.png", caption="Example of Random Forest Feature Importance")
        elif selected_model == "SVR":
            st.latex(r'\min_{w, b, \xi} \frac{1}{2} \|w\|^2 + C \sum_{i=1}^n \xi_i')
            st.write("Support Vector Regression (SVR) finds a function that deviates from y by a value no greater than ε for each training point x.")
    except Exception as e:
        st.error(f"An error occurred during model training: {str(e)}")
        st.write("This error might be due to insufficient data, incompatible data types, or other issues in the dataset.")
        st.write("Try selecting different variables or using a different model.")
def unsupervised_models(df, numeric_columns):
    st.write("Unsupervised Learning Models")
    x_cols = st.multiselect("Select features for clustering", numeric_columns)
    if len(x_cols) == 0:
        st.warning("Please select at least one feature.")
        return
    X = df[x_cols]
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    n_clusters = st.slider("Select number of clusters", 2, 10, 3)
    try:
        kmeans = KMeans(n_clusters=n_clusters, random_state=42)
        cluster_labels = kmeans.fit_predict(X_scaled)
        df_clustered = df.copy()
        df_clustered['Cluster'] = cluster_labels
        if len(x_cols) >= 2:
            fig = px.scatter(df_clustered, x=x_cols[0], y=x_cols[1], color='Cluster', title="K-means Clustering")
            st.plotly_chart(fig, use_container_width=True)
        st.write("Cluster Centers:")
        cluster_centers = scaler.inverse_transform(kmeans.cluster_centers_)
        st.write(pd.DataFrame(cluster_centers, columns=x_cols))
        inertias = []
        k_range = range(1, 11)
        for k in k_range:
            kmeans = KMeans(n_clusters=k, random_state=42)
            kmeans.fit(X_scaled)
            inertias.append(kmeans.inertia_)
        fig = px.line(x=k_range, y=inertias, title="Elbow Method for Optimal k",
                      labels={'x': 'Number of Clusters (k)', 'y': 'Inertia'})
        st.plotly_chart(fig, use_container_width=True)
        st.subheader("Principal Component Analysis (PCA)")
        n_components = st.slider("Select number of components", 2, min(len(x_cols), 10), 2)
        pca = PCA(n_components=n_components)
        pca_result = pca.fit_transform(X_scaled)
        df_pca = pd.DataFrame(data=pca_result, columns=[f'PC{i+1}' for i in range(n_components)])
        fig = px.scatter(df_pca, x='PC1', y='PC2', title="PCA Visualization")
        st.plotly_chart(fig, use_container_width=True)
        explained_variance_ratio = pca.explained_variance_ratio_
        cumulative_variance_ratio = np.cumsum(explained_variance_ratio)
        fig = go.Figure()
        fig.add_trace(go.Bar(x=range(1, n_components+1), y=explained_variance_ratio, name='Individual'))
        fig.add_trace(go.Scatter(x=range(1, n_components+1), y=cumulative_variance_ratio, mode='lines+markers', name='Cumulative'))
        fig.update_layout(title='Explained Variance Ratio', xaxis_title='Principal Components', yaxis_title='Explained Variance Ratio')
        st.plotly_chart(fig, use_container_width=True)
        st.write("Explained Variance Ratio:")
        st.write(pd.DataFrame({'PC': range(1, n_components+1), 'Explained Variance Ratio': explained_variance_ratio, 'Cumulative Variance Ratio': cumulative_variance_ratio}))
        st.subheader("K-means Clustering Formula")
        st.latex(r'\min_{S} \sum_{i=1}^{k} \sum_{x \in S_i} \|x - \mu_i\|^2')
        st.write("Where:")
        st.write("- S is the set of clusters")
        st.write("- k is the number of clusters")
        st.write("- x is a data point")
        st.write("- μ_i is the mean of points in S_i")
        st.subheader("PCA Formula")
        st.latex(r'X = U\Sigma V^T')
        st.write("Where:")
        st.write("- X is the original data matrix")
        st.write("- U is the left singular vectors (eigenvectors of XX^T)")
        st.write("- Σ is a diagonal matrix of singular values")
        st.write("- V^T is the right singular vectors (eigenvectors of X^TX)")
    except Exception as e:
        st.error(f"An error occurred during unsupervised learning: {str(e)}")
        st.write("This error might be due to insufficient data, incompatible data types, or other issues in the dataset.")
        st.write("Try selecting different variables or adjusting the number of clusters/components.")
def advanced_statistics(df, numeric_columns):
    st.subheader("Advanced Statistics")
    column = st.selectbox("Select a column for advanced statistics", numeric_columns)
    st.write("Descriptive Statistics:")
    st.write(df[column].describe())
    st.subheader("Normality Tests")
    shapiro_stat, shapiro_p = stats.shapiro(df[column])
    st.write(f"Shapiro-Wilk Test: statistic = {shapiro_stat:.4f}, p-value = {shapiro_p:.4f}")
    st.write(f"{'Reject' if shapiro_p < 0.05 else 'Fail to reject'} the null hypothesis of normality at 5% significance level.")
    anderson_result = stats.anderson(df[column])
    st.write("Anderson-Darling Test:")
    st.write(f"Statistic: {anderson_result.statistic:.4f}")
    for i in range(len(anderson_result.critical_values)):
        sl, cv = anderson_result.significance_level[i], anderson_result.critical_values[i]
        st.write(f"At {sl}% significance level: critical value = {cv:.4f}")
        if anderson_result.statistic < cv:
            st.write(f"The null hypothesis of normality is not rejected at {sl}% significance level.")
        else:
            st.write(f"The null hypothesis of normality is rejected at {sl}% significance level.")
    jb_stat, jb_p = stats.jarque_bera(df[column])
    st.write(f"Jarque-Bera Test: statistic = {jb_stat:.4f}, p-value = {jb_p:.4f}")
    st.write(f"{'Reject' if jb_p < 0.05 else 'Fail to reject'} the null hypothesis of normality at 5% significance level.")
    fig, ax = plt.subplots()
    stats.probplot(df[column], dist="norm", plot=ax)
    ax.set_title("Q-Q Plot")
    st.pyplot(fig)
    st.subheader("Time Series Analysis")
    adf_result = adfuller(df[column])
    st.write("Augmented Dickey-Fuller Test:")
    st.write(f"ADF Statistic: {adf_result[0]:.4f}")
    st.write(f"p-value: {adf_result[1]:.4f}")
    for key, value in adf_result[4].items():
        st.write(f"Critical Value ({key}): {value:.4f}")
    st.write(f"{'Reject' if adf_result[1] < 0.05 else 'Fail to reject'} the null hypothesis of a unit root at 5% significance level.")
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 10))
    plot_acf(df[column], ax=ax1)
    plot_pacf(df[column], ax=ax2)
    ax1.set_title("Autocorrelation Function (ACF)")
    ax2.set_title("Partial Autocorrelation Function (PACF)")
    st.pyplot(fig)
    st.subheader("Distribution Fitting")
    mu, sigma = stats.norm.fit(df[column])
    x = np.linspace(df[column].min(), df[column].max(), 100)
    y = stats.norm.pdf(x, mu, sigma)
    fig, ax = plt.subplots()
    ax.hist(df[column], density=True, alpha=0.7, bins='auto')
    ax.plot(x, y, 'r-', lw=2, label='Normal fit')
    ax.set_title(f"Distribution Fitting for {column}")
    ax.legend()
    st.pyplot(fig)
    st.write(f"Fitted Normal Distribution: μ = {mu:.4f}, σ = {sigma:.4f}")
    ks_statistic, ks_p_value = stats.kstest(df[column], 'norm', args=(mu, sigma))
    st.write("Kolmogorov-Smirnov Test:")
    st.write(f"Statistic: {ks_statistic:.4f}")
    st.write(f"p-value: {ks_p_value:.4f}")
    st.write(f"{'Reject' if ks_p_value < 0.05 else 'Fail to reject'} the null hypothesis that the data comes from the fitted normal distribution at 5% significance level.")
def create_stats_pdf(stats_data, district):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    title = Paragraph(f"Descriptive Statistics for {district}", styles['Title'])
    elements.append(title)
    data = [['Brand', 'Mean', 'Median', 'Std Dev', 'Min', 'Max', 'Skewness', 'Kurtosis', 'Range', 'IQR']]
    for brand, stats in stats_data.items():
        row = [brand]
        for stat in ['Mean', 'Median', 'Std Dev', 'Min', 'Max', 'Skewness', 'Kurtosis', 'Range', 'IQR']:
            value = stats[stat]
            if isinstance(value, (int, float)):
                row.append(f"{value:.2f}")
            else:
                row.append(str(value))
        data.append(row)
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 12),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)]))
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer
def create_prediction_pdf(prediction_data, district):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    title = Paragraph(f"Price Predictions for {district}", styles['Title'])
    elements.append(title)
    data = [['Brand', 'Predicted Price', 'Lower CI', 'Upper CI']]
    for brand, pred in prediction_data.items():
        row = [brand, f"{pred['forecast']:.2f}", f"{pred['lower_ci']:.2f}", f"{pred['upper_ci']:.2f}"]
        data.append(row)
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 12),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)]))
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer
#st.set_page_config(page_title="WSP Analysis",page_icon="🔬", layout="wide")
st.markdown("""
<style>
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
    }
    .main .block-container {
        padding: 2rem;
        background: rgba(255, 255, 255, 0.9);
        border-radius: 15px;
        box-shadow: 0 8px 32px 0 rgba(31, 38, 135, 0.37);
    }
    h1 {
        color: #2c3e50;
        text-align: center;
        padding: 1.5rem;
        background: rgba(255, 255, 255, 0.95);
        border-radius: 10px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .stSelectbox, .stMultiSelect {
        background: white;
        border-radius: 8px;
        margin-bottom: 1.5rem;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    .stButton > button {
        width: 100%;
        border-radius: 8px;
        background-color: #3498db;
        color: white;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    .stButton > button:hover {
        background-color: #2980b9;
        transform: translateY(-2px);
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .stSlider > div > div > div {
        background-color: #3498db;
    }
    .stCheckbox > label {
        color: #2c3e50;
        font-weight: 500;
    }
    .stSubheader {
        color: #34495e;
        background: rgba(255, 255, 255, 0.9);
        padding: 0.8rem;
        border-radius: 8px;
        margin-top: 1.5rem;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    .uploadedFile {
        background-color: #e8f0fe;
        border-radius: 8px;
        padding: 1rem;
        margin-bottom: 1.5rem;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    .dataframe {
        font-size: 0.8em;
    }
    .dataframe thead tr th {
        background-color: #3498db;
        color: brown;
    }
    .dataframe tbody tr:nth-child(even) {
        background-color: #f2f2f2;
    }
</style>
""", unsafe_allow_html=True)
if 'df' not in st.session_state:
    st.session_state.df = None
if 'week_names_input' not in st.session_state:
    st.session_state.week_names_input = []
if 'desired_diff_input' not in st.session_state:
    st.session_state.desired_diff_input = {}
if 'file_processed' not in st.session_state:
    st.session_state.file_processed = False
if 'diff_week' not in st.session_state:
    st.session_state.diff_week = 0
def transform_data(df, week_names_input):
    brands = ['UTCL', 'JKS', 'JKLC', 'Ambuja', 'Wonder', 'Shree']
    transformed_df = df[['Zone', 'REGION', 'Dist Code', 'Dist Name']].copy()
    region_replacements = {
        '12_Madhya Pradesh(west)': 'Madhya Pradesh(West)',
        '20_Rajasthan': 'Rajasthan', '50_Rajasthan III': 'Rajasthan', '80_Rajasthan II': 'Rajasthan',
        '33_Chhattisgarh(2)': 'Chhattisgarh', '38_Chhattisgarh(3)': 'Chhattisgarh', '39_Chhattisgarh(1)': 'Chhattisgarh',
        '07_Haryana 1': 'Haryana', '07_Haryana 2': 'Haryana',
        '06_Gujarat 1': 'Gujarat', '66_Gujarat 2': 'Gujarat', '67_Gujarat 3': 'Gujarat', '68_Gujarat 4': 'Gujarat', '69_Gujarat 5': 'Gujarat',
        '13_Maharashtra': 'Maharashtra(West)',
        '24_Uttar Pradesh': 'Uttar Pradesh(West)',
        '35_Uttarakhand': 'Uttarakhand',
        '83_UP East Varanasi Region': 'Varanasi',
        '83_UP East Lucknow Region': 'Lucknow',
        '30_Delhi': 'Delhi',
        '19_Punjab': 'Punjab',
        '09_Jammu&Kashmir': 'Jammu&Kashmir',
        '08_Himachal Pradesh': 'Himachal Pradesh',
        '82_Maharashtra(East)': 'Maharashtra(East)',
        '81_Madhya Pradesh': 'Madhya Pradesh(East)',
        '34_Jharkhand': 'Jharkhand',
        '18_ODISHA': 'Odisha',
        '04_Bihar': 'Bihar',
        '27_Chandigarh': 'Chandigarh',
        '82_Maharashtra (East)': 'Maharashtra(East)',
        '25_West Bengal': 'West Bengal'}
    transformed_df['REGION'] = transformed_df['REGION'].replace(region_replacements)
    transformed_df['REGION'] = transformed_df['REGION'].replace(['Delhi', 'Haryana', 'Punjab'], 'North-I')
    transformed_df['REGION'] = transformed_df['REGION'].replace(['Uttar Pradesh(West)','Uttarakhand'], 'North-II')
    zone_replacements = {'EZ_East Zone': 'East Zone','CZ_Central Zone': 'Central Zone','NZ_North Zone': 'North Zone','UPEZ_UP East Zone': 'UP East Zone','upWZ_up West Zone': 'UP West Zone','WZ_West Zone': 'West Zone'}
    transformed_df['Zone'] = transformed_df['Zone'].replace(zone_replacements)
    brand_columns = [col for col in df.columns if any(brand in col for brand in brands)]
    num_weeks = len(brand_columns) // len(brands)
    for i in range(num_weeks):
        start_idx = i * len(brands)
        end_idx = (i + 1) * len(brands)
        week_data = df[brand_columns[start_idx:end_idx]]
        week_name = week_names_input[i]
        week_data = week_data.rename(columns={
            col: f"{brand} ({week_name})"
            for brand, col in zip(brands, week_data.columns)})
        week_data.replace(0, np.nan, inplace=True)
        suffix = f'_{i}'
        transformed_df = pd.merge(transformed_df, week_data, left_index=True, right_index=True, suffixes=('', suffix))
    transformed_df = transformed_df.loc[:, ~transformed_df.columns.str.contains('_\d+$')]
    return transformed_df
def plot_district_graph(df, district_names, benchmark_brands_dict, desired_diff_dict, week_names, diff_week, download_pdf=False):
    brands = ['UTCL', 'JKS', 'JKLC', 'Ambuja', 'Wonder', 'Shree']
    num_weeks = len(df.columns[4:]) // len(brands)
    if download_pdf:
        pdf = matplotlib.backends.backend_pdf.PdfPages("district_plots.pdf")
    for i, district_name in enumerate(district_names):
        fig,ax=plt.subplots(figsize=(10, 8))
        district_df = df[df["Dist Name"] == district_name]
        price_diffs = []
        for brand in brands:
            brand_prices = []
            for week_name in week_names:
                column_name = f"{brand} ({week_name})"
                if column_name in district_df.columns:
                    price = district_df[column_name].iloc[0]
                    brand_prices.append(price)
                else:
                    brand_prices.append(np.nan)
            valid_prices = [p for p in brand_prices if not np.isnan(p)]
            if len(valid_prices) > diff_week:
                price_diff = valid_prices[-1] - valid_prices[diff_week]
            else:
                price_diff = np.nan
            price_diff_label = price_diff
            if np.isnan(price_diff):
               price_diff = 'NA'
            label = f"{brand} ({price_diff if isinstance(price_diff, str) else f'{price_diff:.0f}'})"
            plt.plot(week_names, brand_prices, marker='o', linestyle='-', label=label)
            for week, price in zip(week_names, brand_prices):
                if not np.isnan(price):
                    plt.text(week, price, str(round(price)), fontsize=10)
        plt.grid(False)
        plt.xlabel('Month/Week', weight='bold')
        reference_week = week_names[diff_week]
        last_week = week_names[-1]
        explanation_text = f"***Numbers in brackets next to brand names show the price difference between {reference_week} and {last_week}.***"
        plt.annotate(explanation_text, 
                     xy=(0, -0.23), xycoords='axes fraction', 
                     ha='left', va='center', fontsize=8, style='italic', color='deeppink',
                     bbox=dict(facecolor="#f0f8ff", edgecolor='none', alpha=0.7, pad=3))
        region_name = district_df['REGION'].iloc[0]
        plt.ylabel('Whole Sale Price(in Rs.)', weight='bold')
        region_name = district_df['REGION'].iloc[0]
        if i == 0:
            plt.text(0.5, 1.1, region_name, ha='center', va='center', transform=plt.gca().transAxes, weight='bold', fontsize=16)
            plt.title(f"{district_name} - Brands Price Trend", weight='bold')
        else:
            plt.title(f"{district_name} - Brands Price Trend", weight='bold')
        plt.legend(loc='upper center', bbox_to_anchor=(0.5, -0.15), ncol=6, prop={'weight': 'bold'})
        plt.tight_layout()
        text_str = ''
        if district_name in benchmark_brands_dict:
            brand_texts = []
            max_left_length = 0
            for benchmark_brand in benchmark_brands_dict[district_name]:
                jklc_prices = [district_df[f"JKLC ({week})"].iloc[0] for week in week_names if f"JKLC ({week})" in district_df.columns]
                benchmark_prices = [district_df[f"{benchmark_brand} ({week})"].iloc[0] for week in week_names if f"{benchmark_brand} ({week})" in district_df.columns]
                actual_diff = np.nan
                if jklc_prices and benchmark_prices:
                    for i in range(len(jklc_prices) - 1, -1, -1):
                        if not np.isnan(jklc_prices[i]) and not np.isnan(benchmark_prices[i]):
                            actual_diff = jklc_prices[i] - benchmark_prices[i]
                            break
                desired_diff_str = f" ({desired_diff_dict[district_name][benchmark_brand]:.0f} Rs.)" if district_name in desired_diff_dict and benchmark_brand in desired_diff_dict[district_name] else ""
                brand_text = [f"Benchmark Brand: {benchmark_brand}{desired_diff_str}", f"Actual Diff: {actual_diff:+.0f} Rs."]
                brand_texts.append(brand_text)
                max_left_length = max(max_left_length, len(brand_text[0]))
            num_brands = len(brand_texts)
            if num_brands == 1:
                text_str = "\n".join(brand_texts[0])
            elif num_brands > 1:
                half_num_brands = num_brands // 2
                left_side = brand_texts[:half_num_brands]
                right_side = brand_texts[half_num_brands:]
                lines = []
                for i in range(2):
                    left_text = left_side[0][i] if i < len(left_side[0]) else ""
                    right_text = right_side[0][i] if i < len(right_side[0]) else ""
                    lines.append(f"{left_text.ljust(max_left_length)} \u2502 {right_text.rjust(max_left_length)}")
                text_str = "\n".join(lines)
        plt.text(0.5, -0.3, text_str, weight='bold', ha='center', va='center', transform=plt.gca().transAxes, bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.5'))
        plt.subplots_adjust(bottom=0.25)
        if download_pdf:
            pdf.savefig(fig, bbox_inches='tight')
        st.pyplot(fig)
        buf = BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        b64_data = base64.b64encode(buf.getvalue()).decode()
        st.markdown(f'<a download="district_plot_{district_name}.png" href="data:image/png;base64,{b64_data}">Download Plot as PNG</a>', unsafe_allow_html=True)
        plt.close()
    if download_pdf:
        pdf.close()
        with open("district_plots.pdf", "rb") as f:
            pdf_data = f.read()
        b64_pdf = base64.b64encode(pdf_data).decode()
        st.markdown(f'<a download="{region_name}.pdf" href="data:application/pdf;base64,{b64_pdf}">Download All Plots as PDF</a>', unsafe_allow_html=True)
def update_week_name(index):
    def callback():
        if index < len(st.session_state.week_names_input):
            st.session_state.week_names_input[index] = st.session_state[f'week_{index}']
        else:
            st.warning(f"Attempted to update week {index + 1}, but only {len(st.session_state.week_names_input)} weeks are available.")
        st.session_state.all_weeks_filled = all(st.session_state.week_names_input)
    return callback
def load_lottie_url(url: str):
    try:
        response = requests.get(url)
        response.raise_for_status()  # Raises an HTTPError for bad responses
        return response.json()
    except Exception as e:
        st.warning(f"Failed to load Lottie animation: {str(e)}")
        return None
def Home():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;700&display=swap');
    body {
        font-family: 'Roboto', sans-serif;
        background-color: #f5f7fa;
        color: #333;
    }
    .title {
        font-size: 3.5rem;
        font-weight: 700;
        color: brown;
        text-align: center;
        padding: 2rem 0;
        margin-bottom: 2rem;
        background: linear-gradient(to right, #f0f8ff, #e6f3ff);
    }
    .subtitle {
        font-size: 1.5rem;
        font-weight: 300;
        color: #34495e;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-box {
        background-color: #ffffff;
        border-radius: 8px;
        padding: 2rem;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        transition: all 0.3s ease;
    }
    .section-box:hover {
        transform: translateY(-5px);
        box-shadow: 0 6px 12px rgba(0, 0, 0, 0.15);
    }
    .upload-section {
        background: linear-gradient(120deg, #a1c4fd 0%, #c2e9fb 100%);
        padding: 2rem;
        border-radius: 8px;
        margin-bottom: 2rem;
    }
    .btn-primary {
        background-color: #3498db;
        color: brown;
        padding: 0.5rem 1rem;
        border-radius: 5px;
        border: none;
        cursor: pointer;
        transition: background-color 0.3s ease;
    }
    .btn-primary:hover {
        background-color: #2980b9;
    }
    </style>
    """, unsafe_allow_html=True)
    st.markdown('<h1 class="title">Statistica</h1>', unsafe_allow_html=True)
    st.markdown('<p class="subtitle">Analyze, Visualize, Optimize.</p>', unsafe_allow_html=True)
    lottie_url = "https://assets9.lottiefiles.com/packages/lf20_jcikwtux.json"
    lottie_json = load_lottie_url(lottie_url)
    col1, col2 = st.columns([1, 2])
    with col1:
        if lottie_json:
            from streamlit_lottie import st_lottie
            st_lottie(lottie_json, height=250, key="home_animation")
        else:
            # Fallback content when animation fails to load
            st.image("https://via.placeholder.com/250x250?text=Analytics+Dashboard", 
                    caption="Dashboard Visualization")
    with col2:
        st.markdown("""
        <div class="section-box">
        <h3>Welcome to Your Data Analysis Journey!</h3>
        <p>Our interactive dashboard empowers you to:</p>
        <ul>
            <li>Upload and process your WSP data effortlessly</li>
            <li>Visualize trends across different brands and regions</li>
            <li>Generate descriptive statistics and predictions</li>
            <li>Make data-driven decisions with confidence</li>
        </ul>
        </div>
        """, unsafe_allow_html=True)
    st.markdown("""
    <div class="section-box">
    <h3>How to Use This Dashboard</h3>
    <ol>
        <li><strong>Upload Your Data:</strong> Start by uploading your Excel file containing the WSP data.</li>
        <li><strong>Enter Week Names:</strong> Provide names for each week column in your dataset.</li>
        <li><strong>Choose Your Analysis:</strong> Navigate to either the WSP Analysis Dashboard or Descriptive Statistics and Prediction sections.</li>
        <li><strong>Customize and Explore:</strong> Select your analysis parameters and generate valuable insights!</li>
    </ol>
    </div>
    """, unsafe_allow_html=True)
    st.markdown('<div class="upload-section">', unsafe_allow_html=True)
    st.subheader("Upload Your Data")
    if 'file_processed' not in st.session_state:
        st.session_state.file_processed = False
    if 'file_ready' not in st.session_state:
        st.session_state.file_ready = False
    if 'week_names_input' not in st.session_state:
        st.session_state.week_names_input = []
    uploaded_file = st.file_uploader("Upload Excel File", type=["xlsx"], key="wsp_data")
    if 'edited_df' in st.session_state and 'edited_file_name' in st.session_state and not st.session_state.edited_df.empty:
        st.success(f"Edited file uploaded: {st.session_state.edited_file_name}")
        if st.button("Process Edited File", key="process_edited"):
            process_uploaded_file(st.session_state.edited_df)
    elif uploaded_file:
        st.success(f"File uploaded: {uploaded_file.name}")
        if st.button("Process Uploaded File", key="process_uploaded"):
            process_uploaded_file(uploaded_file)
    if st.session_state.file_ready:
        st.markdown("### Enter Week Names")
        num_weeks = st.session_state.num_weeks
        num_columns = min(4, num_weeks) 
        if len(st.session_state.week_names_input) != num_weeks:
            st.session_state.week_names_input = [''] * num_weeks
        week_cols = st.columns(num_columns)
        for i in range(num_weeks):
            with week_cols[i % num_columns]:
                st.session_state.week_names_input[i] = st.text_input(
                    f'Week {i+1}', 
                    value=st.session_state.week_names_input[i],
                    key=f'week_{i}')
        if st.button("Confirm Week Names", key="confirm_weeks"):
            if all(st.session_state.week_names_input):
                st.session_state.file_processed = True
                st.success("File processed successfully! You can now proceed to the analysis sections.")
            else:
                st.warning("Please fill in all week names before confirming.")
    if st.session_state.file_processed:
        st.success("File processed successfully! You can now proceed to the analysis sections.")
    else:
        st.info("Please upload a file and fill in all week names to proceed with the analysis.")
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class="section-box">
    <h3>Need Assistance?</h3>
    <p>If you have any questions or need help using the dashboard, our support team is here for you. Don't hesitate to reach out!</p>
    <p>Email: prasoon.bajpai@lc.jkmail.com</p>
    <p>Phone: +91-9219393559</p>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("""
    <div style="text-align: center; margin-top: 2rem; padding: 1rem; background-color: #34495e; color: #ecf0f1;">
    <p>© 2024 WSP Analysis Dashboard. All rights reserved.</p>
    </div>
    """, unsafe_allow_html=True)
def process_uploaded_file(uploaded_file):
    if (isinstance(uploaded_file, pd.DataFrame) or uploaded_file) and not st.session_state.file_processed:
        try:
            if isinstance(uploaded_file, pd.DataFrame):
                # Convert DataFrame to Excel file in memory
                buffer = BytesIO()
                uploaded_file.to_excel(buffer, index=False)
                buffer.seek(0)
                file_content = buffer.getvalue()
            else:
                file_content = uploaded_file.read()
            wb = openpyxl.load_workbook(BytesIO(file_content))
            ws = wb.active
            hidden_cols = [idx for idx, col in enumerate(ws.column_dimensions, 1) if ws.column_dimensions[col].hidden]
            df = pd.read_excel(BytesIO(file_content), header=2)
            df = df.dropna(axis=1, how='all')
            df = df.drop(columns=df.columns[hidden_cols], errors='ignore')
            if df.empty:
                st.error("The uploaded file resulted in an empty dataframe. Please check the file content.")
            else:
                st.session_state.df = df
                brands = ['UTCL', 'JKS', 'JKLC', 'Ambuja', 'Wonder', 'Shree']
                brand_columns = [col for col in st.session_state.df.columns if any(brand in str(col) for brand in brands)]
                num_weeks = len(brand_columns) // len(brands)
                if num_weeks > 0:
                    if 'week_names_input' not in st.session_state or len(st.session_state.week_names_input) != num_weeks:
                        st.session_state.week_names_input = [''] * num_weeks
                    st.session_state.num_weeks = num_weeks
                    st.session_state.file_ready = True
                else:
                    st.warning("No weeks detected in the uploaded file. Please check the file content.")
                    st.session_state.week_names_input = []
                    st.session_state.file_processed = False
        except Exception as e:
            st.error(f"Error processing file: {e}")
            st.exception(e)
            st.session_state.file_processed = False
def wsp_analysis_dashboard():
    st.markdown("""
    <style>
    .title {
        font-size: 50px;
        font-weight: bold;
        color: brown;
        text-align: center;
        padding: 20px;
        border-radius: 10px;
        background: linear-gradient(to right, #f0f8ff, #e6f3ff);
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin-bottom: 30px;
        font-family: 'Arial', sans-serif;
    }
    .title span {
        background: linear-gradient(45deg, #3366cc, #6699ff);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .section-box {
        background-color: #f9f9f9;
        border-radius: 10px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        transition: all 0.3s ease;
    }
    .section-box:hover {
        transform: translateY(-5px);
        box-shadow: 0 6px 8px rgba(0, 0, 0, 0.15);
    }
    .stSelectbox, .stMultiSelect {
        background-color: white;
        border-radius: 8px;
        margin-bottom: 10px;
    }
    .stButton>button {
        border-radius: 20px;
        padding: 10px 20px;
        font-weight: bold;
        transition: all 0.3s ease;}
    .stButton>button:hover {
        transform: scale(1.05);}
    </style>
    """, unsafe_allow_html=True)
    st.markdown('<div class="title"><span>WSP Analysis Dashboard</span></div>', unsafe_allow_html=True)
    if not st.session_state.file_processed:
        st.warning("Please upload a file and fill in all week names in the Home section before using this dashboard.")
        return
    st.session_state.df = transform_data(st.session_state.df, st.session_state.week_names_input)
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.subheader("Analysis Settings")
    st.session_state.diff_week = st.slider("Select Week for Difference Calculation",min_value=0,max_value=len(st.session_state.week_names_input) - 1,value=st.session_state.diff_week,key="diff_week_slider") 
    download_pdf = st.checkbox("Download Plots as PDF", value=True)   
    col1, col2 = st.columns(2)
    with col1:
        zone_names = st.session_state.df["Zone"].unique().tolist()
        selected_zone = st.selectbox("Select Zone", zone_names, key="zone_select")
    with col2:
        filtered_df = st.session_state.df[st.session_state.df["Zone"] == selected_zone]
        region_names = filtered_df["REGION"].unique().tolist()
        selected_region = st.selectbox("Select Region", region_names, key="region_select")
    filtered_df = filtered_df[filtered_df["REGION"] == selected_region]
    district_names = filtered_df["Dist Name"].unique().tolist()
    region_recommendations = {"Gujarat": {"districts": ["Ahmadabad", "Mahesana", "Rajkot", "Vadodara", "Surat"],"benchmarks": ["UTCL", "Wonder"],"diffs": {"UTCL": -10.0, "Wonder": 0.0}},
        "Chhattisgarh": {"districts": ["Durg", "Raipur", "Bilaspur", "Raigarh", "Rajnandgaon"],"benchmarks": ["UTCL"],"diffs": {"UTCL": -10.0}},
        "Maharashtra(East)": {"districts": ["Nagpur", "Gondiya"],"benchmarks": ["UTCL"],"diffs": {"UTCL": -10.0}},
        "Odisha": {"districts": ["Cuttack", "Sambalpur", "Khorda"],"benchmarks": ["UTCL"],"diffs": {"UTCL": {"Sambalpur": -25.0, "Cuttack": -15.0, "Khorda": -15.0}}},
        "Rajasthan": {"districts": ["Alwar", "Jodhpur", "Udaipur", "Jaipur", "Kota", "Bikaner"],"benchmarks": [],"diffs": {}},
        "Madhya Pradesh(West)": {"districts": ["Indore", "Neemuch", "Ratlam", "Dhar"],"benchmarks": [],"diffs": {}},
        "Madhya Pradesh(East)": {"districts": ["Jabalpur", "Balaghat", "Chhindwara"],"benchmarks": [],"diffs": {}},
        "North-I": {"districts": ["East", "Gurugram", "Sonipat", "Hisar", "Yamunanagar", "Bathinda"],"benchmarks": [],"diffs": {}},
        "North-II": {"districts": ["Ghaziabad", "Meerut"],"benchmarks": [],"diffs": {}}}
    if selected_region in region_recommendations:
        recommended = region_recommendations[selected_region]
        suggested_districts = [d for d in recommended["districts"] if d in district_names]
        if suggested_districts:
            st.markdown(f"### Suggested Districts for {selected_region}")
            select_all = st.checkbox(f"Select all suggested districts for {selected_region}")
            if select_all:
                selected_districts = st.multiselect("Select District(s)", district_names, default=suggested_districts, key="district_select")
            else:
                selected_districts = st.multiselect("Select District(s)", district_names, key="district_select")
    else:
        selected_districts = st.multiselect("Select District(s)", district_names, key="district_select")
    st.markdown('</div>', unsafe_allow_html=True)
    brands = ['UTCL', 'JKS', 'JKLC', 'Ambuja', 'Wonder', 'Shree']
    benchmark_brands = [brand for brand in brands if brand != 'JKLC']
    benchmark_brands_dict = {}
    desired_diff_dict = {}
    if selected_districts:
        st.markdown("### Benchmark Settings")
        has_recommendations = (
            selected_region in region_recommendations and 
            region_recommendations[selected_region]["benchmarks"])
        if has_recommendations:
            use_recommended_benchmarks = st.checkbox(
                "Use recommended benchmarks and differences", 
                value=False)
        else:
            use_recommended_benchmarks = False
        if use_recommended_benchmarks:
            for district in selected_districts:
                benchmark_brands_dict[district] = region_recommendations[selected_region]["benchmarks"]
                desired_diff_dict[district] = {}
                if selected_region == "Odisha":
                    # Handle Odisha's district-specific differences
                    for brand in benchmark_brands_dict[district]:
                        if brand in region_recommendations[selected_region]["diffs"]:
                            desired_diff_dict[district][brand] = float(
                                region_recommendations[selected_region]["diffs"][brand].get(district, 0.0))
                else:
                    for brand in benchmark_brands_dict[district]:
                        desired_diff_dict[district][brand] = float(
                            region_recommendations[selected_region]["diffs"].get(brand, 0.0))            
        else:
            use_same_benchmarks = st.checkbox("Use same benchmarks for all districts", value=True)
            if use_same_benchmarks:
                selected_benchmarks = st.multiselect("Select Benchmark Brands for all districts", benchmark_brands, key="unified_benchmark_select")
                for district in selected_districts:
                    benchmark_brands_dict[district] = selected_benchmarks
                    desired_diff_dict[district] = {}
                if selected_benchmarks:
                    st.markdown("#### Desired Differences")
                    num_cols = min(len(selected_benchmarks), 3)
                    diff_cols = st.columns(num_cols)
                    for i, brand in enumerate(selected_benchmarks):
                        with diff_cols[i % num_cols]:
                            value = st.number_input(f"{brand}",min_value=-100.0,value=0.0,step=0.1,format="%.1f",key=f"unified_{brand}")
                            for district in selected_districts:
                                desired_diff_dict[district][brand] = float(value)
                else:
                    st.warning("Please select at least one benchmark brand.")
            else:
                for district in selected_districts:
                    st.subheader(f"Settings for {district}")
                    benchmark_brands_dict[district] = st.multiselect(
                        f"Select Benchmark Brands for {district}",
                        benchmark_brands,
                        key=f"benchmark_select_{district}")
                    desired_diff_dict[district] = {}
                    if benchmark_brands_dict[district]:
                        num_cols = min(len(benchmark_brands_dict[district]), 3)
                        diff_cols = st.columns(num_cols)
                        for i, brand in enumerate(benchmark_brands_dict[district]):
                            with diff_cols[i % num_cols]:
                                desired_diff_dict[district][brand] = st.number_input(f"{brand}",min_value=-100.0,value=0.0,step=0.1,format="%.1f",key=f"{district}_{brand}")
                    else:
                        st.warning(f"No benchmark brands selected for {district}.")
    st.markdown("### Generate Analysis")
    if st.button('Generate Plots', key='generate_plots', use_container_width=True):
        with st.spinner('Generating plots...'):
            plot_district_graph(filtered_df, selected_districts, benchmark_brands_dict, desired_diff_dict, st.session_state.week_names_input, st.session_state.diff_week,download_pdf)
            st.success('Plots generated successfully!')
    else:
        st.warning("Please upload a file in the Home section before using this dashboard.")
def descriptive_statistics_and_prediction():
    st.markdown("""
    <style>
    .title {
        font-size: 50px;
        font-weight: bold;
        color: #3366cc;
        text-align: center;
        padding: 20px;
        border-radius: 10px;
        background: linear-gradient(to right, #f0f8ff, #e6f3ff);
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin-bottom: 30px;
        font-family: 'Arial', sans-serif;
    }
    .title span {
        background: linear-gradient(45deg, #3366cc, #6699ff);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .section-box {
        background-color: #f9f9f9;
        border-radius: 10px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        transition: all 0.3s ease;
    }
    .section-box:hover {
        transform: translateY(-5px);
        box-shadow: 0 6px 8px rgba(0, 0, 0, 0.15);
    }
    .stSelectbox, .stMultiSelect {
        background-color: white;
        border-radius: 8px;
        margin-bottom: 10px;
    }
    .stButton>button {
        border-radius: 20px;
        padding: 10px 20px;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        transform: scale(1.05);
    }
    .stats-box {
        background-color: #e6f3ff;
        border-radius: 8px;
        padding: 15px;
        margin-bottom: 15px;
    }
    </style>
    """, unsafe_allow_html=True)
    st.markdown('<div class="title"><span>Descriptive Statistics and Prediction</span></div>', unsafe_allow_html=True)
    if not st.session_state.file_processed:
        st.warning("Please upload a file in the Home section before using this feature.")
        return
    st.session_state.df = transform_data(st.session_state.df, st.session_state.week_names_input)
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.subheader("Analysis Settings")
    col1, col2 = st.columns(2)
    with col1:
        zone_names = st.session_state.df["Zone"].unique().tolist()
        selected_zone = st.selectbox("Select Zone", zone_names, key="stats_zone_select")
    with col2:
        filtered_df = st.session_state.df[st.session_state.df["Zone"] == selected_zone]
        region_names = filtered_df["REGION"].unique().tolist()
        selected_region = st.selectbox("Select Region", region_names, key="stats_region_select")
    filtered_df = filtered_df[filtered_df["REGION"] == selected_region]
    district_names = filtered_df["Dist Name"].unique().tolist()
    if selected_region in ["Rajasthan", "Madhya Pradesh(West)","Madhya Pradesh(East)","Chhattisgarh","Maharashtra(East)","Odisha","North-I","North-II","Gujarat"]:
        suggested_districts = []
        if selected_region == "Rajasthan":
            rajasthan_districts = ["Alwar", "Jodhpur", "Udaipur", "Jaipur", "Kota", "Bikaner"]
            suggested_districts = [d for d in rajasthan_districts if d in district_names]
        elif selected_region == "Madhya Pradesh(West)":
            mp_west_districts = ["Indore", "Neemuch","Ratlam","Dhar"]
            suggested_districts = [d for d in mp_west_districts if d in district_names]
        elif selected_region == "Madhya Pradesh(East)":
            mp_west_districts = ["Jabalpur","Balaghat","Chhindwara"]
            suggested_districts = [d for d in mp_west_districts if d in district_names]
        elif selected_region == "Chhattisgarh":
            mp_west_districts = ["Durg","Raipur","Bilaspur","Raigarh","Rajnandgaon"]
            suggested_districts = [d for d in mp_west_districts if d in district_names]
        elif selected_region == "Maharashtra(East)":
            mp_west_districts = ["Nagpur","Gondiya"]
            suggested_districts = [d for d in mp_west_districts if d in district_names]
        elif selected_region == "Odisha":
            mp_west_districts = ["Cuttack","Sambalpur","Khorda"]
            suggested_districts = [d for d in mp_west_districts if d in district_names]
        elif selected_region == "North-I":
            mp_west_districts = ["East","Gurugram","Sonipat","Hisar","Yamunanagar","Bathinda"]
            suggested_districts = [d for d in mp_west_districts if d in district_names]
        elif selected_region == "North-II":
            mp_west_districts = ["Ghaziabad","Meerut"]
            suggested_districts = [d for d in mp_west_districts if d in district_names]
        elif selected_region == "Gujarat":
            mp_west_districts = ["Ahmadabad","Mahesana","Rajkot","Vadodara","Surat"]
            suggested_districts = [d for d in mp_west_districts if d in district_names]
        if suggested_districts:
            st.markdown(f"### Suggested Districts for {selected_region}")
            select_all = st.checkbox(f"Select all suggested districts for {selected_region}")
            if select_all:
                selected_districts = st.multiselect("Select District(s)", district_names, default=suggested_districts, key="district_select")
            else:
                selected_districts = st.multiselect("Select District(s)", district_names, key="district_select")
        else:
            selected_districts = st.multiselect("Select District(s)", district_names, key="district_select")
    else:
        selected_districts = st.multiselect("Select District(s)", district_names, key="district_select")
    st.markdown('</div>', unsafe_allow_html=True)
    if selected_districts:
        # Add a button to download all stats and predictions in one PDF
        if len(selected_districts) > 1:
            if st.checkbox("Download All Stats and Predictions",value=True):
                all_stats_pdf = BytesIO()
                pdf = SimpleDocTemplate(all_stats_pdf, pagesize=letter)
                elements = []
                for district in selected_districts:
                    elements.append(Paragraph(f"Statistics and Predictions for {district}", getSampleStyleSheet()['Title']))
                    district_df = filtered_df[filtered_df["Dist Name"] == district]
                    brands = ['UTCL', 'JKS', 'JKLC', 'Ambuja', 'Wonder', 'Shree']
                    stats_data = {}
                    prediction_data = {}
                    for brand in brands:
                        brand_data = district_df[[col for col in district_df.columns if brand in col]].values.flatten()
                        brand_data = brand_data[~np.isnan(brand_data)]
                        if len(brand_data) > 0:
                            stats_data[brand] = pd.DataFrame({'Mean': [np.mean(brand_data)],'Median': [np.median(brand_data)],'Std Dev': [np.std(brand_data)],'Min': [np.min(brand_data)],
                                'Max': [np.max(brand_data)],'Skewness': [stats.skew(brand_data)],'Kurtosis': [stats.kurtosis(brand_data)],'Range': [np.ptp(brand_data)],
                                'IQR': [np.percentile(brand_data, 75) - np.percentile(brand_data, 25)]}).iloc[0]
                            if len(brand_data) > 2:
                                model = ARIMA(brand_data, order=(1,1,1))
                                model_fit = model.fit()
                                forecast = model_fit.forecast(steps=1)
                                confidence_interval = model_fit.get_forecast(steps=1).conf_int()
                                prediction_data[brand] = {'forecast': forecast[0],'lower_ci': confidence_interval[0, 0],'upper_ci': confidence_interval[0, 1]}
                    elements.append(Paragraph("Descriptive Statistics", getSampleStyleSheet()['Heading2']))
                    elements.append(create_stats_table(stats_data))
                    elements.append(Paragraph("Price Predictions", getSampleStyleSheet()['Heading2']))
                    elements.append(create_prediction_table(prediction_data))
                    elements.append(PageBreak())
                pdf.build(elements)
                st.download_button(
                    label="Download All Stats and Predictions PDF",
                    data=all_stats_pdf.getvalue(),
                    file_name=f"{selected_districts}stats_and_predictions.pdf",
                    mime="application/pdf")
        st.markdown('<div class="section-box">', unsafe_allow_html=True)
        st.markdown("### Descriptive Statistics")
        for district in selected_districts:
            st.subheader(f"{district}")
            district_df = filtered_df[filtered_df["Dist Name"] == district]
            brands = ['UTCL', 'JKS', 'JKLC', 'Ambuja', 'Wonder', 'Shree']
            stats_data = {}
            prediction_data = {}
            for brand in brands:
                st.markdown(f'<div class="stats-box">', unsafe_allow_html=True)
                st.markdown(f"#### {brand}")
                brand_data = district_df[[col for col in district_df.columns if brand in col]].values.flatten()
                brand_data = brand_data[~np.isnan(brand_data)]
                if len(brand_data) > 0:
                    basic_stats = pd.DataFrame({'Mean': [np.mean(brand_data)],'Median': [np.median(brand_data)],'Std Dev': [np.std(brand_data)],'Min': [np.min(brand_data)],'Max': [np.max(brand_data)],'Skewness': [stats.skew(brand_data)],
                        'Kurtosis': [stats.kurtosis(brand_data)],'Range': [np.ptp(brand_data)],'IQR': [np.percentile(brand_data, 75) - np.percentile(brand_data, 25)]})
                    st.dataframe(basic_stats)
                    stats_data[brand] = basic_stats.iloc[0]
                    if len(brand_data) > 2:  
                        model = ARIMA(brand_data, order=(1,1,1))
                        model_fit = model.fit()
                        forecast = model_fit.forecast(steps=1)
                        confidence_interval = model_fit.get_forecast(steps=1).conf_int()
                        st.markdown(f"Predicted price for next week: {forecast[0]:.2f}")
                        st.markdown(f"95% Confidence Interval: [{confidence_interval[0, 0]:.2f}, {confidence_interval[0, 1]:.2f}]")
                        prediction_data[brand] = {'forecast': forecast[0],'lower_ci': confidence_interval[0, 0],'upper_ci': confidence_interval[0, 1]}
                else:
                    st.warning(f"No data available for {brand} in this district.")
                st.markdown('</div>', unsafe_allow_html=True)
            stats_pdf = create_stats_pdf(stats_data, district)
            predictions_pdf = create_prediction_pdf(prediction_data, district)
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(label="Download Statistics PDF",data=stats_pdf,file_name=f"{district}_statistics.pdf",mime="application/pdf")
            with col2:
                st.download_button(label="Download Predictions PDF",data=predictions_pdf,file_name=f"{district}_predictions.pdf",mime="application/pdf")
        st.markdown('</div>', unsafe_allow_html=True)
def create_stats_table(stats_data):
    data = [['Brand', 'Mean', 'Median', 'Std Dev', 'Min', 'Max', 'Skewness', 'Kurtosis', 'Range', 'IQR']]
    for brand, stats in stats_data.items():
        row = [brand]
        for stat in ['Mean', 'Median', 'Std Dev', 'Min', 'Max', 'Skewness', 'Kurtosis', 'Range', 'IQR']:
            value = stats[stat]
            if isinstance(value, (int, float)):
                row.append(f"{value:.2f}")
            else:
                row.append(str(value))
        data.append(row)
    table = Table(data)
    table.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),('BOTTOMPADDING', (0, 0), (-1, 0), 12),('BACKGROUND', (0, 1), (-1, -1), colors.beige),('TEXTCOLOR', (0, 1), (-1, -1), colors.black),('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),('FONTSIZE', (0, 1), (-1, -1), 10),('TOPPADDING', (0, 1), (-1, -1), 6),('BOTTOMPADDING', (0, 1), (-1, -1), 6),('GRID', (0, 0), (-1, -1), 1, colors.black)]))
    return table
def create_prediction_table(prediction_data):
    data = [['Brand', 'Predicted Price', 'Lower CI', 'Upper CI']]
    for brand, pred in prediction_data.items():
        row = [brand, f"{pred['forecast']:.2f}", f"{pred['lower_ci']:.2f}", f"{pred['upper_ci']:.2f}"]
        data.append(row)
    table = Table(data)
    table.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),('BOTTOMPADDING', (0, 0), (-1, 0), 12),('BACKGROUND', (0, 1), (-1, -1), colors.beige),('TEXTCOLOR', (0, 1), (-1, -1), colors.black),('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),('FONTSIZE', (0, 1), (-1, -1), 10),('TOPPADDING', (0, 1), (-1, -1), 6),('BOTTOMPADDING', (0, 1), (-1, -1), 6),('GRID', (0, 0), (-1, -1), 1, colors.black)]))
    return table
from urllib.parse import quote
@st.cache_data
def load_data(uploaded_file):
    df = pd.read_excel(uploaded_file)
    df = df.fillna(0)
    regions = df['Zone'].unique().tolist()
    brands = df['Brand'].unique().tolist()
    return df, regions, brands
@st.cache_resource
def load_lottie_url(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()
def create_visualization(region_data, region, brand, months):
    fig = plt.figure(figsize=(20, 34))
    gs = fig.add_gridspec(8, 3, height_ratios=[0.5,1, 1, 3, 2.25, 2, 2, 2])
    ax_region = fig.add_subplot(gs[0, :])
    ax_region.axis('off')
    ax_region.text(0.5, 0.5, f'{region} ({brand})', fontsize=28, fontweight='bold', ha='center', va='center')
    ax_current = fig.add_subplot(gs[1, :])
    ax_current.axis('off')
    overall_dec = region_data['Monthly Achievement(Dec)'].iloc[-1]
    trade_dec = region_data['Trade Nov'].iloc[-1]
    non_trade_dec = overall_dec - trade_dec
    table_data_left = [['AGS Target', f"{region_data['AGS Tgt (Dec)'].iloc[-1]:.0f}"],['Plan', f"{region_data['Month Tgt (Dec)'].iloc[-1]:.0f}"],['Trade Target', f"{region_data['Trade Tgt (Dec)'].iloc[-1]:.0f}"],['Non-Trade Target', f"{region_data['Non-Trade Tgt (Dec)'].iloc[-1]:.0f}"]]
    table_data_right = [[f"{overall_dec:.0f}"],[f"{trade_dec:.0f}"],[f"{non_trade_dec:.0f}"]]
    ax_current.text(0.225, 0.9, 'Targets', fontsize=12, fontweight='bold', ha='center')
    ax_current.text(0.35, 0.9, 'Achievement', fontsize=12, fontweight='bold', ha='center')
    table_left = ax_current.table(cellText=table_data_left,cellLoc='center',loc='center',bbox=[0, 0.0, 0.3, 0.8]) 
    table_left.auto_set_font_size(False)
    table_left.set_fontsize(12)
    table_left.scale(1.2, 1.8)
    for i in range(len(table_data_left)):
      cell = table_left[i, 0]
      cell.set_facecolor('#ECF0F1')
      cell.set_text_props(fontweight='bold')
      cell = table_left[i, 1]
      cell.set_facecolor('#F7F9F9')
      cell.set_text_props(fontweight='bold')
    table_right = ax_current.table(cellText=table_data_right,cellLoc='center',loc='center',bbox=[0.3, 0.0, 0.1, 0.8])  
    cell = table_right.add_cell(0, 0,1, 2, text=f"{overall_dec:.0f}",facecolor='#E8F6F3')
    cell.set_text_props(fontweight='bold')
    cell.set_text_props(fontweight='bold')
    cell = table_right.add_cell(1, 0, 1, 1,text=f"{trade_dec:.0f}",facecolor='#E8F6F3')
    cell.set_text_props(fontweight='bold')
    cell = table_right.add_cell(2, 0, 1, 1,text=f"{non_trade_dec:.0f}",facecolor='#E8F6F3')
    cell.set_text_props(fontweight='bold')
    table_right.auto_set_font_size(False)
    table_right.set_fontsize(13)
    table_right.scale(1.2, 1.8)
    ax_current.text(0.2, 1.0, 'December 2024 Performance Metrics', fontsize=16, fontweight='bold', ha='center', va='bottom')
    detailed_metrics = [('Trade', region_data['Trade Dec'].iloc[-1], region_data['Monthly Achievement(Dec)'].iloc[-1], 'Channel'),('Green', region_data['Green Dec'].iloc[-1], region_data['Monthly Achievement(Dec)'].iloc[-1], 'Region'),('Yellow', region_data['Yellow Dec'].iloc[-1], region_data['Monthly Achievement(Dec)'].iloc[-1], 'Region'),('Red', region_data['Red Dec'].iloc[-1], region_data['Monthly Achievement(Dec)'].iloc[-1], 'Region'),('Premium', region_data['Premium Dec'].iloc[-1], region_data['Monthly Achievement(Dec)'].iloc[-1], 'Product'),('Blended', region_data['Blended Dec'].iloc[-1], region_data['Monthly Achievement(Dec)'].iloc[-1], 'Product')]
    colors = ['blue', 'green', '#CDC50A', 'red', 'darkmagenta', 'saddlebrown']
    trade_box = patches.Rectangle((0.45, 0.74), 0.55, 0.125,facecolor='#F0F0F0',edgecolor='black',alpha=1,transform=ax_current.transAxes)
    ax_current.add_patch(trade_box)
    region_box = patches.Rectangle((0.45, 0.35), 0.55, 0.375,facecolor='#F0F0F0',edgecolor='black',alpha=1,transform=ax_current.transAxes)
    ax_current.add_patch(region_box)
    product_box = patches.Rectangle((0.45, 0.08), 0.55, 0.25,facecolor='#F0F0F0',edgecolor='black',alpha=1,transform=ax_current.transAxes)
    ax_current.add_patch(product_box)
    for i, (label, value, total, category) in enumerate(detailed_metrics):
     percentage = (value / total) * 100 if total != 0 else 0
     if i == 0:  
        y_pos = 0.77
     elif i <= 3:  
        y_pos = 0.63 - (i-1) * 0.11
     else:  
        y_pos = 0.24 - (i-4) * 0.11
     if category == 'Region' and value == 0:
        text = f'• {label} region not present'
     else:
        text = f'• {label} {category} has a share of {percentage:.1f}% in total sales, i.e., {value:.0f} MT.'
        ax_current.text(0.50, y_pos, text, fontsize=14, fontweight="bold", color=colors[i])
    ax_current.text(0.50, 0.90, 'Sales Breakown', fontsize=16, fontweight='bold', ha='center', va='bottom')
    ax_table = fig.add_subplot(gs[2, :])
    ax_table.axis('off')
    ax_table.set_title(f"Q-3 FY2024 vs Q-3 FY2025", fontsize=18, fontweight='bold')
    #table_data = [['Overall\nRequirement', 'Trade Channel\nRequirement', 'Premium Product\nRequirement','Blended Product\nRequirement'],
                #[f"{region_data['Q3 2023 Total'].iloc[-1]-region_data['Monthly Achievement(Oct)'].iloc[-1]-region_data['Monthly Achievement(Nov)'].iloc[-1]-region_data['Monthly Achievement(Dec)'].iloc[-1]:.0f}", f"{region_data['Q3 2023 Trade'].iloc[-1]-region_data['Trade Oct'].iloc[-1]-region_data['Trade Nov'].iloc[-1]-region_data['Trade Dec'].iloc[-1]:.0f}",f"{region_data['Q3 2023 Premium'].iloc[-1]-region_data['Premium Oct'].iloc[-1]-region_data['Premium Nov'].iloc[-1]-region_data['Premium Dec'].iloc[-1]:.0f}", 
                 #f"{region_data['Q3 2023 Blended '].iloc[-1]-region_data['Blended Oct'].iloc[-1]-region_data['Blended Nov'].iloc[-1]-region_data['Blended Dec'].iloc[-1]:.0f}"],]
    table_data = [['Inc./Dec. in\nTotal Sales', 'Inc./Dec. in\nTrade Sales', 'Inc./Dec. in\nPremium Sales','Inc./Dec. in\nBlended Sales'],
                [f"{region_data['Q3 2023 Total'].iloc[-1]-region_data['Monthly Achievement(Oct)'].iloc[-1]-region_data['Monthly Achievement(Nov)'].iloc[-1]-region_data['Monthly Achievement(Dec)'].iloc[-1]:.0f}", f"{region_data['Q3 2023 Trade'].iloc[-1]-region_data['Trade Oct'].iloc[-1]-region_data['Trade Nov'].iloc[-1]-region_data['Trade Dec'].iloc[-1]:.0f}",f"{region_data['Q3 2023 Premium'].iloc[-1]-region_data['Premium Oct'].iloc[-1]-region_data['Premium Nov'].iloc[-1]-region_data['Premium Dec'].iloc[-1]:.0f}", 
                 f"{region_data['Q3 2023 Blended '].iloc[-1]-region_data['Blended Oct'].iloc[-1]-region_data['Blended Nov'].iloc[-1]-region_data['Blended Dec'].iloc[-1]:.0f}"],]
    table = ax_table.table(cellText=table_data[1:], colLabels=table_data[0], cellLoc='center', loc='center')
    table.auto_set_font_size(False)
    table.set_fontsize(12)
    table.scale(1, 1.7)
    for (row, col), cell in table.get_celld().items():
                if row == 0:
                    cell.set_text_props(fontweight='bold', color='black')
                    cell.set_facecolor('goldenrod')
                cell.set_edgecolor('brown')
    ax1 = fig.add_subplot(gs[3, :])
    actual_ags = [region_data[f'AGS Tgt ({month})'].iloc[-1] for month in months]
    actual_achievements = [region_data[f'Monthly Achievement({month})'].iloc[-1] for month in months]
    actual_targets = [region_data[f'Month Tgt ({month})'].iloc[-1] for month in months]
    x = np.arange(len(months))
    width = 0.25
    rects1 = ax1.bar(x-width, actual_ags, width, label='AGS Target', color='brown', alpha=0.8)
    rects2 = ax1.bar(x, actual_targets, width, label='Plan', color='purple', alpha=0.8)
    rects3 = ax1.bar(x + width, actual_achievements, width, label='Achievement', color='yellow', alpha=0.8)
    ax1.set_ylabel('Targets and Achievement', fontsize=12, fontweight='bold')
    ax1.set_title(f"Monthly Targets and Achievements for FY 2025", fontsize=18, fontweight='bold')
    ax1.set_xticks(x)
    ax1.set_xticklabels(months)
    ax1.legend()
    def autolabel(rects):
        for rect in rects:
            height = rect.get_height()
            ax1.annotate(f'{height:.0f}',
                        xy=(rect.get_x() + rect.get_width() / 3, height),
                        xytext=(0, 3),
                        textcoords="offset points",
                        ha='center', va='bottom', fontsize=11)
    autolabel(rects1)
    autolabel(rects2)
    autolabel(rects3)
    ax2 = fig.add_subplot(gs[4, :])
    percent_achievements_plan = [((ach / tgt) * 100) for ach, tgt in zip(actual_achievements, actual_targets)]
    percent_achievements_ags = [((ach / ags) * 100) for ach, ags in zip(actual_achievements, actual_ags)]
    line1 = ax2.plot(x, percent_achievements_plan, marker='o', linestyle='-', color='purple', label='Achievement vs Plan')
    line2 = ax2.plot(x, percent_achievements_ags, marker='s', linestyle='-', color='brown', label='Achievement vs AGS')
    ax2.axhline(y=100, color='lightcoral', linestyle='--', alpha=0.7)
    ax2.set_xlabel('Month', fontsize=12, fontweight='bold')
    ax2.set_ylabel('% Achievement', fontsize=12, fontweight='bold')
    ax2.set_xticks(x)
    ax2.set_xticklabels(months)
    ax2.legend(loc='upper right')
    for i, (pct_plan, pct_ags) in enumerate(zip(percent_achievements_plan, percent_achievements_ags)):
        if pct_plan >= pct_ags:
            ax2.annotate(f'{pct_plan:.1f}%',(i, pct_plan),xytext=(0, 10),textcoords='offset points', ha='center',va='bottom',fontsize=12,color='purple')
            ax2.annotate(f'{pct_ags:.1f}%', (i, pct_ags), xytext=(0, -15), textcoords='offset points', ha='center',va='top',fontsize=12,color='brown')
        else:
            ax2.annotate(f'{pct_ags:.1f}%', (i, pct_ags), xytext=(0, 10), textcoords='offset points', ha='center', va='bottom', fontsize=12,color='brown')
            ax2.annotate(f'{pct_plan:.1f}%', (i, pct_plan), xytext=(0, -15), textcoords='offset points', ha='center', va='top', fontsize=12,color='purple')
    ax3 = fig.add_subplot(gs[5, :])
    ax3.axis('off')
    current_year = 2024
    last_year = 2023
    channel_data = [('Trade', region_data['Trade Dec'].iloc[-1], region_data['Trade Dec 2023'].iloc[-1],'Channel'),('Premium', region_data['Premium Dec'].iloc[-1], region_data['Premium Dec 2023'].iloc[-1],'Product'),('Blended', region_data['Blended Dec'].iloc[-1], region_data['Blended Dec 2023'].iloc[-1],'Product')]
    monthly_achievement_dec = region_data['Monthly Achievement(Dec)'].iloc[-1]
    total_dec_current = region_data['Monthly Achievement(Dec)'].iloc[-1]
    total_dec_last = region_data['Total Dec 2023'].iloc[-1]
    ax3.text(0.2, 1, f'December {current_year} Sales Comparison to December 2023:-', fontsize=16, fontweight='bold', ha='center', va='center')
    def get_arrow(value):
        return '↑' if value > 0 else '↓' if value < 0 else '→'
    def get_color(value):
        return 'green' if value > 0 else 'red' if value < 0 else 'black'
    total_change = ((total_dec_current - total_dec_last) / total_dec_last) * 100
    arrow = get_arrow(total_change)
    color = get_color(total_change)
    ax3.text(0.21, 0.9, f"December 2024: {total_dec_current:.0f}", fontsize=14, fontweight='bold', ha='center')
    ax3.text(0.22, 0.85, f"vs December 2023: {total_dec_last:.0f} ({total_change:.1f}% {arrow})", fontsize=12, color=color, ha='center')
    for i, (channel, value_current, value_last,x) in enumerate(channel_data):
        percentage = (value_current / monthly_achievement_dec) * 100
        percentage_last_year = (value_last / total_dec_last) * 100
        change = ((value_current - value_last) / value_last) * 100
        arrow = get_arrow(change)
        color = get_color(change)
        y_pos = 0.75 - i*0.25
        ax3.text(0.15, y_pos, f"{channel}:", fontsize=14, fontweight='bold')
        ax3.text(0.28, y_pos, f"{value_current:.0f}", fontsize=14)
        ax3.text(0.15, y_pos-0.05, f"vs Last Year: {value_last:.0f}", fontsize=12)
        ax3.text(0.28, y_pos-0.05, f"({change:.1f}% {arrow})", fontsize=12, color=color)
        ax3.text(0.12, y_pos-0.1, f"•{channel} {x} has share of {percentage_last_year:.1f}% in Dec. last year as compared to {percentage:.1f}% in Dec. this year.",fontsize=11, color='darkcyan')
    ax4 = fig.add_subplot(gs[5, 2])
    ax4.axis('off')
    channel_data1 = [('Trade', region_data['Trade Dec'].iloc[-1], region_data['Trade Nov'].iloc[-1],'Channel'),('Premium', region_data['Premium Dec'].iloc[-1], region_data['Premium Nov'].iloc[-1],'Product'),('Blended', region_data['Blended Dec'].iloc[-1], region_data['Blended Nov'].iloc[-1],'Product')]
    total_nov_current = region_data['Total Nov'].iloc[-1]
    ax4.text(0.35, 1, f'December {current_year} Sales Comparison to November 2024:-', fontsize=16, fontweight='bold', ha='center', va='center')
    total_change = ((total_dec_current - total_nov_current) / total_nov_current) * 100
    arrow = get_arrow(total_change)
    color = get_color(total_change)
    ax4.text(0.36, 0.9, f"December 2024: {total_dec_current:.0f}", fontsize=14, fontweight='bold', ha='center')
    ax4.text(0.37, 0.85, f"vs November 2024: {total_nov_current:.0f} ({total_change:.1f}% {arrow})", fontsize=12, color=color, ha='center')
    for i, (channel, value_current, value_last,t) in enumerate(channel_data1):
        percentage = (value_current / monthly_achievement_dec) * 100
        percentage_last_month = (value_last / total_nov_current) * 100
        change = ((value_current - value_last) / value_last) * 100
        arrow = get_arrow(change)
        color = get_color(change)
        y_pos = 0.75 - i*0.25
        ax4.text(0.10, y_pos, f"{channel}:", fontsize=14, fontweight='bold')
        ax4.text(0.65, y_pos, f"{value_current:.0f}", fontsize=14)
        ax4.text(0.10, y_pos-0.05, f"vs Last Month: {value_last:.0f}", fontsize=12)
        ax4.text(0.65, y_pos-0.05, f"({change:.1f}% {arrow})", fontsize=12, color=color)
        ax4.text(0.00, y_pos-0.1,f"•{channel} {t} has share of {percentage_last_month:.1f}% in Nov. as compared to {percentage:.1f}% in Dec.",fontsize=11, color='darkcyan')
    def create_pie_data(data_values, labels, colors):
     non_zero_data = []
     non_zero_labels = []
     non_zero_colors = []
     for value, label, color in zip(data_values, labels, colors):
        if value > 0:
            non_zero_data.append(value)
            non_zero_labels.append(label)
            non_zero_colors.append(color)       
     return non_zero_data, non_zero_labels, non_zero_colors
    def make_autopct(values):
     def my_autopct(pct):
        total = sum(values)
        val = int(round(pct*total/100.0))
        return f'{pct:.0f}%\n({val:.0f})'
     return my_autopct
    ax5 = fig.add_subplot(gs[6, 0])
    region_type_data = [region_data['Green Dec'].iloc[-1],region_data['Yellow Dec'].iloc[-1],region_data['Red Dec'].iloc[-1],region_data['Unidentified Dec'].iloc[-1]]
    region_type_labels = ['G', 'Y', 'R', '']
    colors = ['green', 'yellow', 'red', 'gray']
    filtered_data, filtered_labels, filtered_colors = create_pie_data(region_type_data, region_type_labels, colors)
    explode = [0.05] * len(filtered_data)
    ax5.pie(filtered_data, labels=filtered_labels, colors=filtered_colors,
        autopct=make_autopct(filtered_data), startangle=90, explode=explode)
    ax5.set_title('December 2024 Region Type Breakdown:-', fontsize=16, fontweight='bold')
    ax6 = fig.add_subplot(gs[6, 1])
    region_type_data = [region_data['Green Dec 2023'].iloc[-1],region_data['Yellow Dec 2023'].iloc[-1],region_data['Red Dec 2023'].iloc[-1],region_data['Unidentified Dec 2023'].iloc[-1]]
    filtered_data, filtered_labels, filtered_colors = create_pie_data(
    region_type_data, region_type_labels, colors)
    explode = [0.05] * len(filtered_data)
    ax6.pie(filtered_data, labels=filtered_labels, colors=filtered_colors,
        autopct=make_autopct(filtered_data), startangle=90, explode=explode)
    ax6.set_title('December 2023 Region Type Breakdown:-', fontsize=16, fontweight='bold')
    ax7 = fig.add_subplot(gs[6, 2])
    region_type_data = [
    region_data['Green Nov'].iloc[-1],
    region_data['Yellow Nov'].iloc[-1],
    region_data['Red Nov'].iloc[-1],
    region_data['Unidentified Nov'].iloc[-1]]
    filtered_data, filtered_labels, filtered_colors = create_pie_data(region_type_data, region_type_labels, colors)
    explode = [0.05] * len(filtered_data)
    ax7.pie(filtered_data, labels=filtered_labels, colors=filtered_colors,autopct=make_autopct(filtered_data), startangle=90, explode=explode)
    ax7.set_title('November 2024 Region Type Breakdown:-', fontsize=16, fontweight='bold')
    ax_comparison = fig.add_subplot(gs[7, :])
    ax_comparison.axis('off')
    ax_comparison.set_title('Quarterly Performance Analysis (2023 vs 2024)',fontsize=20, fontweight='bold', pad=20)
    def create_modern_quarterly_box(ax, x, y, width, height, q_data, quarter):
        rect = patches.Rectangle((x, y), width, height,facecolor='#f8f9fa',edgecolor='#dee2e6',linewidth=2,alpha=0.9,zorder=1)
        ax.add_patch(rect)
        title_height = height * 0.15
        title_bar = patches.Rectangle((x, y + height - title_height),width,title_height,facecolor='#4a90e2',alpha=0.9,zorder=2)
        ax.add_patch(title_bar)
        ax.text(x + width/2, y + height - title_height/2,f"{quarter} Performance Overview",ha='center', va='center',fontsize=14, fontweight='bold',color='white',zorder=3)
        total_2023, total_2024 = q_data['total_2023'], q_data['total_2024']
        pct_change = ((total_2024 - total_2023) / total_2023) * 100
        trade_2023, trade_2024 = q_data['trade_2023'], q_data['trade_2024']
        trade_pct_change = ((trade_2024 - trade_2023) / trade_2023) * 100
        y_offset = y + height - title_height - 0.1
        ax.text(x + 0.05, y_offset,"Total Sales Comparison:",fontsize=14, fontweight='bold',color='#2c3e50')
        y_offset -= 0.08
        ax.text(x + 0.05, y_offset,f"FY 2024: {total_2023:,.0f}",fontsize=11)
        ax.text(x + width/2, y_offset,f"FY 2025: {total_2024:,.0f}",fontsize=11)
        ax.text(x + 0.375*width, y_offset,f"{pct_change:+.1f}%",fontsize=11,color='green' if pct_change > 0 else 'red')
        y_offset -= 0.12
        ax.text(x + 0.05, y_offset,"Trade Volume:",fontsize=14, fontweight='bold',color='#2c3e50')
        y_offset -= 0.08
        ax.text(x + 0.05, y_offset,f"FY 2024: {trade_2023:,.0f}",fontsize=11)
        ax.text(x + width/2, y_offset,f"FY 2025: {trade_2024:,.0f}",fontsize=11)
        ax.text(x + 0.375*width, y_offset,f"{trade_pct_change:+.1f}%",fontsize=11,color='green' if trade_pct_change > 0 else 'red')
        if pct_change > 0:
            arrow_style = 'fancy,head_length=4,head_width=6'
            arrow_color = 'green'
            start_x = x + 0.13
            end_x = x + width * 0.49
            start_y = y + 0.31
            end_y = y + 0.31
        else:
            arrow_style = 'fancy,head_length=4,head_width=6'
            arrow_color = 'red'
            start_x = x + 0.13
            end_x = x + width * 0.49
            start_y = y + 0.31
            end_y = y + 0.31
        arrow = patches.FancyArrowPatch((start_x, start_y),(end_x, end_y),arrowstyle=arrow_style,color=arrow_color,linewidth=2,zorder=3)
        ax.add_patch(arrow)
        if trade_pct_change > 0:
            arrow_style = 'fancy,head_length=4,head_width=6'
            arrow_color = 'green'
            start_x = x + 0.13
            end_x = x + width * 0.49
            start_y = y + 0.11
            end_y = y + 0.11
        else:
            arrow_style = 'fancy,head_length=4,head_width=6'
            arrow_color = 'red'
            start_x = x + 0.13
            end_x = x + width * 0.49
            start_y = y + 0.11
            end_y = y + 0.11
        arrow = patches.FancyArrowPatch((start_x, start_y),(end_x, end_y),arrowstyle=arrow_style,color=arrow_color,linewidth=2,zorder=3)
        ax.add_patch(arrow)
    box_height = 0.6 
    box_y = 0.2   
    q1_data = {'total_2023': region_data['Q1 2023 Total'].iloc[-1],'total_2024': region_data['Q1 2024 Total'].iloc[-1],'trade_2023': region_data['Q1 2023 Trade'].iloc[-1],'trade_2024': region_data['Q1 2024 Trade'].iloc[-1]}
    q2_data = {'total_2023': region_data['Q2 2023 Total'].iloc[-1],'total_2024': region_data['Q2 2024 Total'].iloc[-1],'trade_2023': region_data['Q2 2023 Trade'].iloc[-1],'trade_2024': region_data['Q2 2024 Trade'].iloc[-1]}
    q3_data = {'total_2023': region_data['Q3 2023 Total'].iloc[-1],'total_2024': region_data['Q3 2024 Total'].iloc[-1],'trade_2023': region_data['Q3 2023 Trade'].iloc[-1],'trade_2024': region_data['Q3 2024 Trade'].iloc[-1]}
    create_modern_quarterly_box(ax_comparison, 0.00, box_y, 0.35, box_height, q1_data, "Q1")
    create_modern_quarterly_box(ax_comparison, 0.35, box_y, 0.35, box_height, q2_data, "Q2")
    create_modern_quarterly_box(ax_comparison, 0.70, box_y, 0.35, box_height, q3_data, "Q3")
    ax_comparison.set_xlim(0, 1)
    ax_comparison.set_ylim(0, 1)
    plt.tight_layout()
    return fig
def load_lottie_url(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()
def generate_full_report(df, regions):
    from matplotlib.backends.backend_pdf import PdfPages
    import matplotlib.pyplot as plt
    from io import BytesIO
    pdf_buffer = BytesIO()
    with PdfPages(pdf_buffer) as pdf:
        for region in regions:
            region_brands = df[df['Zone'] == region]['Brand'].unique().tolist()
            for brand in region_brands:
                region_data = df[(df['Zone'] == region) & (df['Brand'] == brand)]
                months = ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct','Nov','Dec']
                fig = create_visualization(region_data, region, brand, months)
                pdf.savefig(fig)
                plt.close(fig)
    pdf_buffer.seek(0)
    return pdf_buffer
def show_welcome_page():
        st.markdown("# 📈 Sales Review Report Generator")
        st.markdown("""
        ### Transform Your Sales Data into Actionable Insights
        This advanced analytics platform helps you:
        - 📊 Generate comprehensive sales review reports
        - 🎯 Track performance across regions and brands
        - 📈 Visualize key metrics and trends
        - 🔄 Compare historical data
        """)
        st.markdown("""
        <div class='reportBlock'>
        <h3>🚀 Getting Started</h3>
        <p>Upload your Excel file to begin analyzing your sales data:</p>
        </div>
        """, unsafe_allow_html=True)
        uploaded_file = st.file_uploader("Choose your Excel file", type="xlsx", key="Sales_Prediction_uploader")
        if uploaded_file:
            with st.spinner("Processing your data..."):
                progress_bar = st.progress(0)
                for i in range(100):
                    time.sleep(0.01)
                    progress_bar.progress(i + 1)
                df, regions, brands = load_data(uploaded_file)
                st.session_state['df'] = df
                st.session_state['regions'] = regions
                st.session_state['brands'] = brands
                st.success("✅ File processed successfully!")
def show_report_generator():
    st.markdown("# 🎯 Report Generator")
    if st.session_state.get('df') is None:
        st.warning("⚠️ Please upload your data file on the Home page first.")
        return
    df = st.session_state['df']
    regions = st.session_state['regions']
    tab1, tab2 = st.tabs(["📑 Individual Report", "📚 Complete Report"])
    with tab1:
            st.markdown("""
            <div class='reportBlock'>
            <h3>Report Parameters</h3>
            </div>
            """, unsafe_allow_html=True)
            region = st.selectbox("Select Region", regions, key='region_select')
            region_brands = df[df['Zone'] == region]['Brand'].unique().tolist()
            brand = st.selectbox("Select Brand", region_brands, key='brand_select')
            if st.button("🔍 Generate Individual Report", key='individual_report'):
                with st.spinner("Creating your report..."):
                    region_data = df[(df['Zone'] == region) & (df['Brand'] == brand)]
                    months = ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct','Nov','Dec']
                    fig = create_visualization(region_data, region, brand, months)
                    st.pyplot(fig)
                    buf = BytesIO()
                    fig.savefig(buf, format="pdf")
                    buf.seek(0)
                    st.download_button(
                        label="📥 Download Individual Report (PDF)",
                        data=buf,
                        file_name=f"sales_report_{region}_{brand}_{datetime.now().strftime('%Y%m%d')}.pdf",
                        mime="application/pdf")
    with tab2:
        st.markdown("""
        <div class='reportBlock'>
        <h3>Complete Report Generation</h3>
        <p>Generate a comprehensive report covering all regions and brands in your dataset.</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("📊 Generate Complete Report", key='complete_report'):
            with st.spinner("Generating comprehensive report... This may take a few minutes."):
                progress_bar = st.progress(0)
                for i in range(100):
                    time.sleep(0.02)
                    progress_bar.progress(i + 1)
                pdf_buffer = generate_full_report(df, regions)
                st.success("✅ Report generated successfully!")
                st.download_button(
                    label="📥 Download Complete Report (PDF)",
                    data=pdf_buffer,
                    file_name=f"complete_sales_report_{datetime.now().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf")
def show_about_page():
        st.markdown("# ℹ️ About")
        st.markdown("""
        <div class='reportBlock'>
        <h2>Sales Review Report Generator Pro</h2>
        <p>Version 2.0 | Last Updated: October 2024</p>
        <h3>🎯 Purpose</h3>
        Our advanced analytics platform empowers sales teams to:
        - Generate detailed performance reports
        - Track KPIs across regions and brands
        - Identify trends and opportunities
        - Make data-driven decisions
        <h3>🛠️ Features</h3>
        - Automated report generation
        - Interactive visualizations
        - Multi-region analysis
        - Historical comparisons
        - PDF export capabilities
        <h3>📧 Support</h3>
        For technical support or feedback:
        - Email: prasoon.bajpai@lc.jkmail.com
        </div>
        """, unsafe_allow_html=True)
def sales_review_report_generator():
    with st.sidebar:
        st.markdown("# 📊 Navigation")
        selected_page = st.radio("",["🏠 Home", "📈 Report Generator", "ℹ️ About"],key="navigation")
        st.markdown("---")
        st.markdown("### 📅 Current Session")
        st.markdown(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        if 'df' in st.session_state and st.session_state['df'] is not None:
            st.markdown("Status: ✅ Data Loaded")
        else:
            st.markdown("Status: ⚠️ Awaiting Data")
    if selected_page == "🏠 Home":
        show_welcome_page()
    elif selected_page == "📈 Report Generator":
        show_report_generator()
    else:
        show_about_page()
def load_lottie_url(url: str):
    try:
        r = requests.get(url)
        if r.status_code != 200:
            return None
        return r.json()
    except:
        return None
def get_online_editor_url(file_extension):
    extension_mapping = {'.xlsx': 'https://www.office.com/launch/excel?auth=2','.xls': 'https://www.office.com/launch/excel?auth=2','.doc': 'https://www.office.com/launch/word?auth=2',
        '.docx': 'https://www.office.com/launch/word?auth=2','.ppt': 'https://www.office.com/launch/powerpoint?auth=2','.pptx': 'https://www.office.com/launch/powerpoint?auth=2','.pdf': 'https://documentcloud.adobe.com/link/home/'}
    return extension_mapping.get(file_extension.lower(), 'https://www.google.com/drive/')
def folder_menu():
    st.markdown("""
    <style>
    .title {
        font-size: 50px;
        font-weight: bold;
        color: #3366cc;
        text-align: center;
        padding: 20px;
        border-radius: 10px;
        background: linear-gradient(to right, #f0f8ff, #e6f3ff);
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin-bottom: 30px;
        font-family: 'Arial', sans-serif;
    }
    .title span {
        background: linear-gradient(45deg, #3366cc, #6699ff);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .file-box {
        border: 1px solid #ddd;
        padding: 15px;
        margin: 15px 0;
        border-radius: 10px;
        background-color: #f9f9f9;
        transition: all 0.3s ease;
    }
    .file-box:hover {
        box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        transform: translateY(-5px);
    }
    .stButton>button {
        border-radius: 20px;
        padding: 10px 20px;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        transform: scale(1.05);
    }
    .upload-section {
        background-color: #e6f3ff;
        padding: 20px;
        border-radius: 10px;
        margin-bottom: 20px;
    }
    .todo-section {
        background-color: #f0f8ff;
        padding: 20px;
        border-radius: 10px;
        margin-top: 20px;
    }
    .todo-item {
        display: flex;
        align-items: center;
        margin-bottom: 10px;
    }
    .todo-text {
        margin-left: 10px;
    }
    </style>
    """, unsafe_allow_html=True)
    st.markdown('<div class="title"><span>📓 Advanced File Manager</span></div>', unsafe_allow_html=True)
    lottie_urls = [
        "https://assets9.lottiefiles.com/packages/lf20_3vbOcw.json",
        "https://assets9.lottiefiles.com/packages/lf20_5lAtR7.json",
        "https://assets1.lottiefiles.com/packages/lf20_4djadwfo.json",
        "https://assets6.lottiefiles.com/packages/lf20_2a5yxpci.json" ]
    lottie_json = None
    for url in lottie_urls:
        lottie_json = load_lottie_url(url)
        if lottie_json:
            break
    col1, col2 = st.columns([1, 2])
    with col1:
        if lottie_json:
           st_lottie(lottie_json, height=200, key="file_animation")
        else:
           st.image("https://via.placeholder.com/200x200.png?text=File+Manager", use_column_width=True)
    with col2:
        st.markdown("""
        Welcome to the Advanced File Manager! 
        Here you can upload, download, and manage your files with ease. 
        Enjoy the smooth animations, user-friendly interface, and new features like file search and sorting.""")
    if not os.path.exists("uploaded_files"):
        os.makedirs("uploaded_files")
    st.markdown('<div class="upload-section">', unsafe_allow_html=True)
    uploaded_file = st.file_uploader("Upload a file", type=["xlsx", "xls", "doc", "docx", "pdf", "ppt", "pptx", "txt", "csv"])
    if uploaded_file is not None:
        file_details = {"FileName": uploaded_file.name, "FileType": uploaded_file.type, "FileSize": uploaded_file.size}
        with open(os.path.join("uploaded_files", uploaded_file.name), "wb") as f:
            f.write(uploaded_file.getbuffer())
        st.success(f"File {uploaded_file.name} saved successfully!")
    st.markdown('</div>', unsafe_allow_html=True)
    st.subheader("Your Files")
    search_query = st.text_input("Search files", "")
    sort_option = st.selectbox("Sort by", ["Name", "Size", "Date Modified"])
    if 'files_to_delete' not in st.session_state:
        st.session_state.files_to_delete = set()
    files = os.listdir("uploaded_files")
    if search_query:
        files = [f for f in files if search_query.lower() in f.lower()]
    if sort_option == "Name":
        files.sort()
    elif sort_option == "Size":
        files.sort(key=lambda x: os.path.getsize(os.path.join("uploaded_files", x)), reverse=True)
    elif sort_option == "Date Modified":
        files.sort(key=lambda x: os.path.getmtime(os.path.join("uploaded_files", x)), reverse=True)
    for filename in files:
        file_path = os.path.join("uploaded_files", filename)
        file_stats = os.stat(file_path)
        st.markdown(f'<div class="file-box">', unsafe_allow_html=True)
        col1, col2, col3, col4= st.columns([3, 1, 1, 1])
        with col1:
            st.markdown(f"<h3>{filename}</h3>", unsafe_allow_html=True)
            st.text(f"Size: {file_stats.st_size / 1024:.2f} KB")
            st.text(f"Modified: {datetime.fromtimestamp(file_stats.st_mtime).strftime('%Y-%m-%d %H:%M:%S')}")
        with col2:
            if st.button(f"📥 Download", key=f"download_{filename}"):
                with open(file_path, "rb") as file:
                    file_content = file.read()
                    b64 = base64.b64encode(file_content).decode()
                    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Click to download</a>'
                    st.markdown(href, unsafe_allow_html=True)
        with col3:
            if st.button(f"🗑️ Delete", key=f"delete_{filename}"):
                st.session_state.files_to_delete.add(filename)
        with col4:
            file_extension = os.path.splitext(filename)[1]
            editor_url = get_online_editor_url(file_extension)
            st.markdown(f"[🌐 Open Online]({editor_url})")
        st.markdown('</div>', unsafe_allow_html=True)
    files_deleted = False
    for filename in st.session_state.files_to_delete:
        file_path = os.path.join("uploaded_files", filename)
        if os.path.exists(file_path):
            os.remove(file_path)
            st.warning(f"{filename} has been deleted.")
            files_deleted = True
    st.session_state.files_to_delete.clear()
    if files_deleted:
        st.rerun()
    st.info("Note: The 'Open Online' links will redirect you to the appropriate online editor. You may need to manually open your file once there.")
    st.markdown('<div class="todo-section">', unsafe_allow_html=True)
    st.subheader("📝 To-Do List / Diary")
    if 'todo_items' not in st.session_state:
        st.session_state.todo_items = []
    new_item = st.text_input("Add a new to-do item or diary entry")
    if st.button("Add"):
        if new_item:
            st.session_state.todo_items.append({"text": new_item, "done": False, "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S")})
            st.success("Item added successfully!")
    for idx, item in enumerate(st.session_state.todo_items):
        col1, col2, col3 = st.columns([0.1, 3, 1])
        with col1:
            done = st.checkbox("", item["done"], key=f"todo_{idx}")
            if done != item["done"]:
                st.session_state.todo_items[idx]["done"] = done
        with col2:
            st.markdown(f"<div class='todo-text'>{'<s>' if item['done'] else ''}{item['text']}{'</s>' if item['done'] else ''}</div>", unsafe_allow_html=True)
        with col3:
            st.text(item["date"])
        if st.button("Delete", key=f"delete_todo_{idx}"):
            st.session_state.todo_items.pop(idx)
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown("---")
    st.subheader("📚 Fun File Fact")
    fun_facts = ["The first computer virus was created in 1983 and was called the Elk Cloner.",
        "The most common file extension in the world is .dll (Dynamic Link Library).",
        "The largest file size theoretically possible in Windows is 16 exabytes minus 1 KB.",
        "The PDF file format was invented by Adobe in 1993.",
        "The first widely-used image format on the web was GIF, created in 1987.",
        "John McCarthy,an American computer scientist, coined the term Artificial Intelligence in 1956.",
        "About 90% of the World's Currency only exists on Computers.",
        "MyDoom is the most expensive computer virus in history.",
        "The original name of windows was Interface Manager.",
        "The first microprocessor created by Intel was the 4004."]
    st.markdown(f"*{fun_facts[int(os.urandom(1)[0]) % len(fun_facts)]}*")
def load_lottieurl(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()
def sales_dashboard():
    st.title("Sales Dashboard")
    st.markdown("""
    <style>
    .title {
        font-size: 50px;
        font-weight: bold;
        color: #3366cc;
        text-align: center;
        padding: 20px;
        border-radius: 10px;
        background: linear-gradient(to right, #f0f8ff, #e6f3ff);
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin-bottom: 30px;
        font-family: 'Arial', sans-serif;
    }
    .title span {
        background: linear-gradient(45deg, #3366cc, #6699ff);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .section-box {
        background-color: #f9f9f9;
        border-radius: 10px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        transition: all 0.3s ease;
    }
    .section-box:hover {
        transform: translateY(-5px);
        box-shadow: 0 6px 8px rgba(0, 0, 0, 0.15);
    }
    .upload-section {
        background-color: #e6f3ff;
        padding: 20px;
        border-radius: 10px;
        margin-bottom: 20px;
    }
    .stDataFrame {
        font-family: 'Arial', sans-serif;
    }
    </style>
    """, unsafe_allow_html=True)
    lottie_url = "https://assets2.lottiefiles.com/packages/lf20_V9t630.json"  # New interesting animation
    lottie_json = load_lottie_url(lottie_url)
    col1, col2 = st.columns([1, 2])
    with col1:
        st_lottie(lottie_json, height=200, key="home_animation")
    with col2:
        st.markdown("""
        Welcome to our interactive Sales Analysis Dashboard! 
        This powerful tool helps you analyze Sales data for JKLC and UCWL across different regions, districts and channels.
        Let's get started with your data analysis journey!
        """)
    uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx",key="Sales_Dashboard_uploader")
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    if uploaded_file is not None:
        df = pd.read_excel(uploaded_file)
        df = process_dataframe(df)
        regions = df['Region'].unique()
        selected_regions = st.multiselect('Select Regions', regions)
        districts = df[df['Region'].isin(selected_regions)]['Dist Name'].unique()
        selected_districts = st.multiselect('Select Districts', districts)
        channels = ['Overall', 'Trade', 'Non-Trade']
        selected_channels = st.multiselect('Select Channels', channels, default=channels)
        show_whole_region = st.checkbox('Show whole region totals')
        if st.button('Generate Report'):
            display_data(df, selected_regions, selected_districts, selected_channels, show_whole_region)
def process_dataframe(df):
    column_mapping = {pd.to_datetime('2024-09-23 00:00:00'): '23-Sep',pd.to_datetime('2024-08-23 00:00:00'): '23-Aug',pd.to_datetime('2024-07-23 00:00:00'): '23-Jul',pd.to_datetime('2024-06-23 00:00:00'): '23-Jun',pd.to_datetime('2024-05-23 00:00:00'): '23-May',pd.to_datetime('2024-04-23 00:00:00'): '23-Apr',pd.to_datetime('2024-08-24 00:00:00'): '24-Aug',pd.to_datetime('2024-07-24 00:00:00'): '24-Jul',pd.to_datetime('2024-06-24 00:00:00'): '24-Jun',pd.to_datetime('2024-05-24 00:00:00'): '24-May',pd.to_datetime('2024-04-24 00:00:00'): '24-Apr'}
    df = df.rename(columns=column_mapping)
    df['FY 2024 till Aug'] = df['24-Apr'] + df['24-May'] + df['24-Jun'] + df['24-Jul'] + df['24-Aug']
    df['FY 2023 till Aug'] = df['23-Apr'] + df['23-May'] + df['23-Jun'] + df['23-Jul'] + df['23-Aug']
    df['Quarterly Requirement'] = df['23-Jul'] + df['23-Aug'] + df['23-Sep'] - df['24-Jul'] - df['24-Aug']
    df['Growth/Degrowth(MTD)'] = (df['24-Aug'] - df['23-Aug']) / df['23-Aug'] * 100
    df['Growth/Degrowth(YTD)'] = (df['FY 2024 till Aug'] - df['FY 2023 till Aug']) / df['FY 2023 till Aug'] * 100
    df['Q3 2023'] = df['23-Jul'] + df['23-Aug'] + df['23-Sep']
    df['Q3 2024 till August'] = df['24-Jul'] + df['24-Aug']
    for month in ['Sep', 'Aug', 'Jul', 'Jun', 'May', 'Apr']:
        df[f'23-{month} Non-Trade'] = df[f'23-{month}'] - df[f'23-{month} Trade']
        if month != 'Sep':
            df[f'24-{month} Non-Trade'] = df[f'24-{month}'] - df[f'24-{month} Trade']
    df['FY 2024 till Aug Trade'] = df['24-Apr Trade'] + df['24-May Trade'] + df['24-Jun Trade'] + df['24-Jul Trade'] + df['24-Aug Trade']
    df['FY 2023 till Aug Trade'] = df['23-Apr Trade'] + df['23-May Trade'] + df['23-Jun Trade'] + df['23-Jul Trade'] + df['23-Aug Trade']
    df['Quarterly Requirement Trade'] = df['23-Jul Trade'] + df['23-Aug Trade'] + df['23-Sep Trade'] - df['24-Jul Trade'] - df['24-Aug Trade']
    df['Growth/Degrowth(MTD) Trade'] = (df['24-Aug Trade'] - df['23-Aug Trade']) / df['23-Aug Trade'] * 100
    df['Growth/Degrowth(YTD) Trade'] = (df['FY 2024 till Aug Trade'] - df['FY 2023 till Aug Trade']) / df['FY 2023 till Aug Trade'] * 100
    df['Q3 2023 Trade'] = df['23-Jul Trade'] + df['23-Aug Trade'] + df['23-Sep Trade']
    df['Q3 2024 till August Trade'] = df['24-Jul Trade'] + df['24-Aug Trade']
    df['FY 2024 till Aug Non-Trade'] = df['24-Apr Non-Trade'] + df['24-May Non-Trade'] + df['24-Jun Non-Trade'] + df['24-Jul Non-Trade'] + df['24-Aug Non-Trade']
    df['FY 2023 till Aug Non-Trade'] = df['23-Apr Non-Trade'] + df['23-May Non-Trade'] + df['23-Jun Non-Trade'] + df['23-Jul Non-Trade'] + df['23-Aug Non-Trade']
    df['Quarterly Requirement Non-Trade'] = df['23-Jul Non-Trade'] + df['23-Aug Non-Trade'] + df['23-Sep Non-Trade'] - df['24-Jul Non-Trade'] - df['24-Aug Non-Trade']
    df['Growth/Degrowth(MTD) Non-Trade'] = (df['24-Aug Non-Trade'] - df['23-Aug Non-Trade']) / df['23-Aug Non-Trade'] * 100
    df['Growth/Degrowth(YTD) Non-Trade'] = (df['FY 2024 till Aug Non-Trade'] - df['FY 2023 till Aug Non-Trade']) / df['FY 2023 till Aug Non-Trade'] * 100
    df['Q3 2023 Non-Trade'] = df['23-Jul Non-Trade'] + df['23-Aug Non-Trade'] + df['23-Sep Non-Trade']
    df['Q3 2024 till August Non-Trade'] = df['24-Jul Non-Trade'] + df['24-Aug Non-Trade']
    return df
    pass
def display_data(df, selected_regions, selected_districts, selected_channels, show_whole_region):
    def color_growth(val):
        try:
            value = float(val.strip('%'))
            color = 'green' if value > 0 else 'red' if value < 0 else 'black'
            return f'color: {color}'
        except:
            return 'color: black'
    if show_whole_region:
        filtered_data = df[df['Region'].isin(selected_regions)].copy()
        sum_columns = ['24-Apr','24-May','24-Jun','24-Jul','24-Aug','23-Apr','23-May','23-Jun','23-Jul', '23-Aug', 'FY 2024 till Aug', 'FY 2023 till Aug', 'Q3 2023', 'Q3 2024 till August','24-Apr Trade','24-May Trade','24-Jun Trade','24-Jul Trade', 
                        '24-Aug Trade','23-Apr Trade','23-May Trade','23-Jun Trade','23-Jul Trade', '23-Aug Trade', 'FY 2024 till Aug Trade', 'FY 2023 till Aug Trade', 
                        'Q3 2023 Trade', 'Q3 2024 till August Trade','24-Apr Non-Trade','24-May Non-Trade','24-Jun Non-Trade','24-Jul Non-Trade',
                        '24-Aug Non-Trade','23-Apr Non-Trade','23-May Non-Trade','23-Jun Non-Trade','23-Jul Non-Trade', '23-Aug Non-Trade', 'FY 2024 till Aug Non-Trade', 'FY 2023 till Aug Non-Trade', 
                        'Q3 2023 Non-Trade', 'Q3 2024 till August Non-Trade']
        grouped_data = filtered_data.groupby('Region')[sum_columns].sum().reset_index()
        grouped_data['Growth/Degrowth(MTD)'] = (grouped_data['24-Aug'] - grouped_data['23-Aug']) / grouped_data['23-Aug'] * 100
        grouped_data['Growth/Degrowth(YTD)'] = (grouped_data['FY 2024 till Aug'] - grouped_data['FY 2023 till Aug']) / grouped_data['FY 2023 till Aug'] * 100
        grouped_data['Quarterly Requirement'] = grouped_data['Q3 2023'] - grouped_data['Q3 2024 till August']
        grouped_data['Growth/Degrowth(MTD) Trade'] = (grouped_data['24-Aug Trade'] - grouped_data['23-Aug Trade']) / grouped_data['23-Aug Trade'] * 100
        grouped_data['Growth/Degrowth(YTD) Trade'] = (grouped_data['FY 2024 till Aug Trade'] - grouped_data['FY 2023 till Aug Trade']) / grouped_data['FY 2023 till Aug Trade'] * 100
        grouped_data['Quarterly Requirement Trade'] = grouped_data['Q3 2023 Trade'] - grouped_data['Q3 2024 till August Trade']
        grouped_data['Growth/Degrowth(MTD) Non-Trade'] = (grouped_data['24-Aug Non-Trade'] - grouped_data['23-Aug Non-Trade']) / grouped_data['23-Aug Non-Trade'] * 100
        grouped_data['Growth/Degrowth(YTD) Non-Trade'] = (grouped_data['FY 2024 till Aug Non-Trade'] - grouped_data['FY 2023 till Aug Non-Trade']) / grouped_data['FY 2023 till Aug Non-Trade'] * 100
        grouped_data['Quarterly Requirement Non-Trade'] = grouped_data['Q3 2023 Non-Trade'] - grouped_data['Q3 2024 till August Non-Trade']
    else:
        if selected_districts:
            filtered_data = df[df['Dist Name'].isin(selected_districts)].copy()
        else:
            filtered_data = df[df['Region'].isin(selected_regions)].copy()
        grouped_data = filtered_data
    for selected_channel in selected_channels:
        if selected_channel == 'Trade':
            columns_to_display = ['24-Aug Trade','23-Aug Trade','Growth/Degrowth(MTD) Trade','FY 2024 till Aug Trade', 'FY 2023 till Aug Trade','Growth/Degrowth(YTD) Trade','Q3 2023 Trade','Q3 2024 till August Trade', 'Quarterly Requirement Trade']
            suffix = ' Trade'
        elif selected_channel == 'Non-Trade':
            columns_to_display = ['24-Aug Non-Trade','23-Aug Non-Trade','Growth/Degrowth(MTD) Non-Trade','FY 2024 till Aug Non-Trade', 'FY 2023 till Aug Non-Trade','Growth/Degrowth(YTD) Non-Trade','Q3 2023 Non-Trade','Q3 2024 till August Non-Trade', 'Quarterly Requirement Non-Trade']
            suffix = ' Non-Trade'
        else: 
            columns_to_display = ['24-Aug','23-Aug','Growth/Degrowth(MTD)','FY 2024 till Aug', 'FY 2023 till Aug','Growth/Degrowth(YTD)','Q3 2023','Q3 2024 till August', 'Quarterly Requirement']
            suffix = ''
        display_columns = ['Region' if show_whole_region else 'Dist Name'] + columns_to_display  
        st.subheader(f"{selected_channel} Sales Data")
        display_df = grouped_data[display_columns].copy()
        display_df.set_index('Region' if show_whole_region else 'Dist Name', inplace=True)
        styled_df = display_df.style.format({
            col: '{:,.0f}' if 'Growth' not in col else '{:.2f}%' for col in columns_to_display
        }).applymap(color_growth, subset=[col for col in columns_to_display if 'Growth' in col])
        st.dataframe(styled_df)
        fig = go.Figure(data=[
            go.Bar(name='FY 2023', x=grouped_data['Region' if show_whole_region else 'Dist Name'], y=grouped_data[f'FY 2023 till Aug{suffix}']),
            go.Bar(name='FY 2024', x=grouped_data['Region' if show_whole_region else 'Dist Name'], y=grouped_data[f'FY 2024 till Aug{suffix}']),])
        fig.update_layout(barmode='group', title=f'{selected_channel} YTD Comparison')
        st.plotly_chart(fig)
        months = ['Apr', 'May', 'Jun', 'Jul', 'Aug']
        fig_trend = go.Figure()
        for year in ['23', '24']:
            y_values = []
            for month in months:
                column_name = f'{year}-{month}{suffix}'
                if column_name in grouped_data.columns:
                    y_values.append(grouped_data[column_name].sum())
                else:
                    y_values.append(None)
            fig_trend.add_trace(go.Scatter(
                x=months, 
                y=y_values, 
                mode='lines+markers+text',
                name=f'FY 20{year}',
                text=[f'{y:,.0f}' if y is not None else '' for y in y_values],
                textposition='top center'
            ))
        fig_trend.update_layout(
            title=f'{selected_channel} Monthly Trends', 
            xaxis_title='Month', 
            yaxis_title='Sales',
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
        )
        st.plotly_chart(fig_trend)
def load_lottieurl(url: str):
    try:
        r = requests.get(url)
        if r.status_code != 200:
            return None
        return r.json()
    except:
        return None
def normal():
 lottie_analysis = load_lottieurl("https://assets4.lottiefiles.com/packages/lf20_qp1q7mct.json")
 lottie_upload = load_lottieurl("https://assets9.lottiefiles.com/packages/lf20_ABViugg1T8.json")
 with st.sidebar:
    selected = option_menu(
        menu_title="Navigation",
        options=["Home", "Product-Mix Analysis", "About"],
        icons=["house", "graph-up", "info-circle"],
        menu_icon="cast",
        default_index=0,)
 def create_pdf_report(region, df, region_subset=None):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    def add_page_number(canvas):
      canvas.saveState()
      canvas.setFont('Helvetica', 10)
      page_number_text = f"Page {canvas.getPageNumber()}"
      canvas.drawString(width - 100, 30, page_number_text)
      canvas.restoreState()
    def add_header(page_number):
        c.setFillColorRGB(0.2, 0.2, 0.7)  # Dark blue color for header
        c.rect(0, height - 50, width, 50, fill=True)
        c.setFillColorRGB(1, 1, 1)  # White color for text
        c.setFont("Helvetica-Bold", 18)
        header_text = f"Product Mix Analysis Report: {region}"
        if region_subset:
            header_text += f" ({region_subset})"
        c.drawString(30, height - 35, header_text)
    def add_front_page():
        c.setFillColorRGB(0.4,0.5,0.3)
        c.rect(0, 0, width, height, fill=True)
        c.setFillColorRGB(1, 1, 1)
        c.setFont("Helvetica-Bold", 36)
        c.drawCentredString(width / 2, height - 200, "Product Mix Analysis Report")
        c.setFont("Helvetica", 24)
        report_title = f"Region: {region}"
        if region_subset:
            report_title += f" ({region_subset})"
        c.drawCentredString(width / 2, height - 250, report_title)
        c.setFont("Helvetica", 18)
        c.drawCentredString(width / 2, height - 300, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        add_page_number(c)
        c.showPage()
    def draw_graph(fig, x, y, width, height):
        img_buffer = BytesIO()
        fig.write_image(img_buffer, format="png",scale=2)
        img_buffer.seek(0)
        img = ImageReader(img_buffer)
        c.drawImage(img, x, y, width, height)
    def draw_table(data, x, y, col_widths):
        table = Table(data, colWidths=col_widths)
        table.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),('BOTTOMPADDING', (0, 0), (-1, 0), 6),('BACKGROUND', (0, 1), (-1, -1), colors.beige),('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),('FONTSIZE', (0, 1), (-1, -1), 6), ('TOPPADDING', (0, 1), (-1, -1), 3),('BOTTOMPADDING', (0, 1), (-1, -1), 3),('GRID', (0, 0), (-1, -1), 1, colors.black)]))
        w, h = table.wrapOn(c, width, height)
        table.drawOn(c, x, y - h)
    def add_tutorial_page():
        c.setFont("Helvetica-Bold", 24)
        c.drawString(inch, height - inch, "Understanding the Product Mix Analysis")
        drawing = Drawing(400, 200)
        lc = HorizontalLineChart()
        lc.x = 40
        lc.y = 50
        lc.height = 125
        lc.width = 300
        lc.data = [[random.randint(2000, 3000) for _ in range(12)],[random.randint(1500, 2500) for _ in range(12)],[random.randint(1800, 2800) for _ in range(12)],[random.randint(2200, 3200) for _ in range(12)],]
        lc.lines[0].strokeColor = colors.green
        lc.lines[1].strokeColor = colors.blue
        lc.lines[2].strokeColor = colors.pink
        lc.lines[3].strokeColor = colors.brown
        legend = Legend()
        legend.alignment = 'right'
        legend.x = 330
        legend.y = 150
        legend.colorNamePairs = [(colors.green, 'Normal EBITDA'),(colors.blue, 'Premium EBITDA'),(colors.crimson, 'Overall EBITDA'),(colors.brown, 'Imaginary EBITDA'),]
        drawing.add(lc)
        drawing.add(legend)
        renderPDF.draw(drawing, c, inch, height - 300)
        c.setFont("Helvetica-Bold", 18)
        c.drawString(inch, height - 350, "Key Concepts:")
        concepts = [("Overall EBITDA:", "Weighted average of Normal and Premium EBITDA based on their actual quantities."),("Imaginary EBITDA:", "Calculated by adjusting shares based on the following rules:"),("", "• If both (Normal and Premium) are present: Premium +5%, Normal -5%"),("", "• If only one is present: No change"),("Adjusted Shares:", "These adjustments aim to model potential improvements in product mix."),]
        text_object = c.beginText(inch, height - 380)
        for title, description in concepts:
            if title:
                text_object.setFont("Helvetica-Bold", 12)
                text_object.setFillColorRGB(0.7, 0.3, 0.1) 
                text_object.textLine(title)
                text_object.setFont("Helvetica", 12)
                text_object.setFillColorRGB(0, 0, 0)  
            text_object.textLine(description)
            if not title:
                text_object.textLine("")
        c.drawText(text_object)
        add_page_number(c)
        c.showPage()
    def add_appendix():
        c.setFont("Helvetica-Bold", 24)
        c.drawString(inch, height - inch, "Appendix")
        sections = [("Graph Interpretation:", "Each line represents a different metric over time. The differences between metrics are shown below\n each month."),("Tables:", "The descriptive statistics table provides a summary of the data. The monthly share distribution table\n shows the proportion of Normal and Premium Product for each month."),("Importance:", "These visualizations help identify trends, compare performance across product categories, and\n understand the potential impact of changing product distributions."),]
        text_object = c.beginText(inch, height - 1.5*inch)
        text_object.setFont("Helvetica-Bold", 14)
        for title, content in sections:
            text_object.textLine(title)
            text_object.setFont("Helvetica", 12)
            text_object.textLines(content)
            text_object.textLine("")
            text_object.setFont("Helvetica-Bold", 14)
        c.drawText(text_object)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(inch, height - 4*inch, "Suggestions for Improvement:")
        suggestions = ["Increase the share of Premium Product , which typically have higher EBITDA.","Analyze factors contributing to higher EBITDA in Premium Channel,and apply insights to Normal.","Regularly review and adjust pricing strategies to optimize EBITDA across all channels.","Invest in product innovation to expand Premium Product offerings.",]
        text_object = c.beginText(inch, height - 4.3*inch)
        text_object.setFont("Helvetica", 12)
        for suggestion in suggestions:
            text_object.textLine(f"• {suggestion}")
        c.drawText(text_object)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(inch, height - 5.2*inch, "Limitations:")
        limitations = ["This analysis is based on historical data and may not predict future market changes.","External factors such as economic conditions are not accounted for in this report.","This report analyzes the EBIDTA for Normal and Premium Product ceteris paribus.",]
        text_object = c.beginText(inch, height - 5.5*inch)
        text_object.setFont("Helvetica", 12)
        for limitation in limitations:
            text_object.textLine(f"• {limitation}")
        c.drawText(text_object)
        c.setFont("Helvetica", 12)
        c.drawString(inch, 2*inch, "We are currently working on including all other factors which impact the EBIDTA across products,")
        c.drawString(inch, 1.8*inch, "regions which will make this analysis more robust and helpful. We will also include NSR and") 
        c.drawString(inch,1.6*inch,"Contribution in our next report.")
        c.setFont("Helvetica-Bold", 14)
        c.drawString(inch, inch, "Thank You.")
        c.showPage()
    add_front_page()
    add_tutorial_page()
    brands = df['Brand'].unique()
    types = df['Type'].unique()
    region_subsets = df['Region subsets'].unique()
    page_number = 1
    for brand in brands:
        for product_type in types:
            for region_subset in region_subsets:
                filtered_df = df[(df['Region'] == region) & (df['Brand'] == brand) &
                                 (df['Type'] == product_type) & (df['Region subsets'] == region_subset)].copy()
                if not filtered_df.empty:
                    add_header(c)
                    cols = ['Normal EBITDA', 'Premium EBITDA']
                    overall_col = 'Overall EBITDA'
                    total_quantity = filtered_df['Normal'] + filtered_df['Premium']
                    filtered_df[overall_col] = (
                        (filtered_df['Normal'] * filtered_df['Normal EBITDA'] +
                         filtered_df['Premium'] * filtered_df['Premium EBITDA'])/ total_quantity)
                    filtered_df['Average Normal Share'] = filtered_df['Normal'] / total_quantity
                    filtered_df['Average Premium Share'] = filtered_df['Premium'] / total_quantity
                    def adjust_shares(row):
                        normal = row['Average Normal Share']
                        premium = row['Average Premium Share']
                        if normal == 1 or premium == 1 :
                            return normal,premium
                        else:
                            premium = min(premium + 0.05, 1)
                            normal = max(normal - 0.05, 1 - premium)
                        return normal,premium
                    filtered_df['Adjusted Normal Share'], filtered_df['Adjusted Premium Share'] = zip(*filtered_df.apply(adjust_shares, axis=1))
                    filtered_df['Imaginary EBITDA'] = (
                        filtered_df['Adjusted Normal Share'] * filtered_df['Normal EBITDA'] +
                        filtered_df['Adjusted Premium Share'] * filtered_df['Premium EBITDA'])
                    filtered_df['P-N Difference'] = filtered_df['Premium EBITDA'] - filtered_df['Normal EBITDA']
                    filtered_df['I-O Difference'] = filtered_df['Imaginary EBITDA'] - filtered_df[overall_col]
                    fig = go.Figure()
                    fig = make_subplots(rows=2, cols=1, row_heights=[0.58, 0.42], vertical_spacing=0.18)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['Normal EBITDA'],mode='lines+markers', name='Normal EBITDA', line=dict(color='green')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['Premium EBITDA'],mode='lines+markers', name='Premium EBITDA', line=dict(color='blue')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[overall_col],mode='lines+markers', name=overall_col, line=dict(color='crimson', dash='dash')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['Imaginary EBITDA'],mode='lines+markers', name='Imaginary EBITDA',line=dict(color='brown', dash='dot')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['I-O Difference'],mode='lines+markers+text', name='I-O Difference',text=filtered_df['I-O Difference'].round(2),textposition='top center',textfont=dict(size=8,weight="bold"),line=dict(color='fuchsia')), row=2, col=1)
                    mean_diff = filtered_df['I-O Difference'].mean()
                    if not np.isnan(mean_diff):
                        mean_diff=round(mean_diff)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=[mean_diff] * len(filtered_df),mode='lines', name=f'Mean I-O Difference[{mean_diff}]',line=dict(color='black', dash='dash')), row=2, col=1)
                    x_labels = [f"{month}<br>(P-N: {g_r:.0f})<br>(I-O: {g_y:.0f}))" 
                                for month, g_r, g_y in 
                                zip(filtered_df['Month'], filtered_df['P-N Difference'],filtered_df['I-O Difference'])]
                    fig.update_layout(title=f"EBITDA Analysis for {brand}(Type:-{product_type}) in {region}({region_subset})",legend_title='Metrics',plot_bgcolor='cornsilk',paper_bgcolor='lightcyan',height=710,)
                    fig.update_xaxes(tickmode='array', tickvals=list(range(len(x_labels))), ticktext=x_labels, row=1, col=1)
                    fig.update_xaxes(title_text='Months', row=2, col=1)
                    fig.update_yaxes(title_text='EBITDA(Rs./MT)', row=1, col=1)
                    fig.update_yaxes(title_text='I-O Difference(Rs./MT)', row=2, col=1)
                    # Add new page if needed
                    #if page_number > 1:
                        #c.showPage()
                    # Draw the graph
                    draw_graph(fig, 50, height - 410, 500, 350)
                    c.setFillColorRGB(0.2, 0.2, 0.7)  # Dark grey color for headers
                    c.setFont("Helvetica-Bold", 10)  # Reduced font size
                    c.drawString(50, height - 425, "Descriptive Statistics")
                    desc_stats = filtered_df[['Normal','Premium']+cols + [overall_col, 'Imaginary EBITDA']].describe().reset_index()
                    desc_stats = desc_stats[desc_stats['index'] != 'count'].round(2)  # Remove 'count' row
                    table_data = [['Metric'] + list(desc_stats.columns[1:])] + desc_stats.values.tolist()
                    draw_table(table_data, 50, height - 435, [45,45,45] + [75] * (len(desc_stats.columns) - 4))  # Reduced column widths
                    c.setFont("Helvetica-Bold", 10)  # Reduced font size
                    c.drawString(50, height - 600, "Average Share Distribution")
                    average_shares = filtered_df[['Average Normal Share', 'Average Premium Share']].mean()
                    share_fig = px.pie(values=average_shares.values,names=average_shares.index,color=average_shares.index,color_discrete_map={'Average Normal Share': 'green', 'Average Premium Share': 'blue'},title="",hole=0.3)
                    share_fig.update_layout(width=475, height=475, margin=dict(l=0, r=0, t=0, b=0))  
                    draw_graph(share_fig, 80, height - 810, 200, 200)  # Adjusted position and size
                    c.setFont("Helvetica-Bold", 10)
                    c.drawString(330, height - 600, "Monthly Share Distribution")
                    share_data = [['Month', 'Normal', 'Premium']]
                    for _, row in filtered_df[['Month', 'Normal', 'Premium','Average Normal Share', 'Average Premium Share']].iterrows():
                        share_data.append([row['Month'],f"{row['Normal']:.0f} ({row['Average Normal Share']:.2%})",f"{row['Premium']:.0f} ({row['Average Premium Share']:.2%})"])
                    draw_table(share_data, 330, height - 620, [40, 60, 60, 60])
                    add_page_number(c)
                    c.showPage()
    for i in range(c.getPageNumber()):
        c.setPageSize((width, height))
        add_page_number(c)         
    add_appendix()
    c.save()
    buffer.seek(0)
    return buffer
 if selected == "Home":
    st.title("🔍 Advanced Product Mix Analysis")
    st.markdown("Welcome to our advanced data analysis platform. Upload your Excel file to get started with interactive visualizations and insights.")
    st.markdown("<div class='upload-section'>", unsafe_allow_html=True)
    col1, col2 = st.columns([2, 1])
    with col1:
        uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx",key='NormalvsPremiumuploader')
        if uploaded_file is not None:
            st.session_state.uploaded_file = uploaded_file
            st.success("File successfully uploaded! Please go to the Analysis page to view results.")
    with col2:
        if lottie_upload:
            st_lottie(lottie_upload, height=150, key="upload")
        else:
            st.image("https://cdn-icons-png.flaticon.com/512/4503/4503700.png", width=150)
    st.markdown("</div>", unsafe_allow_html=True)
 elif selected == "Product-Mix Analysis":
    st.title("📈 Product Mix Dashboard")
    if 'uploaded_file' not in st.session_state or st.session_state.uploaded_file is None:
        st.warning("Please upload an Excel file on the Home page to begin the analysis.")
    else:
        df = pd.read_excel(st.session_state.uploaded_file)
        st.markdown("<div class='analysis-section'>", unsafe_allow_html=True)
        if lottie_analysis:
            st_lottie(lottie_analysis, height=200, key="analysis")
        else:
            st.image("https://cdn-icons-png.flaticon.com/512/2756/2756778.png", width=200)
        st.sidebar.header("Filter Options")
        region = st.sidebar.selectbox("Select Region", options=df['Region'].unique(), key="region_select")
        st.sidebar.subheader(f"Download Report for {region}")
        download_choice = st.sidebar.radio(
            "Choose report type:",
            ('Full Region', 'Region Subset'))
        if download_choice == 'Full Region':
            if st.sidebar.button(f"Download Full Report for {region}"):
                subset_df = df[(df['Region'] == region) & (df['Type'] != 'PPC Premium')]
                pdf_buffer = create_pdf_report(region, subset_df)
                pdf_bytes = pdf_buffer.getvalue()
                b64 = base64.b64encode(pdf_bytes).decode()
                href = f'<a href="data:application/pdf;base64,{b64}" download="Product_Mix_Analysis_Report_{region}.pdf">Download Full Region PDF Report</a>'
                st.sidebar.markdown(href, unsafe_allow_html=True)
        else:
            region_subsets = df[df['Region'] == region]['Region subsets'].unique()
            selected_subset = st.sidebar.selectbox("Select Region Subset", options=region_subsets)
            if st.sidebar.button(f"Download Report for {region} - {selected_subset}"):
                subset_df = df[(df['Region'] == region) & (df['Region subsets'] == selected_subset) & (df['Type'] != 'PPC Premium')]
                pdf_buffer = create_pdf_report(region, subset_df, selected_subset)
                pdf_bytes = pdf_buffer.getvalue()
                b64 = base64.b64encode(pdf_bytes).decode()
                href = f'<a href="data:application/pdf;base64,{b64}" download="Product_Mix_Analysis_Report_{region}_{selected_subset}.pdf">Download Region Subset PDF Report</a>'
                st.sidebar.markdown(href, unsafe_allow_html=True)
        brand = st.sidebar.selectbox("Select Brand", options=df[df['Region']==region]['Brand'].unique(), key="brand_select")
        product_type = st.sidebar.selectbox("Select Type", options=df[df['Region']==region]['Type'].unique(), key="type_select")
        region_subset = st.sidebar.selectbox("Select Region Subset", options=df[df['Region']==region]['Region subsets'].unique(), key="region_subset_select")
        st.sidebar.header("Analysis on")
        analysis_options = ["NSR Analysis", "Contribution Analysis", "EBITDA Analysis"]
        if 'analysis_type' not in st.session_state:
            st.session_state.analysis_type = "EBITDA Analysis"
        analysis_type = st.sidebar.radio("Select Analysis Type", analysis_options, index=analysis_options.index(st.session_state.analysis_type), key="analysis_type_radio")
        st.session_state.analysis_type = analysis_type
        premium_share = st.sidebar.slider("Adjust Premium Share (%)", 0, 100, 50)
        filtered_df = df[(df['Region'] == region) & (df['Brand'] == brand) &(df['Type'] == product_type) & (df['Region subsets'] == region_subset)].copy()
        if not filtered_df.empty:
            if analysis_type == 'NSR Analysis':
                cols = ['Normal NSR', 'Premium NSR']
                overall_col = 'Overall NSR'
            elif analysis_type == 'Contribution Analysis':
                cols = ['Normal Contribution', 'Premium Contribution']
                overall_col = 'Overall Contribution'
            elif analysis_type == 'EBITDA Analysis':
                cols = ['Normal EBITDA', 'Premium EBITDA']
                overall_col = 'Overall EBITDA'
            filtered_df[overall_col] = (filtered_df['Normal'] * filtered_df[cols[0]] + filtered_df['Premium'] * filtered_df[cols[1]]) / (filtered_df['Normal'] + filtered_df['Premium'])
            imaginary_col = f'Imaginary {overall_col}'
            filtered_df[imaginary_col] = ((1 - premium_share/100) * filtered_df[cols[0]] +
                                          (premium_share/100) * filtered_df[cols[1]])
            filtered_df['Difference'] = filtered_df[cols[1]] - filtered_df[cols[0]]
            filtered_df['Imaginary vs Overall Difference'] = filtered_df[imaginary_col] - filtered_df[overall_col]
            fig = go.Figure()
            for col in cols:
                fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[col],mode='lines+markers', name=col))
            fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[overall_col],mode='lines+markers', name=overall_col, line=dict(dash='dash')))
            fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[imaginary_col],mode='lines+markers', name=f'Imaginary {overall_col} ({premium_share}% Premium)',line=dict(color='brown', dash='dot')))
            x_labels = [f"{month}<br>(P-N: {diff:.2f})<br>(I-O: {i_diff:.2f})" for month, diff, i_diff in zip(filtered_df['Month'], filtered_df['Difference'], filtered_df['Imaginary vs Overall Difference'])]
            fig.update_layout(title=analysis_type,xaxis_title='Month (P-N: Premium - Normal, I-O: Imaginary - Overall)',yaxis_title='Value',legend_title='Metrics',hovermode="x unified",xaxis=dict(tickmode='array', tickvals=list(range(len(x_labels))), ticktext=x_labels))
            st.plotly_chart(fig, use_container_width=True)
            st.subheader("Descriptive Statistics")
            desc_stats = filtered_df[cols + [overall_col, imaginary_col]].describe()
            st.dataframe(desc_stats.style.format("{:.2f}"), use_container_width=True)
            st.subheader("Share of Normal and Premium Products")
            total_quantity = filtered_df['Normal'] + filtered_df['Premium']
            normal_share = (filtered_df['Normal'] / total_quantity * 100).round(2)
            premium_share = (filtered_df['Premium'] / total_quantity * 100).round(2)
            share_df = pd.DataFrame({'Month': filtered_df['Month'],'Premium Share (%)': premium_share,'Normal Share (%)': normal_share})     
            fig_pie = px.pie(share_df, values=[normal_share.mean(), premium_share.mean()],names=['Normal', 'Premium'], title='Average Share Distribution',color=["N","P"],color_discrete_map={"N":"green","P":"blue"},hole=0.5)
            st.plotly_chart(fig_pie, use_container_width=True)  
            st.dataframe(share_df.set_index('Month').style.format("{:.2f}").background_gradient(cmap='RdYlGn'), use_container_width=True)
        else:
            st.warning("No data available for the selected combination.")
        st.markdown("</div>", unsafe_allow_html=True)
 elif selected == "About":
    st.title("About the Product Mix Analysis App")
    st.markdown("""This advanced data analysis application is designed to provide insightful visualizations and statistics for your Product(Normal and Premium) Mix data. 
    Key features include:
    - Interactive data filtering
    - Multiple analysis types (NSR, Contribution, EBITDA)
    - Dynamic visualizations with Plotly
    - Descriptive statistics and share analysis
    - Customizable Premium share adjustments""")
def load_lottieurl(url: str):
    try:
        r = requests.get(url)
        if r.status_code != 200:
            return None
        return r.json()
    except:
        return None
def trade():
 lottie_analysis = load_lottieurl("https://assets4.lottiefiles.com/packages/lf20_qp1q7mct.json")
 lottie_upload = load_lottieurl("https://assets9.lottiefiles.com/packages/lf20_ABViugg1T8.json")
 with st.sidebar:
    selected = option_menu(
        menu_title="Navigation",
        options=["Home", "Segment-Mix Analysis", "About"],
        icons=["house", "graph-up", "info-circle"],
        menu_icon="cast",
        default_index=0,)
 def create_pdf_report(region, df, region_subset=None):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    def add_page_number(canvas):
      canvas.saveState()
      canvas.setFont('Helvetica', 10)
      page_number_text = f"Page {canvas.getPageNumber()}"
      canvas.drawString(width - 100, 30, page_number_text)
      canvas.restoreState()
    def add_header(page_number):
        c.setFillColorRGB(0.2, 0.2, 0.7)  # Dark blue color for header
        c.rect(0, height - 50, width, 50, fill=True)
        c.setFillColorRGB(1, 1, 1)  # White color for text
        c.setFont("Helvetica-Bold", 18)
        header_text = f"Segment Mix Analysis Report: {region}"
        if region_subset:
            header_text += f" ({region_subset})"
        c.drawString(30, height - 35, header_text)
    def add_front_page():
        c.setFillColorRGB(0.4,0.5,0.3)
        c.rect(0, 0, width, height, fill=True)
        c.setFillColorRGB(1, 1, 1)
        c.setFont("Helvetica-Bold", 36)
        c.drawCentredString(width / 2, height - 200, "Segment Mix Analysis Report")
        c.setFont("Helvetica", 24)
        report_title = f"Region: {region}"
        if region_subset:
            report_title += f" ({region_subset})"
        c.drawCentredString(width / 2, height - 250, report_title)
        c.setFont("Helvetica", 18)
        c.drawCentredString(width / 2, height - 300, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        add_page_number(c)
        c.showPage()
    def draw_graph(fig, x, y, width, height):
        img_buffer = BytesIO()
        fig.write_image(img_buffer, format="png",scale=2)
        img_buffer.seek(0)
        img = ImageReader(img_buffer)
        c.drawImage(img, x, y, width, height)
    def draw_table(data, x, y, col_widths):
        table = Table(data, colWidths=col_widths)
        table.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8), ('BOTTOMPADDING', (0, 0), (-1, 0), 6), ('BACKGROUND', (0, 1), (-1, -1), colors.beige),('TEXTCOLOR', (0, 1), (-1, -1), colors.black),('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),('FONTSIZE', (0, 1), (-1, -1), 6), ('TOPPADDING', (0, 1), (-1, -1), 3), ('BOTTOMPADDING', (0, 1), (-1, -1), 3),('GRID', (0, 0), (-1, -1), 1, colors.black)]))
        w, h = table.wrapOn(c, width, height)
        table.drawOn(c, x, y - h)
    def add_tutorial_page():
        c.setFont("Helvetica-Bold", 24)
        c.drawString(inch, height - inch, "Understanding the Segment Mix Analysis")
        drawing = Drawing(400, 200)
        lc = HorizontalLineChart()
        lc.x = 40
        lc.y = 50
        lc.height = 125
        lc.width = 300
        lc.data = [[random.randint(2000, 3000) for _ in range(12)],[random.randint(1500, 2500) for _ in range(12)], [random.randint(1800, 2800) for _ in range(12)],[random.randint(2200, 3200) for _ in range(12)],]
        lc.lines[0].strokeColor = colors.green
        lc.lines[1].strokeColor = colors.blue
        lc.lines[2].strokeColor = colors.pink
        lc.lines[3].strokeColor = colors.brown
        legend = Legend()
        legend.alignment = 'right'
        legend.x = 330
        legend.y = 150
        legend.colorNamePairs = [(colors.green, 'Trade EBITDA'),(colors.blue, 'Non-Trade EBITDA'),(colors.crimson, 'Overall EBITDA'),(colors.brown, 'Imaginary EBITDA'),]
        drawing.add(lc)
        drawing.add(legend)
        renderPDF.draw(drawing, c, inch, height - 300)
        c.setFont("Helvetica-Bold", 18)
        c.drawString(inch, height - 350, "Key Concepts:")
        concepts = [("Overall EBITDA:", "Weighted average of Trade and Non-Trade EBITDA based on their actual quantities."),("Imaginary EBITDA:", "Calculated by adjusting shares based on the following rules:"),("", "• If both (Trade,Non-Trade) are present: Trade +5%, Non-Trade -5%"),("", "• If only one is present: No change"),("Adjusted Shares:", "These adjustments aim to model potential improvements in product mix."),]
        text_object = c.beginText(inch, height - 380)
        for title, description in concepts:
            if title:
                text_object.setFont("Helvetica-Bold", 12)
                text_object.setFillColorRGB(0.7, 0.3, 0.1)
                text_object.textLine(title)
                text_object.setFont("Helvetica", 12)
                text_object.setFillColorRGB(0, 0, 0)  
            text_object.textLine(description)
            if not title:
                text_object.textLine("")
        c.drawText(text_object)
        add_page_number(c)
        c.showPage()
    def add_appendix():
        c.setFont("Helvetica-Bold", 24)
        c.drawString(inch, height - inch, "Appendix")
        sections = [("Graph Interpretation:", "Each line represents a different metric over time. The differences between metrics are shown below\n each month."),("Tables:", "The descriptive statistics table provides a summary of the data. The monthly share distribution table\n shows the proportion of Trade and Non-Trade Channel for each month."),("Importance:", "These visualizations help identify trends, compare performance across product categories, and\n understand the potential impact of changing product distributions."),]
        text_object = c.beginText(inch, height - 1.5*inch)
        text_object.setFont("Helvetica-Bold", 14)
        for title, content in sections:
            text_object.textLine(title)
            text_object.setFont("Helvetica", 12)
            text_object.textLines(content)
            text_object.textLine("")
            text_object.setFont("Helvetica-Bold", 14)
        c.drawText(text_object)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(inch, height - 4*inch, "Suggestions for Improvement:")
        suggestions = ["Increase the share of Trade Channel specifically for PPC, which typically have higher EBIDTA.","Analyze factors contributing to higher EBIDTA in Trade Channel,and apply insights to Non-Trade.","Regularly review and adjust pricing strategies to optimize EBITDA across all channels.","Invest in product innovation to expand Trade Channel offerings.",]
        text_object = c.beginText(inch, height - 4.3*inch)
        text_object.setFont("Helvetica", 12)
        for suggestion in suggestions:
            text_object.textLine(f"• {suggestion}")
        c.drawText(text_object)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(inch, height - 5.2*inch, "Limitations:")
        limitations = ["This analysis is based on historical data and may not predict future market changes.","External factors such as economic conditions are not accounted for in this report.","This report analyzes the EBIDTA for Trade and Non-Trade channel ceteris paribus.",]
        text_object = c.beginText(inch, height - 5.5*inch)
        text_object.setFont("Helvetica", 12)
        for limitation in limitations:
            text_object.textLine(f"• {limitation}")
        c.drawText(text_object)
        c.setFont("Helvetica", 12)
        c.drawString(inch, 2*inch, "We are currently working on including all other factors which impact the EBIDTA across GYR,")
        c.drawString(inch, 1.8*inch, "regions which will make this analysis more robust and helpful. We will also include NSR and") 
        c.drawString(inch,1.6*inch,"Contribution in our next report.")
        c.setFont("Helvetica-Bold", 14)
        c.drawString(inch, inch, "Thank You.")
        c.showPage()
    add_front_page()
    add_tutorial_page()
    brands = df['Brand'].unique()
    types = df['Type'].unique()
    region_subsets = df['Region subsets'].unique()
    page_number = 1
    for brand in brands:
        for product_type in types:
            for region_subset in region_subsets:
                filtered_df = df[(df['Region'] == region) & (df['Brand'] == brand) &(df['Type'] == product_type) & (df['Region subsets'] == region_subset)].copy()
                if not filtered_df.empty:
                    add_header(c)
                    cols = ['Trade EBITDA', 'Non-Trade EBITDA']
                    overall_col = 'Overall EBITDA'
                    total_quantity = filtered_df['Trade'] + filtered_df['Non-Trade']
                    filtered_df[overall_col] = ((filtered_df['Trade'] * filtered_df['Trade EBITDA'] + filtered_df['Non-Trade'] * filtered_df['Non-Trade EBITDA'])/ total_quantity)
                    filtered_df['Average Trade Share'] = filtered_df['Trade'] / total_quantity
                    filtered_df['Average Non-Trade Share'] = filtered_df['Non-Trade'] / total_quantity
                    def adjust_shares(row):
                        trade = row['Average Trade Share']
                        nontrade = row['Average Non-Trade Share']
                        if trade == 1 or nontrade == 1 :
                            return trade,nontrade
                        else:
                            trade = min(trade + 0.05, 1)
                            nontrade = min(nontrade - 0.05, 1 - trade)
                        return trade,nontrade
                    filtered_df['Adjusted Trade Share'], filtered_df['Adjusted Non-Trade Share'] = zip(*filtered_df.apply(adjust_shares, axis=1))
                    filtered_df['Imaginary EBITDA'] = (filtered_df['Adjusted Trade Share'] * filtered_df['Trade EBITDA'] +filtered_df['Adjusted Non-Trade Share'] * filtered_df['Non-Trade EBITDA'])
                    filtered_df['T-NT Difference'] = filtered_df['Trade EBITDA'] - filtered_df['Non-Trade EBITDA']
                    filtered_df['I-O Difference'] = filtered_df['Imaginary EBITDA'] - filtered_df[overall_col]
                    fig = go.Figure()
                    fig = make_subplots(rows=2, cols=1, row_heights=[0.58, 0.42], vertical_spacing=0.18)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['Trade EBITDA'],mode='lines+markers', name='Trade EBIDTA', line=dict(color='green')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['Non-Trade EBITDA'],mode='lines+markers', name='Non-Trade EBIDTA', line=dict(color='blue')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[overall_col],mode='lines+markers', name=overall_col, line=dict(color='crimson', dash='dash')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['Imaginary EBITDA'],mode='lines+markers', name='Imaginary EBIDTA',line=dict(color='brown', dash='dot')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['I-O Difference'],mode='lines+markers+text', name='I-O Difference',text=filtered_df['I-O Difference'].round(2),textposition='top center',textfont=dict(size=8,weight="bold"),line=dict(color='fuchsia')), row=2, col=1)
                    mean_diff = filtered_df['I-O Difference'].mean()
                    if not np.isnan(mean_diff):
                        mean_diff=round(mean_diff)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=[mean_diff] * len(filtered_df),mode='lines', name=f'Mean I-O Difference[{mean_diff}]',line=dict(color='black', dash='dash')), row=2, col=1)
                    x_labels = [f"{month}<br>(T-NT: {g_r:.0f})<br>(I-O: {g_y:.0f}))" for month, g_r, g_y in zip(filtered_df['Month'],filtered_df['T-NT Difference'],filtered_df['I-O Difference'])]
                    fig.update_layout(title=f"EBITDA Analysis for {brand}(Type:-{product_type}) in {region}({region_subset})",legend_title='Metrics',plot_bgcolor='cornsilk',paper_bgcolor='lightcyan',height=710,)
                    fig.update_xaxes(tickmode='array', tickvals=list(range(len(x_labels))), ticktext=x_labels, row=1, col=1)
                    fig.update_xaxes(title_text='Months', row=2, col=1)
                    fig.update_yaxes(title_text='EBITDA(Rs./MT)', row=1, col=1)
                    fig.update_yaxes(title_text='I-O Difference(Rs./MT)', row=2, col=1)
                    # Add new page if needed
                    #if page_number > 1:
                        #c.showPage()
                    # Draw the graph
                    draw_graph(fig, 50, height - 410, 500, 350)
                    c.setFillColorRGB(0.2, 0.2, 0.7) 
                    c.setFont("Helvetica-Bold", 10) 
                    c.drawString(50, height - 425, "Descriptive Statistics")
                    desc_stats = filtered_df[['Trade','Non-Trade']+cols + [overall_col, 'Imaginary EBITDA']].describe().reset_index()
                    desc_stats = desc_stats[desc_stats['index'] != 'count'].round(2)  # Remove 'count' row
                    table_data = [['Metric'] + list(desc_stats.columns[1:])] + desc_stats.values.tolist()
                    draw_table(table_data, 50, height - 435, [45,45,45] + [75] * (len(desc_stats.columns) - 4))
                    c.setFont("Helvetica-Bold", 10)
                    c.drawString(50, height - 600, "Average Share Distribution")
                    average_shares = filtered_df[['Average Trade Share', 'Average Non-Trade Share']].mean()
                    share_fig = px.pie(values=average_shares.values,names=average_shares.index,color=average_shares.index,color_discrete_map={'Average Trade Share': 'green', 'Average Non-Trade Share': 'blue'},title="",hole=0.3)
                    share_fig.update_layout(width=475, height=475, margin=dict(l=0, r=0, t=0, b=0)) 
                    draw_graph(share_fig, 80, height - 810, 200, 200)  # Adjusted position and size
                    c.setFont("Helvetica-Bold", 10)
                    c.drawString(330, height - 600, "Monthly Share Distribution")
                    share_data = [['Month', 'Trade', 'Non-Trade']]
                    for _, row in filtered_df[['Month', 'Trade', 'Non-Trade','Average Trade Share', 'Average Non-Trade Share']].iterrows():
                        share_data.append([row['Month'],f"{row['Trade']:.0f} ({row['Average Trade Share']:.2%})",f"{row['Non-Trade']:.0f} ({row['Average Non-Trade Share']:.2%})"])
                    draw_table(share_data, 330, height - 620, [40, 60, 60, 60])
                    add_page_number(c)
                    c.showPage()
    for i in range(c.getPageNumber()):
        c.setPageSize((width, height))
        add_page_number(c)         
    add_appendix()
    c.save()
    buffer.seek(0)
    return buffer
 if selected == "Home":
    st.title("🔍 Advanced Segment Mix Analysis")
    st.markdown("Welcome to our advanced data analysis platform. Upload your Excel file to get started with interactive visualizations and insights.")
    st.markdown("<div class='upload-section'>", unsafe_allow_html=True)
    col1, col2 = st.columns([2, 1])
    with col1:
        uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx",key='TradevsNontradeuploader')
        if uploaded_file is not None:
            st.session_state.uploaded_file = uploaded_file
            st.success("File successfully uploaded! Please go to the Analysis page to view results.")
    with col2:
        if lottie_upload:
            st_lottie(lottie_upload, height=150, key="upload")
        else:
            st.image("https://cdn-icons-png.flaticon.com/512/4503/4503700.png", width=150)
    st.markdown("</div>", unsafe_allow_html=True)
 elif selected == "Segment-Mix Analysis":
    st.title("📈 Segment Mix Dashboard")
    if 'uploaded_file' not in st.session_state or st.session_state.uploaded_file is None:
        st.warning("Please upload an Excel file on the Home page to begin the analysis.")
    else:
        df = pd.read_excel(st.session_state.uploaded_file)
        st.markdown("<div class='analysis-section'>", unsafe_allow_html=True)
        if lottie_analysis:
            st_lottie(lottie_analysis, height=200, key="analysis")
        else:
            st.image("https://cdn-icons-png.flaticon.com/512/2756/2756778.png", width=200)
        st.sidebar.header("Filter Options")
        region = st.sidebar.selectbox("Select Region", options=df['Region'].unique(), key="region_select")
        st.sidebar.subheader(f"Download Report for {region}")
        download_choice = st.sidebar.radio("Choose report type:",('Full Region', 'Region Subset'))
        if download_choice == 'Full Region':
            if st.sidebar.button(f"Download Full Report for {region}"):
                subset_df = df[(df['Region'] == region) & (df['Type'] != 'PPC Premium')]
                pdf_buffer = create_pdf_report(region, subset_df)
                pdf_bytes = pdf_buffer.getvalue()
                b64 = base64.b64encode(pdf_bytes).decode()
                href = f'<a href="data:application/pdf;base64,{b64}" download="Segment_Mix_Analysis_Report_{region}.pdf">Download Full Region PDF Report</a>'
                st.sidebar.markdown(href, unsafe_allow_html=True)
        else:
            region_subsets = df[df['Region'] == region]['Region subsets'].unique()
            selected_subset = st.sidebar.selectbox("Select Region Subset", options=region_subsets)
            if st.sidebar.button(f"Download Report for {region} - {selected_subset}"):
                subset_df = df[(df['Region'] == region) & (df['Region subsets'] == selected_subset) & (df['Type'] != 'PPC Premium')]
                pdf_buffer = create_pdf_report(region, subset_df, selected_subset)
                pdf_bytes = pdf_buffer.getvalue()
                b64 = base64.b64encode(pdf_bytes).decode()
                href = f'<a href="data:application/pdf;base64,{b64}" download="GYR_Analysis_Report_{region}_{selected_subset}.pdf">Download Region Subset PDF Report</a>'
                st.sidebar.markdown(href, unsafe_allow_html=True)
        brand = st.sidebar.selectbox("Select Brand", options=df[df['Region']==region]['Brand'].unique(), key="brand_select")
        product_type = st.sidebar.selectbox("Select Type", options=df[df['Region']==region]['Type'].unique(), key="type_select")
        region_subset = st.sidebar.selectbox("Select Region Subset", options=df[df['Region']==region]['Region subsets'].unique(), key="region_subset_select")
        st.sidebar.header("Analysis on")
        analysis_options = ["NSR Analysis", "Contribution Analysis", "EBITDA Analysis"]
        if 'analysis_type' not in st.session_state:
            st.session_state.analysis_type = "EBITDA Analysis"  
        analysis_type = st.sidebar.radio("Select Analysis Type", analysis_options, index=analysis_options.index(st.session_state.analysis_type), key="analysis_type_radio")
        st.session_state.analysis_type = analysis_type
        trade_share = st.sidebar.slider("Adjust Trade Share (%)", 0, 100, 50)
        filtered_df = df[(df['Region'] == region) & (df['Brand'] == brand) &(df['Type'] == product_type) & (df['Region subsets'] == region_subset)].copy()
        if not filtered_df.empty:
            if analysis_type == 'NSR Analysis':
                cols = ['Trade NSR', 'Non-Trade NSR']
                overall_col = 'Overall NSR'
            elif analysis_type == 'Contribution Analysis':
                cols = ['Trade Contribution', 'Non-Trade Contribution']
                overall_col = 'Overall Contribution'
            elif analysis_type == 'EBITDA Analysis':
                cols = ['Trade EBITDA', 'Non-Trade EBITDA']
                overall_col = 'Overall EBITDA'
            filtered_df[overall_col] = (filtered_df['Trade'] * filtered_df[cols[0]] +filtered_df['Non-Trade'] * filtered_df[cols[1]]) / (filtered_df['Trade'] + filtered_df['Non-Trade'])
            imaginary_col = f'Imaginary {overall_col}'
            filtered_df[imaginary_col] = ((1 - trade_share/100) * filtered_df[cols[1]] +(trade_share/100) * filtered_df[cols[0]])
            filtered_df['Difference'] = filtered_df[cols[0]] - filtered_df[cols[1]]
            filtered_df['Imaginary vs Overall Difference'] = filtered_df[imaginary_col] - filtered_df[overall_col]
            fig = go.Figure()
            for col in cols:
                fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[col],mode='lines+markers', name=col))
            fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[overall_col],mode='lines+markers', name=overall_col, line=dict(dash='dash')))
            fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[imaginary_col],mode='lines+markers', name=f'Imaginary {overall_col} ({trade_share}% Trade)',line=dict(color='brown', dash='dot')))
            x_labels = [f"{month}<br>(T-NT: {diff:.2f})<br>(I-O: {i_diff:.2f})" for month, diff, i_diff in zip(filtered_df['Month'], filtered_df['Difference'], filtered_df['Imaginary vs Overall Difference'])]
            fig.update_layout(title=analysis_type,xaxis_title='Month (T-NT: Trade - Non-Trade, I-O: Imaginary - Overall)',yaxis_title='Value',legend_title='Metrics',hovermode="x unified",xaxis=dict(tickmode='array', tickvals=list(range(len(x_labels))), ticktext=x_labels))
            st.plotly_chart(fig, use_container_width=True)
            st.subheader("Descriptive Statistics")
            desc_stats = filtered_df[cols + [overall_col, imaginary_col]].describe()
            st.dataframe(desc_stats.style.format("{:.2f}"), use_container_width=True)
            st.subheader("Share of Trade and Non-Trade Channel")
            total_quantity = filtered_df['Trade'] + filtered_df['Non-Trade']
            trade_share = (filtered_df['Trade'] / total_quantity * 100).round(2)
            nontrade_share = (filtered_df['Non-Trade'] / total_quantity * 100).round(2)
            share_df = pd.DataFrame({'Month': filtered_df['Month'],'Trade Share (%)': trade_share,'Non-Trade Share (%)': nontrade_share}) 
            fig_pie = px.pie(share_df, values=[trade_share.mean(), nontrade_share.mean()],names=['Trade', 'Non-Trade'], title='Average Share Distribution',color=["T","NT"],color_discrete_map={"T":"green","NT":"blue"},hole=0.5)
            st.plotly_chart(fig_pie, use_container_width=True)        
            st.dataframe(share_df.set_index('Month').style.format("{:.2f}").background_gradient(cmap='RdYlGn'), use_container_width=True)
        else:
            st.warning("No data available for the selected combination.")
        st.markdown("</div>", unsafe_allow_html=True)
 elif selected == "About":
    st.title("About the Segment Mix Analysis App")
    st.markdown("""
    This advanced data analysis application is designed to provide insightful visualizations and statistics for your Segment(Trade,Non-Trade) Mix data. 
    Key features include:
    - Interactive data filtering
    - Multiple analysis types (NSR, Contribution, EBITDA)
    - Dynamic visualizations with Plotly
    - Descriptive statistics and share analysis
    - Customizable Trade share adjustments
    """)
def load_lottieurl(url: str):
    try:
        r = requests.get(url)
        if r.status_code != 200:
            return None
        return r.json()
    except:
        return None
def green():
 with st.sidebar:
    selected = option_menu(menu_title="Navigation",options=["Home", "Geo-Mix Analysis", "About"],icons=["house", "graph-up", "info-circle"],menu_icon="cast",default_index=0,)
 lottie_analysis = load_lottieurl("https://assets4.lottiefiles.com/packages/lf20_qp1q7mct.json")
 lottie_upload = load_lottieurl("https://assets9.lottiefiles.com/packages/lf20_ABViugg1T8.json")
 def create_pdf_report(region, df, region_subset=None):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    def add_page_number(canvas):
      canvas.saveState()
      canvas.setFont('Helvetica', 10)
      page_number_text = f"Page {canvas.getPageNumber()}"
      canvas.drawString(width - 100, 30, page_number_text)
      canvas.restoreState()
    def add_header(page_number):
        c.setFillColorRGB(0.2, 0.2, 0.7)  # Dark blue color for header
        c.rect(0, height - 50, width, 50, fill=True)
        c.setFillColorRGB(1, 1, 1)  # White color for text
        c.setFont("Helvetica-Bold", 18)
        header_text = f"GYR Analysis Report: {region}"
        if region_subset:
            header_text += f" ({region_subset})"
        c.drawString(30, height - 35, header_text)
    def add_front_page():
        c.setFillColorRGB(0.4,0.5,0.3)
        c.rect(0, 0, width, height, fill=True)
        c.setFillColorRGB(1, 1, 1)
        c.setFont("Helvetica-Bold", 36)
        c.drawCentredString(width / 2, height - 200, "GYR Analysis Report")
        c.setFont("Helvetica", 24)
        report_title = f"Region: {region}"
        if region_subset:
            report_title += f" ({region_subset})"
        c.drawCentredString(width / 2, height - 250, report_title)
        c.setFont("Helvetica", 18)
        c.drawCentredString(width / 2, height - 300, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        add_page_number(c)
        c.showPage()
    def draw_graph(fig, x, y, width, height):
        img_buffer = BytesIO()
        fig.write_image(img_buffer, format="png",scale=2)
        img_buffer.seek(0)
        img = ImageReader(img_buffer)
        c.drawImage(img, x, y, width, height)
    def draw_table(data, x, y, col_widths):
        table = Table(data, colWidths=col_widths)
        table.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),('FONTSIZE', (0, 0), (-1, 0), 8),  ('BOTTOMPADDING', (0, 0), (-1, 0), 6), ('BACKGROUND', (0, 1), (-1, -1), colors.beige),('TEXTCOLOR', (0, 1), (-1, -1), colors.black),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),('FONTSIZE', (0, 1), (-1, -1), 6), ('TOPPADDING', (0, 1), (-1, -1), 3), ('BOTTOMPADDING', (0, 1), (-1, -1), 3), ('GRID', (0, 0), (-1, -1), 1, colors.black)]))
        w, h = table.wrapOn(c, width, height)
        table.drawOn(c, x, y - h)
    def add_tutorial_page():
        c.setFont("Helvetica-Bold", 24)
        c.drawString(inch, height - inch, "Understanding the GYR Analysis")
        drawing = Drawing(400, 200)
        lc = HorizontalLineChart()
        lc.x = 40
        lc.y = 50
        lc.height = 125
        lc.width = 300
        lc.data = [[random.randint(2000, 3000) for _ in range(12)],[random.randint(1500, 2500) for _ in range(12)],[random.randint(1000, 2000) for _ in range(12)], [random.randint(1800, 2800) for _ in range(12)],[random.randint(2200, 3200) for _ in range(12)], ]
        lc.lines[0].strokeColor = colors.green
        lc.lines[1].strokeColor = colors.yellow
        lc.lines[2].strokeColor = colors.red
        lc.lines[3].strokeColor = colors.blue
        lc.lines[4].strokeColor = colors.purple
        legend = Legend()
        legend.alignment = 'right'
        legend.x = 330
        legend.y = 150
        legend.colorNamePairs = [(colors.green, 'Green EBITDA'),(colors.yellow, 'Yellow EBITDA'),(colors.red, 'Red EBITDA'),(colors.blue, 'Overall EBITDA'),(colors.purple, 'Imaginary EBITDA'),]
        drawing.add(lc)
        drawing.add(legend)
        renderPDF.draw(drawing, c, inch, height - 300)
        c.setFont("Helvetica-Bold", 18)
        c.drawString(inch, height - 350, "Key Concepts:")
        concepts = [("Overall EBITDA:", "Weighted average of Green, Yellow, and Red EBITDA based on their actual quantities."),("Imaginary EBITDA:", "Calculated by adjusting shares based on the following rules:"),("", "• If all three (Green, Yellow, Red) are present: Green +5%, Yellow +2.5%, Red -7.5%"),("", "• If only two are present: Superior one (Green in GR or GY, Yellow in YR) +5%, other -5%"),("", "• If only one is present: No change"),("Adjusted Shares:", "These adjustments aim to model potential improvements in product mix."),]
        text_object = c.beginText(inch, height - 380)
        for title, description in concepts:
            if title:
                text_object.setFont("Helvetica-Bold", 12)
                text_object.setFillColorRGB(0.7, 0.3, 0.1) 
                text_object.textLine(title)
                text_object.setFont("Helvetica", 12)
                text_object.setFillColorRGB(0, 0, 0) 
            text_object.textLine(description)
            if not title:
                text_object.textLine("")
        c.drawText(text_object)
        add_page_number(c)
        c.showPage()
    def add_appendix():
        c.setFont("Helvetica-Bold", 24)
        c.drawString(inch, height - inch, "Appendix")
        sections = [("Graph Interpretation:", "Each line represents a different metric over time. The differences between metrics are shown below\n each month."),("Tables:", "The descriptive statistics table provides a summary of the data. The monthly share distribution table\n shows the proportion of Green, Yellow, and Red products for each month."),("Importance:", "These visualizations help identify trends, compare performance across product categories, and\n understand the potential impact of changing product distributions."),]
        text_object = c.beginText(inch, height - 1.5*inch)
        text_object.setFont("Helvetica-Bold", 14)
        for title, content in sections:
            text_object.textLine(title)
            text_object.setFont("Helvetica", 12)
            text_object.textLines(content)
            text_object.textLine("")
            text_object.setFont("Helvetica-Bold", 14)
        c.drawText(text_object)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(inch, height - 4*inch, "Suggestions for Improvement:")
        suggestions = ["Increase the share of Green Region products, which typically have higher EBIDTA margins.","Analyze factors contributing to higher EBIDTA in Green zone,and apply insights to Red zone.","Regularly review and adjust pricing strategies to optimize EBITDA across all product categories.","Invest in product innovation to expand Green and Yellow region offerings.",]
        text_object = c.beginText(inch, height - 4.3*inch)
        text_object.setFont("Helvetica", 12)
        for suggestion in suggestions:
            text_object.textLine(f"• {suggestion}")
        c.drawText(text_object)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(inch, height - 5.2*inch, "Limitations:")
        limitations = ["This analysis is based on historical data and may not predict future market changes.","External factors such as economic conditions are not accounted for in this report.","This report analyzes the EBIDTA for GYR keeping everything else constant.",]
        text_object = c.beginText(inch, height - 5.5*inch)
        text_object.setFont("Helvetica", 12)
        for limitation in limitations:
            text_object.textLine(f"• {limitation}")
        c.drawText(text_object)
        c.setFont("Helvetica", 12)
        c.drawString(inch, 2*inch, "We are currently working on including all other factors which impact the EBIDTA across GYR,")
        c.drawString(inch, 1.8*inch, "regions which will make this analysis more robust and helpful. We will also include NSR and") 
        c.drawString(inch,1.6*inch,"Contribution in our next report.")
        c.setFont("Helvetica-Bold", 14)
        c.drawString(inch, inch, "Thank You.")
        c.showPage()
    add_front_page()
    add_tutorial_page()
    brands = df['Brand'].unique()
    types = df['Type'].unique()
    region_subsets = df['Region subsets'].unique()
    page_number = 1
    for brand in brands:
        for product_type in types:
            for region_subset in region_subsets:
                filtered_df = df[(df['Region'] == region) & (df['Brand'] == brand) &(df['Type'] == product_type) & (df['Region subsets'] == region_subset)].copy()
                if not filtered_df.empty:
                    add_header(c)
                    cols = ['Green EBITDA', 'Yellow EBITDA', 'Red EBITDA']
                    overall_col = 'Overall EBITDA'
                    total_quantity = filtered_df['Green'] + filtered_df['Yellow'] + filtered_df['Red']
                    filtered_df[overall_col] = ((filtered_df['Green'] * filtered_df['Green EBITDA'] +filtered_df['Yellow'] * filtered_df['Yellow EBITDA'] + filtered_df['Red'] * filtered_df['Red EBITDA']) / total_quantity)
                    filtered_df['Average Green Share'] = filtered_df['Green'] / total_quantity
                    filtered_df['Average Yellow Share'] = filtered_df['Yellow'] / total_quantity
                    filtered_df['Average Red Share'] = filtered_df['Red'] / total_quantity
                    def adjust_shares(row):
                        green = row['Average Green Share']
                        yellow = row['Average Yellow Share']
                        red = row['Average Red Share']
                        if green == 1 or yellow == 1 or red == 1:
                            return green, yellow, red
                        elif red == 0:
                            green = min(green +0.05, 1)
                            yellow = max(1-green, 0)
                        elif green == 0 and yellow == 0:
                            return green, yellow, red
                        elif green == 0:
                            yellow = min(yellow + 0.05, 1)
                            red = max(1 - yellow, 0)
                        elif yellow == 0:
                            green = min(green + 0.05, 1)
                            red = max(1 - green, 0)
                        else:
                            green = min(green + 0.05, 1)
                            yellow = min(yellow + 0.025, 1 - green)
                            red = max(1 - green - yellow, 0)
                        return green, yellow, red
                    filtered_df['Adjusted Green Share'], filtered_df['Adjusted Yellow Share'], filtered_df['Adjusted Red Share'] = zip(*filtered_df.apply(adjust_shares, axis=1))
                    filtered_df['Imaginary EBITDA'] = (filtered_df['Adjusted Green Share'] * filtered_df['Green EBITDA'] +filtered_df['Adjusted Yellow Share'] * filtered_df['Yellow EBITDA'] +filtered_df['Adjusted Red Share'] * filtered_df['Red EBITDA'])
                    filtered_df['G-R Difference'] = filtered_df['Green EBITDA'] - filtered_df['Red EBITDA']
                    filtered_df['G-Y Difference'] = filtered_df['Green EBITDA'] - filtered_df['Yellow EBITDA']
                    filtered_df['Y-R Difference'] = filtered_df['Yellow EBITDA'] - filtered_df['Red EBITDA']
                    filtered_df['I-O Difference'] = filtered_df['Imaginary EBITDA'] - filtered_df[overall_col]
                    fig = go.Figure()
                    fig = make_subplots(rows=2, cols=1, row_heights=[0.58, 0.42], vertical_spacing=0.18)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['Green EBITDA'],mode='lines+markers', name='Green EBIDTA', line=dict(color='green')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['Yellow EBITDA'],mode='lines+markers', name='Yellow EBIDTA', line=dict(color='yellow')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['Red EBITDA'],mode='lines+markers', name='Red EBIDTA', line=dict(color='red')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[overall_col],mode='lines+markers', name=overall_col, line=dict(color='blue', dash='dash')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['Imaginary EBITDA'],mode='lines+markers', name='Imaginary EBIDTA',line=dict(color='purple', dash='dot')), row=1, col=1)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df['I-O Difference'],mode='lines+markers+text', name='I-O Difference',text=filtered_df['I-O Difference'].round(2),textposition='top center',textfont=dict(size=8,weight="bold"),line=dict(color='fuchsia')), row=2, col=1)
                    mean_diff = filtered_df['I-O Difference'].mean()
                    if not np.isnan(mean_diff):
                        mean_diff=round(mean_diff)
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=[mean_diff] * len(filtered_df),mode='lines', name=f'Mean I-O Difference[{mean_diff}]',line=dict(color='black', dash='dash')), row=2, col=1)
                    x_labels = [f"{month}<br>(G-R: {g_r:.0f})<br>(G-Y: {g_y:.0f})<br>(Y-R: {y_r:.0f})" for month, g_r, g_y, y_r, i_o in zip(filtered_df['Month'], filtered_df['G-R Difference'],filtered_df['G-Y Difference'],filtered_df['Y-R Difference'],filtered_df['I-O Difference'])]
                    fig.update_layout(title=f"EBITDA Analysis for {brand}({product_type}) in {region}({region_subset})",legend_title='Metrics',plot_bgcolor='cornsilk',paper_bgcolor='lightcyan',height=710,)
                    fig.update_xaxes(tickmode='array', tickvals=list(range(len(x_labels))), ticktext=x_labels, row=1, col=1)
                    fig.update_xaxes(title_text='Months', row=2, col=1)
                    fig.update_yaxes(title_text='EBITDA(Rs./MT)', row=1, col=1)
                    fig.update_yaxes(title_text='I-O Difference(Rs./MT)', row=2, col=1)
                    # Add new page if needed
                    #if page_number > 1:
                        #c.showPage()
                    # Draw the graph
                    draw_graph(fig, 50, height - 410, 500, 350)
                    c.setFillColorRGB(0.2, 0.2, 0.7)  # Dark grey color for headers
                    c.setFont("Helvetica-Bold", 10)  # Reduced font size
                    c.drawString(50, height - 425, "Descriptive Statistics")
                    desc_stats = filtered_df[['Green','Yellow','Red']+cols + [overall_col, 'Imaginary EBITDA']].describe().reset_index()
                    desc_stats = desc_stats[desc_stats['index'] != 'count'].round(2)
                    table_data = [['Metric'] + list(desc_stats.columns[1:])] + desc_stats.values.tolist()
                    draw_table(table_data, 50, height - 435, [40,40,40,40] + [75] * (len(desc_stats.columns) - 4))  
                    c.setFont("Helvetica-Bold", 10) 
                    c.drawString(50, height - 600, "Average Share Distribution")
                    average_shares = filtered_df[['Average Green Share', 'Average Yellow Share', 'Average Red Share']].mean()
                    share_fig = px.pie(values=average_shares.values,names=average_shares.index,color=average_shares.index,color_discrete_map={'Average Green Share': 'green', 'Average Yellow Share': 'yellow', 'Average Red Share': 'red'},title="",hole=0.3)
                    share_fig.update_layout(width=475, height=475, margin=dict(l=0, r=0, t=0, b=0)) 
                    draw_graph(share_fig, 80, height - 810, 200, 200) 
                    c.setFont("Helvetica-Bold", 10)
                    c.drawString(330, height - 600, "Monthly Share Distribution")
                    share_data = [['Month', 'Green', 'Yellow', 'Red']]
                    for _, row in filtered_df[['Month', 'Green', 'Yellow', 'Red', 'Average Green Share', 'Average Yellow Share', 'Average Red Share']].iterrows():
                        share_data.append([row['Month'],f"{row['Green']:.0f} ({row['Average Green Share']:.2%})",f"{row['Yellow']:.0f} ({row['Average Yellow Share']:.2%})",f"{row['Red']:.0f} ({row['Average Red Share']:.2%})"])
                    draw_table(share_data, 330, height - 620, [40, 60, 60, 60])
                    add_page_number(c)
                    c.showPage()
    for i in range(c.getPageNumber()):
        c.setPageSize((width, height))
        add_page_number(c)         
    add_appendix()
    c.save()
    buffer.seek(0)
    return buffer
 if selected == "Home":
    st.title("🔍 Advanced Geo Mix Analysis")
    st.markdown("Welcome to our advanced data analysis platform. Upload your Excel file to get started with interactive visualizations and insights.")
    st.markdown("<div class='upload-section'>", unsafe_allow_html=True)
    col1, col2 = st.columns([2, 1])
    with col1:
        uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx",key="gyruploader")
        if uploaded_file is not None:
            st.session_state.uploaded_file = uploaded_file
            st.success("File successfully uploaded! Please go to the Analysis page to view results.")
    with col2:
        if lottie_upload:
            st_lottie(lottie_upload, height=150, key="upload")
        else:
            st.image("https://cdn-icons-png.flaticon.com/512/4503/4503700.png", width=150)
    st.markdown("</div>", unsafe_allow_html=True)
 elif selected == "Geo-Mix Analysis":
    st.title("📈 Geo Mix Dashboard")
    if 'uploaded_file' not in st.session_state or st.session_state.uploaded_file is None:
        st.warning("Please upload an Excel file on the Home page to begin the analysis.")
    else:
        df = pd.read_excel(st.session_state.uploaded_file)
        st.markdown("<div class='analysis-section'>", unsafe_allow_html=True)
        if lottie_analysis:
            st_lottie(lottie_analysis, height=200, key="analysis")
        else:
            st.image("https://cdn-icons-png.flaticon.com/512/2756/2756778.png", width=200)
        st.sidebar.header("Filter Options")
        region = st.sidebar.selectbox("Select Region", options=df['Region'].unique(), key="region_select")
        st.sidebar.subheader(f"Download Report for {region}")
        download_choice = st.sidebar.radio("Choose report type:",('Full Region', 'Region Subset'))
        if download_choice == 'Full Region':
            if st.sidebar.button(f"Download Full Report for {region}"):
                pdf_buffer = create_pdf_report(region, df)
                pdf_bytes = pdf_buffer.getvalue()
                b64 = base64.b64encode(pdf_bytes).decode()
                href = f'<a href="data:application/pdf;base64,{b64}" download="GYR_Analysis_Report_{region}.pdf">Download Full Region PDF Report</a>'
                st.sidebar.markdown(href, unsafe_allow_html=True)
        else:
            region_subsets = df[df['Region'] == region]['Region subsets'].unique()
            selected_subset = st.sidebar.selectbox("Select Region Subset", options=region_subsets)
            if st.sidebar.button(f"Download Report for {region} - {selected_subset}"):
                subset_df = df[(df['Region'] == region) & (df['Region subsets'] == selected_subset)]
                pdf_buffer = create_pdf_report(region, subset_df, selected_subset)
                pdf_bytes = pdf_buffer.getvalue()
                b64 = base64.b64encode(pdf_bytes).decode()
                href = f'<a href="data:application/pdf;base64,{b64}" download="GYR_Analysis_Report_{region}_{selected_subset}.pdf">Download Region Subset PDF Report</a>'
                st.sidebar.markdown(href, unsafe_allow_html=True)
        brand = st.sidebar.selectbox("Select Brand", options=df[df['Region']==region]['Brand'].unique(), key="brand_select")
        product_type = st.sidebar.selectbox("Select Type", options=df[df['Region']==region]['Type'].unique(), key="type_select")
        region_subset = st.sidebar.selectbox("Select Region Subset", options=df[df['Region']==region]['Region subsets'].unique(), key="region_subset_select")
        st.sidebar.header("Analysis on")
        analysis_options = ["NSR Analysis", "Contribution Analysis", "EBITDA Analysis"]
        if 'analysis_type' not in st.session_state:
            st.session_state.analysis_type = "EBITDA Analysis"
        analysis_type = st.sidebar.radio("Select Analysis Type", analysis_options, index=analysis_options.index(st.session_state.analysis_type), key="analysis_type_radio")
        st.session_state.analysis_type = analysis_type
        green_share = st.sidebar.slider("Adjust Green Share (%)", 0, 99, 50, key="green_share_slider")
        yellow_share = st.sidebar.slider("Adjust Yellow Share (%)", 0, 100-green_share, 0, key="yellow_share_slider")
        red_share = 100 - green_share - yellow_share
        st.sidebar.text(f"Red Share: {red_share}%")
        filtered_df = df[(df['Region'] == region) & (df['Brand'] == brand) &(df['Type'] == product_type) & (df['Region subsets'] == region_subset)].copy()
        if not filtered_df.empty:
            if analysis_type == 'NSR Analysis':
                cols = ['Green NSR', 'Yellow NSR', 'Red NSR']
                overall_col = 'Overall NSR'
            elif analysis_type == 'Contribution Analysis':
                cols = ['Green Contribution', 'Yellow Contribution','Red Contribution']
                overall_col = 'Overall Contribution'
            elif analysis_type == 'EBITDA Analysis':
                cols = ['Green EBITDA', 'Yellow EBITDA','Red EBITDA']
                overall_col = 'Overall EBITDA'
            filtered_df[overall_col] = (filtered_df['Green'] * filtered_df[cols[0]] +filtered_df['Yellow'] * filtered_df[cols[1]] + filtered_df['Red']*filtered_df[cols[2]]) / (filtered_df['Green'] + filtered_df['Yellow']+filtered_df['Red'])
            imaginary_col = f'Imaginary {overall_col}'
            filtered_df[imaginary_col] = ((1 - (green_share+yellow_share)/100) * filtered_df[cols[2]] +(green_share/100) * filtered_df[cols[0]] + (yellow_share/100) * filtered_df[cols[1]])
            filtered_df['G-Y Difference'] = filtered_df[cols[0]] - filtered_df[cols[1]]
            filtered_df['G-R Difference'] = filtered_df[cols[0]] - filtered_df[cols[2]]
            filtered_df['Y-R Difference'] = filtered_df[cols[1]] - filtered_df[cols[2]]
            filtered_df['Imaginary vs Overall Difference'] = filtered_df[imaginary_col] - filtered_df[overall_col]
            fig = go.Figure()
            if cols[0] in cols:
                  fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[cols[0]],mode='lines+markers', name=cols[0],line_color="green"))
            if cols[1] in cols:
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[cols[1]],mode='lines+markers', name=cols[1],line_color="yellow"))
            if cols[2] in cols:
                    fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[cols[2]],mode='lines+markers', name=cols[2],line_color="red"))
            fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[overall_col],mode='lines+markers', name=overall_col, line=dict(dash='dash')))
            fig.add_trace(go.Scatter(x=filtered_df['Month'], y=filtered_df[imaginary_col],mode='lines+markers', name=f'Imaginary {overall_col} ({green_share}% Green & {yellow_share}% Yellow)',line=dict(color='brown', dash='dot')))
            x_labels = [f"{month}<br>(G-Y: {diff:.2f})<br>(G-R: {i_diff:.2f})<br>(Y-R: {j_diff:.2f})<br>(I-O: {k_diff:.2f})" for month, diff, i_diff, j_diff, k_diff in 
                        zip(filtered_df['Month'], filtered_df['G-Y Difference'], filtered_df['G-R Difference'], filtered_df['Y-R Difference'], filtered_df['Imaginary vs Overall Difference'])]
            fig.update_layout(title=analysis_type,xaxis_title='Month (G-Y: Green - Red,G-R: Green - Red,Y-R: Yellow - Red, I-O: Imaginary - Overall)',yaxis_title='Value',legend_title='Metrics',hovermode="x unified",xaxis=dict(tickmode='array', tickvals=list(range(len(x_labels))), ticktext=x_labels))
            st.plotly_chart(fig, use_container_width=True)
            st.subheader("Descriptive Statistics")
            desc_stats = filtered_df[cols + [overall_col, imaginary_col]].describe()
            st.dataframe(desc_stats.style.format("{:.2f}").background_gradient(cmap='Blues'), use_container_width=True)
            st.subheader("Share of Green, Yellow, and Red Products")
            total_quantity = filtered_df['Green'] + filtered_df['Yellow'] + filtered_df['Red']
            green_share = (filtered_df['Green'] / total_quantity * 100).round(2)
            yellow_share = (filtered_df['Yellow'] / total_quantity * 100).round(2)
            red_share = (filtered_df['Red'] / total_quantity * 100).round(2)
            share_df = pd.DataFrame({'Month': filtered_df['Month'],'Green Share (%)': green_share,'Yellow Share (%)': yellow_share,'Red Share (%)': red_share})
            fig_pie = px.pie(share_df, values=[green_share.mean(), yellow_share.mean(), red_share.mean()],names=['Green', 'Yellow', 'Red'], title='Average Share Distribution',color=["G","Y","R"],color_discrete_map={"G":"green","Y":"yellow","R":"red"},hole=0.5)
            st.plotly_chart(fig_pie, use_container_width=True)  
            st.dataframe(share_df.set_index('Month').style.format("{:.2f}").background_gradient(cmap='RdYlGn'), use_container_width=True)
        else:
            st.warning("No data available for the selected combination.")
        st.markdown("</div>", unsafe_allow_html=True)
 elif selected == "About":
    st.title("About the GYR Analysis App")
    st.markdown("""
    This advanced data analysis application is designed to provide insightful visualizations and statistics for your GYR (Green, Yellow, Red) data. 
    Key features include:
    - Interactive data filtering
    - Multiple analysis types (NSR, Contribution, EBITDA)
    - Dynamic visualizations with Plotly
    - Descriptive statistics and share analysis
    - Customizable Green and Yellow share adjustments""")
def projection():
 def get_cookie_password():
    if 'cookie_password' not in st.session_state:
        st.session_state.cookie_password = secrets.token_hex(16)
    return st.session_state.cookie_password
 cookies = EncryptedCookieManager(prefix="sales_predictor_",password=get_cookie_password())
 CORRECT_PASSWORD = "prasoonA1@"
 MAX_ATTEMPTS = 5
 LOCKOUT_DURATION = 3600 
 def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()
 def check_password():
    if not cookies.ready():
        st.warning("Initializing cookies...")
        return False
    st.markdown("""<style>.stTextInput > div > div > input {background-color: #f0f0f0;color: #333;border: 2px solid #4a69bd;border-radius: 5px;padding: 10px;font-size: 16px;}.stButton > button {background-color: #4a69bd;color: white;border: none;border-radius: 5px;padding: 10px 20px;font-size: 16px;cursor: pointer;transition: background-color 0.3s;}.stButton > button:hover {background-color: #82ccdd;}.attempt-text {color: #ff4b4b;font-size: 14px;margin-top: 5px;text-align: center;}</style>""", unsafe_allow_html=True)
    lockout_time = cookies.get('lockout_time')
    if lockout_time is not None and time.time() < float(lockout_time):
        remaining_time = int(float(lockout_time) - time.time())
        st.error(f"Too many incorrect attempts. Please try again in {remaining_time // 60} minutes and {remaining_time % 60} seconds.")
        return False
    if 'login_attempts' not in st.session_state:
     login_attempts = cookies.get('login_attempts')
     st.session_state.login_attempts = int(login_attempts) if login_attempts is not None else 0
    def password_entered():
        if hash_password(st.session_state["password"]) == hash_password(CORRECT_PASSWORD):
            st.session_state["password_correct"] = True
            st.session_state.login_attempts = 0
            del st.session_state["password"]  # don't store password
        else:
            st.session_state["password_correct"] = False
            st.session_state.login_attempts += 1
            if st.session_state.login_attempts >= MAX_ATTEMPTS:
                cookies['lockout_time'] = str(time.time() + LOCKOUT_DURATION)
        cookies['login_attempts'] = str(st.session_state.login_attempts)
        cookies.save()
    if "password_correct" not in st.session_state:
        st.markdown("<h1 style='text-align: center; color: #4a69bd;'>Sales Prediction Simulator</h1>", unsafe_allow_html=True)
        st.markdown("<h3 style='text-align: center; color: #333;'>Please enter your password to access the application</h3>", unsafe_allow_html=True)
        st.text_input("Password", type="password", key="password")
        if st.button("Login"):
            password_entered()
        if st.session_state.login_attempts > 0:
            st.markdown(f"<p class='attempt-text'>Incorrect password. Attempt {st.session_state.login_attempts} of {MAX_ATTEMPTS}.</p>", unsafe_allow_html=True)
        return False
    elif st.session_state.get("password_correct", False):
        return True
    else:
        st.markdown("<h1 style='text-align: center; color: #4a69bd;'>Sales Prediction Simulator</h1>", unsafe_allow_html=True)
        st.markdown("<h3 style='text-align: center; color: #333;'>Please enter your password to access the application</h3>", unsafe_allow_html=True)
        st.text_input("Password", type="password", key="password")
        if st.button("Login"):
            password_entered()
        if st.session_state.login_attempts > 0:
            st.markdown(f"<p class='attempt-text'>Incorrect password. Attempt {st.session_state.login_attempts} of {MAX_ATTEMPTS}.</p>", unsafe_allow_html=True)
        return False
 if check_password():
  st.markdown("""<style>body {background-color: #0e1117;color: #ffffff;}.stApp {background-image: linear-gradient(45deg, #1e3799, #0c2461);}.big-font {font-size: 48px !important;font-weight: bold;color: lime;text-align: center;text-shadow: 2px 2px 4px #000000;}.subheader {font-size: 24px;color: moccasin;text-align: center;}.stButton>button {background-color: #4a69bd;color: white;border-radius: 20px;border: 2px solid #82ccdd;padding: 10px 24px;font-size: 16px;transition: all 0.3s;}.stButton>button:hover {background-color: #82ccdd;color: #0c2461;transform: scale(1.05);}.stProgress > div > div > div > div {background-color: #4a69bd;}.stSelectbox {background-color: #1e3799;}.stDataFrame {background-color: #0c2461;}.metric-value {color: gold !important;font-size: 24px !important;font-weight: bold !important;}.metric-label {color: white !important;}h3 {color: #ff9f43 !important;font-size: 28px !important;font-weight: bold !important;text-shadow: 1px 1px 2px #000000;}/* Updated styles for file uploader */.stFileUploader {background-color: rgba(255, 255, 255, 0.1);border-radius: 10px;padding: 20px;margin-bottom: 20px;}.custom-file-upload {display: inline-block;padding: 10px 20px;cursor: pointer;background-color: #4a69bd;color: #ffffff;border-radius: 5px;transition: all 0.3s;}.custom-file-upload:hover {background-color: #82ccdd;color: #0c2461;}.file-upload-text {font-size: 18px;color: fuchsia;font-weight: bold;margin-bottom: 10px;}/* Style for uploaded file name */.uploaded-filename {background-color: rgba(255, 255, 255, 0.2);color: cyan;padding: 10px;border-radius: 5px;margin-top: 10px;font-weight: bold;}</style>""", unsafe_allow_html=True)
  def custom_file_uploader(label, type):
    st.markdown(f'<p class="file-upload-text">{label}</p>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader("Choose file", type=type, key="file_uploader", label_visibility="collapsed")
    return uploaded_file
  @st.cache_data
  def load_data(file):
    data = pd.read_excel(file)
    return data
  @st.cache_resource
  def train_model(X, y):
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
    model = RandomForestRegressor(n_estimators=100, random_state=42)
    model.fit(X_train, y_train)
    return model, X_test, y_test
  def create_monthly_performance_graph(data):
    months = ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct','Nov']
    colors = px.colors.qualitative.Pastel
    fig = go.Figure()
    for i, month in enumerate(months):
        if month != 'Nov':
            target = data[f'Month Tgt ({month})'].iloc[0]
            achievement = data[f'Monthly Achievement({month})'].iloc[0]
            percentage = (achievement / target * 100) if target != 0 else 0
            fig.add_trace(go.Bar(x=[f"{month} Tgt", f"{month} Ach"],y=[target, achievement],name=month,marker_color=colors[i],text=[f"{target:,.0f}", f"{achievement:,.0f}<br>{percentage:.1f}%"],textposition='auto'))
        else:
            target = data['Month Tgt (Nov)'].iloc[0]
            projection = data['Predicted Nov 2024'].iloc[0]
            percentage = (projection / target * 100) if target != 0 else 0
            fig.add_trace(go.Bar(x=[f"{month} Tgt", f"{month} Proj"],y=[target, projection],name=month,marker_color=[colors[i], 'red'],text=[f"{target:,.0f}", f"{projection:,.0f}<br><span style='color:black'>{percentage:.1f}%</span>"],textposition='auto'))
    fig.update_layout(title='Monthly Performance',plot_bgcolor='rgba(255,255,255,0.1)',paper_bgcolor='rgba(0,0,0,0)',font_color='burlywood',title_font_color='burlywood',xaxis_title_font_color='burlywood',yaxis_title_font_color='burlywood',legend_font_color='burlywood',height=500,width=800,barmode='group')
    fig.update_xaxes(tickfont_color='peru')
    fig.update_yaxes(title_text='Sales', tickfont_color='peru')
    fig.update_traces(textfont_color='black')
    return fig
  def create_target_vs_projected_graph(data):
    fig = go.Figure()
    fig.add_trace(go.Bar(x=data['Zone'], y=data['Month Tgt (Nov)'], name='Month Target (Nov)', marker_color='#4a69bd'))
    fig.add_trace(go.Bar(x=data['Zone'], y=data['Predicted Nov 2024'], name='Projected Sales (Nov)', marker_color='#82ccdd'))
    fig.update_layout(title='October 2024: Target vs Projected Sales',barmode='group',plot_bgcolor='rgba(255,255,255,0.1)',paper_bgcolor='rgba(0,0,0,0)',font_color='burlywood',title_font_color='burlywood',xaxis_title_font_color='burlywood',yaxis_title_font_color='burlywood',legend_font_color='burlywood',height=500)
    fig.update_xaxes(title_text='Zone', tickfont_color='peru')
    fig.update_yaxes(title_text='Sales', tickfont_color='peru')
    return fig
  def prepare_data_for_pdf(data):
    excluded_zones = ['Bihar', 'J&K', 'North-I', 'Punjab,HP and J&K', 'U.P.+U.K.', 'Odisha+Jharkhand+Bihar']
    filtered_data = data[~data['Zone'].isin(excluded_zones)]
    filtered_data = filtered_data[filtered_data['Brand'].isin(['LC', 'PHD'])]
    lc_data = filtered_data[filtered_data['Brand'] == 'LC']
    phd_data = filtered_data[filtered_data['Brand'] == 'PHD']
    lc_phd_data = filtered_data
    totals = []
    for brand_data, brand_name in [(lc_data, 'LC'), (phd_data, 'PHD'), (lc_phd_data, 'LC+PHD')]:
        total_month_tgt_nov = brand_data['Month Tgt (Nov)'].sum()
        total_predicted_nov_2024 = brand_data['Predicted Nov 2024'].sum()
        total_nov_2023 = brand_data['Total Nov 2023'].sum()
        total_yoy_growth = (total_predicted_nov_2024 - total_nov_2023) / total_nov_2023 * 100
        totals.append({'Zone': 'All India Total','Brand': brand_name,'Month Tgt (Nov)': total_month_tgt_nov,'Predicted Oct 2024': total_predicted_nov_2024,'Total Oct 2023': total_nov_2023,'YoY Growth': total_yoy_growth})
    final_data = pd.concat([filtered_data, pd.DataFrame(totals)], ignore_index=True)
    final_data['Month Tgt (Nov)'] = final_data['Month Tgt (Nov)'].round().astype(int)
    final_data['Predicted Nov 2024'] = final_data['Predicted Nov 2024'].round().astype(int)
    final_data['Total Nov 2023'] = final_data['Total Nov 2023'].round().astype(int)
    final_data['YoY Growth'] = final_data['YoY Growth'].round(2)
    return final_data
  def create_pdf(data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,rightMargin=inch, leftMargin=inch,topMargin=0.2*inch, bottomMargin=0.5*inch)
    elements = []
    styles = getSampleStyleSheet()
    title_style = styles['Heading1']
    title_style.alignment = 1
    title = Paragraph("Sales Predictions for Novemeber 2024", title_style)
    elements.append(title)
    elements.append(Spacer(1, 0.15*inch))
    elements.append(Paragraph("<br/><br/>", styles['Normal']))
    pdf_data = prepare_data_for_pdf(data)
    table_data = [['Zone', 'Brand', 'Month Tgt (Nov)', 'Predicted Nov 2024', 'Total Nov 2023', 'YoY Growth']]
    for _, row in pdf_data.iterrows():
        table_data.append([row['Zone'],row['Brand'],f"{row['Month Tgt (Nov)']:,}",f"{row['Predicted Nov 2024']:,}",f"{row['Total Nov 2023']:,}",f"{row['YoY Growth']:.2f}%"])
    table_data[0][-1] = table_data[0][-1] + "*"  
    table = Table(table_data, colWidths=[1.25*inch, 0.80*inch, 1.5*inch, 1.75*inch, 1.5*inch, 1.20*inch],rowHeights=[0.60*inch] + [0.38*inch] * (len(table_data) - 1))
    style = TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4A708B')),('BACKGROUND', (0, len(table_data) - 3), (-1, len(table_data) - 1), colors.orange),('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),('FONTSIZE', (0, 0), (-1, 0), 12),('BOTTOMPADDING', (0, 0), (-1, 0), 10),('BACKGROUND', (0, 1), (-1, -4), colors.white),('TEXTCOLOR', (0, 1), (-1, -1), colors.black),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),('FONTSIZE', (0, 1), (-1, -1), 10),('GRID', (0, 0), (-1, -1), 1, colors.lightgrey)])
    table.setStyle(style)
    elements.append(table)
    footnote_style = getSampleStyleSheet()['Normal']
    footnote_style.fontSize = 8
    footnote_style.leading = 10 
    footnote_style.alignment = 0
    footnote = Paragraph("*YoY Growth is calculated using November 2023 sales and predicted November 2024 sales.", footnote_style)
    indented_footnote = Indenter(left=-0.75*inch)
    elements.append(Spacer(1, 0.15*inch))
    elements.append(indented_footnote)
    elements.append(footnote)
    elements.append(Indenter(left=0.5*inch))
    doc.build(elements)
    buffer.seek(0)
    return buffer
  def style_dataframe(df):
    styler = df.style
    for col in df.columns:
        if df[col].dtype in ['float64', 'int64']:
            styler.apply(lambda x: ['background-color: #f0f0f0'] * len(x), subset=[col])
        else:
            styler.apply(lambda x: ['background-color: #f0f0f0'] * len(x), subset=[col])
    numeric_format = {'November 2024 Target': '{:.0f}','November Projection': '{:.2f}','November 2023 Sales': '{:.0f}','YoY Growth(Projected)': '{:.2f}%'}
    styler.format(numeric_format)
    return styler
  def main():
    st.markdown('<p class="big-font">Sales Prediction Simulator</p>', unsafe_allow_html=True)
    st.markdown('<p class="subheader">Upload your data and unlock the future of sales!</p>', unsafe_allow_html=True)
    uploaded_file = custom_file_uploader("Choose your sales data file (Excel format)", ["xlsx"])
    if uploaded_file is not None:
        st.markdown(f'<div class="uploaded-filename">Uploaded file: {uploaded_file.name}</div>', unsafe_allow_html=True)
        data = load_data(uploaded_file)
        features = ['Month Tgt (Nov)', 'Monthly Achievement(Oct)', 'Total Oct 2023', 'Total Nov 2023','Monthly Achievement(Apr)', 'Monthly Achievement(May)', 'Monthly Achievement(June)','Monthly Achievement(July)', 'Monthly Achievement(Aug)','Monthly Achievement(Sep)']
        X = data[features]
        y = data['Total Oct 2023']
        model, X_test, y_test = train_model(X, y)
        st.sidebar.header("Control Panel")
        if 'selected_brands' not in st.session_state:
            st.session_state.selected_brands = []
        if 'selected_zones' not in st.session_state:
            st.session_state.selected_zones = []
        st.sidebar.subheader("Select Brands")
        for brand in data['Brand'].unique():
            if st.sidebar.checkbox(brand, key=f"brand_{brand}"):
                if brand not in st.session_state.selected_brands:
                    st.session_state.selected_brands.append(brand)
            elif brand in st.session_state.selected_brands:
                st.session_state.selected_brands.remove(brand)
        st.sidebar.subheader("Select Zones")
        for zone in data['Zone'].unique():
            if st.sidebar.checkbox(zone, key=f"zone_{zone}"):
                if zone not in st.session_state.selected_zones:
                    st.session_state.selected_zones.append(zone)
            elif zone in st.session_state.selected_zones:
                st.session_state.selected_zones.remove(zone)
        if st.session_state.selected_brands and st.session_state.selected_zones:
            filtered_data = data[data['Brand'].isin(st.session_state.selected_brands) & data['Zone'].isin(st.session_state.selected_zones)]
        else:
            filtered_data = data
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("<h3>Model Performance Metrics</h3>", unsafe_allow_html=True)
            y_pred = model.predict(X_test)
            mse = mean_squared_error(y_test, y_pred)
            r2 = r2_score(y_test, y_pred)
            st.markdown(f'<div class="metric-label">Accuracy Score</div><div class="metric-value">{r2:.2f}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="metric-label">Error Margin</div><div class="metric-value">{np.sqrt(mse):.2f}</div>', unsafe_allow_html=True)
            feature_importance = pd.DataFrame({'feature': features,'importance': model.feature_importances_}).sort_values('importance', ascending=False)
            fig_importance = px.bar(feature_importance, x='importance', y='feature', orientation='h',title='Feature Impact Analysis', labels={'importance': 'Impact', 'feature': 'Feature'})
            fig_importance.update_layout(plot_bgcolor='rgba(255,255,255,0.1)', paper_bgcolor='rgba(0,0,0,0)', font_color='burlywood',title_font_color='burlywood',xaxis_title_font_color='burlywood',yaxis_title_font_color='burlywood',legend_font_color='burlywood')
            fig_importance.update_xaxes(tickfont_color='peru')
            fig_importance.update_yaxes(tickfont_color='peru')
            filtered_data['FY 2025 Till Oct']= filtered_data['Monthly Achievement(Apr)']+filtered_data['Monthly Achievement(May)']+filtered_data['Monthly Achievement(June)']+filtered_data['Monthly Achievement(July)']+filtered_data['Monthly Achievement(Aug)']+filtered_data['Monthly Achievement(Sep)']+filtered_data['Monthly Achievement(Oct)']
            fig_predictions1 = go.Figure()
            fig_predictions1.add_trace(go.Bar(x=filtered_data['Zone'], y=filtered_data['FY 2025 Till Oct'], name='Till OctSales', marker_color='#4a69bd'))
            fig_predictions1.update_layout(title='FY 2025 Till Oct',barmode='group',plot_bgcolor='rgba(255,255,255,0.1)',paper_bgcolor='rgba(0,0,0,0)',font_color='burlywood',xaxis_title_font_color='burlywood',yaxis_title_font_color='burlywood',title_font_color='burlywood',legend_font_color='burlywood')
            fig_predictions1.update_xaxes(title_text='Zone', tickfont_color='peru')
            fig_predictions1.update_yaxes(title_text='Sales', tickfont_color='peru')
            st.plotly_chart(fig_importance, use_container_width=True)
            st.plotly_chart(fig_predictions1, use_container_width=True)
        with col2:
            st.markdown("<h3>Sales Forecast Visualization</h3>", unsafe_allow_html=True)
            X_2024 = filtered_data[features].copy()
            X_2024['Total Nov 2023'] = filtered_data['Total Nov 2023']
            predictions_2024 = model.predict(X_2024)
            filtered_data['Predicted Nov 2024'] = predictions_2024
            filtered_data['YoY Growth'] = (filtered_data['Predicted Nov 2024'] - filtered_data['Total Nov 2023']) / filtered_data['Total Nov 2023'] * 100
            fig_predictions = go.Figure()
            fig_predictions.add_trace(go.Bar(x=filtered_data['Zone'], y=filtered_data['Total Nov 2023'], name='Nov 2023 Sales', marker_color='#4a69bd'))
            fig_predictions.add_trace(go.Bar(x=filtered_data['Zone'], y=filtered_data['Predicted Nov 2024'], name='Predicted Nov 2024 Sales', marker_color='#82ccdd'))
            fig_predictions.update_layout(title='Sales Projection: 2023 vs 2024',barmode='group',plot_bgcolor='rgba(255,255,255,0.1)',paper_bgcolor='rgba(0,0,0,0)',font_color='burlywood',xaxis_title_font_color='burlywood',yaxis_title_font_color='burlywood',title_font_color='burlywood',legend_font_color='burlywood')
            fig_predictions.update_xaxes(title_text='Zone', tickfont_color='peru')
            fig_predictions.update_yaxes(title_text='Sales', tickfont_color='peru')
            st.plotly_chart(fig_predictions, use_container_width=True)
            fig_target_vs_projected = create_target_vs_projected_graph(filtered_data)
            st.plotly_chart(fig_target_vs_projected, use_container_width=True)
        st.markdown("<h3>Monthly Performance by Zone and Brand</h3>", unsafe_allow_html=True)
        col_zone, col_brand = st.columns(2)
        with col_zone:
            selected_zone = st.selectbox("Select Zone", options=filtered_data['Zone'].unique())
        with col_brand:
            selected_brand = st.selectbox("Select Brand", options=filtered_data[filtered_data['Zone']==selected_zone]['Brand'].unique())
        selected_data = filtered_data[(filtered_data['Zone'] == selected_zone) & (filtered_data['Brand']==selected_brand)]
        if not selected_data.empty:
            fig_monthly_performance = create_monthly_performance_graph(selected_data)
            months = ['Apr', 'May', 'June', 'July', 'Aug', 'Sep', 'Oct','Nov']
            for i, month in enumerate(months):
                if month != 'Nov':
                    fig_monthly_performance.data[i].y = [selected_data[f'Month Tgt ({month})'].iloc[0],selected_data[f'Monthly Achievement({month})'].iloc[0]]
                else:
                    fig_monthly_performance.data[i].y = [selected_data['Month Tgt (Nov)'].iloc[0],selected_data['Predicted Nov 2024'].iloc[0]]
            st.plotly_chart(fig_monthly_performance, use_container_width=True)
        else:
            st.warning("No data available for the selected Zone and Brand combination.")
        st.markdown("<h3>Detailed Sales Forecast</h3>", unsafe_allow_html=True)
        share_df = pd.DataFrame({'Zone': filtered_data['Zone'],'Brand': filtered_data['Brand'],'November 2024 Target': filtered_data['Month Tgt (Nov)'],'November Projection': filtered_data['Predicted Nov 2024'],'November 2023 Sales': filtered_data['Total Nov 2023'],'YoY Growth(Projected)': filtered_data['YoY Growth']})
        styled_df = style_dataframe(share_df)
        st.dataframe(styled_df, use_container_width=True,hide_index=True)
        pdf_buffer = create_pdf(filtered_data)
        st.download_button(label="Download Forecast Report",data=pdf_buffer,file_name="sales_forecast_2024.pdf",mime="application/pdf")
    else:
        st.info("Upload your sales data to begin the simulation!")
  if __name__ == "__main__":
    main()
 else:
    st.stop()
def market_share():
    THEME = {'PRIMARY': '#2563eb','SECONDARY': '#64748b','SUCCESS': '#10b981','WARNING': '#f59e0b','DANGER': '#ef4444','BACKGROUND': '#ffffff','SIDEBAR': '#f8fafc','TEXT': '#1e293b','HEADER': '#0f172a'}
    st.markdown("""<style>/* Global Styles */.stApp {background-color: #ffffff;}/* Main Content Area */.main {background-color: #f8fafc;padding: 2rem;border-radius: 1rem;box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1);}/* Headers */h1 {color: #0f172a;font-size: 2.25rem !important;font-weight: 700 !important;margin-bottom: 1.5rem !important;padding-bottom: 0.5rem;border-bottom: 2px solid #e2e8f0;}h2 {color: #1e293b;font-size: 1.875rem !important;font-weight: 600 !important;margin-top: 2rem !important;}h3 {color: #334155;font-size: 1.5rem !important;font-weight: 600 !important;}/* Sidebar */.css-1d391kg {background-color: #f8fafc;padding: 2rem 1.5rem;border-right: 1px solid #e2e8f0;}/* Cards */.stMetric {background-color: white;padding: 1rem;border-radius: 0.75rem;box-shadow: 0 1px 3px 0 rgb(0 0 0 / 0.1);transition: transform 0.2s;}.stMetric:hover {transform: translateY(-2px);}/* Buttons */.stButton>button {background-color: #2563eb;color: white;border: none;padding: 0.5rem 1.25rem;border-radius: 0.5rem;font-weight: 500;transition: all 0.2s;box-shadow: 0 2px 4px rgba(37, 99, 235, 0.2);}.stButton>button:hover {background-color: #1d4ed8;transform: translateY(-1px);box-shadow: 0 4px 6px rgba(37, 99, 235, 0.3);}/* Select Boxes */.stSelectbox>div>div {background-color: white;border-radius: 0.5rem;border: 1px solid #e2e8f0;}/* Expander */.streamlit-expanderHeader {background-color: white;border-radius: 0.5rem;border: 1px solid #e2e8f0;padding: 0.75rem 1rem;}/* Plots */.stPlot {background-color: white;padding: 1rem;border-radius: 0.75rem;box-shadow: 0 1px 3px 0 rgb(0 0 0 / 0.1);}/* Loading Animation */.stSpinner {text-align: center;color: #2563eb;}/* Tooltips */.tooltip {position: relative;display: inline-block;border-bottom: 1px dotted #64748b;} .tooltip .tooltiptext {visibility: hidden;background-color: #1e293b;color: white;text-align: center;padding: 0.5rem 1rem;border-radius: 0.375rem;position: absolute;z-index: 1;bottom: 125%;left: 50%;transform: translateX(-50%);opacity: 0;transition: opacity 0.2s;}.tooltip:hover .tooltiptext {visibility: visible;opacity: 1;}</style>""", unsafe_allow_html=True)
    COMPANY_COLORS = {}
    @st.cache_data
    def generate_distinct_color(existing_colors):
        if existing_colors:
            return distinctipy.get_colors(1, existing_colors)[0]
        return distinctipy.get_colors(1)[0]
    @st.cache_data
    def get_company_color(company):
     if 'company_colors' not in st.session_state:
        st.session_state.company_colors = {}
     if company not in st.session_state.company_colors:
        existing_colors = list(st.session_state.company_colors.values())
        st.session_state.company_colors[company] = generate_distinct_color(existing_colors)
     return st.session_state.company_colors[company]
    @st.cache_data
    def load_and_process_data(uploaded_file):
        xl = pd.ExcelFile(uploaded_file)
        states = xl.sheet_names
        state_dfs = {state: pd.read_excel(uploaded_file, sheet_name=state) for state in states}
        all_companies = set()
        for df in state_dfs.values():
            all_companies.update(df['Company'].unique())
        for company in all_companies:
            get_company_color(company)
        return state_dfs, states
    def get_available_months(df):
     share_cols = [col for col in df.columns if col.startswith('Share_')]
     months = [col.split('_')[1] for col in share_cols]
     month_order = {'Jan': 1, 'Feb': 2, 'Mar': 3, 'Apr': 4, 'May': 5, 'Jun': 6,'Jul': 7, 'Aug': 8, 'Sep': 9, 'Oct': 10, 'Nov': 11, 'Dec': 12}
     sorted_months = sorted(months, key=lambda x: month_order[x])
     return sorted_months
    @st.cache_data
    def create_share_plot_with_state(df, month, state_name):
     fig = create_share_plot(df, month)
     plt.figtext(0.15, 0.90,state_name,rotation=0,fontsize=14,fontweight='bold',color='#2c3e50',ha='center',va='center',bbox=dict(facecolor='#f8f9fa',edgecolor='#bdc3c7',boxstyle='round,pad=0.5',alpha=0.9))
     plt.subplots_adjust(left=0.1, right=0.82, bottom=0.2, top=0.88)
     return fig
    def create_all_states_report(state_dfs, selected_months):
     figs = []
     for state_name, df in state_dfs.items():
        for month in selected_months:
            fig = create_share_plot_with_state(df, month, state_name)
            figs.append(fig)
            plt.close(fig)
     return figs
    def create_share_plot(df, month):
     def create_stripe_pattern(spacing=5):
        return patches.PathPatch(Path([(0., 0.), (1., 0.),(1., 1.), (0., 1.),(0., 0.)],[Path.MOVETO] + [Path.LINETO] * 3 + [Path.CLOSEPOLY]),transform=None, clip_on=True,facecolor='none', edgecolor='none', alpha=1.)
     def draw_curly_brace(ax, x, y1, y2):
        mid_y = (y1 + y2) / 2
        width = 0.03
        brace_points = [[x, y1],[x + width, y1],[x + width, y1],[x + width, (y1 + mid_y)/2],[x, mid_y],[x + width, (mid_y + y2)/2],[x + width, y2],[x + width, y2],[x, y2]]
        for i in range(len(brace_points)-1):
            line = Line2D([brace_points[i][0], brace_points[i+1][0]],[brace_points[i][1], brace_points[i+1][1]],color='#2c3e50',linewidth=1.5)
        return mid_y
     def cascade_label_positions(positions, y_max, min_gap=12):
        if not positions:
            return [], {}
        x_groups = {}
        for vol, original_y, color, x_pos in positions:
            if x_pos not in x_groups:
                x_groups[x_pos] = []
            x_groups[x_pos].append((vol, original_y, color, x_pos))
        x_positions = sorted(x_groups.keys(), reverse=True)
        y_range = y_max * 0.9
        min_allowed_y = y_max * 0.0001
        total_price_ranges = len(x_positions)
        height_per_range = y_range / total_price_ranges
        adjusted = []
        group_info = {}  # Store information about each group for brackets
        for i, x_pos in enumerate(x_positions):
            group = x_groups[x_pos]
            n_labels = len(group)
            top_y = y_range - (i * height_per_range)
            bottom_y = top_y - height_per_range
            group_gap = min(min_gap, height_per_range / (n_labels + 1))
            group = sorted(group, key=lambda x: x[1])
            group_volumes = []
            group_positions = []
            for j, (vol, original_y, color, x) in enumerate(group):
                label_y = top_y - ((j + 1) * group_gap)
                label_y = max(label_y, min_allowed_y)
                adjusted.append((vol, original_y, label_y, color, x_pos))
                group_volumes.append(vol)
                group_positions.append(label_y)
            if len(group) >= 1:
                group_info[x_pos] = {'total_volume': sum(group_volumes),'top_y': max(group_positions),'bottom_y': min(group_positions)}
        return adjusted, group_info
     def adjust_label_positions(positions, y_max, min_gap=12):
        if not positions:
            return positions
        positions = sorted(positions, key=lambda x: x[1])
        y_range = y_max * 0.9  # Use 90% of the plot height for labels
        n_labels = len(positions)
        optimal_gap = min(min_gap, y_range / (n_labels + 1))
        adjusted = []
        used_positions = set()
        min_allowed_y = y_max * 0.0001
        for vol, original_y, color, x_pos in positions:
            label_y = max(original_y, min_allowed_y)
            while any(abs(label_y - used_y) < optimal_gap for used_y in used_positions):
                label_y += optimal_gap
                if label_y > y_range:
                    label_y = min_allowed_y
                    while any(abs(label_y - used_y) < optimal_gap for used_y in used_positions):
                        label_y += optimal_gap
                        if label_y > y_range:
                         optimal_gap *= 0.8
                         label_y = max(original_y, min_allowed_y)
            used_positions.add(label_y)
            adjusted.append((vol, original_y, label_y, color, x_pos))
        return adjusted
     def check_overlap(y1, y2, height=10):  # height is the estimated text height in points
        return abs(y1 - y2) < height
     def adjust_positions(positions, min_gap=10):
        if not positions:
            return positions
        positions = sorted(positions, key=lambda x: x[1])  
        adjusted_positions = [positions[0]]  
        for vol, y_pos, color, x_pos in positions[1:]:
            prev_y = adjusted_positions[-1][1]
            if check_overlap(y_pos, prev_y):
                new_y = prev_y + min_gap
            else:
                new_y = y_pos
            adjusted_positions.append((vol, new_y, color, x_pos))
        return adjusted_positions
     plt.style.use('seaborn-v0_8-whitegrid')
     plt.rcParams.update({'font.family': 'sans-serif','font.size': 10,'axes.labelweight': 'bold','axes.titleweight': 'bold','figure.facecolor': 'white','axes.facecolor': '#f8f9fa','grid.alpha': 0.2,'grid.color': '#b4b4b4','figure.dpi': 120,'axes.spines.top': False,'axes.spines.right': False,'axes.linewidth': 1.5})
     month_data = df[['Company', f'Share_{month}', f'WSP_{month}', f'Vol_{month}']].copy()
     month_data.columns = ['Company', 'Share', 'WSP', 'Volume']
     total_market_size = month_data['Volume'].sum()
     companies_without_price = month_data[month_data['WSP'].isna() & (month_data['Volume'] > 0)]
     month_data_with_price = month_data.dropna(subset=['WSP'])
     min_price = (month_data_with_price['WSP'].min() // 10) * 10
     max_price = (month_data_with_price['WSP'].max() // 10 + 1) * 10
     price_ranges = pd.interval_range(start=min_price, end=max_price, freq=10)
     month_data_with_price['Price_Range'] = pd.cut(month_data_with_price['WSP'], bins=price_ranges)
     pivot_df = pd.pivot_table(month_data_with_price,values=['Share', 'Volume'],index='Price_Range',columns='Company',aggfunc='sum',fill_value=0)
     share_df = pivot_df['Share']
     volume_df = pivot_df['Volume']
     share_df = share_df.loc[:, (share_df != 0).any(axis=0)]
     volume_df = volume_df.loc[:, (volume_df != 0).any(axis=0)]
     company_wsps = {company: month_data_with_price[month_data_with_price['Company'] == company]['WSP'].iloc[0]for company in share_df.columns}
     sorted_companies = sorted(company_wsps.keys(), key=lambda x: company_wsps[x], reverse=True)
     company_colors = {}
     for i, company in enumerate(sorted_companies):
        if company == 'JK Lakshmi':
            company_colors[company] = '#FF6B6B'
        else:
            company_colors[company] = get_company_color(company)
     share_df = share_df[sorted_companies]
     volume_df = volume_df[sorted_companies]
     fig, ax1 = plt.subplots(figsize=(14, 9), dpi=120)
     ax2 = ax1.twinx()
     bottom = np.zeros(len(share_df))
     volume_positions = []
     total_shares = share_df.sum(axis=1)
     total_volumes = volume_df.sum(axis=1)
     pattern = create_stripe_pattern()
     for company in sorted_companies:
        values = share_df[company].values
        bar_container = ax1.bar(range(len(share_df)),values,bottom=bottom,label=company,color=company_colors[company],alpha=0.95,edgecolor='white',linewidth=0.5)
        if company == 'JK Lakshmi':
            for bar in bar_container:
                x, y = bar.get_xy()
                w, h = bar.get_width(), bar.get_height()
                glow = patches.Rectangle((x, y), w, h,facecolor='none',edgecolor='#FFD700',linewidth=2,alpha=0.6)
                ax1.add_patch(glow)
                bar.set_hatch('///')
                edge = patches.Rectangle((x, y), w, h,facecolor='none',edgecolor='#FF4136',linewidth=1.5)
                ax1.add_patch(edge)
        for i, v in enumerate(values):
            if v > 0:
                center = bottom[i] + v/2
                if v > 0.2:
                    text_color = 'white' if company != 'JK Lakshmi' else '#000000'
                    ax1.text(i, center, f'{v:.1f}%',ha='center', va='center',fontsize=8,color=text_color,fontweight='bold',zorder=10)
                vol = volume_df.loc[share_df.index[i], company]
                if vol > 0:
                    volume_positions.append((vol, center, company_colors[company], i))
        bottom += values
     max_total_share = total_shares.max()
     y_max = max_total_share * 1.15
     ax1.set_ylim(0, y_max)
     for i, total in enumerate(total_shares):
        ax1.text(i, total + (y_max * 0.02), f'Total: {total:.1f}%',ha='center', va='bottom',fontsize=12,fontweight='bold',color='#2c3e50')
     adjusted_positions, group_info = cascade_label_positions(volume_positions, y_max)
     for vol, line_y, label_y, color, x_pos in adjusted_positions:
        if abs(label_y - line_y) > 0.5:
            mid_x = x_pos + (len(share_df)-0.15 - x_pos) * 0.7
            ax1.plot([x_pos, mid_x, len(share_df)-0.15], [line_y, label_y, label_y],color=color, linestyle='--', alpha=1, linewidth=1)
        else:
            ax1.plot([x_pos, len(share_df)-0.15], [line_y, line_y],color=color, linestyle='--', alpha=1, linewidth=1)
        label = f'{vol:,.2f} lakh MT'
        ax2.text(0.98, label_y, label,
                transform=ax1.get_yaxis_transform(),va='center', ha='left',color=color,fontsize=11,fontweight='bold',bbox=dict(facecolor='white',edgecolor='none',alpha=1,pad=1))
     for x_pos, info in group_info.items():
        brace_x = 1.095 
        mid_y = draw_curly_brace(ax2, brace_x, info['top_y'], info['bottom_y'])
        total_label = f'Total: {info["total_volume"]:,.2f} lakh MT'
        ax2.text(brace_x, mid_y, total_label,transform=ax1.get_yaxis_transform(),va='center', ha='left',color='#2c3e50',fontsize=11,fontweight='bold',bbox=dict(facecolor='white',edgecolor='#bdc3c7',boxstyle='round,pad=0.5',alpha=0.9))
     plt.subplots_adjust(right=0.75)
     x_labels = [f'₹{interval.left:.0f}-{interval.right:.0f}'for interval in share_df.index]
     ax1.set_xticks(range(len(x_labels)))
     ax1.set_xticklabels(x_labels, ha='center')
     plt.suptitle('Market Share Distribution by Price Range',fontsize=16, y=1.05,color='#2c3e50',fontweight='bold')
     plt.title(f'{month.capitalize()}',fontsize=14,pad=15,color='#34495e',y=1.11)
     ax1.set_xlabel('WSP Price Range (₹)',fontsize=11,labelpad=15,color='#2c3e50')
     ax1.set_ylabel('Market Share (%)',fontsize=11,labelpad=10,color='#2c3e50')
     legend_elements = []
     for company in sorted_companies:
        marker = patches.Rectangle((0, 0), 1, 1, facecolor=company_colors[company])
        if company == 'JK Lakshmi':
            marker.set_hatch('///')
            marker.set_edgecolor('#FF4136')
            marker.set_linewidth(1.5)
        else:
            marker.set_edgecolor('white')
            marker.set_linewidth(0.5)
        legend_elements.append(marker)
     legend_labels = [f'{company} (WSP: ₹{company_wsps[company]:.0f})'for company in sorted_companies]
     legend = ax1.legend(legend_elements,legend_labels,bbox_to_anchor=(1.28, 0.8),loc='upper left',fontsize=9,frameon=True,facecolor='white',edgecolor='brown',title='Companies',title_fontsize=10,borderpad=1)
     legend_texts = legend.get_texts()
     for text in legend_texts:
        if 'JK Lakshmi' in text.get_text():
            text.set_color('#FF6B6B')
            text.set_fontweight('bold')
     legend.get_frame().set_alpha(1)
     ax2.set_yticks([])
     plt.figtext(0.45, 0.925,f'Total Market Size: {total_market_size:,.2f} lakh MT',ha='center', va='center',bbox=dict(facecolor='#f8f9fa',edgecolor='#bdc3c7',boxstyle='round,pad=0.7',alpha=1),fontsize=11,fontweight='bold',color='#2c3e50')
     if not companies_without_price.empty:
        company_info = companies_without_price.apply(lambda row: f"{row['Company']} ({row['Share']:.1f}%, {row['Volume']:,.2f} lakh MT)",axis=1).tolist()
        if len(company_info) == 1:
            note = f"Note: Price not reported for {company_info[0]}."
        elif len(company_info) == 2:
            note = f"Note: Price not reported for {company_info[0]} and {company_info[1]}."
        else:
            *first_companies, last_company = company_info
            note = f"Note: Price not reported for {', '.join(first_companies)}, and {last_company}."
        total_missing_share = companies_without_price['Share'].sum()
        total_missing_volume = companies_without_price['Volume'].sum()
        summary = f"\nTotal market share and volume without price data: {total_missing_share:.1f}% ({total_missing_volume:,.2f} lakh MT)"
        note += summary
        plt.figtext(0.1, 0.001, note,ha='left', va='bottom',fontsize=10,color='#2c3e50',style='italic',bbox=dict(facecolor='#f8f9fa',edgecolor='#bdc3c7',boxstyle='round,pad=0.5',alpha=0.9))
     plt.tight_layout()
     plt.subplots_adjust(right=0.82, bottom=0.2, top=0.88)
     return fig
    @st.cache_data
    def calculate_share_changes(shares, months):
        sequential_changes = []
        for i in range(1, len(shares)):
            change = (shares[i] - shares[i-1])/shares[i]*100
            sequential_changes.append(change)
        total_change = (shares[-1] - shares[0])/shares[0]*100
        return sequential_changes, total_change
    @st.cache_data
    def create_trend_line_plot(_df, selected_companies, state_name):
     df = _df.copy()
     share_cols = [col for col in df.columns if col.startswith('Share_')]
     months = [col.split('_')[1] for col in share_cols]
     month_order = {'Jan': 1, 'Feb': 2, 'Mar': 3, 'Apr': 4, 'May': 5, 'Jun': 6,'Jul': 7, 'Aug': 8, 'Sep': 9, 'Oct': 10, 'Nov': 11, 'Dec': 12}
     month_col_pairs = [(col, month_order[month]) for col, month in zip(share_cols, months)]
     sorted_pairs = sorted(month_col_pairs, key=lambda x: x[1])
     sorted_share_cols = [pair[0] for pair in sorted_pairs]
     sorted_months = [col.split('_')[1] for col in sorted_share_cols]
     fig, ax = plt.subplots(figsize=(14, 8))
     lines = []
     legend_labels = []
     for company in selected_companies:
        color = get_company_color(company)
        company_shares = df[df['Company'] == company][sorted_share_cols].iloc[0].values
        avg_share = company_shares.mean()
        line = ax.plot(range(len(sorted_months)), company_shares, marker='o', linewidth=2, color=color,label=company)[0]
        lines.append(line)
        ax.axhline(y=avg_share, color=color, linestyle='--', alpha=0.3)
        for i, share in enumerate(company_shares):
            ax.annotate(f'{share:.1f}%', (i, share),xytext=(0, 0),textcoords='offset points',ha='center',va='bottom',fontsize=8)
        sequential_changes, total_change = calculate_share_changes(company_shares, sorted_months)
        for i, change in enumerate(sequential_changes):
            mid_x = (i + 0.5)
            mid_y = (company_shares[i] + company_shares[i + 1]) / 2
            arrow_color = 'green' if change > 0 else 'red'
            arrow_symbol = '↑' if change > 0 else '↓'
            ax.annotate(f'{arrow_symbol}{abs(change):.1f}%',(mid_x, mid_y),xytext=(0, 0 if i % 2 == 0 else 0),textcoords='offset points',ha='center',va='center',color=arrow_color,fontsize=8,bbox=dict(facecolor='white', edgecolor='none',alpha=0.7,pad=0.5))
        change_symbol = '↑' if total_change > 0 else '↓'
        legend_labels.append(f"{company} (Avg: {avg_share:.1f}% | Total Change: {change_symbol}{abs(total_change):.1f}%)")
     plt.title(f'Market Share Trends Over Time - {state_name}', fontsize=20, pad=20, fontweight='bold')
     plt.xlabel('Months', fontsize=12, fontweight='bold')
     plt.ylabel('Market Share (%)', fontsize=12, fontweight='bold')
     plt.xticks(range(len(sorted_months)), sorted_months, rotation=45)
     plt.grid(True, linestyle='--', alpha=0.3)
     ax.legend(lines, legend_labels,bbox_to_anchor=(1.15, 1),loc='upper left',borderaxespad=0.,frameon=True,fontsize=10,title='Companies with Average Share & Total Change',title_fontsize=12,edgecolor='gray')
     ax.set_facecolor('#f8f9fa')
     fig.patch.set_facecolor('#ffffff')
     plt.tight_layout()
     return fig
    def create_title_page(state_name):
     fig, ax = plt.subplots(figsize=(11.7, 8.3))  # A4 size
     ax.axis('off')
     ax.text(0.5, 0.6, 'Price Band Analysis Report', horizontalalignment='center',fontsize=24,fontweight='bold')
     ax.text(0.5, 0.5, f'State: {state_name}',horizontalalignment='center',fontsize=20)
     current_date = datetime.now().strftime("%d %B %Y")
     ax.text(0.5, 0.4, f'Generated on: {current_date}',horizontalalignment='center',fontsize=16)
     fig.patch.set_facecolor('#ffffff')
     return fig
    def create_dashboard_header():
        st.markdown("""<div style='padding: 1.5rem; background: linear-gradient(90deg, #2563eb 0%, #3b82f6 100%);border-radius: 1rem; margin-bottom: 2rem; color: white;'><h1 style='color: brown; margin: 0; border: none;'>Market Share Analysis Dashboard</h1><p style='margin: 0.5rem 0 0 0; opacity: 0.9;'>Comprehensive market analysis and visualization tool</p></div>""", unsafe_allow_html=True)
    def create_metric_card(title, value, delta=None, help_text=None):
            st.metric(label=title,value=value,delta=delta,help=help_text)
    def export_to_pdf(figs, filename):
        with PdfPages(filename) as pdf:
            for fig in figs:
                pdf.savefig(fig, bbox_inches='tight')
                plt.close(fig)
    def main():
     if 'computed_figures' not in st.session_state:
        st.session_state.computed_figures = {}
     create_dashboard_header()
     col1, col2 = st.columns([1, 4])
     with col1:
        st.markdown("### 🎯 Analysis Controls")
        uploaded_file = st.file_uploader("Upload Excel File",type=['xlsx', 'xls'],help="Upload your market share data file")
        if uploaded_file:
            state_dfs, states = load_and_process_data(uploaded_file)
            st.markdown("### 🎯 Settings")
            selected_state = st.selectbox("Select State",states,index=0,help="Choose the state for analysis")
            available_months = get_available_months(state_dfs[list(state_dfs.keys())[0]])
            selected_months = st.multiselect("Select Months",available_months,default=[available_months[0]],help="Choose months for comparison")
            if selected_months:
                all_states_pdf_buf = io.BytesIO()
                with st.spinner("Generating report for all states..."):
                    figs = create_all_states_report(state_dfs, selected_months)
                    with PdfPages(all_states_pdf_buf) as pdf:
                        for fig in figs:
                            pdf.savefig(fig, bbox_inches='tight')
                            plt.close(fig)
                all_states_pdf_buf.seek(0)
                st.download_button(label="📊 Download All States Report",data=all_states_pdf_buf,file_name=f'market_share_all_states_{"-".join(selected_months)}.pdf',mime='application/pdf',help="Download a PDF report containing market share graphs for all states")
            all_companies = state_dfs[selected_state]['Company'].unique()
            default_companies = ['JK Lakshmi', 'Ultratech', 'Ambuja', 'Wonder', 'Shree', 'JK Cement (N)']
            available_defaults = [company for company in default_companies if company in all_companies]
            selected_companies = st.multiselect("Select Companies for Trend Analysis",all_companies,default=available_defaults,help="Choose companies to show in the trend line graph")
     with col2:
        if uploaded_file and selected_companies:
            st.markdown("### Market Share Trends")
            df_hash = hash(str(state_dfs[selected_state]))
            trend_key = f"trend_{selected_state}_{'-'.join(sorted(selected_companies))}_{df_hash}"
            if trend_key not in st.session_state.computed_figures:
                st.session_state.computed_figures[trend_key] = create_trend_line_plot(state_dfs[selected_state],selected_companies,selected_state)
            st.pyplot(st.session_state.computed_figures[trend_key])
            st.markdown("---")
        if uploaded_file and selected_months:
            st.markdown("### 📊 Key Metrics")
            metric_cols = st.columns(len(selected_months))
            for idx, month in enumerate(selected_months):
                df = state_dfs[selected_state]
                with metric_cols[idx]:
                    create_metric_card(f"{month.capitalize()}",f"{len(df[df[f'Share_{month}'] > 0])} Companies",f"Avg WSP: ₹{df[f'WSP_{month}'].mean():.0f}","Number of active companies and average wholesale price")
            st.markdown("---")
            for month in selected_months:
                plot_key = f"share_{selected_state}_{month}"
                if plot_key not in st.session_state.computed_figures:
                    with st.spinner(f"📊 Generating visualization for {month.capitalize()}..."):
                        st.session_state.computed_figures[plot_key] = create_share_plot(state_dfs[selected_state],month)
                st.pyplot(st.session_state.computed_figures[plot_key])
                col1, col2, col3 = st.columns([1, 1, 2])
                with col1:
                    buf = io.BytesIO()
                    st.session_state.computed_figures[plot_key].savefig(buf,format='png',dpi=300,bbox_inches='tight')
                    buf.seek(0)
                    st.download_button(label="📥 Download PNG",data=buf,file_name=f'market_share_{selected_state}_{month}.png',mime='image/png',key=f"download_png_{month}")
                with col2:
                    pdf_buf = io.BytesIO()
                    with PdfPages(pdf_buf) as pdf:
                        pdf.savefig(st.session_state.computed_figures[plot_key], bbox_inches='tight')
                    pdf_buf.seek(0)
                    st.download_button(label="📄 Download PDF",data=pdf_buf,file_name=f'market_share_{selected_state}_{month}.pdf',mime='application/pdf',key=f"download_pdf_{month}")
                st.markdown("---")
            if st.session_state.computed_figures:
             st.markdown("### 📑 Download Complete Report")
             all_pdf_buf = io.BytesIO()
             with PdfPages(all_pdf_buf) as pdf:
                title_page = create_title_page(selected_state)
                pdf.savefig(title_page, bbox_inches='tight')
                plt.close(title_page)
                df_hash = hash(str(state_dfs[selected_state]))
                trend_key = f"trend_{selected_state}_{'-'.join(sorted(selected_companies))}_{df_hash}"
                if trend_key in st.session_state.computed_figures:
                    pdf.savefig(st.session_state.computed_figures[trend_key], bbox_inches='tight')
                for month in selected_months:
                    plot_key = f"share_{selected_state}_{month}"
                    if plot_key in st.session_state.computed_figures:
                        pdf.savefig(st.session_state.computed_figures[plot_key], bbox_inches='tight')
             all_pdf_buf.seek(0)
             st.download_button(label="📥 Download Complete Report (PDF)",data=all_pdf_buf,file_name=f'market_share_{selected_state}_complete_report.pdf',mime='application/pdf',key="download_complete_pdf")
    if __name__ == "__main__":
        main()
def discount():
 warnings.filterwarnings('ignore')
 st.markdown("""<style>/* Global Styles */[data-testid="stSidebar"] {background-color: #f8fafc;border-right: 1px solid #e2e8f0;}.stButton button {background-color: #3b82f6;color: white;border-radius: 6px;padding: 0.5rem 1rem;border: none;transition: all 0.2s;}.stButton button:hover {background-color: #2563eb;box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);}
    /* Ticker Animation */@keyframes ticker {0% { transform: translateX(100%); }100% { transform: translateX(-100%); }}.ticker-container {background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);color: white;padding: 16px;overflow: hidden;white-space: nowrap;position: relative;margin-bottom: 24px;border-radius: 12px;box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);}
    .ticker-content {display: inline-block;animation: ticker 2500s linear infinite;animation-delay: -1250s;padding-right: 100%;will-change: transform;}.ticker-content:hover {animation-play-state: paused;}
    .ticker-item {display: inline-block;margin-right: 80px;font-size: 16px;padding: 8px 16px;opacity: 1;transition: opacity 0.3s;background: rgba(255, 255, 255, 0.1);border-radius: 8px;}
    /* Enhanced Metrics */.state-name {color: #10B981;font-weight: 600;}.month-name {color: #60A5FA;font-weight: 600;}.discount-value {color: #FBBF24;font-weight: 600;}/* Card Styles */
    .metric-card {background: white;padding: 1.5rem;border-radius: 12px;box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);transition: transform 0.2s;border: 1px solid #e2e8f0;}
    .metric-card:hover {transform: translateY(-2px);box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1);}.metric-value {font-size: 2rem;font-weight: 600;color: #1e293b;}.metric-label {color: #64748b;font-size: 0.875rem;margin-top: 0.5rem;}
    /* Chart Container */.chart-container {background: white;padding: 1.5rem;border-radius: 12px;box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);margin: 1rem 0;border: 1px solid #e2e8f0;}
    /* Selectbox Styling */.stSelectbox {background: white;border-radius: 8px;border: 1px solid #e2e8f0;}
    /* Custom Header */.dashboard-header {padding: 1.5rem;background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);color: white;border-radius: 12px;margin-bottom: 2rem;text-align: center;}
    .dashboard-title {font-size: 2rem;font-weight: 600;margin-bottom: 0.5rem;}.dashboard-subtitle {color: #94a3b8;font-size: 1rem;}</style>""", unsafe_allow_html=True)
 @st.cache_data(ttl=3600)
 def process_excel_file(file_content, excluded_sheets):
    excel_data = io.BytesIO(file_content)
    excel_file = pd.ExcelFile(excel_data)
    processed_data = {}
    for sheet in excel_file.sheet_names:
        if not any(excluded_sheet in sheet for excluded_sheet in excluded_sheets):
            df = pd.read_excel(excel_data, sheet_name=sheet, usecols=range(22))
            cash_discount_patterns = ['CASH DISCOUNT', 'Cash Discount', 'CD']
            start_idx = None
            for idx, value in enumerate(df.iloc[:, 0]):
                if isinstance(value, str):
                    if any(pattern.lower() in value.lower() for pattern in cash_discount_patterns):
                        start_idx = idx
                        break
            if start_idx is not None:
                df = df.iloc[start_idx:].reset_index(drop=True)
            g_total_idx = None
            for idx, value in enumerate(df.iloc[:, 0]):
                if isinstance(value, str) and 'G. TOTAL' in value:
                    g_total_idx = idx
                    break
            if g_total_idx is not None:
                df = df.iloc[:g_total_idx].copy()
            processed_data[sheet] = df
    return processed_data
 class HorizontalLine(Flowable):
    def __init__(self, width, thickness=1):
        Flowable.__init__(self)
        self.width = width
        self.thickness = thickness
    def draw(self):
        self.canv.setLineWidth(self.thickness)
        self.canv.line(0, 0, self.width, 0)
 class LineChart(Drawing):
    def __init__(self, width=500, height=250, data=None, months=None):
        Drawing.__init__(self, width, height)
        self._months = months or []
        self.add(LinePlot(), name='chart')
        self.chart.width = width * 0.8
        self.chart.height = height * 0.7
        self.chart.x = width * 0.1
        self.chart.y = height * 0.15
        self.chart.lines[0].strokeColor = HexColor('#3b82f6')
        self.chart.lines[1].strokeColor = HexColor('#ef4444')
        self.chart.lines[0].strokeWidth = 0.5
        self.chart.lines[1].strokeWidth = 0.5
        if months:
            self._months = list(dict.fromkeys(months))
        if data:
            self.chart.data = data
        self.chart.xValueAxis.labelTextFormat = self._month_formatter
        self.chart.xValueAxis.labels.boxAnchor = 'n'
        self.chart.xValueAxis.labels.angle = 0
        self.chart.xValueAxis.labels.dy = -20
        self.chart.xValueAxis.tickDown = 5
        self.chart.xValueAxis.tickUp = 0
        self.chart.xValueAxis.visibleGrid = False
        self.chart.xValueAxis.valueMin = 0
        self.chart.xValueAxis.valueMax = len(self._months)-1
        self.chart.yValueAxis.labelTextFormat = 'Rs.%.1f'
        self.chart.yValueAxis.gridStrokeColor = HexColor('#e2e8f0')
        self.chart.yValueAxis.gridStrokeWidth = 0.5
        self.chart.lines[0].symbol = makeMarker('Circle')
        self.chart.lines[1].symbol = makeMarker('Circle')
        self.chart.lines[0].symbol.strokeColor = HexColor('#3b82f6')
        self.chart.lines[1].symbol.strokeColor = HexColor('#ef4444')
        self.chart.lines[0].symbol.fillColor = HexColor('#ffffff')
        self.chart.lines[1].symbol.fillColor = HexColor('#ffffff')
        self.chart.lines[0].symbol.size = 6
        self.chart.lines[1].symbol.size = 6
        if data and len(data) >= 2: 
            approved_values = [point[1] for point in data[0]]
            actual_values = [point[1] for point in data[1]]
            min_y = min(min(approved_values), min(actual_values))
            max_y = max(max(approved_values), max(actual_values))
            y_range = max_y - min_y if max_y != min_y else 1
            x_scale = self.chart.width / (len(self._months) - 1)
            y_scale = self.chart.height / y_range
            for i, (_, y) in enumerate(data[0]):
                x_pos = self.chart.x + (i * x_scale)
                y_pos = self.chart.y + ((y - min_y) * y_scale)
                label = String(x_pos,y_pos + 15,  f'Rs.{y:.1f}',fontSize=8,fillColor=HexColor('#3b82f6'),textAnchor='middle')
                self.add(label)
            for i, (_, y) in enumerate(data[1]): 
                x_pos = self.chart.x + (i * x_scale)
                y_pos = self.chart.y + ((y - min_y) * y_scale)
                label = String(x_pos,y_pos - 15, f'Rs.{y:.1f}',fontSize=8,fillColor=HexColor('#ef4444'),textAnchor='middle')
                self.add(label)
    def _month_formatter(self, value):
        try:
            index = int(round(value))
            if 0 <= index < len(self._months):
                return self._months[index]
        except (ValueError, TypeError):
            pass
        return ''
 class DiscountReportGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.setup_custom_styles()
    def setup_custom_styles(self):
        self.styles.add(ParagraphStyle(name='ChartTitle',fontName='Helvetica-Bold',fontSize=14,spaceBefore=30,spaceAfter=10,alignment=1,textColor=HexColor('#1e293b')))
        self.styles.add(ParagraphStyle(name='ChartLegend',fontName='Helvetica',fontSize=10,spaceAfter=20,alignment=1,textColor=HexColor('#475569')))
        self.styles.add(ParagraphStyle(name='TitlePageFooter',fontName='Helvetica-Bold',fontSize=11,spaceBefore=30,alignment=1,textColor=HexColor('#64748b')))
        self.styles.add(ParagraphStyle(name='MonthHeader',fontName='Helvetica-Bold',fontSize=20,spaceAfter=10,alignment=1,textColor=HexColor('#1e293b')))
        self.styles.add(ParagraphStyle(name='QuantityInfo',fontName='Helvetica',fontSize=12,spaceAfter=15,alignment=1,textColor=HexColor('#475569')))
        self.styles.add(ParagraphStyle(name='FooterNote',fontName='Helvetica-Bold',fontSize=10,spaceBefore=20,spaceAfter=10,alignment=1,textColor=HexColor('#64748b')))
    def is_valid_discount(self, approved, actual):
        if pd.isna(approved) and pd.isna(actual):
            return False
        if approved == 0 and actual == 0:
            return False
        return True
    def create_monthly_page(self, month_data, month_name):
        elements = []
        elements.append(Paragraph(month_name, self.styles['MonthHeader']))
        elements.append(Paragraph(f"Total Quantity: {month_data['quantity']:,.2f} MT",self.styles['QuantityInfo']))
        table_data = [['Discount Type', 'Approved Rate', 'Actual Rate', 'Difference', 'Total Impact']]
        valid_discounts = []
        for discount_name, values in month_data['discounts'].items():
            approved = values['approved']
            actual = values['actual']
            if not self.is_valid_discount(approved, actual):
                continue
            diff = approved - actual
            total_diff = diff * month_data['quantity']*20
            valid_discounts.append({'name': discount_name,'approved': approved,'actual': actual,'diff': diff,'total_diff': total_diff})
        valid_discounts.sort(key=lambda x: x['approved'])
        for discount in valid_discounts:
            diff_text = f"{'↓' if discount['diff'] >= 0 else '↑'} Rs.{abs(discount['diff']):,.2f}"
            impact_text = f"Rs.{abs(discount['total_diff']):,.2f}"
            row = [discount['name'],f"Rs.{discount['approved']:,.2f}",f"Rs.{discount['actual']:,.2f}",diff_text,impact_text]
            table_data.append(row)
        if len(table_data) > 1:
            table = Table(table_data, colWidths=[2.2*inch, 1.3*inch, 1.3*inch, 1.1*inch, 1.3*inch])
            style = TableStyle([('BACKGROUND', (0, 0), (-1, 0), HexColor('#f1f5f9')),('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#334155')),('ALIGN', (0, 0), (-1, -1), 'CENTER'),('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),('FONTSIZE', (0, 0), (-1, 0), 10),('BOTTOMPADDING', (0, 0), (-1, 0), 12),('TOPPADDING', (0, 0), (-1, 0), 12),('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),('FONTSIZE', (0, 1), (-1, -1), 9),('GRID', (0, 0), (-1, -1), 1, HexColor('#e2e8f0')),('ALIGN', (1, 1), (-1, -1), 'RIGHT'),('ALIGN', (0, 0), (0, -1), 'LEFT'),('BOTTOMPADDING', (0, 1), (-1, -1), 8),('TOPPADDING', (0, 1), (-1, -1), 8),])
            for i in range(1, len(table_data)):
                if '↓' in table_data[i][3]: 
                    style.add('TEXTCOLOR', (3, i), (4, i), HexColor('#10b981'))
                else:  # Excess
                    style.add('TEXTCOLOR', (3, i), (4, i), HexColor('#ef4444'))
            table.setStyle(style)
            elements.append(table)
        else:
            elements.append(Paragraph("No valid discount data available for this month",self.styles['Normal']))
        elements.append(Paragraph("This report currently presents data for Q1 2024. The forthcoming update will incorporate comprehensive data for Q2 2024, providing a more extensive analysis of discount trends.",self.styles['FooterNote']))
        return elements
    def get_highest_discount_data(self, data):
        months = []
        approved_values = []
        actual_values = []
        for month, month_data in data.items():
            valid_discounts = []
            for discount_name, values in month_data['discounts'].items():
                if self.is_valid_discount(values['approved'], values['actual']):
                    valid_discounts.append({'name': discount_name,'approved': values['approved'],'actual': values['actual']})
            if valid_discounts:
                highest_discount = max(valid_discounts, key=lambda x: x['approved'])
                months.append(month)
                approved_values.append(highest_discount['approved'])
                actual_values.append(highest_discount['actual'])
        return months, approved_values, actual_values
    def generate_report(self, state, data):
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer,pagesize=A4,rightMargin=36,leftMargin=36,topMargin=36,bottomMargin=36,title=f"Discount Report - {state}") 
        story = []
        story.append(Paragraph(f"Discount Analysis Report<br/>{state}",ParagraphStyle('Title',parent=self.styles['Title'],fontSize=24,spaceAfter=30,alignment=1)))
        story.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}",ParagraphStyle('Date',parent=self.styles['Normal'],alignment=1,fontSize=12,textColor=HexColor('#64748b'))))
        story.append(Spacer(1, 20))
        story.append(HorizontalLine(540, 2))  # 540 points = 7.5 inches (standard page width minus margins)
        story.append(Spacer(1, 30))
        story.append(Paragraph("Discount Rate Trend(Grand Total)",self.styles['ChartTitle']))
        months, approved_values, actual_values = self.get_highest_discount_data(data)
        chart_data = [list(zip(range(len(months)), approved_values)), list(zip(range(len(months)), actual_values))]
        drawing = LineChart(500, 300, chart_data, months)
        story.append(drawing)
        story.append(Paragraph("""<para alignment="center"><font color="#3b82f6">--Approved Rate</font>  <font color="#ef4444">--Actual Rate</font></para>""",self.styles['ChartLegend']))
        story.append(Paragraph("Find detailed month-wise analyses in the following pages, where each discount type is meticulously examined and presented in ascending order of approved rates.",self.styles['TitlePageFooter']))
        story.append(PageBreak())
        for month, month_data in data.items():
            story.extend(self.create_monthly_page(month_data, month))
            if month != list(data.keys())[-1]:
                story.append(PageBreak())
        doc.build(story)
        buffer.seek(0)
        return buffer
 def add_pdf_download(analytics_instance, data, selected_state):
    pdf_data = {}
    for month, cols in analytics_instance.month_columns.items():
        month_data = {'quantity': 0,'discounts': {}}
        df = data[selected_state]
        discount_types = analytics_instance.get_discount_types(df, selected_state)
        first_discount = discount_types[0]
        if first_discount == analytics_instance.combined_discount_name:
            combined_data = analytics_instance.get_combined_data(df, cols, selected_state)
            month_data['quantity'] = combined_data['quantity']
        else:
            mask = df.iloc[:, 0].fillna('').astype(str).str.strip() == first_discount.strip()
            filtered_df = df[mask]
            if len(filtered_df) > 0:
                month_data['quantity'] = filtered_df.iloc[0, cols['quantity']]
        for discount in discount_types:
            if discount == analytics_instance.combined_discount_name:
                combined_data = analytics_instance.get_combined_data(df, cols, selected_state)
                month_data['discounts'][discount] = {'approved': combined_data['approved'],'actual': combined_data['actual']}
            else:
                mask = df.iloc[:, 0].fillna('').astype(str).str.strip() == discount.strip()
                filtered_df = df[mask]
                if len(filtered_df) > 0:
                    month_data['discounts'][discount] = {'approved': filtered_df.iloc[0, cols['approved']],'actual': filtered_df.iloc[0, cols['actual']]}
        pdf_data[month] = month_data
    report_generator = DiscountReportGenerator()
    pdf_buffer = report_generator.generate_report(selected_state, pdf_data)
    st.download_button(label="📄 Download Detailed Report",data=pdf_buffer,file_name=f"discount_report_{selected_state}_{datetime.now().strftime('%Y%m%d')}.pdf",mime="application/pdf",key=f"pdf_download_{selected_state}",help="Download a detailed PDF report for this state")
 class DiscountAnalytics:
    def __init__(self):
        self.excluded_discounts = ['Sub Total','TOTAL OF DP PAYOUT','TOTAL OF STS & RD','Other (Please specify',]
        self.discount_mappings = {'group1': {'states': ['HP', 'JMU', 'PUN'],'discounts': ['CASH DISCOUNT', 'ADVANCE CD & NIL OS']},'group2': {'states': ['UP (W)'],'discounts': ['CD', 'Adv CD']}}
        self.combined_discount_name = 'CD and Advance CD'
        self.month_columns = {'April': {'quantity': 1,'approved': 2,'actual': 4},'May': {'quantity': 8,'approved': 9,'actual': 11},'June': {'quantity': 15,'approved': 16,'actual': 18}}
        self.total_patterns = ['G. TOTAL', 'G.TOTAL', 'G. Total', 'G.Total', 'GRAND TOTAL',"G. Total (STD + STS)"]
        self.excluded_states = ['MP (JK)', 'MP (U)','East']
    def create_ticker(self, data):
     ticker_items = []
     last_month = "June"
     month_cols = self.month_columns[last_month]
     for state in data.keys():
        df = data[state]
        if not df.empty:
            state_text = f"<span class='state-name'>📍 {state}</span>"
            month_text = f"<span class='month-name'>📅 {last_month}</span>"
            state_group = next((group for group, config in self.discount_mappings.items()if state in config['states']),None)
            discount_items = []
            if state_group:
                relevant_discounts = self.discount_mappings[state_group]['discounts']
                combined_data = self.get_combined_data(df, month_cols, state)
                if combined_data:
                    actual = combined_data.get('actual', 0)
                    discount_items.append(f"{self.combined_discount_name}: <span class='discount-value'>₹{actual:,.2f}</span>")
                for discount in self.get_discount_types(df, state):
                    if discount != self.combined_discount_name:
                        mask = df.iloc[:, 0].fillna('').astype(str).str.strip() == discount.strip()
                        filtered_df = df[mask]
                        if len(filtered_df) > 0:
                            actual = filtered_df.iloc[0, month_cols['actual']]
                            discount_items.append(
                                f"{discount}: <span class='discount-value'>₹{actual:,.2f}</span>")
            else:
                for discount in self.get_discount_types(df, state):
                    mask = df.iloc[:, 0].fillna('').astype(str).str.strip() == discount.strip()
                    filtered_df = df[mask]
                    if len(filtered_df) > 0:
                        actual = filtered_df.iloc[0, month_cols['actual']]
                        discount_items.append(
                            f"{discount}: <span class='discount-value'>₹{actual:,.2f}</span>")
            if discount_items:
                full_text = f"{state_text} | {month_text} | {' | '.join(discount_items)}"
                ticker_items.append(f"<span class='ticker-item'>{full_text}</span>")
     ticker_items = ticker_items * 3
     ticker_html = f"""<div class="ticker-container"><div class="ticker-content">{' '.join(ticker_items)}</div></div>"""
     st.markdown(ticker_html, unsafe_allow_html=True)
    def create_summary_metrics(self, data):
        total_states = len(data)
        total_discounts = sum(len(self.get_discount_types(df)) for df in data.values())
        avg_discount = np.mean([df.iloc[0, 4] for df in data.values() if not df.empty])
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total States", total_states, "Active")
        with col2:
            st.metric("Total Discount Types", total_discounts, "Available")
        with col3:
            st.metric("Average Discount Rate", f"₹{avg_discount:,.2f}", "Per Bag")
    def create_monthly_metrics(self, data, selected_state, selected_discount):
        df = data[selected_state]
        if selected_discount == self.combined_discount_name:
            monthly_data = {
                month: self.get_combined_data(df, cols, selected_state)
                for month, cols in self.month_columns.items()}
        else:
            mask = df.iloc[:, 0].fillna('').astype(str).str.strip() == selected_discount.strip()
            filtered_df = df[mask]
            if len(filtered_df) > 0:
                monthly_data = {month: {'actual': filtered_df.iloc[0, cols['actual']],'approved': filtered_df.iloc[0, cols['approved']],'quantity': filtered_df.iloc[0, cols['quantity']]}for month, cols in self.month_columns.items()}
        for month, data in monthly_data.items():
            st.markdown(f"""<div style='text-align: center; margin-bottom: 10px;'><h3 style='color: #1e293b; margin-bottom: 15px;'>{month}</h3></div>""", unsafe_allow_html=True)
            col1, col2, col3 = st.columns(3)
            with col1:
                quantity = data.get('quantity', 0)
                st.metric("Quantity Sold",f"{quantity:,.2f}",delta=None,help=f"Total quantity sold in {month}")
            with col2:
                approved = data.get('approved', 0)
                st.metric("Approved Payout",f"₹{approved:,.2f}",delta=None,help=f"Approved discount rate for {month}")
            with col3:
                actual = data.get('actual', 0)
                difference = approved - actual
                delta_color = "normal" if difference >= 0 else "inverse"
                st.metric("Actual Payout",f"₹{actual:,.2f}",delta=f"₹{abs(difference):,.2f}" + (" under approved" if difference >= 0 else " over approved"),delta_color=delta_color,help=f"Actual discount rate for {month}")
            st.markdown("---")
    def process_excel(self, uploaded_file):
        return process_excel_file(uploaded_file.getvalue(), ['MP (U)', 'MP (JK)'])
    def create_trend_chart(self, data, selected_state, selected_discount):
        df = data[selected_state]
        if selected_discount == self.combined_discount_name:
            monthly_data = {
                month: self.get_combined_data(df, cols, selected_state)
                for month, cols in self.month_columns.items()}
        else:
            mask = df.iloc[:, 0].fillna('').astype(str).str.strip() == selected_discount.strip()
            filtered_df = df[mask]
            if len(filtered_df) > 0:
                monthly_data = {month: {'actual': filtered_df.iloc[0, cols['actual']],'approved': filtered_df.iloc[0, cols['approved']]}for month, cols in self.month_columns.items()}
        months = list(monthly_data.keys())
        actual_values = [data['actual'] for data in monthly_data.values()]
        approved_values = [data['approved'] for data in monthly_data.values()]
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=months,y=actual_values,name='Actual',line=dict(color='#10B981', width=3)))
        fig.add_trace(go.Scatter(x=months,y=approved_values,name='Approved',line=dict(color='#3B82F6', width=3)))
        fig.update_layout(title=f'Discount Trends - {selected_state}',xaxis_title='Month',yaxis_title='Discount Rate (₹/Bag)',template='plotly_white',height=400,margin=dict(t=50, b=50, l=50, r=50))
        st.plotly_chart(fig, use_container_width=True)
        self.create_difference_chart(months, approved_values, actual_values, selected_state)
    def create_difference_chart(self, months, approved_values, actual_values, selected_state):
        differences = [approved - actual for approved, actual in zip(approved_values, actual_values)]
        fig = go.Figure()
        for i in range(len(months)):
            color = '#10B981' if differences[i] >= 0 else '#EF4444'  # Green for positive, red for negative
            fig.add_trace(go.Scatter(x=[months[i], months[i]],y=[0, differences[i]],mode='lines',line=dict(color=color, width=3),showlegend=False))
        fig.add_trace(go.Scatter(x=months,y=differences,mode='markers',marker=dict(size=8,color=['#10B981' if d >= 0 else '#EF4444' for d in differences],line=dict(width=2, color='white')),name='Difference'))
        fig.add_shape(type='line',x0=months[0],x1=months[-1],y0=0,y1=0,line=dict(color='gray', width=1, dash='dash'))
        fig.update_layout(title=f'Approved vs Actual Difference - {selected_state}',xaxis_title='Month',yaxis_title='Difference in Discount Rate (₹/Bag)',template='plotly_white',height=300,margin=dict(t=50, b=50, l=50, r=50))
        st.plotly_chart(fig, use_container_width=True)
    def get_discount_types(self, df, state=None):
     first_col = df.iloc[:, 0]
     valid_discounts = []
     if state:
        state_group = next((group for group, config in self.discount_mappings.items()if state in config['states']),None)
        if state_group:
            relevant_discounts = self.discount_mappings[state_group]['discounts']
            if any(d in first_col.values for d in relevant_discounts):
                valid_discounts.append(self.combined_discount_name)
            for d in first_col.unique():
                if (isinstance(d, str) and d.strip() not in self.excluded_discounts and d.strip() not in relevant_discounts):
                    valid_discounts.append(d)
        else:
            valid_discounts = [d for d in first_col.unique() if isinstance(d, str) and d.strip() not in self.excluded_discounts]
     else:
        valid_discounts = [d for d in first_col.unique() if isinstance(d, str) and d.strip() not in self.excluded_discounts]
     return sorted(valid_discounts)
    def get_combined_data(self, df, month_cols, state):
     combined_data = {'actual': np.nan, 'approved': np.nan,'quantity': np.nan}
     state_group = next((group for group, config in self.discount_mappings.items()if state in config['states']),None)
     if state_group:
        relevant_discounts = self.discount_mappings[state_group]['discounts']
        mask = df.iloc[:, 0].fillna('').astype(str).str.strip().isin(relevant_discounts)
        filtered_df = df[mask]
        if len(filtered_df) > 0:
            combined_data['approved'] = filtered_df.iloc[:, month_cols['approved']].sum()
            combined_data['actual'] = filtered_df.iloc[:, month_cols['actual']].sum()
            total_quantity = filtered_df.iloc[:, month_cols['quantity']].sum()
            combined_data['quantity'] = total_quantity / 2  # Divide summed quantity by 2
     return combined_data
 def main():
    processor = DiscountAnalytics()
    with st.sidebar:
        st.markdown("""<div style='text-align: center; padding: 1rem;'><h2 style='color: #1e293b;'>Dashboard Controls</h2></div>""", unsafe_allow_html=True)
        uploaded_file = st.file_uploader("Upload Excel File", type=['xlsx', 'xls'])
    st.markdown("""<div class='dashboard-header'><div class='dashboard-title'>Discount Analytics Dashboard</div><div class='dashboard-subtitle'>Monitor and analyze discount performance across states</div></div>""", unsafe_allow_html=True)
    if uploaded_file is not None:
        with st.spinner('Processing data...'):
            data = processor.process_excel(uploaded_file)
            processor.create_ticker(data)
        st.markdown("""<div style='margin: 2rem 0;'><h3 style='color: #1e293b; margin-bottom: 1rem;'>Key Performance Indicators</h3></div>""", unsafe_allow_html=True)
        processor.create_summary_metrics(data)
        st.markdown("""<div style='margin: 2rem 0;'><h3 style='color: #1e293b; margin-bottom: 1rem;'>Detailed Analysis</h3></div>""", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            selected_state = st.selectbox("Select State", list(data.keys()))
        if selected_state:
            with col2:
                discount_types = processor.get_discount_types(data[selected_state], selected_state)
                selected_discount = st.selectbox("Select Discount Type", discount_types)
            st.markdown("<div class='chart-container'>", unsafe_allow_html=True)
            processor.create_monthly_metrics(data, selected_state, selected_discount)
            st.markdown("</div>", unsafe_allow_html=True)
            st.markdown("<div class='chart-container'>", unsafe_allow_html=True)
            processor.create_trend_chart(data, selected_state, selected_discount)
            st.markdown("</div>", unsafe_allow_html=True)
            add_pdf_download(processor, data, selected_state)
    else:
        st.markdown("""<div style='text-align: center; padding: 3rem; background: white; border-radius: 12px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);'><h2 style='color: #1e293b; margin-bottom: 1rem;'>Welcome to Discount Analytics</h2><p style='color: #64748b; margin-bottom: 2rem;'>Please upload an Excel file to begin your analysis.</p></div>""", unsafe_allow_html=True)
        st.markdown("<div style='margin-top: 2rem;'>", unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total States", "0", "Waiting")
        with col2:
            st.metric("Total Discount Types", "0", "Waiting")
        with col3:
            st.metric("Average Discount Rate", "₹0.00", "Waiting")
        st.markdown("</div>", unsafe_allow_html=True)
 if __name__ == "__main__":
    main()
def load_visit_data():
    try:
        with open('visit_data.json', 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        return {'total_visits': 0, 'daily_visits': {}}
def save_visit_data(data):
    with open('visit_data.json', 'w') as f:
        json.dump(data, f)
def update_visit_count():
    visit_data = load_visit_data()
    today = datetime.now().strftime('%Y-%m-%d')
    visit_data['total_visits'] += 1
    visit_data['daily_visits'][today] = visit_data['daily_visits'].get(today, 0) + 1
    save_visit_data(visit_data)
    return visit_data['total_visits'], visit_data['daily_visits'][today]
def load_visit_data():
    try:
        with open('visit_data.json', 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        return {'total_visits': 0, 'daily_visits': {}}
def save_visit_data(data):
    with open('visit_data.json', 'w') as f:
        json.dump(data, f)
def update_visit_count():
    visit_data = load_visit_data()
    today = datetime.now().strftime('%Y-%m-%d')
    visit_data['total_visits'] += 1
    visit_data['daily_visits'][today] = visit_data['daily_visits'].get(today, 0) + 1
    save_visit_data(visit_data)
    return visit_data['total_visits'], visit_data['daily_visits'][today]
def main():
    st.markdown("""<style>.sidebar .sidebar-content {background-image: linear-gradient(180deg, #2e7bcf 25%, #4527A0 100%);color: white;}.sidebar-text {color: white !important;}.stButton>button {width: 100%;border-radius: 20px;background-color: #4CAF50;color: white;border: none;padding: 10px;font-weight: bold;transition: all 0.3s ease;}
    .stButton>button:hover {background-color: #45a049;box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);}.stProgress .st-bo {background-color: #4CAF50;}.stProgress .st-bp {background-color: #E0E0E0;}.settings-container {background-color: rgba(255, 255, 255, 0.1);backdrop-filter: blur(10px);padding: 20px;border-radius: 10px;margin-top: 20px;box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);}
    .visit-counter {background-color: rgba(255, 228, 225, 0.7);border-radius: 10px;padding: 15px;margin-top: 20px;text-align: center;box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);}
    .visit-counter h3 {color: #FFD700;font-size: 18px;margin-bottom: 10px;}.visit-counter p {color: #8B4513;font-size: 14px;margin: 5px 0;}.user-info {background-color: rgba(255, 255, 255, 0.1);border-radius: 10px;padding: 10px;margin-bottom: 20px;}</style>""", unsafe_allow_html=True)
    st.sidebar.title("Analytics Dashboard")
    if 'username' not in st.session_state:
        st.session_state.username = "Guest"
    st.sidebar.markdown(f"""<div class="user-info"><i class="fas fa-user"></i> Logged in as: {st.session_state.username}<br><small>Last login: {datetime.now().strftime('%Y-%m-%d %H:%M')}</small></div>""", unsafe_allow_html=True)
    with st.sidebar:
        selected = option_menu(menu_title="Main Menu",options=["Home", "Data Management", "Sales Volume Analysis","Price Analysis", "Predictions", "Settings"],icons=["house-fill", "database-fill-gear", "graph-up-arrow","lightbulb-fill","gear-fill"],menu_icon="cast",default_index=0,styles={"container": {"padding": "0!important", "background-color": "transparent"},"icon": {"color": "orange", "font-size": "20px"}, "nav-link": {"font-size": "16px", "text-align": "left", "margin":"0px", "--hover-color": "#eee"},"nav-link-selected": {"background-color": "rgba(255, 255, 255, 0.2)"},})
    if selected == "Home":
        Home()
    elif selected == "Data Management":
        data_management_menu = option_menu(menu_title="Data Management",options=["Editor", "File Manager","Anil Maheswari EBITDA Data Processor"],icons=["pencil-square", "folder"],orientation="horizontal",)
        if data_management_menu == "Editor":
            excel_editor_and_analyzer()
        elif data_management_menu == "File Manager":
            folder_menu()
        elif data_management_menu =="Anil Maheswari EBITDA Data Processor":
            geo()
    elif selected == "Price Analysis":
        analysis_menu1 = option_menu(menu_title ="Price Analysis",options=["WSP Analysis","Discount Analysis","Price Trend","Price Input"],icons=["clipboard-data","cash","shuffle","globe"],orientation="horizontal",)
        if analysis_menu1 == "WSP Analysis":
            wsp_analysis_dashboard()
        elif analysis_menu1 == "Discount Analysis":
              discount()
        elif analysis_menu1 == "Price Trend":
              price()
        elif analysis_menu1 == "Price Input":
             price_input()
    elif selected == "Sales Volume Analysis":
        analysis_menu = option_menu(menu_title="Volume Analysis Dashboards",options=["Sales Dashboard","Sales Review Report","Market Share Analysis","Product-Mix", "Segment-Mix","Geo-Mix"],icons=["clipboard-data", "cash","bar-chart", "arrow-up-right", "shuffle", "globe"],orientation="horizontal",)
        if analysis_menu == "Sales Dashboard":
            sales_dashboard()
        elif analysis_menu == "Sales Review Report":
            sales_review_report_generator()
        elif analysis_menu == "Product-Mix":
            normal()
        elif analysis_menu == "Segment-Mix":
            trade()
        elif analysis_menu == "Market Share Analysis":
            market_share()
        elif analysis_menu == "Geo-Mix":
            green()
    elif selected == "Predictions":
        prediction_menu = option_menu(menu_title="Predictions",options=["WSP Projection","Sales Projection(Old Model)","Sales Projection(New Model)"],icons=["bar-chart", "graph-up-arrow"],orientation="horizontal",)
        if prediction_menu == "WSP Projection":
            descriptive_statistics_and_prediction()
        elif prediction_menu == "Sales Projection(Old Model)":
            projection()
        elif prediction_menu == "Sales Projection(New Model)":
            pro()
    elif selected == "Settings":
        st.title("Settings")
        st.markdown('<div class="settings-container">', unsafe_allow_html=True)
        st.subheader("User Settings")
        username = st.text_input("Username", value=st.session_state.username)
        email = st.text_input("Email", value="johndoe@example.com")
        if st.button("Update Profile"):
            st.session_state.username = username
            st.success("Profile updated successfully!")
        st.subheader("Appearance")
        theme = st.selectbox("Theme", ["Light", "Dark", "System Default"])
        chart_color = st.color_picker("Default Chart Color", "#2e7bcf")
        st.subheader("Notifications")
        email_notifications = st.checkbox("Receive Email Notifications", value=True)
        notification_frequency = st.select_slider("Notification Frequency", options=["Daily", "Weekly", "Monthly"])
        if st.button("Save Settings"):
            st.success("Settings saved successfully!")
        st.markdown('</div>', unsafe_allow_html=True)
    st.sidebar.markdown("---")
    st.sidebar.subheader("📢 Feedback")
    feedback = st.sidebar.text_area("Share your thoughts:")
    if st.sidebar.button("Submit Feedback", key="submit_feedback"):
        st.sidebar.success("Thank you for your valuable feedback!")
    total_visits, daily_visits = update_visit_count()
    st.sidebar.markdown(f"""<div class="visit-counter"><h3>📊 Visit Statistics</h3><p>Total Visits: <span class="count">{total_visits}</span></p><p>Visits Today: <span class="count">{daily_visits}</span></p></div><script>const countElements = document.querySelectorAll('.count');countElements.forEach(element => {{const target = parseInt(element.innerText);let count = 0;const timer = setInterval(() => {{element.innerText = count;if (count === target) {{clearInterval(timer);}}count++;}}, 20);}});</script>""", unsafe_allow_html=True)
if __name__ == "__main__":
    main()
